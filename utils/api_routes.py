"""
API routes for the reporting system
"""
import asyncio
import json
from datetime import datetime
import os
import httpx
from fastapi import APIRouter, Query, HTTPException, Request, UploadFile, File, Form
from fastapi.responses import Response, FileResponse, StreamingResponse
from typing import Optional

from services import APIService, PDFService, ExcelService, DatabaseService
from services.bank_check_service import BankCheckService
from services.enhanced_bank_check_service import EnhancedBankCheckService
try:
    from services.dashboard_activity_service import DashboardActivityService
except Exception:
    DashboardActivityService = None  # type: ignore
from export_utils import get_default_header_config
from models import ExportRequest, ExportResponse

# Initialize services
api_service = APIService()
pdf_service = PDFService()
excel_service = ExcelService()
db_service = DatabaseService() if DatabaseService else None
dashboard_activity_service = DashboardActivityService() if DashboardActivityService else None
bank_check_service = BankCheckService()
enhanced_bank_check_service = EnhancedBankCheckService()

# Create router
router = APIRouter()
@router.get("/api/grc/incidents")
async def get_incidents_dashboard(
    startDate: str = Query(None),
    endDate: str = Query(None)
):
    """Return incidents dashboard data including statusOverview (details list)"""
    try:
        data = await api_service.get_incidents_data(startDate, endDate)
        # Ensure aggregates exist; if missing, fill from SQL
        if not db_service:
            raise HTTPException(status_code=503, detail="Database service unavailable")
        if not data.get('categoryDistribution'):
            data['categoryDistribution'] = await db_service.get_incidents_by_category(startDate, endDate)
        if not data.get('incidentsByCategory') and data.get('categoryDistribution'):
            data['incidentsByCategory'] = data['categoryDistribution']
        if not data.get('statusDistribution'):
            data['statusDistribution'] = await db_service.get_incidents_by_status(startDate, endDate)
        if not data.get('incidentsByStatus') and data.get('statusDistribution'):
            data['incidentsByStatus'] = data['statusDistribution']
        if not data.get('monthlyTrend'):
            data['monthlyTrend'] = await db_service.get_incidents_monthly_trend(startDate, endDate)

        # Always include statusOverview list for table (like controls)
        status_rows = await db_service.get_incidents_list(startDate, endDate)
        data['statusOverview'] = status_rows
        data['overallStatuses'] = status_rows

        # Derive totals if missing
        if 'totalIncidents' not in data:
            data['totalIncidents'] = len(status_rows)

        return data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load incidents dashboard: {str(e)}")
# Incidents list (JSON) for frontend table fallback
@router.get("/api/grc/incidents/list")
async def incidents_list(
    startDate: str = Query(None),
    endDate: str = Query(None),
    page: int = Query(1),
    limit: int = Query(1000)
):
    try:
        rows = await db_service.get_incidents_list(startDate, endDate)
        # Simple pagination (frontend expects large page anyway)
        total = len(rows)
        start = max(0, (page - 1) * limit)
        end = start + limit
        return {
            "data": rows[start:end],
            "pagination": {
                "page": page,
                "limit": limit,
                "total": total,
                "totalPages": (total + limit - 1) // limit,
                "hasNext": end < total,
                "hasPrev": start > 0
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load incidents list: {str(e)}")
# Incidents: PDF export
@router.get("/api/grc/incidents/export-pdf")
async def export_incidents_pdf(
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    cardType: str = Query(None),
    onlyCard: str = Query("False"),
    onlyChart: str = Query("False"),
    chartType: str = Query(None),
    onlyOverallTable: str = Query("False")
):
    try:
        # Convert string parameters to boolean
        onlyCard = onlyCard.lower() in ['true', '1', 'yes']
        onlyChart = onlyChart.lower() in ['true', '1', 'yes']
        onlyOverallTable = onlyOverallTable.lower() in ['true', '1', 'yes']
        
        header_config = {}
        if headerConfig:
            try:
                header_config = json.loads(headerConfig)
            except json.JSONDecodeError:
                header_config = {}
        default_config = get_default_header_config("incidents")
        header_config = {**default_config, **header_config}

        # Normalize chart/table params to cardType
        if onlyChart and chartType in ['byCategory', 'byStatus', 'monthlyTrend', 'netLossAndRecovery', 'topFinancialImpacts']:
            cardType = chartType
            onlyCard = True
        if onlyOverallTable:
            cardType = 'overallStatuses'
            onlyCard = True

        incidents_data = await api_service.get_incidents_data(startDate, endDate)

        # Card-specific fallbacks
        if onlyCard and cardType:
            # totals/list
            if cardType in ['totalIncidents']:
                card_rows = await api_service.get_incidents_card_data('totalIncidents', startDate, endDate)
                if not card_rows:
                    card_rows = await db_service.get_incidents_list(startDate, endDate)
                # reuse risks naming: allRisks style -> use 'list'
                incidents_data['list'] = card_rows
            elif cardType == 'overallStatuses':
                if not incidents_data.get('list'):
                    incidents_data['list'] = await db_service.get_incidents_list(startDate, endDate)
            elif cardType == 'byCategory':
                dist = incidents_data.get('categoryDistribution') or []
                if not dist:
                    incidents_data['categoryDistribution'] = await db_service.get_incidents_by_category(startDate, endDate)
            elif cardType == 'byStatus':
                dist = incidents_data.get('statusDistribution') or []
                if not dist:
                    incidents_data['statusDistribution'] = await db_service.get_incidents_by_status(startDate, endDate)
            elif cardType == 'monthlyTrend':
                trend = incidents_data.get('monthlyTrend') or []
                if not trend:
                    incidents_data['monthlyTrend'] = await db_service.get_incidents_monthly_trend(startDate, endDate)
            elif cardType == 'topFinancialImpacts':
                impacts = incidents_data.get('topFinancialImpacts') or []
                if not impacts:
                    incidents_data['topFinancialImpacts'] = await db_service.get_incidents_top_financial_impacts(startDate, endDate)
            elif cardType == 'netLossAndRecovery':
                net_loss = incidents_data.get('netLossAndRecovery') or []
                if not net_loss:
                    incidents_data['netLossAndRecovery'] = await db_service.get_incidents_net_loss_recovery(startDate, endDate)

        from services import PDFService
        pdf_service_local = PDFService()
        
        # Debug logging
        print(f"DEBUG: onlyCard={onlyCard} (type: {type(onlyCard)})")
        print(f"DEBUG: cardType={cardType} (type: {type(cardType)})")
        print(f"DEBUG: onlyCard and cardType = {onlyCard and cardType}")
        
        pdf_bytes = await pdf_service_local.generate_incidents_pdf(
            incidents_data=incidents_data, 
            start_date=startDate, 
            end_date=endDate, 
            header_config=header_config, 
            card_type=cardType, 
            only_card=onlyCard
        )

        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"incidents_{cardType+'_' if onlyCard and cardType else ''}{ts}.pdf"
        return Response(content=pdf_bytes, media_type='application/pdf', headers={'Content-Disposition': f'attachment; filename="{filename}"'})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# Incidents: Excel export
@router.get("/api/grc/incidents/export-excel")
async def export_incidents_excel(
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    cardType: str = Query(None),
    onlyCard: str = Query("False"),
    onlyChart: str = Query("False"),
    chartType: str = Query(None),
    onlyOverallTable: str = Query("False")
):
    try:
        # Convert string parameters to boolean
        onlyCard = onlyCard.lower() in ['true', '1', 'yes']
        onlyChart = onlyChart.lower() in ['true', '1', 'yes']
        onlyOverallTable = onlyOverallTable.lower() in ['true', '1', 'yes']
        
        header_config = {}
        if headerConfig:
            try:
                header_config = json.loads(headerConfig)
            except json.JSONDecodeError:
                header_config = {}
        default_config = get_default_header_config("incidents")
        header_config = {**default_config, **header_config}

        # Normalize chart/table params to cardType
        if onlyChart and chartType in ['byCategory', 'byStatus', 'monthlyTrend', 'netLossAndRecovery', 'topFinancialImpacts']:
            cardType = chartType
            onlyCard = True
        if onlyOverallTable:
            cardType = 'overallStatuses'
            onlyCard = True

        incidents_data = await api_service.get_incidents_data(startDate, endDate)

        # Card-specific fallbacks
        if onlyCard and cardType:
            if cardType in ['totalIncidents']:
                card_rows = await api_service.get_incidents_card_data('totalIncidents', startDate, endDate)
                if not card_rows:
                    card_rows = await db_service.get_incidents_list(startDate, endDate)
                incidents_data['list'] = card_rows
            elif cardType == 'overallStatuses':
                if not incidents_data.get('list'):
                    incidents_data['list'] = await db_service.get_incidents_list(startDate, endDate)
            elif cardType == 'byCategory':
                dist = incidents_data.get('categoryDistribution') or []
                if not dist:
                    incidents_data['categoryDistribution'] = await db_service.get_incidents_by_category(startDate, endDate)
            elif cardType == 'byStatus':
                dist = incidents_data.get('statusDistribution') or []
                if not dist:
                    incidents_data['statusDistribution'] = await db_service.get_incidents_by_status(startDate, endDate)
            elif cardType == 'monthlyTrend':
                trend = incidents_data.get('monthlyTrend') or []
                if not trend:
                    incidents_data['monthlyTrend'] = await db_service.get_incidents_monthly_trend(startDate, endDate)
            elif cardType == 'topFinancialImpacts':
                impacts = incidents_data.get('topFinancialImpacts') or []
                if not impacts:
                    incidents_data['topFinancialImpacts'] = await db_service.get_incidents_top_financial_impacts(startDate, endDate)
            elif cardType == 'netLossAndRecovery':
                net_loss = incidents_data.get('netLossAndRecovery') or []
                if not net_loss:
                    incidents_data['netLossAndRecovery'] = await db_service.get_incidents_net_loss_recovery(startDate, endDate)

        from services import ExcelService
        excel_service_local = ExcelService()
        excel_bytes = await excel_service_local.generate_incidents_excel(
            incidents_data=incidents_data, 
            start_date=startDate, 
            end_date=endDate, 
            header_config=header_config, 
            card_type=cardType, 
            only_card=onlyCard
        )

        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"incidents_{cardType+'_' if onlyCard and cardType else ''}{ts}.xlsx"
        return Response(content=excel_bytes, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={'Content-Disposition': f'attachment; filename="{filename}"'})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.get("/api/grc/risks/export-pdf")
async def export_risks_pdf(
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    cardType: str = Query(None),
    onlyCard: bool = Query(False),
    onlyChart: bool = Query(False),
    chartType: str = Query(None),
    onlyOverallTable: bool = Query(False)
):
    """Export risks report in PDF format"""
    try:
        # Parse header config
        header_config = {}
        if headerConfig:
            try:
                header_config = json.loads(headerConfig)
            except json.JSONDecodeError as e:
                header_config = {}
        
        # Get default header config for risks
        default_config = get_default_header_config("risks")
        header_config = {**default_config, **header_config}
        
        # Normalize short card names to canonical keys
        if cardType in ['low', 'medium', 'high']:
            cardType = {'low': 'lowRisk', 'medium': 'mediumRisk', 'high': 'highRisk'}[cardType]
        
        # Get risks data
        risks_data = await api_service.get_risks_data(startDate, endDate)
        
        # Resolve a usable "all risks" list from multiple possible keys
        all_risks: list = []
        for key in ['allRisks', 'risks', 'list', 'items', 'data']:
            value = risks_data.get(key)
            if isinstance(value, list) and (len(value) == 0 or isinstance(value[0], dict)):
                all_risks = value
                break
        
        # DB fallback if Node data missing
        if not all_risks and db_service:
            db_rows = await db_service.get_risks_data(startDate, endDate)
            all_risks = [r.to_dict() for r in db_rows]
            risks_data['allRisks'] = all_risks
        
        # For card-specific exports, filter data from resolved list
        if onlyCard and cardType:
            if cardType in ['highRisk', 'mediumRisk', 'lowRisk']:
                def normalize_level(val):
                    if val is None:
                        return None
                    s = str(val).strip().lower()
                    # numeric mapping: 1–3=low, 4=medium, 5+=high
                    if s.isdigit():
                        n = int(s)
                        if n >= 5:
                            return 'high'
                        if n >= 4:
                            return 'medium'
                        return 'low'
                    if 'high' in s:
                        return 'high'
                    if 'medium' in s:
                        return 'medium'
                    if 'low' in s:
                        return 'low'
                    return None

                filtered_risks = []
                for risk in all_risks:
                    level = normalize_level(
                        risk.get('inherent_value')
                        or risk.get('inherent')
                        or risk.get('inherentLevel')
                        or risk.get('inherent_rating')
                    )
                    if (cardType == 'highRisk' and level == 'high') or \
                       (cardType == 'mediumRisk' and level == 'medium') or \
                       (cardType == 'lowRisk' and level == 'low'):
                        filtered_risks.append(risk)
                
                # If still empty, try Node card endpoint
                if not filtered_risks:
                    fetched = await api_service.get_risks_card_data(cardType, startDate, endDate)
                    if isinstance(fetched, list) and fetched:
                        filtered_risks = fetched
                
                risks_data[cardType] = filtered_risks
        
        # Generate PDF
        pdf_content = await pdf_service.generate_risks_pdf(
            risks_data, startDate, endDate, header_config, cardType, onlyCard
        )
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if onlyCard and cardType:
            filename = f"risks_{cardType}_{timestamp}.pdf"
        else:
            filename = f"risks_report_{timestamp}.pdf"
        
        return Response(
            content=pdf_content,
            media_type='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.get("/api/grc/risks/export-excel")
async def export_risks_excel(
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    cardType: str = Query(None),
    onlyCard: bool = Query(False),
    onlyChart: bool = Query(False),
    chartType: str = Query(None),
    onlyOverallTable: bool = Query(False)
):
    """Export risks report in Excel format"""
    try:
        # Parse header config
        header_config = {}
        if headerConfig:
            try:
                header_config = json.loads(headerConfig)
            except json.JSONDecodeError as e:
                header_config = {}
        
        # Get default header config for risks
        default_config = get_default_header_config("risks")
        header_config = {**default_config, **header_config}

        # Normalize short card names to canonical keys
        if cardType in ['low', 'medium', 'high']:
            cardType = {'low': 'lowRisk', 'medium': 'mediumRisk', 'high': 'highRisk'}[cardType]
        
        # Get risks data
        risks_data = await api_service.get_risks_data(startDate, endDate)
        
        # Resolve a usable "all risks" list from multiple possible keys
        all_risks: list = []
        for key in ['allRisks', 'risks', 'list', 'items', 'data']:
            value = risks_data.get(key)
            if isinstance(value, list) and (len(value) == 0 or isinstance(value[0], dict)):
                all_risks = value
                break
        
        # DB fallback if Node data missing
        if not all_risks and db_service:
            db_rows = await db_service.get_risks_data(startDate, endDate)
            all_risks = [r.to_dict() for r in db_rows]
            risks_data['allRisks'] = all_risks
        
        # For card-specific exports, filter data from resolved list
        if onlyCard and cardType:
            if cardType in ['highRisk', 'mediumRisk', 'lowRisk']:
                def normalize_level(val):
                    if val is None:
                        return None
                    s = str(val).strip().lower()
                    if s.isdigit():
                        n = int(s)
                        if n >= 5:
                            return 'high'
                        if n >= 4:
                            return 'medium'
                        return 'low'
                    if 'high' in s:
                        return 'high'
                    if 'medium' in s:
                        return 'medium'
                    if 'low' in s:
                        return 'low'
                    return None

                filtered_risks = []
                for risk in all_risks:
                    level = normalize_level(
                        risk.get('inherent_value')
                        or risk.get('inherent')
                        or risk.get('inherentLevel')
                        or risk.get('inherent_rating')
                    )
                    if (cardType == 'highRisk' and level == 'high') or \
                       (cardType == 'mediumRisk' and level == 'medium') or \
                       (cardType == 'lowRisk' and level == 'low'):
                        filtered_risks.append(risk)

                # If still empty, try Node card endpoint
                if not filtered_risks:
                    fetched = await api_service.get_risks_card_data(cardType, startDate, endDate)
                    if isinstance(fetched, list) and fetched:
                        filtered_risks = fetched

                risks_data[cardType] = filtered_risks
        
        # Generate Excel
        excel_content = await excel_service.generate_risks_excel(
            risks_data, startDate, endDate, header_config, cardType, onlyCard
        )
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if onlyCard and cardType:
            filename = f"risks_{cardType}_{timestamp}.xlsx"
        else:
            filename = f"risks_report_{timestamp}.xlsx"
        
        return Response(
            content=excel_content,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.get("/api/grc/controls/export-pdf")
async def export_controls_pdf(
    request: Request,
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    cardType: str = Query(None),
    onlyCard: str = Query("False"),
    onlyChart: str = Query("False"),
    chartType: str = Query(None),
    onlyOverallTable: str = Query("False")
):
    """Export controls report in PDF format"""
    print("=== PDF EXPORT CALLED ===")
    print(f"onlyChart={onlyChart}, chartType={chartType}")
    print("=== TESTING CHART EXPORT ===")
    try:
        # Parse header config
        header_config = {}
        if headerConfig:
            try:
                header_config = json.loads(headerConfig)
            except json.JSONDecodeError as e:
                header_config = {}
        
        # Get default header config for controls
        default_config = get_default_header_config("controls")
        header_config = {**default_config, **header_config}
        
        # Get table type from query params
        table_type = request.query_params.get('tableType', 'overallStatuses')
        
        # Convert string parameters to boolean
        onlyCard = onlyCard.lower() in ['true', '1', 'yes']
        onlyChart = onlyChart.lower() in ['true', '1', 'yes']
        onlyOverallTable = onlyOverallTable.lower() in ['true', '1', 'yes']
        
        # Respect explicit cardType from client; only derive when not provided
        if not cardType:
            if onlyChart and chartType in ['department', 'risk', 'quarterlyControlCreationTrend', 'controlsByType', 'antiFraudDistribution', 'controlsPerLevel', 'controlExecutionFrequency', 'numberOfControlsByIcofrStatus', 'numberOfFocusPointsPerPrinciple', 'numberOfFocusPointsPerComponent', 'numberOfControlsPerComponent', 'actionPlansStatus']:
                cardType = chartType
                onlyCard = True
                onlyChart = True
            elif onlyOverallTable:
                if table_type == 'controlsTestingApprovalCycle':
                    cardType = 'controlsTestingApprovalCycle'
                elif table_type == 'keyNonKeyControlsPerDepartment':
                    cardType = 'keyNonKeyControlsPerDepartment'
                elif table_type == 'keyNonKeyControlsPerProcess':
                    cardType = 'keyNonKeyControlsPerProcess'
                elif table_type == 'keyNonKeyControlsPerBusinessUnit':
                    cardType = 'keyNonKeyControlsPerBusinessUnit'
                elif table_type == 'controlCountByAssertionName':
                    cardType = 'controlCountByAssertionName'
                elif table_type == 'icofrControlCoverageByCoso':
                    cardType = 'icofrControlCoverageByCoso'
                elif table_type == 'actionPlanForAdequacy':
                    cardType = 'actionPlanForAdequacy'
                elif table_type == 'actionPlanForEffectiveness':
                    cardType = 'actionPlanForEffectiveness'
                elif table_type == 'controlSubmissionStatusByQuarterFunction':
                    cardType = 'controlSubmissionStatusByQuarterFunction'
                elif table_type == 'functionsWithFullyTestedControlTests':
                    cardType = 'functionsWithFullyTestedControlTests'
                elif table_type == 'controlsNotMappedToAssertions':
                    cardType = 'controlsNotMappedToAssertions'
                elif table_type == 'controlsNotMappedToPrinciples':
                    cardType = 'controlsNotMappedToPrinciples'
                else:
                    cardType = 'overallStatuses'
                    onlyCard = True
            elif table_type == 'controlCountByAssertionName':
                cardType = 'controlCountByAssertionName'
            elif table_type == 'icofrControlCoverageByCoso':
                cardType = 'icofrControlCoverageByCoso'
            elif table_type == 'actionPlanForAdequacy':
                cardType = 'actionPlanForAdequacy'
            elif table_type == 'actionPlanForEffectiveness':
                cardType = 'actionPlanForEffectiveness'
            elif table_type == 'functionsWithFullyTestedControlTests':
                cardType = 'functionsWithFullyTestedControlTests'
            elif table_type == 'functionsWithFullySubmittedControlTests':
                cardType = 'functionsWithFullySubmittedControlTests'
            else:
                cardType = 'overallStatuses'
                onlyCard = True

        # Short-circuit: skip heavy dashboard fetch for card-only exports
        if onlyCard and cardType:
            controls_data = {}
        else:
            try:
                print("DEBUG: About to call get_controls_data")
                controls_data = await asyncio.wait_for(
                    api_service.get_controls_data(startDate, endDate), 
                    timeout=45.0
                )
                print(f"DEBUG: Controls data keys from API: {list(controls_data.keys())}")
                print(f"DEBUG: Controls data type: {type(controls_data)}")
                # Check if any values are Response objects
                for key, value in controls_data.items():
                    if hasattr(value, 'status'):
                        print(f"DEBUG: WARNING - {key} is a Response object: {type(value)}")
                    elif isinstance(value, list) and len(value) > 0:
                        print(f"DEBUG: {key} is a list with {len(value)} items")
                        if hasattr(value[0], 'status'):
                            print(f"DEBUG: WARNING - {key}[0] is a Response object: {type(value[0])}")
            except asyncio.TimeoutError:
                print("DEBUG: Controls data API timeout - using empty data")
                controls_data = {}
            except Exception as e:
                print(f"DEBUG: Controls data API error: {e}")
                controls_data = {}
        # Ensure specific data exists for card-only routes (and add SQL fallbacks)
        if onlyCard and cardType:
            if cardType in ['totalControls', 'unmappedControls', 'pendingPreparer', 'pendingChecker', 'pendingReviewer', 'pendingAcceptance', 'testsPendingPreparer', 'testsPendingChecker', 'testsPendingReviewer', 'testsPendingAcceptance', 'unmappedIcofrControls', 'unmappedNonIcofrControls']:
                card_data = await api_service.get_controls_card_data(cardType, startDate, endDate)
                if not card_data:
                    if cardType == 'unmappedControls':
                        card_data = await db_service.get_unmapped_controls(startDate, endDate)
                    elif cardType == 'pendingPreparer':
                        card_data = await db_service.get_pending_controls('preparer', startDate, endDate)
                    elif cardType == 'pendingChecker':
                        card_data = await db_service.get_pending_controls('checker', startDate, endDate)
                    elif cardType == 'pendingReviewer':
                        card_data = await db_service.get_pending_controls('reviewer', startDate, endDate)
                    elif cardType == 'pendingAcceptance':
                        card_data = await db_service.get_pending_controls('acceptance', startDate, endDate)
                    elif cardType == 'testsPendingPreparer':
                        card_data = await db_service.get_tests_pending_controls('preparer', startDate, endDate)
                    elif cardType == 'testsPendingChecker':
                        card_data = await db_service.get_tests_pending_controls('checker', startDate, endDate)
                    elif cardType == 'testsPendingReviewer':
                        card_data = await db_service.get_tests_pending_controls('reviewer', startDate, endDate)
                    elif cardType == 'testsPendingAcceptance':
                        card_data = await db_service.get_tests_pending_controls('acceptance', startDate, endDate)
                    elif cardType == 'unmappedIcofrControls':
                        card_data = await db_service.get_unmapped_icofr_controls(startDate, endDate)
                    elif cardType == 'unmappedNonIcofrControls':
                        card_data = await db_service.get_unmapped_non_icofr_controls(startDate, endDate)
                    elif cardType == 'testsPendingPreparer':
                        card_data = await db_service.get_tests_pending_controls('preparer', startDate, endDate)
                    elif cardType == 'testsPendingChecker':
                        card_data = await db_service.get_tests_pending_controls('checker', startDate, endDate)
                    elif cardType == 'testsPendingReviewer':
                        card_data = await db_service.get_tests_pending_controls('reviewer', startDate, endDate)
                    elif cardType == 'testsPendingAcceptance':
                        card_data = await db_service.get_tests_pending_controls('acceptance', startDate, endDate)
                    else:
                        card_data = []
                controls_data[f'{cardType}'] = card_data
            elif cardType == 'overallStatuses':
                # Prefer API aggregate; else DB fallback
                if not controls_data.get('overallStatuses'):
                    rows = await db_service.execute_query(
                        """
                        SELECT 
                            c.code as code,
                            c.name as name,
                            c.preparerStatus,
                            c.checkerStatus,
                            c.reviewerStatus,
                            c.acceptanceStatus
                        FROM dbo.[Controls] c
                        WHERE c.isDeleted = 0
                        ORDER BY c.createdAt DESC
                        """
                    )
                    controls_data['overallStatuses'] = rows

    # Generate PDF
        try:
            print(f"DEBUG: About to call generate_controls_pdf with cardType={cardType}, onlyOverallTable={onlyOverallTable}")
            pdf_content = await pdf_service.generate_controls_pdf(
                controls_data, startDate, endDate, header_config, cardType, onlyCard, onlyOverallTable, onlyChart
            )
            print(f"DEBUG: PDF generated successfully, type={type(pdf_content)}")
        except Exception as pdf_error:
            print(f"DEBUG: Error in generate_controls_pdf: {pdf_error}")
            import traceback
            traceback.print_exc()
            raise
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        print(f"DEBUG: Filename generation - onlyCard={onlyCard}, onlyChart={onlyChart}, onlyOverallTable={onlyOverallTable}")
        print(f"DEBUG: Filename generation - cardType={cardType}, chartType={chartType}, table_type={table_type}")
        if onlyCard and cardType:
            filename = f"controls_{cardType}_{timestamp}.pdf"
            print(f"DEBUG: Using card filename: {filename}")
        elif onlyOverallTable and table_type:
            filename = f"controls_{table_type}_{timestamp}.pdf"
            print(f"DEBUG: Using table filename: {filename}")
        elif onlyChart and chartType:
            filename = f"controls_{chartType}_{timestamp}.pdf"
            print(f"DEBUG: Using chart filename: {filename}")
        else:
            filename = f"controls_report_{timestamp}.pdf"
            print(f"DEBUG: Using default filename: {filename}")
        
        return Response(
            content=pdf_content,
            media_type='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.get("/api/grc/controls/tests/pending-acceptance")
async def get_tests_pending_acceptance(
    startDate: str = Query(None, description="Start date for filtering"),
    endDate: str = Query(None, description="End date for filtering")
):
    """Get tests pending acceptance data for modal display"""
    try:
        if not db_service:
            raise HTTPException(status_code=503, detail="Database service unavailable")
        
        # Get the data using the same method as the card
        data = await db_service.get_tests_pending_controls('acceptance', startDate, endDate)
        
        # Transform the data to match frontend expectations
        transformed_data = []
        for item in data:
            transformed_data.append({
                "id": item.get("id"),
                "control_code": item.get("code"),
                "control_name": item.get("control_name"),
                "preparerStatus": item.get("status"),
                "function_name": item.get("function_name")
            })
        
        return {
            "success": True,
            "data": transformed_data,
            "count": len(transformed_data)
        }
    except Exception as e:
        print(f"Error in get_tests_pending_acceptance: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/api/grc/controls/export-excel")
async def export_controls_excel(
    request: Request,
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    cardType: str = Query(None),
    onlyCard: str = Query("False"),
    onlyChart: str = Query("False"),
    chartType: str = Query(None),
    onlyOverallTable: str = Query("False")
):
    """Export controls report in Excel format"""
    try:
        # Parse header config
        header_config = {}
        if headerConfig:
            try:
                header_config = json.loads(headerConfig)
            except json.JSONDecodeError as e:
                header_config = {}
        
        # Get default header config for controls
        default_config = get_default_header_config("controls")
        header_config = {**default_config, **header_config}
        
        # Get table type from query params
        table_type = request.query_params.get('tableType', 'overallStatuses')
        
        # Convert string parameters to boolean
        onlyCard = onlyCard.lower() in ['true', '1', 'yes']
        onlyChart = onlyChart.lower() in ['true', '1', 'yes']
        onlyOverallTable = onlyOverallTable.lower() in ['true', '1', 'yes']
        
        # Normalize table/chart params → card params for backend handling
        # This lets URLs with onlyChart=true&chartType=... or onlyOverallTable work
        
        # First check if we have a specific cardType from the URL parameter
        if onlyCard and cardType:
            # Keep the cardType as provided, don't override it
            pass
        elif onlyChart and chartType in ['department', 'risk', 'quarterlyControlCreationTrend', 'controlsByType', 'antiFraudDistribution', 'controlsPerLevel', 'controlExecutionFrequency', 'numberOfControlsByIcofrStatus', 'numberOfFocusPointsPerPrinciple', 'numberOfFocusPointsPerComponent', 'numberOfControlsPerComponent', 'actionPlansStatus']:
            cardType = chartType
            onlyCard = True
        elif onlyOverallTable:
            # Check if specific table type is requested
            if table_type == 'controlsTestingApprovalCycle':
                cardType = 'controlsTestingApprovalCycle'
            elif table_type == 'keyNonKeyControlsPerDepartment':
                cardType = 'keyNonKeyControlsPerDepartment'
            elif table_type == 'keyNonKeyControlsPerProcess':
                cardType = 'keyNonKeyControlsPerProcess'
            elif table_type == 'keyNonKeyControlsPerBusinessUnit':
                cardType = 'keyNonKeyControlsPerBusinessUnit'
            elif table_type == 'controlCountByAssertionName':
                cardType = 'controlCountByAssertionName'
            elif table_type == 'icofrControlCoverageByCoso':
                cardType = 'icofrControlCoverageByCoso'
            elif table_type == 'actionPlanForAdequacy':
                cardType = 'actionPlanForAdequacy'
            elif table_type == 'actionPlanForEffectiveness':
                cardType = 'actionPlanForEffectiveness'
            elif table_type == 'controlSubmissionStatusByQuarterFunction':
                cardType = 'controlSubmissionStatusByQuarterFunction'
            elif table_type == 'functionsWithFullyTestedControlTests':
                cardType = 'functionsWithFullyTestedControlTests'
            elif table_type == 'controlsNotMappedToAssertions':
                cardType = 'controlsNotMappedToAssertions'
            elif table_type == 'controlsNotMappedToPrinciples':
                cardType = 'controlsNotMappedToPrinciples'
            else:
                cardType = 'overallStatuses'
                onlyCard = True
        elif table_type == 'controlCountByAssertionName':
            cardType = 'controlCountByAssertionName'
        elif table_type == 'icofrControlCoverageByCoso':
            cardType = 'icofrControlCoverageByCoso'
        elif table_type == 'actionPlanForAdequacy':
            cardType = 'actionPlanForAdequacy'
        elif table_type == 'actionPlanForEffectiveness':
            cardType = 'actionPlanForEffectiveness'
        elif table_type == 'functionsWithFullyTestedControlTests':
            cardType = 'functionsWithFullyTestedControlTests'
        elif table_type == 'functionsWithFullySubmittedControlTests':
            cardType = 'functionsWithFullySubmittedControlTests'
        elif table_type == 'controlsByFunction':
            cardType = 'controlsByFunction'
        else:
            # Only set default if no cardType was provided
            if not cardType:
                cardType = 'overallStatuses'
                onlyCard = True

        # Get controls data with timeout handling
        try:
            controls_data = await asyncio.wait_for(
                api_service.get_controls_data(startDate, endDate), 
                timeout=10.0  # Reduced timeout
            )
        except asyncio.TimeoutError:
            print("DEBUG: Controls data API timeout - using database fallback")
            # Use database fallback instead of empty data
            try:
                controls_data = await db_service.get_controls_dashboard_data(startDate, endDate)
            except Exception as db_error:
                print(f"DEBUG: Database fallback error: {db_error}")
            controls_data = {}
        except Exception as e:
            print(f"DEBUG: Controls data API error: {e} - using database fallback")
            try:
                controls_data = await db_service.get_controls_dashboard_data(startDate, endDate)
            except Exception as db_error:
                print(f"DEBUG: Database fallback error: {db_error}")
            controls_data = {}
        
        # Construct date filter for SQL queries
        date_filter = ""
        if startDate and endDate:
            try:
                # Parse dates and create SQL date filter
                start_date = datetime.strptime(startDate, '%Y-%m-%d').strftime('%Y-%m-%d')
                end_date = datetime.strptime(endDate, '%Y-%m-%d').strftime('%Y-%m-%d')
                date_filter = f"AND c.createdAt >= '{start_date}' AND c.createdAt <= '{end_date}'"
            except ValueError:
                print(f"DEBUG: Invalid date format - startDate: {startDate}, endDate: {endDate}")
                date_filter = ""
        
        # For card-specific exports, fetch specific card data (large page size)
        if onlyCard and cardType:
            # Handle chart types first
            if cardType in ['quarterlyControlCreationTrend', 'controlsByType', 'numberOfControlsPerComponent', 'department', 'risk', 'antiFraudDistribution', 'controlsPerLevel', 'controlExecutionFrequency', 'numberOfControlsByIcofrStatus', 'numberOfFocusPointsPerPrinciple', 'numberOfFocusPointsPerComponent', 'actionPlansStatus']:
                # Fetch chart data directly from database
                try:
                    if cardType == 'quarterlyControlCreationTrend':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CONCAT('Q', DATEPART(QUARTER, c.createdAt), ' ', YEAR(c.createdAt)) AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY YEAR(c.createdAt), DATEPART(QUARTER, c.createdAt)
                            ORDER BY YEAR(c.createdAt), DATEPART(QUARTER, c.createdAt)
                            """
                        )
                    elif cardType == 'controlsByType':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN c.type IS NULL OR c.type = '' THEN 'Not Specified'
                                    ELSE c.type
                                END AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY c.type
                            ORDER BY COUNT(c.id) DESC
                            """
                        )
                    elif cardType == 'numberOfControlsPerComponent':
                        print("DEBUG: Fetching numberOfControlsPerComponent data from database")
                        # Use the exact SQL query from Node.js frontend
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                cc.name AS name,
                                COUNT(DISTINCT c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            JOIN {db_service.get_fully_qualified_table_name('ControlCosos')} ccx ON c.id = ccx.control_id
                            JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} cp ON ccx.coso_id = cp.id
                            JOIN {db_service.get_fully_qualified_table_name('CosoPrinciples')} pr ON cp.principle_id = pr.id
                            JOIN {db_service.get_fully_qualified_table_name('CosoComponents')} cc ON pr.component_id = cc.id
                            WHERE c.isDeleted = 0 
                                AND ccx.deletedAt IS NULL 
                                AND cp.deletedAt IS NULL 
                                AND pr.deletedAt IS NULL 
                                AND cc.deletedAt IS NULL {date_filter}
                            GROUP BY cc.name
                            ORDER BY COUNT(DISTINCT c.id) DESC
                            """
                        )
                        print(f"DEBUG: numberOfControlsPerComponent query returned {len(chart_data) if chart_data else 0} rows")
                        if chart_data and len(chart_data) > 0:
                            print(f"DEBUG: Sample data: {chart_data[0]}")
                    elif cardType == 'department':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                f.name as name,
                                COUNT(c.id) as value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            JOIN {db_service.get_fully_qualified_table_name('ControlFunctions')} cf ON c.id = cf.control_id
                            JOIN {db_service.get_fully_qualified_table_name('Functions')} f ON cf.function_id = f.id
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY f.name
                            ORDER BY COUNT(c.id) DESC, f.name
                            """
                        )
                    elif cardType == 'risk':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                c.risk_response as name,
                                COUNT(c.id) as value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY c.risk_response
                            """
                        )
                    elif cardType == 'antiFraudDistribution':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN c.AntiFraud = 1 THEN 'Anti-Fraud'
                                    WHEN c.AntiFraud = 0 THEN 'Non-Anti-Fraud'
                                    ELSE 'Unknown'
                                END AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY c.AntiFraud
                            ORDER BY COUNT(c.id) DESC
                            """
                        )
                    elif cardType == 'controlsPerLevel':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN c.entityLevel IS NULL OR c.entityLevel = '' THEN 'Not Specified'
                                    ELSE c.entityLevel
                                END AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY c.entityLevel
                            ORDER BY COUNT(c.id) DESC
                            """
                        )
                    elif cardType == 'controlExecutionFrequency':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN c.frequency = 'Daily' THEN 'Daily'
                                    WHEN c.frequency = 'Event Base' THEN 'Event Base'
                                    WHEN c.frequency = 'Weekly' THEN 'Weekly'
                                    WHEN c.frequency = 'Monthly' THEN 'Monthly'
                                    WHEN c.frequency = 'Quarterly' THEN 'Quarterly'
                                    WHEN c.frequency = 'Semi Annually' THEN 'Semi Annually'
                                    WHEN c.frequency = 'Annually' THEN 'Annually'
                                    WHEN c.frequency IS NULL OR c.frequency = '' THEN 'Not Specified'
                                    ELSE c.frequency
                                END AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY c.frequency
                            ORDER BY COUNT(c.id) DESC
                            """
                        )
                    elif cardType == 'numberOfControlsByIcofrStatus':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN a.id IS NULL THEN 'Non-ICOFR'
                                    WHEN (a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1)
                                         AND (a.account_type IN ('Balance Sheet', 'Income Statement')) 
                                      THEN 'ICOFR' 
                                    ELSE 'Non-ICOFR' 
                                END AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            LEFT JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id AND a.isDeleted = 0
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY 
                                CASE 
                                    WHEN a.id IS NULL THEN 'Non-ICOFR'
                                    WHEN (a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1)
                                         AND (a.account_type IN ('Balance Sheet', 'Income Statement')) 
                                      THEN 'ICOFR' 
                                    ELSE 'Non-ICOFR' 
                                END
                            ORDER BY COUNT(c.id) DESC
                            """
                        )
                    elif cardType == 'numberOfFocusPointsPerPrinciple':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                prin.name AS name,
                                COUNT(point.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('CosoPrinciples')} prin
                            LEFT JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} point ON prin.id = point.principle_id
                            WHERE prin.deletedAt IS NULL {date_filter}
                            GROUP BY prin.name
                            ORDER BY COUNT(point.id) DESC, prin.name
                            """
                        )
                    elif cardType == 'numberOfFocusPointsPerComponent':
                        # Use the exact same SQL query as Node.js backend
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                comp.name AS name,
                                COUNT(point.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('CosoComponents')} comp
                            JOIN {db_service.get_fully_qualified_table_name('CosoPrinciples')} prin ON prin.component_id = comp.id
                            LEFT JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} point ON point.principle_id = prin.id
                            WHERE comp.deletedAt IS NULL AND prin.deletedAt IS NULL {date_filter}
                            GROUP BY comp.name
                            ORDER BY COUNT(point.id) DESC
                            """
                        )
                        # Store the data in controls_data for Excel service
                        controls_data[cardType] = chart_data if chart_data else []
                    elif cardType == 'numberOfControlsPerComponent':
                        # Use the exact SQL query provided by user
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                cc.name AS name,
                                COUNT(DISTINCT c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            JOIN {db_service.get_fully_qualified_table_name('ControlCosos')} ccx ON c.id = ccx.control_id
                            JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} cp ON ccx.coso_id = cp.id
                            JOIN {db_service.get_fully_qualified_table_name('CosoPrinciples')} pr ON cp.principle_id = pr.id
                            JOIN {db_service.get_fully_qualified_table_name('CosoComponents')} cc ON pr.component_id = cc.id
                            WHERE c.isDeleted = 0 
                                AND ccx.deletedAt IS NULL 
                                AND cp.deletedAt IS NULL 
                                AND pr.deletedAt IS NULL 
                                AND cc.deletedAt IS NULL
                            GROUP BY cc.name
                            ORDER BY COUNT(DISTINCT c.id) DESC
                            """
                        )
                        print(f"DEBUG: numberOfControlsPerComponent query returned {len(chart_data) if chart_data else 0} rows")
                        if chart_data and len(chart_data) > 0:
                            print(f"DEBUG: Sample data: {chart_data[0]}")
                            # Store the data in controls_data for Excel service
                            controls_data[cardType] = chart_data
                        else:
                            print("DEBUG: No data returned from numberOfControlsPerComponent query")
                            controls_data[cardType] = []
                    elif cardType == 'actionPlansStatus':
                        chart_data = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN a.done = 0 AND a.implementation_date < GETDATE() THEN 'Overdue'
                                    ELSE 'Not Overdue'
                                END AS name,
                                COUNT(a.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Actionplans')} a
                            WHERE a.deletedAt IS NULL {date_filter}
                            GROUP BY 
                                CASE 
                                    WHEN a.done = 0 AND a.implementation_date < GETDATE() THEN 'Overdue'
                                    ELSE 'Not Overdue'
                                END
                            ORDER BY COUNT(a.id) DESC
                            """
                        )
                    else:
                        # Fallback: create a simple chart with basic data
                        chart_data = [
                            {'name': 'Sample Data', 'value': 1},
                            {'name': 'Another Sample', 'value': 2}
                        ]
                    
                    controls_data[cardType] = chart_data
                    print(f"DEBUG: Fetched {len(chart_data)} rows for chart type {cardType}")
                    if chart_data:
                        print(f"DEBUG: Sample data: {chart_data[0] if chart_data else 'Empty'}")
                except Exception as e:
                    print(f"DEBUG: Error fetching chart data for {cardType}: {e}")
                    # Fallback data on error
                    controls_data[cardType] = [
                        {'name': 'Error - No Data', 'value': 0}
                    ]
            elif cardType in ['totalControls', 'unmappedControls', 'pendingPreparer', 'pendingChecker', 'pendingReviewer', 'pendingAcceptance', 'testsPendingPreparer', 'testsPendingChecker', 'testsPendingReviewer', 'testsPendingAcceptance', 'unmappedIcofrControls', 'unmappedNonIcofrControls']:
                # Try Node API first
                card_data = await api_service.get_controls_card_data(cardType, startDate, endDate)
                # If empty, fallback to direct SQL for reliability
                if not card_data:
                    if cardType == 'unmappedControls':
                        card_data = await db_service.get_unmapped_controls(startDate, endDate)
                    elif cardType == 'pendingPreparer':
                        card_data = await db_service.get_pending_controls('preparer', startDate, endDate)
                    elif cardType == 'pendingChecker':
                        card_data = await db_service.get_pending_controls('checker', startDate, endDate)
                    elif cardType == 'pendingReviewer':
                        card_data = await db_service.get_pending_controls('reviewer', startDate, endDate)
                    elif cardType == 'pendingAcceptance':
                        card_data = await db_service.get_pending_controls('acceptance', startDate, endDate)
                    elif cardType == 'unmappedIcofrControls':
                        card_data = await db_service.get_unmapped_icofr_controls(startDate, endDate)
                    elif cardType == 'unmappedNonIcofrControls':
                        card_data = await db_service.get_unmapped_non_icofr_controls(startDate, endDate)
                    elif cardType == 'testsPendingPreparer':
                        card_data = await db_service.get_tests_pending_controls('preparer', startDate, endDate)
                    elif cardType == 'testsPendingChecker':
                        card_data = await db_service.get_tests_pending_controls('checker', startDate, endDate)
                    elif cardType == 'testsPendingReviewer':
                        card_data = await db_service.get_tests_pending_controls('reviewer', startDate, endDate)
                    elif cardType == 'testsPendingAcceptance':
                        card_data = await db_service.get_tests_pending_controls('acceptance', startDate, endDate)
                    else:
                        card_data = []
                controls_data[f'{cardType}'] = card_data
            elif cardType == 'overallStatuses':
                # Fix: populate the expected 'overallStatuses' key (not 'statusOverview')
                if not controls_data.get('overallStatuses'):
                    rows = await db_service.execute_query(
                        """
                        SELECT 
                            c.code as code,
                            c.name as name,
                            c.preparerStatus,
                            c.checkerStatus,
                            c.reviewerStatus,
                            c.acceptanceStatus
                        FROM dbo.[Controls] c
                        WHERE c.isDeleted = 0
                        ORDER BY c.createdAt DESC
                        """
                    )
                    controls_data['overallStatuses'] = rows
            elif cardType == 'controlsTestingApprovalCycle':
                if not controls_data.get('controlsTestingApprovalCycle'):
                    print("DEBUG: Fetching controlsTestingApprovalCycle data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                c.name AS [Control Name],
                                c.createdAt AS [Created At],
                                c.id AS [Control ID],
                                c.code AS [Code],
                                t.preparerStatus AS [Preparer Status],
                                t.checkerStatus AS [Checker Status],
                                t.reviewerStatus AS [Reviewer Status],
                                t.acceptanceStatus AS [Acceptance Status],
                                f.name AS [Business Unit]
                            FROM {db_service.get_fully_qualified_table_name('ControlDesignTests')} AS t
                            INNER JOIN {db_service.get_fully_qualified_table_name('Controls')} AS c ON t.control_id = c.id
                            INNER JOIN {db_service.get_fully_qualified_table_name('Functions')} AS f ON t.function_id = f.id
                            WHERE c.isDeleted = 0 AND (t.deletedAt IS NULL) AND t.function_id IS NOT NULL
                            ORDER BY c.createdAt DESC, c.name
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlsTestingApprovalCycle")
                        if rows:
                            print(f"DEBUG: First row sample: {rows[0]}")
                        else:
                            print("DEBUG: No rows returned from query")
                        controls_data['controlsTestingApprovalCycle'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlsTestingApprovalCycle data: {e}")
                        controls_data['controlsTestingApprovalCycle'] = []
            elif cardType == 'quarterlyControlCreationTrend':
                print(f"DEBUG: Processing quarterlyControlCreationTrend, current data: {controls_data.get('quarterlyControlCreationTrend', 'NOT_FOUND')}")
                if not controls_data.get('quarterlyControlCreationTrend'):
                    print("DEBUG: Fetching quarterlyControlCreationTrend data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                CONCAT('Q', DATEPART(QUARTER, c.createdAt), ' ', YEAR(c.createdAt)) AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0
                            GROUP BY YEAR(c.createdAt), DATEPART(QUARTER, c.createdAt)
                            ORDER BY YEAR(c.createdAt), DATEPART(QUARTER, c.createdAt)
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for quarterlyControlCreationTrend")
                        controls_data['quarterlyControlCreationTrend'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching quarterlyControlCreationTrend data: {e}")
                        controls_data['quarterlyControlCreationTrend'] = []
            elif cardType == 'controlsByType':
                print(f"DEBUG: Processing controlsByType, current data: {controls_data.get('controlsByType', 'NOT_FOUND')}")
                if not controls_data.get('controlsByType'):
                    print("DEBUG: Fetching controlsByType data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN c.type IS NULL OR c.type = '' THEN 'Not Specified'
                                    ELSE c.type
                                END AS name,
                                COUNT(c.id) AS value
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0
                            GROUP BY c.type
                            ORDER BY COUNT(c.id) DESC
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlsByType")
                        controls_data['controlsByType'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlsByType data: {e}")
                        controls_data['controlsByType'] = []
            elif cardType == 'keyNonKeyControlsPerDepartment':
                if not controls_data.get('keyNonKeyControlsPerDepartment'):
                    print("DEBUG: Fetching keyNonKeyControlsPerDepartment data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                COALESCE(jt.name, 'Unassigned Department') AS [Department],
                                SUM(CASE WHEN c.keyControl = 1 THEN 1 ELSE 0 END) AS [Key Controls],
                                SUM(CASE WHEN c.keyControl = 0 THEN 1 ELSE 0 END) AS [Non-Key Controls],
                                COUNT(c.id) AS [Total Controls]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            LEFT JOIN {db_service.get_fully_qualified_table_name('JobTitles')} jt ON c.departmentId = jt.id
                            WHERE c.isDeleted = 0
                            GROUP BY COALESCE(jt.name, 'Unassigned Department'), c.departmentId
                            ORDER BY COUNT(c.id) DESC, COALESCE(jt.name, 'Unassigned Department')
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for keyNonKeyControlsPerDepartment")
                        controls_data['keyNonKeyControlsPerDepartment'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching keyNonKeyControlsPerDepartment data: {e}")
                        controls_data['keyNonKeyControlsPerDepartment'] = []
            elif cardType == 'keyNonKeyControlsPerProcess':
                if not controls_data.get('keyNonKeyControlsPerProcess'):
                    print("DEBUG: Fetching keyNonKeyControlsPerProcess data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                CASE 
                                    WHEN p.name IS NULL THEN 'Unassigned Process'
                                    ELSE p.name
                                END AS [Process],
                                SUM(CASE WHEN c.keyControl = 1 THEN 1 ELSE 0 END) AS [Key Controls],
                                SUM(CASE WHEN c.keyControl = 0 THEN 1 ELSE 0 END) AS [Non-Key Controls],
                                COUNT(c.id) AS [Total Controls]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            LEFT JOIN {db_service.get_fully_qualified_table_name('ControlProcesses')} cp ON c.id = cp.control_id
                            LEFT JOIN {db_service.get_fully_qualified_table_name('Processes')} p ON cp.process_id = p.id
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY 
                                CASE 
                                    WHEN p.name IS NULL THEN 'Unassigned Process'
                                    ELSE p.name
                                END
                            ORDER BY COUNT(c.id) DESC, 
                                CASE 
                                    WHEN p.name IS NULL THEN 'Unassigned Process'
                                    ELSE p.name
                                END
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for keyNonKeyControlsPerProcess")
                        controls_data['keyNonKeyControlsPerProcess'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching keyNonKeyControlsPerProcess data: {e}")
                        controls_data['keyNonKeyControlsPerProcess'] = []
            elif cardType == 'keyNonKeyControlsPerBusinessUnit':
                if not controls_data.get('keyNonKeyControlsPerBusinessUnit'):
                    print("DEBUG: Fetching keyNonKeyControlsPerBusinessUnit data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                f.name AS [Business Unit],
                                SUM(CASE WHEN c.keyControl = 1 THEN 1 ELSE 0 END) AS [Key Controls],
                                SUM(CASE WHEN c.keyControl = 0 THEN 1 ELSE 0 END) AS [Non-Key Controls],
                                COUNT(c.id) AS [Total Controls]
                            FROM {db_service.get_fully_qualified_table_name('ControlFunctions')} cf
                            JOIN {db_service.get_fully_qualified_table_name('Functions')} f ON cf.function_id = f.id
                            JOIN {db_service.get_fully_qualified_table_name('Controls')} c ON cf.control_id = c.id
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY f.name
                            ORDER BY COUNT(c.id) DESC, f.name
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for keyNonKeyControlsPerBusinessUnit")
                        controls_data['keyNonKeyControlsPerBusinessUnit'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching keyNonKeyControlsPerBusinessUnit data: {e}")
                        controls_data['keyNonKeyControlsPerBusinessUnit'] = []
            elif cardType == 'controlCountByAssertionName':
                if not controls_data.get('controlCountByAssertionName'):
                    print("DEBUG: Fetching controlCountByAssertionName data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                COALESCE(a.name, 'Unassigned Assertion') AS [Assertion Name],
                                COALESCE(a.account_type, 'Not Specified') AS [Type],
                                COUNT(c.id) AS [Control Count]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            LEFT JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id AND a.isDeleted = 0
                            WHERE c.isDeleted = 0 {date_filter}
                            GROUP BY a.name, a.account_type
                            ORDER BY COUNT(c.id) DESC, a.name
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlCountByAssertionName")
                        controls_data['controlCountByAssertionName'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlCountByAssertionName data: {e}")
                        controls_data['controlCountByAssertionName'] = []
            elif cardType == 'icofrControlCoverageByCoso':
                if not controls_data.get('icofrControlCoverageByCoso'):
                    print("DEBUG: Fetching icofrControlCoverageByCoso data from database")
                    try:
                        # Try simplified query first for better performance
                        rows = await db_service.execute_query(
                            """
                            SELECT 
                                'Control Environment' AS [Component], 
                                'ICOFR' AS [IcofrStatus], 
                                COUNT(*) AS [Control Count]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 AND c.icof_id IS NOT NULL
                            UNION ALL
                            SELECT 
                                'Control Environment' AS [Component], 
                                'Non-ICOFR' AS [IcofrStatus], 
                                COUNT(*) AS [Control Count]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            WHERE c.isDeleted = 0 AND c.icof_id IS NULL
                            ORDER BY [Component], [IcofrStatus]
                            """
                        )
                        if not rows:
                            # Fallback to complex query if simplified doesn't work
                            rows = await db_service.execute_query(
                                """
                                SELECT 
                                    comp.name AS [Component], 
                                    CASE 
                                      WHEN c.icof_id IS NOT NULL 
                                        AND (a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1) 
                                        AND (a.account_type IN ('Balance Sheet', 'Income Statement')) 
                                      THEN 'ICOFR' 
                                      WHEN c.icof_id IS NULL 
                                        OR ((a.C IS NULL OR a.C = 0) AND (a.E IS NULL OR a.E = 0) AND (a.A IS NULL OR a.A = 0) 
                                            AND (a.V IS NULL OR a.V = 0) AND (a.O IS NULL OR a.O = 0) AND (a.P IS NULL OR a.P = 0)) 
                                        OR a.account_type NOT IN ('Balance Sheet', 'Income Statement')
                                      THEN 'Non-ICOFR'
                                      ELSE 'Other'
                                    END AS [IcofrStatus], 
                                    COUNT(DISTINCT c.id) AS [Control Count]
                                FROM {db_service.get_fully_qualified_table_name('Controls')} c
                                LEFT JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id AND (a.isDeleted = 0 OR a.id IS NULL)
                                JOIN {db_service.get_fully_qualified_table_name('ControlCosos')} ccx ON c.id = ccx.control_id AND ccx.deletedAt IS NULL
                                JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} point ON ccx.coso_id = point.id AND point.deletedAt IS NULL
                                JOIN {db_service.get_fully_qualified_table_name('CosoPrinciples')} prin ON point.principle_id = prin.id AND prin.deletedAt IS NULL
                                JOIN {db_service.get_fully_qualified_table_name('CosoComponents')} comp ON prin.component_id = comp.id AND comp.deletedAt IS NULL
                                WHERE c.isDeleted = 0
                                GROUP BY comp.name, 
                                  CASE 
                                    WHEN c.icof_id IS NOT NULL 
                                      AND (a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1) 
                                      AND (a.account_type IN ('Balance Sheet', 'Income Statement')) 
                                    THEN 'ICOFR' 
                                    WHEN c.icof_id IS NULL 
                                      OR ((a.C IS NULL OR a.C = 0) AND (a.E IS NULL OR a.E = 0) AND (a.A IS NULL OR a.A = 0) 
                                          AND (a.V IS NULL OR a.V = 0) AND (a.O IS NULL OR a.O = 0) AND (a.P IS NULL OR a.P = 0)) 
                                      OR a.account_type NOT IN ('Balance Sheet', 'Income Statement')
                                    THEN 'Non-ICOFR'
                                    ELSE 'Other'
                                  END
                                ORDER BY comp.name, [IcofrStatus]
                                """
                            )
                        print(f"DEBUG: Fetched {len(rows)} rows for icofrControlCoverageByCoso")
                        controls_data['icofrControlCoverageByCoso'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching icofrControlCoverageByCoso data: {e}")
                        controls_data['icofrControlCoverageByCoso'] = []
            elif cardType == 'actionPlanForAdequacy':
                if not controls_data.get('actionPlanForAdequacy'):
                    print("DEBUG: Fetching actionPlanForAdequacy data from database")
                    try:
                        rows = await db_service.execute_query(
                            """
                            SELECT 
                                COALESCE(c.name, 'N/A') AS [Control Name], 
                                COALESCE(f.name, 'N/A') AS [Function Name], 
                                ap.factor AS [Factor], 
                                ap.riskType AS [Risk Treatment], 
                                ap.control_procedure AS [Control Procedure], 
                                ap.[type] AS [Control Procedure Type], 
                                ap.responsible AS [Action Plan Owner], 
                                ap.expected_cost AS [Expected Cost], 
                                ap.business_unit AS [Business Unit Status], 
                                ap.meeting_date AS [Meeting Date], 
                                ap.implementation_date AS [Expected Implementation Date], 
                                ap.not_attend AS [Did Not Attend]
                            FROM {db_service.get_fully_qualified_table_name('Actionplans')} ap
                            LEFT JOIN {db_service.get_fully_qualified_table_name('ControlDesignTests')} cdt ON ap.controlDesignTest_id = cdt.id AND cdt.deletedAt IS NULL
                            LEFT JOIN {db_service.get_fully_qualified_table_name('Controls')} c ON cdt.control_id = c.id AND c.isDeleted = 0
                            LEFT JOIN {db_service.get_fully_qualified_table_name('Functions')} f ON cdt.function_id = f.id AND f.deletedAt IS NULL
                            WHERE ap.[from] = 'adequacy' 
                                AND ap.deletedAt IS NULL
                            ORDER BY ap.createdAt DESC
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for actionPlanForAdequacy")
                        controls_data['actionPlanForAdequacy'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching actionPlanForAdequacy data: {e}")
                        controls_data['actionPlanForAdequacy'] = []
            elif cardType == 'actionPlanForEffectiveness':
                if not controls_data.get('actionPlanForEffectiveness'):
                    print("DEBUG: Fetching actionPlanForEffectiveness data from Node.js API")
                    try:
                        # Use the Node.js API data directly since it's working
                        node_data = await api_service.get_controls_data(startDate, endDate)
                        if node_data.get('actionPlanForEffectiveness'):
                            controls_data['actionPlanForEffectiveness'] = node_data['actionPlanForEffectiveness']
                            print(f"DEBUG: Fetched {len(controls_data['actionPlanForEffectiveness'])} rows for actionPlanForEffectiveness from Node.js API")
                        else:
                            print("DEBUG: No actionPlanForEffectiveness data in Node.js API response")
                            controls_data['actionPlanForEffectiveness'] = []
                    except Exception as e:
                        print(f"DEBUG: Error fetching actionPlanForEffectiveness data from Node.js API: {e}")
                        controls_data['actionPlanForEffectiveness'] = []
            elif cardType == 'controlSubmissionStatusByQuarterFunction':
                if not controls_data.get('controlSubmissionStatusByQuarterFunction'):
                    print("DEBUG: Fetching controlSubmissionStatusByQuarterFunction data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                c.name AS [Control Name], 
                                f.name AS [Function Name], 
                                CASE WHEN cdt.quarter = 'quarterOne' THEN 1 
                                     WHEN cdt.quarter = 'quarterTwo' THEN 2 
                                     WHEN cdt.quarter = 'quarterThree' THEN 3 
                                     WHEN cdt.quarter = 'quarterFour' THEN 4 
                                     ELSE NULL END AS [Quarter], 
                                cdt.year AS [Year], 
                                -- Submitted? (Control-level full approval cycle)
                                CASE WHEN ( c.preparerStatus = 'sent' AND c.acceptanceStatus = 'approved' ) 
                                     THEN CAST(1 AS bit) ELSE CAST(0 AS bit) END AS [Control Submitted?], 
                                -- Approved? (ControlDesignTests-level full approval cycle)
                                CASE WHEN ( cdt.preparerStatus = 'sent' AND cdt.acceptanceStatus = 'approved' ) 
                                     THEN CAST(1 AS bit) ELSE CAST(0 AS bit) END AS [Test Approved?] 
                            FROM {db_service.get_fully_qualified_table_name('ControlDesignTests')} cdt 
                            JOIN {db_service.get_fully_qualified_table_name('Controls')} c ON cdt.control_id = c.id 
                            JOIN {db_service.get_fully_qualified_table_name('Functions')} f ON cdt.function_id = f.id 
                            WHERE c.isDeleted = 0 AND cdt.deletedAt IS NULL
                            ORDER BY c.createdAt DESC
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlSubmissionStatusByQuarterFunction")
                        controls_data['controlSubmissionStatusByQuarterFunction'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlSubmissionStatusByQuarterFunction data: {e}")
                        controls_data['controlSubmissionStatusByQuarterFunction'] = []
            elif cardType == 'functionsWithFullyTestedControlTests':
                if not controls_data.get('functionsWithFullyTestedControlTests'):
                    print("DEBUG: Fetching functionsWithFullyTestedControlTests data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                f.name AS [Function Name],
                                CASE WHEN cdt.quarter = 'quarterOne' THEN 1 
                                     WHEN cdt.quarter = 'quarterTwo' THEN 2 
                                     WHEN cdt.quarter = 'quarterThree' THEN 3 
                                     WHEN cdt.quarter = 'quarterFour' THEN 4 
                                     ELSE NULL END AS [Quarter],
                                cdt.year AS [Year],
                                COUNT(DISTINCT c.id) AS [Total Controls],
                                COUNT(DISTINCT CASE WHEN (c.preparerStatus = 'sent' AND c.acceptanceStatus = 'approved') THEN c.id END) AS [Controls Submitted],
                                COUNT(DISTINCT CASE WHEN (cdt.preparerStatus = 'sent' AND cdt.acceptanceStatus = 'approved') THEN c.id END) AS [Tests Approved]
                            FROM {db_service.get_fully_qualified_table_name('Functions')} AS f 
                            JOIN {db_service.get_fully_qualified_table_name('ControlFunctions')} AS cf ON f.id = cf.function_id 
                            JOIN {db_service.get_fully_qualified_table_name('Controls')} AS c ON cf.control_id = c.id AND c.isDeleted = 0 
                            LEFT JOIN {db_service.get_fully_qualified_table_name('ControlDesignTests')} AS cdt ON cdt.control_id = c.id AND cdt.deletedAt IS NULL 
                            GROUP BY f.name, cdt.quarter, cdt.year
                            ORDER BY f.name, cdt.year, cdt.quarter
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for functionsWithFullyTestedControlTests")
                        controls_data['functionsWithFullyTestedControlTests'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching functionsWithFullyTestedControlTests data: {e}")
                        controls_data['functionsWithFullyTestedControlTests'] = []
            elif cardType == 'functionsWithFullySubmittedControlTests':
                if not controls_data.get('functionsWithFullySubmittedControlTests'):
                    print("DEBUG: Fetching functionsWithFullySubmittedControlTests data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                f.name AS [Function Name],
                                CASE WHEN cdt.quarter = 'quarterOne' THEN 1 
                                     WHEN cdt.quarter = 'quarterTwo' THEN 2 
                                     WHEN cdt.quarter = 'quarterThree' THEN 3 
                                     WHEN cdt.quarter = 'quarterFour' THEN 4 
                                     ELSE NULL END AS [Quarter],
                                cdt.year AS [Year],
                                COUNT(DISTINCT c.id) AS [Total Controls],
                                COUNT(DISTINCT CASE WHEN (c.preparerStatus = 'sent' AND c.acceptanceStatus = 'approved') THEN c.id END) AS [Controls Submitted],
                                COUNT(DISTINCT CASE WHEN (cdt.preparerStatus = 'sent' AND cdt.acceptanceStatus = 'approved') THEN c.id END) AS [Tests Approved]
                            FROM {db_service.get_fully_qualified_table_name('Functions')} AS f 
                            JOIN {db_service.get_fully_qualified_table_name('ControlFunctions')} AS cf ON f.id = cf.function_id 
                            JOIN {db_service.get_fully_qualified_table_name('Controls')} AS c ON cf.control_id = c.id AND c.isDeleted = 0 
                            LEFT JOIN {db_service.get_fully_qualified_table_name('ControlDesignTests')} AS cdt ON cdt.control_id = c.id AND cdt.deletedAt IS NULL 
                            GROUP BY f.name, cdt.quarter, cdt.year
                            ORDER BY f.name, cdt.year, cdt.quarter
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for functionsWithFullySubmittedControlTests")
                        controls_data['functionsWithFullySubmittedControlTests'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching functionsWithFullySubmittedControlTests data: {e}")
                        controls_data['functionsWithFullySubmittedControlTests'] = []
            elif cardType == 'controlsNotMappedToAssertions':
                if not controls_data.get('controlsNotMappedToAssertions'):
                    print("DEBUG: Fetching controlsNotMappedToAssertions data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                c.name AS [Control Name], 
                                c.departmentId AS [Department]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            LEFT JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id AND a.isDeleted = 0
                            WHERE a.id IS NULL AND c.isDeleted = 0 {date_filter}
                            ORDER BY c.createdAt DESC
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlsNotMappedToAssertions")
                        controls_data['controlsNotMappedToAssertions'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlsNotMappedToAssertions data: {e}")
                        controls_data['controlsNotMappedToAssertions'] = []
            elif cardType == 'controlsNotMappedToPrinciples':
                if not controls_data.get('controlsNotMappedToPrinciples'):
                    print("DEBUG: Fetching controlsNotMappedToPrinciples data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                c.name AS [Control Name], 
                                c.departmentId AS [Department]
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            LEFT JOIN {db_service.get_fully_qualified_table_name('ControlCosos')} ccx ON ccx.control_id = c.id AND ccx.deletedAt IS NULL
                            LEFT JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} point ON point.id = ccx.coso_id
                            LEFT JOIN {db_service.get_fully_qualified_table_name('CosoPrinciples')} prin ON prin.id = point.principle_id
                            WHERE prin.id IS NULL AND c.isDeleted = 0 {date_filter}
                            ORDER BY c.createdAt DESC
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlsNotMappedToPrinciples")
                        controls_data['controlsNotMappedToPrinciples'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlsNotMappedToPrinciples data: {e}")
                        controls_data['controlsNotMappedToPrinciples'] = []
            elif cardType == 'controlsByFunction':
                if not controls_data.get('controlsByFunction'):
                    print("DEBUG: Fetching controlsByFunction data from database")
                    try:
                        rows = await db_service.execute_query(
                            f"""
                            SELECT 
                                f.name as function_name,
                                c.id as control_id,
                                c.name as control_name,
                                c.code as control_code
                            FROM {db_service.get_fully_qualified_table_name('Controls')} c
                            JOIN {db_service.get_fully_qualified_table_name('ControlFunctions')} cf ON c.id = cf.control_id
                            JOIN {db_service.get_fully_qualified_table_name('Functions')} f ON cf.function_id = f.id
                            WHERE c.isDeleted = 0 {date_filter}
                            ORDER BY c.createdAt DESC, f.name, c.name
                            """
                        )
                        print(f"DEBUG: Fetched {len(rows)} rows for controlsByFunction")
                        controls_data['controlsByFunction'] = rows
                    except Exception as e:
                        print(f"DEBUG: Error fetching controlsByFunction data: {e}")
                        controls_data['controlsByFunction'] = []
        
        # Ensure chart data is present for chart-only exports
        if onlyCard and cardType in ['department', 'risk', 'quarterlyControlCreationTrend', 'controlsByType', 'antiFraudDistribution', 'controlsPerLevel', 'controlExecutionFrequency']:
            # department chart expects departmentDistribution: [{name, value}]
            if cardType == 'department':
                dist = controls_data.get('departmentDistribution', []) or []
                if not dist:
                    rows = await db_service.get_controls_by_department(startDate, endDate)
                    # rows fields: department_name, control_count
                    controls_data['departmentDistribution'] = [
                        { 'name': (r.get('department_name') or r.get('name') or 'Unknown'), 'value': r.get('control_count', 0) }
                        for r in rows
                    ]
            # risk response chart expects statusDistribution: [{name, value}]
            if cardType == 'risk':
                dist = controls_data.get('statusDistribution', []) or []
                if not dist:
                    rows = await db_service.get_controls_by_risk_response(startDate, endDate)
                    # rows fields: risk_response_type, control_count
                    controls_data['statusDistribution'] = [
                        { 'name': (r.get('risk_response_type') or r.get('name') or 'Unknown'), 'value': r.get('control_count', 0) }
                        for r in rows
                    ]
            # antiFraudDistribution chart expects antiFraudDistribution: [{name, value}]
            if cardType == 'antiFraudDistribution':
                dist = controls_data.get('antiFraudDistribution', []) or []
                if not dist:
                    rows = await db_service.execute_query(
                        """
                        SELECT 
                            CASE 
                                WHEN c.AntiFraud = 1 THEN 'Anti-Fraud'
                                WHEN c.AntiFraud = 0 THEN 'Non-Anti-Fraud'
                                ELSE 'Unknown'
                            END AS name,
                            COUNT(c.id) AS value
                        FROM {db_service.get_fully_qualified_table_name('Controls')} c
                        WHERE c.isDeleted = 0
                        GROUP BY c.AntiFraud
                        ORDER BY COUNT(c.id) DESC
                        """
                    )
                    controls_data['antiFraudDistribution'] = [
                        { 'name': (r.get('name') or 'Unknown'), 'value': r.get('value', 0) }
                        for r in rows
                    ]
            # controlsPerLevel chart expects controlsPerLevel: [{name, value}]
            if cardType == 'controlsPerLevel':
                dist = controls_data.get('controlsPerLevel', []) or []
                if not dist:
                    rows = await db_service.execute_query(
                        """
                        SELECT 
                            CASE 
                                WHEN c.entityLevel IS NULL OR c.entityLevel = '' THEN 'Not Specified'
                                ELSE c.entityLevel
                            END AS name,
                            COUNT(c.id) AS value
                        FROM {db_service.get_fully_qualified_table_name('Controls')} c
                        WHERE c.isDeleted = 0
                        GROUP BY c.entityLevel
                        ORDER BY COUNT(c.id) DESC
                        """
                    )
                    controls_data['controlsPerLevel'] = [
                        { 'name': (r.get('name') or 'Unknown'), 'value': r.get('value', 0) }
                        for r in rows
                    ]
            # controlExecutionFrequency chart expects controlExecutionFrequency: [{name, value}]
            if cardType == 'controlExecutionFrequency':
                dist = controls_data.get('controlExecutionFrequency', []) or []
                if not dist:
                    rows = await db_service.execute_query(
                        """
                        SELECT 
                            CASE 
                                WHEN c.frequency = 'Daily' THEN 'Daily'
                                WHEN c.frequency = 'Event Base' THEN 'Event Base'
                                WHEN c.frequency = 'Weekly' THEN 'Weekly'
                                WHEN c.frequency = 'Monthly' THEN 'Monthly'
                                WHEN c.frequency = 'Quarterly' THEN 'Quarterly'
                                WHEN c.frequency = 'Semi Annually' THEN 'Semi Annually'
                                WHEN c.frequency = 'Annually' THEN 'Annually'
                                WHEN c.frequency IS NULL OR c.frequency = '' THEN 'Not Specified'
                                ELSE c.frequency
                            END AS name,
                            COUNT(c.id) AS value
                        FROM {db_service.get_fully_qualified_table_name('Controls')} c
                        WHERE c.isDeleted = 0
                        GROUP BY c.frequency
                        ORDER BY COUNT(c.id) DESC
                        """
                    )
                    controls_data['controlExecutionFrequency'] = [
                        { 'name': (r.get('name') or 'Unknown'), 'value': r.get('value', 0) }
                        for r in rows
                    ]
            # numberOfControlsByIcofrStatus chart expects numberOfControlsByIcofrStatus: [{name, value}]
            if cardType == 'numberOfControlsByIcofrStatus':
                dist = controls_data.get('numberOfControlsByIcofrStatus', []) or []
                if not dist:
                    rows = await db_service.execute_query(
                        """
                        SELECT 
                            CASE 
                                WHEN a.id IS NULL THEN 'Non-ICOFR'
                                WHEN (a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1) 
                                     AND (a.account_type IN ('Balance Sheet', 'Income Statement')) 
                                  THEN 'ICOFR' 
                                ELSE 'Non-ICOFR' 
                            END AS name,
                            COUNT(c.id) AS value
                        FROM {db_service.get_fully_qualified_table_name('Controls')} c
                        LEFT JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id AND a.isDeleted = 0
                        WHERE c.isDeleted = 0
                        GROUP BY 
                            CASE 
                                WHEN a.id IS NULL THEN 'Non-ICOFR'
                                WHEN (a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1) 
                                     AND (a.account_type IN ('Balance Sheet', 'Income Statement')) 
                                  THEN 'ICOFR' 
                                ELSE 'Non-ICOFR' 
                            END
                        ORDER BY COUNT(c.id) DESC
                        """
                    )
                    controls_data['numberOfControlsByIcofrStatus'] = [
                        { 'name': (r.get('name') or 'Unknown'), 'value': r.get('value', 0) }
                        for r in rows
                    ]
            # numberOfFocusPointsPerPrinciple chart expects numberOfFocusPointsPerPrinciple: [{name, value}]
            if cardType == 'numberOfFocusPointsPerPrinciple':
                dist = controls_data.get('numberOfFocusPointsPerPrinciple', []) or []
                if not dist:
                    rows = await db_service.execute_query(
                        f"""
                        SELECT 
                            prin.name AS name,
                            COUNT(point.id) AS value
                        FROM {db_service.get_fully_qualified_table_name('CosoPrinciples')} prin
                        LEFT JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} point ON prin.id = point.principle_id
                        WHERE prin.deletedAt IS NULL
                        GROUP BY prin.name
                        ORDER BY COUNT(point.id) DESC, prin.name
                        """
                    )
                    controls_data['numberOfFocusPointsPerPrinciple'] = [
                        { 'name': (r.get('name') or 'Unknown'), 'value': r.get('value', 0) }
                        for r in rows
                    ]
            # numberOfFocusPointsPerComponent chart expects numberOfFocusPointsPerComponent: [{name, value}]
            if cardType == 'numberOfFocusPointsPerComponent':
                dist = controls_data.get('numberOfFocusPointsPerComponent', []) or []
                if not dist:
                    rows = await db_service.execute_query(
                        f"""
                        SELECT 
                            comp.name AS name,
                            COUNT(point.id) AS value
                        FROM {db_service.get_fully_qualified_table_name('CosoComponents')} comp
                        JOIN {db_service.get_fully_qualified_table_name('CosoPrinciples')} prin ON prin.component_id = comp.id
                        LEFT JOIN {db_service.get_fully_qualified_table_name('CosoPoints')} point ON point.principle_id = prin.id
                        WHERE comp.deletedAt IS NULL AND prin.deletedAt IS NULL
                        GROUP BY comp.name
                        ORDER BY COUNT(point.id) DESC
                        """
                    )
                    controls_data['numberOfFocusPointsPerComponent'] = [
                        { 'name': (r.get('name') or 'Unknown'), 'value': r.get('value', 0) }
                        for r in rows
                    ]
        
        # Add ICOFR metrics if not present
        if not controls_data.get('unmappedIcofrControls'):
            try:
                rows = await db_service.execute_query(
                    """
                    SELECT COUNT(*) AS total 
                    FROM {db_service.get_fully_qualified_table_name('Controls')} c 
                    JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id 
                    WHERE c.isDeleted = 0 AND c.icof_id IS NOT NULL 
                    AND NOT EXISTS (
                        SELECT 1 FROM {db_service.get_fully_qualified_table_name('ControlCosos')} ccx 
                        WHERE ccx.control_id = c.id AND ccx.deletedAt IS NULL
                    ) 
                    AND ((a.C = 1 OR a.E = 1 OR a.A = 1 OR a.V = 1 OR a.O = 1 OR a.P = 1) 
                         AND a.account_type IN ('Balance Sheet', 'Income Statement')) 
                    AND a.isDeleted = 0
                    """
                )
                controls_data['unmappedIcofrControls'] = rows[0].get('total', 0) if rows else 0
            except Exception as e:
                print(f"DEBUG: Error fetching unmappedIcofrControls: {e}")
                controls_data['unmappedIcofrControls'] = 0
        
        if not controls_data.get('unmappedNonIcofrControls'):
            try:
                rows = await db_service.execute_query(
                    """
                    SELECT COUNT(*) AS total 
                    FROM {db_service.get_fully_qualified_table_name('Controls')} c 
                    LEFT JOIN {db_service.get_fully_qualified_table_name('Assertions')} a ON c.icof_id = a.id 
                    WHERE c.isDeleted = 0 
                    AND NOT EXISTS (
                        SELECT 1 FROM {db_service.get_fully_qualified_table_name('ControlCosos')} ccx 
                        WHERE ccx.control_id = c.id AND ccx.deletedAt IS NULL
                    ) 
                    AND (c.icof_id IS NULL OR ((a.C IS NULL OR a.C = 0) AND (a.E IS NULL OR a.E = 0) AND (a.A IS NULL OR a.A = 0) 
                         AND (a.V IS NULL OR a.V = 0) AND (a.O IS NULL OR a.O = 0) AND (a.P IS NULL OR a.P = 0) 
                         OR a.account_type NOT IN ('Balance Sheet', 'Income Statement'))) 
                    AND (a.isDeleted = 0 OR a.id IS NULL)
                    """
                )
                controls_data['unmappedNonIcofrControls'] = rows[0].get('total', 0) if rows else 0
            except Exception as e:
                print(f"DEBUG: Error fetching unmappedNonIcofrControls: {e}")
                controls_data['unmappedNonIcofrControls'] = 0

        # Generate Excel
        excel_content = await excel_service.generate_controls_excel(
            controls_data, startDate, endDate, header_config, cardType, onlyCard, onlyOverallTable, onlyChart
        )
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if onlyCard and cardType:
            filename = f"controls_{cardType}_{timestamp}.xlsx"
        elif onlyOverallTable and table_type:
            filename = f"controls_{table_type}_{timestamp}.xlsx"
        elif onlyChart and chartType:
            filename = f"controls_{chartType}_{timestamp}.xlsx"
        else:
            filename = f"controls_report_{timestamp}.xlsx"
        
        return Response(
            content=excel_content,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# Dashboard Activity Endpoints
@router.get("/api/dashboard-activity")
async def get_dashboard_activities(user_id: str = Query("default_user")):
    """Get dashboard activity data from database"""
    try:
        # Initialize table if it doesn't exist
        await dashboard_activity_service.create_activity_table()
        
        # Initialize default activities if none exist
        await dashboard_activity_service.initialize_default_activities()
        
        # Get activities from database
        activities = await dashboard_activity_service.get_dashboard_activities(user_id)
        
        return activities
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch dashboard activities: {str(e)}")

@router.post("/api/dashboard-activity")
async def update_dashboard_activity(request: Request):
    """Update dashboard activity in database"""
    try:
        # Parse JSON body
        body = await request.json()
        
        # Extract parameters from request body
        dashboard_id = body.get('dashboard_id')
        user_id = body.get('user_id', 'default_user')
        card_count = body.get('card_count', 0)
        
        if not dashboard_id:
            raise HTTPException(status_code=400, detail="dashboard_id is required")
        
        # Initialize table if it doesn't exist
        await dashboard_activity_service.create_activity_table()
        
        # Update activity
        activity = await dashboard_activity_service.update_dashboard_activity(
            dashboard_id, user_id, card_count
        )
        
        return activity
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update dashboard activity: {str(e)}")

# KRI Export Endpoints
@router.get("/api/grc/kris")
async def get_kris_dashboard(
    startDate: str = Query(None),
    endDate: str = Query(None)
):
    """Return KRIs dashboard data"""
    try:
        data = await api_service.get_kris_data(startDate, endDate)
        return data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch KRIs data: {str(e)}")

@router.get("/api/grc/kris/export/pdf")
async def export_kris_pdf(
    onlyCard: str = Query(None),
    onlyChart: str = Query(None),
    onlyOverallTable: str = Query(None),
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    chartType: str = Query(None)
):
    """Export KRIs dashboard to PDF"""
    try:
        # Convert string parameters to boolean
        only_card = onlyCard and onlyCard.lower() in ['true', '1', 'yes']
        only_chart = onlyChart and onlyChart.lower() in ['true', '1', 'yes']
        only_overall_table = onlyOverallTable and onlyOverallTable.lower() in ['true', '1', 'yes']
        
        # Normalize onlyChart parameter for KRI cards
        if only_chart and chartType:
            if chartType in ['krisByStatus', 'krisByLevel', 'breachedKRIsByDepartment', 'kriAssessmentCount']:
                only_chart = True
            else:
                only_chart = False
        
        # Get header configuration
        config = get_default_header_config("kris")
        if headerConfig:
            try:
                user_config = json.loads(headerConfig)
                config.update(user_config)
            except:
                pass
        
        # Get data
        data = await api_service.get_kris_data(startDate, endDate)
        
        # Generate PDF
        pdf_buffer = await pdf_service.generate_kris_pdf(
            data=data,
            only_card=only_card,
            only_chart=only_chart,
            only_overall_table=only_overall_table,
            card_type=chartType,
            start_date=startDate,
            end_date=endDate,
            header_config=config
        )
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"kris_{timestamp}.pdf"
        
        return Response(
            content=pdf_buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

@router.get("/api/grc/kris/export/excel")
async def export_kris_excel(
    onlyCard: str = Query(None),
    onlyChart: str = Query(None),
    onlyOverallTable: str = Query(None),
    startDate: str = Query(None),
    endDate: str = Query(None),
    headerConfig: str = Query(None),
    chartType: str = Query(None)
):
    """Export KRIs dashboard to Excel"""
    try:
        # Convert string parameters to boolean
        only_card = onlyCard and onlyCard.lower() in ['true', '1', 'yes']
        only_chart = onlyChart and onlyChart.lower() in ['true', '1', 'yes']
        only_overall_table = onlyOverallTable and onlyOverallTable.lower() in ['true', '1', 'yes']
        
        # Get header configuration
        config = get_default_header_config("kris")
        if headerConfig:
            try:
                user_config = json.loads(headerConfig)
                config.update(user_config)
            except:
                pass
        
        # Get data
        data = await api_service.get_kris_data(startDate, endDate)
        
        # Generate Excel
        excel_buffer = await excel_service.generate_kris_excel(
            kris_data=data,
            only_card=only_card,
            only_chart=only_chart,
            only_overall_table=only_overall_table,
            card_type=chartType,
            start_date=startDate,
            end_date=endDate,
            header_config=config
        )
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"kris_{timestamp}.xlsx"
        
        return Response(
            content=excel_buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Excel: {str(e)}")

# -------------------------------
# Export logging (Excel/PDF) APIs
# -------------------------------
@router.post("/api/exports/log")
async def log_report_export(request: Request):
    """Log an export (excel/pdf/word/zip) with title and src for later download listing."""
    try:
        import pyodbc
        from config import get_database_connection_string

        body = await request.json()
        title = (body.get("title") or "").strip() or "Untitled Report"
        src = (body.get("src") or "").strip()
        fmt = (body.get("format") or "").strip().lower() or "unknown"
        dashboard = (body.get("dashboard") or "").strip() or "general"

        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            # Ensure table exists
            cursor.execute(
                """
                IF NOT EXISTS (
                  SELECT * FROM INFORMATION_SCHEMA.TABLES 
                  WHERE TABLE_NAME = 'report_exports' AND TABLE_SCHEMA='dbo'
                )
                BEGIN
                  CREATE TABLE dbo.report_exports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    title NVARCHAR(255) NOT NULL,
                    src NVARCHAR(1024) NULL,
                    format NVARCHAR(20) NOT NULL,
                    dashboard NVARCHAR(100) NULL,
                    created_at DATETIME2 DEFAULT GETDATE()
                  )
                END
                """
            )
            conn.commit()

            cursor.execute(
                """
                INSERT INTO dbo.report_exports (title, src, format, dashboard)
                VALUES (?, ?, ?, ?)
                """,
                (title, src, fmt, dashboard)
            )
            conn.commit()

            new_id = cursor.execute("SELECT @@IDENTITY").fetchone()[0]
            return {"success": True, "id": int(new_id)}
        finally:
            cursor.close()
            conn.close()
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.get("/api/exports/recent")
async def list_recent_exports(limit: int = Query(50), page: int = Query(1), search: str = Query("")):
    """Return recent report exports (newest first) with simple pagination."""
    try:
        import pyodbc
        from config import get_database_connection_string

        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            cursor.execute(
                """
                IF NOT EXISTS (
                  SELECT * FROM INFORMATION_SCHEMA.TABLES 
                  WHERE TABLE_NAME = 'report_exports' AND TABLE_SCHEMA='dbo'
                )
                BEGIN
                  CREATE TABLE dbo.report_exports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    title NVARCHAR(255) NOT NULL,
                    src NVARCHAR(1024) NULL,
                    format NVARCHAR(20) NOT NULL,
                    dashboard NVARCHAR(100) NULL,
                    created_at DATETIME2 DEFAULT GETDATE()
                  )
                END
                """
            )
            conn.commit()

            # Build search condition
            search_condition = ""
            search_params = []
            if search and search.strip():
                search_condition = "WHERE title LIKE ?"
                search_params.append(f"%{search.strip()}%")

            # Total count with search
            count_query = f"SELECT COUNT(*) FROM dbo.report_exports {search_condition}"
            cursor.execute(count_query, search_params)
            total_count = int(cursor.fetchone()[0])

            # Pagination via OFFSET/FETCH
            safe_limit = max(1, min(200, int(limit)))
            safe_page = max(1, int(page))
            offset = (safe_page - 1) * safe_limit
            
            select_query = f"""
                SELECT id, title, src, format, dashboard, created_at
                FROM dbo.report_exports
                {search_condition}
                ORDER BY created_at DESC, id DESC
                OFFSET ? ROWS FETCH NEXT ? ROWS ONLY
            """
            cursor.execute(select_query, search_params + [offset, safe_limit])
            rows = cursor.fetchall()
            exports = [
                {
                    "id": int(r[0]),
                    "title": r[1],
                    "src": r[2],
                    "format": r[3],
                    "dashboard": r[4],
                    "created_at": r[5].isoformat() if hasattr(r[5], 'isoformat') else str(r[5])
                }
                for r in rows
            ]
            return {
                "success": True,
                "exports": exports,
                "pagination": {
                    "page": safe_page,
                    "limit": safe_limit,
                    "total": total_count,
                    "totalPages": (total_count + safe_limit - 1) // safe_limit,
                    "hasNext": offset + safe_limit < total_count,
                    "hasPrev": safe_page > 1
                }
            }
        finally:
            cursor.close()
            conn.close()
    except Exception as e:
        return {"success": False, "error": str(e), "exports": [], "pagination": {}}

@router.delete("/api/exports/{export_id}")
async def delete_export(export_id: int):
    """Delete an export row and its file if present"""
    try:
        import pyodbc
        from config import get_database_connection_string
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            cursor.execute("SELECT src FROM dbo.report_exports WHERE id = ?", export_id)
            row = cursor.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Export not found")
            src = row[0]

            # Delete DB row
            cursor.execute("DELETE FROM dbo.report_exports WHERE id = ?", export_id)
            conn.commit()

            # Delete file if exists
            if src:
                try:
                    import os
                    file_path = src if os.path.isabs(src) else os.path.join(os.getcwd(), src)
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except Exception as fe:
                    return {"success": True, "deleted": True, "fileDeleted": False, "warning": str(fe)}

            return {"success": True, "deleted": True, "fileDeleted": True}
        finally:
            cursor.close()
            conn.close()
    except HTTPException:
        raise
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.post("/api/reports/dynamic")
async def generate_dynamic_report(request: Request):
    """Generate dynamic report based on table selection, joins, columns, and conditions"""
    try:
        body = await request.json()
        tables = body.get('tables', [])
        joins = body.get('joins', [])
        columns = body.get('columns', [])
        where_conditions = body.get('whereConditions', [])
        time_filter = body.get('timeFilter')
        format_type = body.get('format', 'excel')
        
        if not tables or not columns:
            raise HTTPException(status_code=400, detail="Tables and columns are required")
        
        # Build SQL query
        sql_query = build_dynamic_sql_query(tables, joins, columns, where_conditions, time_filter)
        
        # Execute query and get data
        import pyodbc
        from config import get_database_connection_string
        
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            cursor.execute(sql_query)
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data_rows = []
            for row in rows:
                data_rows.append([str(cell) if cell is not None else '' for cell in row])
            
            # Get header configuration from request body
            header_config = body.get('headerConfig', {})
            if header_config:
                from export_utils import get_default_header_config
                default_config = get_default_header_config("dynamic")
                merged_config = {**default_config, **header_config}
            else:
                from export_utils import get_default_header_config
                merged_config = get_default_header_config("dynamic")
            
            # Generate report based on format
            if format_type == 'excel':
                return generate_excel_report(columns, data_rows, merged_config)
            elif format_type == 'word':
                return generate_word_report(columns, data_rows, merged_config)
            elif format_type == 'pdf':
                return generate_pdf_report(columns, data_rows, merged_config)
            else:
                raise HTTPException(status_code=400, detail="Unsupported format")
                
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate dynamic report: {str(e)}")

def build_dynamic_sql_query(tables, joins, columns, where_conditions, time_filter):
    """Build SQL query from dynamic report configuration"""
    # Start with SELECT clause
    select_columns = ', '.join(columns) if columns else '*'
    sql = f"SELECT {select_columns}"
    
    # Add FROM clause with first table
    if not tables:
        raise ValueError("At least one table is required")
    
    sql += f" FROM {tables[0]}"
    
    # Add JOINs
    for join in joins:
        if join.get('leftTable') and join.get('rightTable') and join.get('leftColumn') and join.get('rightColumn'):
            join_type = join.get('type', 'INNER')
            sql += f" {join_type} JOIN {join['rightTable']} ON {join['leftTable']}.{join['leftColumn']} = {join['rightTable']}.{join['rightColumn']}"
    
    # Add WHERE clause
    where_clauses = []
    
    # Add time filter
    if time_filter and time_filter.get('column') and time_filter.get('startDate') and time_filter.get('endDate'):
        where_clauses.append(f"{time_filter['column']} BETWEEN '{time_filter['startDate']}' AND '{time_filter['endDate']}'")
    
    # Add custom WHERE conditions
    for i, condition in enumerate(where_conditions):
        if condition.get('column') and condition.get('value'):
            operator = condition.get('operator', '=')
            value = condition.get('value', '')
            logical_op = condition.get('logicalOperator', 'AND') if i > 0 else ''
            
            if logical_op:
                where_clauses.append(f" {logical_op} {condition['column']} {operator} '{value}'")
            else:
                where_clauses.append(f"{condition['column']} {operator} '{value}'")
    
    if where_clauses:
        sql += " WHERE " + " ".join(where_clauses)
    
    return sql

def generate_excel_report(columns, data_rows, header_config=None):
    """Generate Excel report from dynamic data with full header configuration support"""
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    import os
    import base64
    
    # Get default header config if none provided
    if not header_config:
        from export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    
    wb = Workbook()
    ws = wb.active
    ws.title = header_config.get('title', 'Dynamic Report')
    
    # Extract ALL configuration values from header modal configuration
    # Basic report settings
    include_header = header_config.get("includeHeader", True)
    title = header_config.get("title", "Dynamic Report")
    subtitle = header_config.get("subtitle", "")
    icon = header_config.get("icon", "chart-line")
    
    # Date and time settings
    show_date = header_config.get("showDate", True)
    footer_show_date = header_config.get("footerShowDate", True)
    
    # Bank information settings
    show_bank_info = header_config.get("showBankInfo", True)
    bank_info_location = header_config.get("bankInfoLocation", "top")  # top, bottom, none
    bank_info_align = (
        header_config.get("bankInfoAlign")
        or header_config.get("logoPosition", "center")
        or "center"
    ).lower()  # left, center, right
    bank_name = header_config.get("bankName", "")
    bank_address = header_config.get("bankAddress", "")
    bank_phone = header_config.get("bankPhone", "")
    bank_website = header_config.get("bankWebsite", "")
    
    # Logo settings
    show_logo = header_config.get("showLogo", True)
    logo_base64 = header_config.get("logoBase64", "")
    logo_position = header_config.get("logoPosition", "left")
    logo_height = header_config.get("logoHeight", 36)
    logo_file = header_config.get("logoFile", None)
    
    # Color and styling settings
    font_color = header_config.get("fontColor", "#1F4E79")
    table_header_bg_color = header_config.get("tableHeaderBgColor", "#1F4E79")
    table_body_bg_color = header_config.get("tableBodyBgColor", "#FFFFFF")
    background_color = header_config.get("backgroundColor", "#FFFFFF")
    border_style = header_config.get("borderStyle", "solid")
    border_color = header_config.get("borderColor", "#E5E7EB")
    border_width = header_config.get("borderWidth", 1)
    
    # Font and size settings
    font_size = header_config.get("fontSize", "medium")
    padding = header_config.get("padding", 20)
    margin = header_config.get("margin", 72)  # 1 inch = 72 points
    
    # Excel specific settings
    excel_auto_fit_columns = header_config.get("excelAutoFitColumns", True)
    excel_zebra_stripes = header_config.get("excelZebraStripes", True)
    excel_fit_to_width = header_config.get("excelFitToWidth", True)
    excel_freeze_top_row = header_config.get("excelFreezeTopRow", True)
    
    # Watermark settings
    watermark_enabled = header_config.get("watermarkEnabled", False)
    watermark_text = header_config.get("watermarkText", "CONFIDENTIAL")
    watermark_opacity = header_config.get("watermarkOpacity", 10)
    watermark_diagonal = header_config.get("watermarkDiagonal", True)
    
    # Footer settings
    footer_show_confidentiality = header_config.get("footerShowConfidentiality", True)
    footer_confidentiality_text = header_config.get("footerConfidentialityText", "Confidential Report - Internal Use Only")
    footer_show_page_numbers = header_config.get("footerShowPageNumbers", True)
    footer_align = header_config.get("footerAlign", "center")
    
    # Page settings
    show_page_numbers = header_config.get("showPageNumbers", True)
    location = header_config.get("location", "top")
    
    # Convert hex colors to RGB for openpyxl
    def hex_to_rgb(hex_color):
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        return hex_color
    
    font_color_rgb = hex_to_rgb(font_color)
    header_bg_rgb = hex_to_rgb(table_header_bg_color)
    body_bg_rgb = hex_to_rgb(table_body_bg_color)
    background_rgb = hex_to_rgb(background_color)
    border_rgb = hex_to_rgb(border_color)
    
    current_row = 1
    
    # Only add header if includeHeader is True
    if include_header:
        # Add logo if available
        if show_logo and logo_base64:
            try:
                import base64
                from PIL import Image as PILImage
                from openpyxl.drawing.image import Image as XLImage
                
                logo_bytes = base64.b64decode(logo_base64.split(',')[-1])
                logo_buf = BytesIO(logo_bytes)
                pil_img = PILImage.open(logo_buf)
                
                # Resize logo based on configuration
                desired_h = min(logo_height, 64)
                w, h = pil_img.size
                if h > 0:
                    scale = desired_h / h
                    new_w = int(w * scale)
                    max_w = 180
                    if new_w > max_w:
                        new_w = max_w
                        scale = new_w / w
                        desired_h = int(h * scale)
                    pil_img = pil_img.resize((new_w, desired_h), PILImage.Resampling.LANCZOS)
                
                # Save to BytesIO for openpyxl
                logo_buf = BytesIO()
                pil_img.save(logo_buf, format='PNG')
                logo_buf.seek(0)
                
                # Create openpyxl image
                xl_image = XLImage(logo_buf)
                
                # Position logo based on logoPosition
                if logo_position == 'left':
                    ws.add_image(xl_image, 'A1')
                elif logo_position == 'center':
                    ws.add_image(xl_image, 'C1')
                elif logo_position == 'right':
                    ws.add_image(xl_image, 'E1')
                
                current_row += 2  # Leave space for logo
            except Exception as e:
                pass  # Continue without logo if there's an error
        
        # Add bank information at top if configured
        if show_bank_info and bank_name and bank_info_location == "top":
            # Resolve Excel horizontal alignment from bank_info_align
            _xl_bank_align = 'center'
            if bank_info_align == 'left':
                _xl_bank_align = 'left'
            elif bank_info_align == 'right':
                _xl_bank_align = 'right'
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=bank_name)
            ws.cell(row=current_row, column=1).font = Font(size=12, bold=True, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal=_xl_bank_align)
            current_row += 1
            
            if bank_address:
                ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
                ws.cell(row=current_row, column=1, value=bank_address)
                ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
                ws.cell(row=current_row, column=1).alignment = Alignment(horizontal=_xl_bank_align)
                current_row += 1
                
            if bank_phone or bank_website:
                contact_info = []
                if bank_phone:
                    contact_info.append(f"Tel: {bank_phone}")
                if bank_website:
                    contact_info.append(f"Web: {bank_website}")
                
                ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
                ws.cell(row=current_row, column=1, value=" | ".join(contact_info))
                ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
                ws.cell(row=current_row, column=1).alignment = Alignment(horizontal=_xl_bank_align)
                current_row += 1
        
        # Add title
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
        ws.cell(row=current_row, column=1, value=title)
        ws.cell(row=current_row, column=1).font = Font(size=16, bold=True, color=font_color_rgb)
        ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center', vertical='center')
        current_row += 1
        
        # Add subtitle
        if subtitle:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=subtitle)
            ws.cell(row=current_row, column=1).font = Font(size=12, italic=True, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1
        
        # Add generation date
        if show_date:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            ws.cell(row=current_row, column=1).font = Font(size=10, italic=True, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1
        
        # Add empty row for spacing
        current_row += 1
    
    # Table headers
    header_row = current_row
    for idx, col in enumerate(columns, start=1):
        cell = ws.cell(row=header_row, column=idx, value=col)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color=header_bg_rgb, end_color=header_bg_rgb, fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Data rows
    for row_idx, row_data in enumerate(data_rows, start=header_row + 1):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(vertical='top')
            
            # Apply zebra stripes if enabled
            if excel_zebra_stripes and (row_idx - header_row) % 2 == 0:
                cell.fill = PatternFill(start_color=body_bg_rgb, end_color=body_bg_rgb, fill_type='solid')

    # Optional footer totals row
    footer_totals_cols = header_config.get("tableFooterTotals", []) or []
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        name_to_index = {str(col): idx for idx, col in enumerate(columns)}
        totals = [None] * len(columns)
        for col_name in footer_totals_cols:
            if str(col_name) in name_to_index:
                idx = name_to_index[str(col_name)]
                total_value = 0.0
                for row_data in data_rows:
                    try:
                        val = row_data[idx]
                        if val is None or val == "":
                            continue
                        total_value += float(str(val).replace(',', ''))
                    except Exception:
                        pass
                totals[idx] = total_value
        totals_row = ws.max_row + 1
        for i in range(len(columns)):
            cell = ws.cell(row=totals_row, column=i + 1)
            if i == 0:
                cell.value = "Total"
                cell.alignment = Alignment(horizontal='left', vertical='center')
            elif totals[i] is not None:
                cell.value = totals[i]
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='right', vertical='center')
            # Style totals row
            cell.fill = PatternFill(start_color=header_bg_rgb, end_color=header_bg_rgb, fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFF')
    
    # Freeze top row if enabled
    if excel_freeze_top_row:
        ws.freeze_panes = f"A{header_row + 1}"
    
    # Generate chart if chart_data is provided (right side)
    chart_data = header_config.get('chart_data')
    chart_type = header_config.get('chart_type', 'bar')
    if chart_data and chart_data.get('labels') and chart_data.get('values'):
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            from io import BytesIO
            from openpyxl.drawing.image import Image as XLImage
            
            # Create chart
            fig, ax = plt.subplots(figsize=(8, 5))
            
            if chart_type == 'bar':
                ax.barh(chart_data['labels'], chart_data['values'], color='#4472C4')
                ax.set_xlabel('Controls Count')
                ax.set_ylabel('Component')
            elif chart_type == 'pie':
                ax.pie(chart_data['values'], labels=chart_data['labels'], autopct='%1.1f%%')
            
            ax.set_title(title if title else 'Chart')
            plt.tight_layout()
            
            # Save chart to buffer
            chart_buffer = BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=150, bbox_inches='tight')
            chart_buffer.seek(0)
            plt.close()
            
            # Add chart to Excel on the right side
            img = XLImage(chart_buffer)
            img.width = 500
            img.height = 300
            # Position chart to the right of the table (column F)
            chart_col = len(columns) + 2  # Start after the table columns + 1 space
            ws.add_image(img, f'{get_column_letter(chart_col)}{header_row}')
        except Exception as e:
            print(f"Error generating chart for Excel: {e}")
            pass
    
    # Auto-fit columns if enabled
    if excel_auto_fit_columns:
        for col_idx in range(1, len(columns) + 1):
            max_length = 0
            column_letter = get_column_letter(col_idx)
            for row_idx in range(1, ws.max_row + 1):
                try:
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if hasattr(cell, 'value') and cell.value:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length
                except:
                    pass
            adjusted_width = min(max_length + 3, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    else:
        # Set default column width
        for i in range(1, len(columns) + 1):
            ws.column_dimensions[get_column_letter(i)].width = 15
    
    # Apply fit to width if enabled
    if excel_fit_to_width:
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
    
    # Add watermark if enabled
    if watermark_enabled:
        try:
            from export_utils import add_watermark_to_excel_sheet
            add_watermark_to_excel_sheet(ws, header_config)
        except Exception as e:
            pass  # Continue without watermark if there's an error
    
    # Add bank information at bottom if configured
    if show_bank_info and bank_name and bank_info_location == "bottom":
        # Add some spacing before bank info
        current_row += 2
        
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
        ws.cell(row=current_row, column=1, value=bank_name)
        ws.cell(row=current_row, column=1).font = Font(size=12, bold=True, color=font_color_rgb)
        ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
        current_row += 1
        
        if bank_address:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=bank_address)
            ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1
            
        if bank_phone or bank_website:
            contact_info = []
            if bank_phone:
                contact_info.append(f"Tel: {bank_phone}")
            if bank_website:
                contact_info.append(f"Web: {bank_website}")
            
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=" | ".join(contact_info))
            ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1

    # Configure page setup and headers/footers
    if show_page_numbers:
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE if len(columns) > 6 else ws.ORIENTATION_PORTRAIT
    
    # Add footer if configured
    if footer_show_date or footer_show_confidentiality or footer_show_page_numbers:
        footer_text = []
        if footer_show_date:
            footer_text.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if footer_show_confidentiality:
            footer_text.append(footer_confidentiality_text)
        if footer_show_page_numbers:
            footer_text.append("Page &P of &N")
        
        if footer_align == "center":
            ws.HeaderFooter.oddFooter.center.text = " | ".join(footer_text)
        elif footer_align == "left":
            ws.HeaderFooter.oddFooter.left.text = " | ".join(footer_text)
        elif footer_align == "right":
            ws.HeaderFooter.oddFooter.right.text = " | ".join(footer_text)
    
    # Save to file
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    date_folder = datetime.now().strftime('%Y-%m-%d')
    os.makedirs(f"reports_export/{date_folder}", exist_ok=True)
    filename = f"dynamic_report_{ts}.xlsx"
    file_path = f"reports_export/{date_folder}/{filename}"
    
    wb.save(file_path)
    
    # Save export record to database (best-effort)
    try:
        import pyodbc
        from config import get_database_connection_string
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            cursor.execute(
                """
                IF NOT EXISTS (
                  SELECT * FROM INFORMATION_SCHEMA.TABLES 
                  WHERE TABLE_NAME = 'report_exports' AND TABLE_SCHEMA='dbo'
                )
                BEGIN
                  CREATE TABLE dbo.report_exports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    title NVARCHAR(255) NOT NULL,
                    src NVARCHAR(1024) NULL,
                    format NVARCHAR(20) NOT NULL,
                    dashboard NVARCHAR(100) NULL,
                    created_at DATETIME2 DEFAULT GETDATE()
                  )
                END
                """
            )
            conn.commit()
            export_title = header_config.get("title", "Dynamic Report") if header_config else "Dynamic Report"
            export_title = f"{export_title} - {datetime.now().strftime('%Y-%m-%d')}"
            export_dashboard = header_config.get("dashboard", "dynamic") if header_config else "dynamic"
            cursor.execute(
                """
                INSERT INTO dbo.report_exports (title, src, format, dashboard)
                VALUES (?, ?, ?, ?)
                """,
                (export_title, file_path, 'excel', export_dashboard)
            )
            conn.commit()
        finally:
            cursor.close()
            conn.close()
    except Exception:
        pass

    with open(file_path, 'rb') as f:
        content = f.read()
    
    return content

def generate_word_report(columns, data_rows, header_config=None):
    """Generate Word report from dynamic data with full header configuration support"""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os
    
    # Get default header config if none provided
    if not header_config:
        from export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    
    doc = Document()
    
    # Extract ALL configuration values from header modal configuration
    # Basic report settings
    include_header = header_config.get("includeHeader", True)
    title = header_config.get("title", "Dynamic Report")
    subtitle = header_config.get("subtitle", "")
    icon = header_config.get("icon", "chart-line")
    
    # Date and time settings
    show_date = header_config.get("showDate", True)
    footer_show_date = header_config.get("footerShowDate", True)
    
    # Bank information settings
    show_bank_info = header_config.get("showBankInfo", True)
    bank_info_location = header_config.get("bankInfoLocation", "top")  # top, bottom, none
    bank_info_align = (
        header_config.get("bankInfoAlign")
        or header_config.get("logoPosition", "center")
        or "center"
    ).lower()  # left, center, right
    bank_name = header_config.get("bankName", "")
    bank_address = header_config.get("bankAddress", "")
    bank_phone = header_config.get("bankPhone", "")
    bank_website = header_config.get("bankWebsite", "")
    
    # Logo settings
    show_logo = header_config.get("showLogo", True)
    logo_base64 = header_config.get("logoBase64", "")
    logo_position = header_config.get("logoPosition", "left")
    logo_height = header_config.get("logoHeight", 36)
    logo_file = header_config.get("logoFile", None)
    
    # Color and styling settings
    font_color = header_config.get("fontColor", "#1F4E79")
    table_header_bg_color = header_config.get("tableHeaderBgColor", "#1F4E79")
    table_body_bg_color = header_config.get("tableBodyBgColor", "#FFFFFF")
    background_color = header_config.get("backgroundColor", "#FFFFFF")
    border_style = header_config.get("borderStyle", "solid")
    border_color = header_config.get("borderColor", "#E5E7EB")
    border_width = header_config.get("borderWidth", 1)
    
    # Font and size settings
    font_size = header_config.get("fontSize", "medium")
    padding = header_config.get("padding", 20)
    margin = header_config.get("margin", 72)  # 1 inch = 72 points
    
    # Watermark settings
    watermark_enabled = header_config.get("watermarkEnabled", False)
    watermark_text = header_config.get("watermarkText", "CONFIDENTIAL")
    watermark_opacity = header_config.get("watermarkOpacity", 10)
    watermark_diagonal = header_config.get("watermarkDiagonal", True)
    
    # Footer settings
    footer_show_confidentiality = header_config.get("footerShowConfidentiality", True)
    footer_confidentiality_text = header_config.get("footerConfidentialityText", "Confidential Report - Internal Use Only")
    footer_show_page_numbers = header_config.get("footerShowPageNumbers", True)
    footer_align = header_config.get("footerAlign", "center")
    
    # Page settings
    show_page_numbers = header_config.get("showPageNumbers", True)
    location = header_config.get("location", "top")
    
    # Convert hex colors to RGB for python-docx
    def hex_to_rgb(hex_color):
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        except:
            return RGBColor(31, 78, 121)  # Default color
    
    font_color_rgb = hex_to_rgb(font_color)
    header_bg_color_rgb = hex_to_rgb(table_header_bg_color)
    body_bg_color_rgb = hex_to_rgb(table_body_bg_color)
    background_color_rgb = hex_to_rgb(background_color)
    border_color_rgb = hex_to_rgb(border_color)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(margin)
        section.bottom_margin = Pt(margin)
        section.left_margin = Pt(margin)
        section.right_margin = Pt(margin)
    
    # Only add header if includeHeader is True
    if include_header:
        # Add logo image into document header if configured
        if show_logo and (logo_base64 or logo_file):
            try:
                img_bytes_io = None
                if logo_base64:
                    b64_data = logo_base64
                    if ',' in b64_data:
                        b64_data = b64_data.split(',')[1]
                    img_bytes_io = BytesIO(base64.b64decode(b64_data))
                elif logo_file and os.path.exists(logo_file):
                    with open(logo_file, 'rb') as lf:
                        img_bytes_io = BytesIO(lf.read())
                if img_bytes_io:
                    header = doc.sections[0].header
                    # Use existing header paragraph if available, else add one
                    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    run = header_para.add_run()
                    # Preserve aspect ratio by setting only height
                    run.add_picture(img_bytes_io, height=Pt(float(logo_height)))
                    if logo_position == 'left':
                        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif logo_position == 'right':
                        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    else:
                        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass
        # Add bank information at top if configured
        if show_bank_info and bank_name and bank_info_location == "top":
            bank_para = doc.add_paragraph()
            bank_run = bank_para.add_run(f"🏦 {bank_name}")
            bank_run.font.size = Pt(12)
            bank_run.font.bold = True
            bank_run.font.color.rgb = font_color_rgb
            if bank_info_align == 'left':
                bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif bank_info_align == 'right':
                bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if bank_address:
                address_para = doc.add_paragraph()
                address_run = address_para.add_run(bank_address)
                address_run.font.size = Pt(10)
                address_run.font.color.rgb = font_color_rgb
                if bank_info_align == 'left':
                    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif bank_info_align == 'right':
                    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if bank_phone or bank_website:
                contact_info = []
                if bank_phone:
                    contact_info.append(f"Tel: {bank_phone}")
                if bank_website:
                    contact_info.append(f"Web: {bank_website}")
                
                contact_para = doc.add_paragraph()
                contact_run = contact_para.add_run(" | ".join(contact_info))
                contact_run.font.size = Pt(10)
                contact_run.font.color.rgb = font_color_rgb
                if bank_info_align == 'left':
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif bank_info_align == 'right':
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add spacing
            doc.add_paragraph()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = font_color_rgb
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add subtitle
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.italic = True
            subtitle_run.font.color.rgb = font_color_rgb
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add generation date
        if show_date:
            date_para = doc.add_paragraph()
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            date_run = date_para.add_run(f"Generated on: {current_date}")
            date_run.font.size = Pt(10)
            date_run.font.italic = True
            date_run.font.color.rgb = font_color_rgb
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    # Create table with configuration-based styling
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Helper function to shade table cells
    def shade_cell(cell, fill_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color.replace('#', ''))
        tcPr.append(shd)
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = str(col)
        # Style header cell
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Apply header background color
        shade_cell(hdr_cells[i], table_header_bg_color)
    
    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            # Style data cell
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Apply alternating row colors if needed
            if row_idx % 2 == 0:
                shade_cell(row_cells[i], table_body_bg_color)

    # Optional footer totals row for Word
    footer_totals_cols = header_config.get("tableFooterTotals", []) or []
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        totals_row = table.add_row().cells
        for i in range(len(columns)):
            totals_row[i].text = ""
        totals_row[0].text = "Total"
        name_to_index = {str(col): idx for idx, col in enumerate(columns)}
        for col_name in footer_totals_cols:
            if str(col_name) in name_to_index:
                idx = name_to_index[str(col_name)]
                s = 0.0
                for r in data_rows:
                    try:
                        val = r[idx]
                        if val is None or val == "":
                            continue
                        s += float(str(val).replace(',', ''))
                    except Exception:
                        pass
                totals_row[idx].text = f"{s:,.2f}"
        # Bold and align totals row
        for i in range(len(columns)):
            for paragraph in totals_row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if i > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add bank information at bottom if configured
    if show_bank_info and bank_name and bank_info_location == "bottom":
        doc.add_paragraph()  # Add spacing
        
        bank_para = doc.add_paragraph()
        bank_run = bank_para.add_run(f"🏦 {bank_name}")
        bank_run.font.size = Pt(12)
        bank_run.font.bold = True
        bank_run.font.color.rgb = font_color_rgb
        if bank_info_align == 'left':
            bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif bank_info_align == 'right':
            bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if bank_address:
            address_para = doc.add_paragraph()
            address_run = address_para.add_run(bank_address)
            address_run.font.size = Pt(10)
            address_run.font.color.rgb = font_color_rgb
            if bank_info_align == 'left':
                address_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif bank_info_align == 'right':
                address_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                address_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if bank_phone or bank_website:
            contact_info = []
            if bank_phone:
                contact_info.append(f"Tel: {bank_phone}")
            if bank_website:
                contact_info.append(f"Web: {bank_website}")
            
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_info))
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = font_color_rgb
            if bank_info_align == 'left':
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif bank_info_align == 'right':
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add footer if configured
    if footer_show_date or footer_show_confidentiality or footer_show_page_numbers:
        doc.add_paragraph()  # Add spacing
        
        footer_text = []
        if footer_show_date:
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            footer_text.append(f"Generated on: {current_date}")
        if footer_show_confidentiality:
            footer_text.append(footer_confidentiality_text)
        if footer_show_page_numbers:
            footer_text.append("Page &P of &N")
        
        if footer_text:
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run(" | ".join(footer_text))
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            
            if footer_align == "center":
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif footer_align == "left":
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif footer_align == "right":
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add watermark if enabled (single centered line)
    if watermark_enabled:
        try:
            watermark_para = doc.add_paragraph()
            if watermark_diagonal:
                spaced_text = " ".join(watermark_text)
                watermark_run = watermark_para.add_run(spaced_text)
            else:
                watermark_run = watermark_para.add_run(watermark_text)
            watermark_run.font.size = Pt(48)
            watermark_run.font.color.rgb = RGBColor(200, 200, 200)
            watermark_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            watermark_para.space_after = Pt(0)
            watermark_para.space_before = Pt(0)
        except Exception:
            pass
    
    # Save to file
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    date_folder = datetime.now().strftime('%Y-%m-%d')
    os.makedirs(f"reports_export/{date_folder}", exist_ok=True)
    filename = f"dynamic_report_{ts}.docx"
    file_path = f"reports_export/{date_folder}/{filename}"
    
    doc.save(file_path)
    
    # Save export record to database (best-effort)
    try:
        import pyodbc
        from config import get_database_connection_string
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            cursor.execute(
                """
                IF NOT EXISTS (
                  SELECT * FROM INFORMATION_SCHEMA.TABLES 
                  WHERE TABLE_NAME = 'report_exports' AND TABLE_SCHEMA='dbo'
                )
                BEGIN
                  CREATE TABLE dbo.report_exports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    title NVARCHAR(255) NOT NULL,
                    src NVARCHAR(1024) NULL,
                    format NVARCHAR(20) NOT NULL,
                    dashboard NVARCHAR(100) NULL,
                    created_at DATETIME2 DEFAULT GETDATE()
                  )
                END
                """
            )
            conn.commit()
            export_title = header_config.get("title", "Dynamic Report") if header_config else "Dynamic Report"
            export_title = f"{export_title} - {datetime.now().strftime('%Y-%m-%d')}"
            export_dashboard = header_config.get("dashboard", "dynamic") if header_config else "dynamic"
            cursor.execute(
                """
                INSERT INTO dbo.report_exports (title, src, format, dashboard)
                VALUES (?, ?, ?, ?)
                """,
                (export_title, file_path, 'word', export_dashboard)
            )
            conn.commit()
        finally:
            cursor.close()
            conn.close()
    except Exception:
        pass

    with open(file_path, 'rb') as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'X-Export-Src': file_path
        }
    )

def generate_pdf_report(columns, data_rows, header_config=None):
    """Generate PDF report from dynamic data with full header configuration support"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    import os
    import base64
    from io import BytesIO
    try:
        from PIL import Image as PILImage
    except Exception:
        PILImage = None
    
    # Import Arabic text support
    from shared_pdf_utils import shape_text_for_arabic, ARABIC_FONT_NAME
    
    # Get default header config if none provided
    if not header_config:
        from export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    date_folder = datetime.now().strftime('%Y-%m-%d')
    os.makedirs(f"reports_export/{date_folder}", exist_ok=True)
    filename = f"dynamic_report_{ts}.pdf"
    file_path = f"reports_export/{date_folder}/{filename}"
    
    # Extract ALL configuration values from header modal configuration
    # Basic report settings
    include_header = header_config.get("includeHeader", True)
    title = header_config.get("title", "Dynamic Report")
    subtitle = header_config.get("subtitle", "")
    icon = header_config.get("icon", "chart-line")
    
    # Date and time settings
    show_date = header_config.get("showDate", True)
    footer_show_date = header_config.get("footerShowDate", True)
    
    # Bank information settings
    show_bank_info = header_config.get("showBankInfo", True)
    bank_info_location = header_config.get("bankInfoLocation", "top")  # top, bottom, none
    bank_info_align = (
        header_config.get("bankInfoAlign")
        or header_config.get("logoPosition", "center")
        or "center"
    ).lower()
    bank_name = header_config.get("bankName", "")
    bank_address = header_config.get("bankAddress", "")
    bank_phone = header_config.get("bankPhone", "")
    bank_website = header_config.get("bankWebsite", "")
    
    # Logo settings
    show_logo = header_config.get("showLogo", True)
    logo_base64 = header_config.get("logoBase64", "")
    logo_position = header_config.get("logoPosition", "left")
    logo_height = header_config.get("logoHeight", 36)
    logo_file = header_config.get("logoFile", None)
    
    # Color and styling settings
    font_color = header_config.get("fontColor", "#1F4E79")
    table_header_bg_color = header_config.get("tableHeaderBgColor", "#1F4E79")
    table_body_bg_color = header_config.get("tableBodyBgColor", "#FFFFFF")
    background_color = header_config.get("backgroundColor", "#FFFFFF")
    border_style = header_config.get("borderStyle", "solid")
    border_color = header_config.get("borderColor", "#E5E7EB")
    border_width = header_config.get("borderWidth", 1)
    
    # Font and size settings
    font_size = header_config.get("fontSize", "medium")
    padding = header_config.get("padding", 20)
    margin = header_config.get("margin", 72)  # 1 inch = 72 points
    
    # Watermark settings
    watermark_enabled = header_config.get("watermarkEnabled", False)
    watermark_text = header_config.get("watermarkText", "CONFIDENTIAL")
    watermark_opacity = header_config.get("watermarkOpacity", 10)
    watermark_diagonal = header_config.get("watermarkDiagonal", True)
    
    # Footer settings
    footer_show_confidentiality = header_config.get("footerShowConfidentiality", True)
    footer_confidentiality_text = header_config.get("footerConfidentialityText", "Confidential Report - Internal Use Only")
    footer_show_page_numbers = header_config.get("footerShowPageNumbers", True)
    footer_align = header_config.get("footerAlign", "center")
    
    # Page settings
    show_page_numbers = header_config.get("showPageNumbers", True)
    location = header_config.get("location", "top")
    
    # Convert hex colors to ReportLab colors
    def hex_to_color(hex_color):
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        try:
            r = int(hex_color[0:2], 16) / 255.0
            g = int(hex_color[2:4], 16) / 255.0
            b = int(hex_color[4:6], 16) / 255.0
            return colors.Color(r, g, b)
        except:
            return colors.HexColor(f"#{hex_color}")
    
    font_color_rl = hex_to_color(font_color)
    header_bg_color_rl = hex_to_color(table_header_bg_color)
    body_bg_color_rl = hex_to_color(table_body_bg_color)
    background_color_rl = hex_to_color(background_color)
    border_color_rl = hex_to_color(border_color)
    
    # Choose page size based on number of columns
    page_size = A4 if len(columns) <= 6 else letter
    
    # Create document with margins
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=page_size,
        rightMargin=margin,
        leftMargin=margin,
        topMargin=margin,
        bottomMargin=margin
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Create custom styles based on configuration
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        textColor=font_color_rl,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=font_color_rl,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    # Resolve text alignment for bank info
    _pdf_align = TA_CENTER
    if bank_info_align == 'left':
        _pdf_align = TA_LEFT
    elif bank_info_align == 'right':
        _pdf_align = TA_RIGHT

    bank_style = ParagraphStyle(
        'BankInfo',
        parent=styles['Normal'],
        fontSize=10,
        textColor=font_color_rl,
        alignment=_pdf_align,
        spaceAfter=3
    )
    
    date_style = ParagraphStyle(
        'DateInfo',
        parent=styles['Normal'],
        fontSize=10,
        textColor=font_color_rl,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    # Only add header if includeHeader is True
    if include_header:
        # Add logo image if configured
        if show_logo and (logo_base64 or logo_file):
            try:
                img_stream = None
                if logo_base64:
                    b64_data = logo_base64
                    if ',' in b64_data:
                        b64_data = b64_data.split(',')[1]
                    img_bytes = base64.b64decode(b64_data)
                    img_stream = BytesIO(img_bytes)
                elif logo_file and os.path.exists(logo_file):
                    with open(logo_file, 'rb') as lf:
                        img_stream = BytesIO(lf.read())
                if img_stream:
                    width_arg = None
                    height_arg = None
                    if PILImage is not None:
                        img_stream.seek(0)
                        pil_img = PILImage.open(img_stream)
                        orig_w, orig_h = pil_img.size
                        if orig_h > 0:
                            scale = float(logo_height) / float(orig_h)
                            width_arg = orig_w * scale
                            height_arg = logo_height
                        img_stream.seek(0)
                    rl_img = RLImage(img_stream, width=width_arg, height=height_arg)
                    if logo_position == 'left':
                        rl_img.hAlign = 'LEFT'
                    elif logo_position == 'right':
                        rl_img.hAlign = 'RIGHT'
                    else:
                        rl_img.hAlign = 'CENTER'
                    story.append(rl_img)
                    story.append(Spacer(1, 6))
            except Exception:
                pass
        # Add bank information at top if configured
        if show_bank_info and bank_name and bank_info_location == "top":
            story.append(Paragraph(f"🏦 {bank_name}", bank_style))
            if bank_address:
                story.append(Paragraph(bank_address, bank_style))
            if bank_phone or bank_website:
                contact_info = []
                if bank_phone:
                    contact_info.append(f"Tel: {bank_phone}")
                if bank_website:
                    contact_info.append(f"Web: {bank_website}")
                story.append(Paragraph(" | ".join(contact_info), bank_style))
            story.append(Spacer(1, 12))
        
        # Add title
        story.append(Paragraph(title, title_style))
        
        # Add subtitle
        if subtitle:
            story.append(Paragraph(subtitle, subtitle_style))
        
        # Add generation date
        if show_date:
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            story.append(Paragraph(f"Generated on: {current_date}", date_style))
        
        story.append(Spacer(1, 20))
    
    # Generate chart if chart_data is provided
    chart_data = header_config.get('chart_data')
    chart_type = header_config.get('chart_type', 'bar')
    if chart_data and chart_data.get('labels') and chart_data.get('values'):
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            from io import BytesIO
            
            # Create chart
            fig, ax = plt.subplots(figsize=(8, 5))
            
            if chart_type == 'bar':
                ax.barh(chart_data['labels'], chart_data['values'], color='#4472C4')
                ax.set_xlabel('Controls Count')
                ax.set_ylabel('Component')
            elif chart_type == 'pie':
                ax.pie(chart_data['values'], labels=chart_data['labels'], autopct='%1.1f%%')
            
            ax.set_title(title if title else 'Chart')
            plt.tight_layout()
            
            # Save chart to buffer
            chart_buffer = BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=150, bbox_inches='tight')
            chart_buffer.seek(0)
            plt.close()
            
            # Add chart to PDF
            chart_img = RLImage(chart_buffer, width=6*inch, height=3.5*inch)
            chart_img.hAlign = 'CENTER'
            story.append(chart_img)
            story.append(Spacer(1, 20))
        except Exception as e:
            print(f"Error generating chart: {e}")
            pass
    
    # Table data with Arabic text support and multi-line text handling
    # Create styles for table cells
    styles = getSampleStyleSheet()
    
    # Header style
    header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontSize=12,
        fontName=ARABIC_FONT_NAME or 'Helvetica-Bold',
        alignment=TA_CENTER,
        textColor=colors.whitesmoke,
        spaceAfter=6,
        spaceBefore=6
    )
    
    # Data cell style
    data_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontSize=10,
        fontName=ARABIC_FONT_NAME or 'Helvetica',
        alignment=TA_CENTER,
        spaceAfter=4,
        spaceBefore=4,
        leading=12
    )
    
    # Process columns with Paragraph objects for multi-line support
    processed_columns = [Paragraph(shape_text_for_arabic(str(col)), header_style) for col in columns]
    
    # Process data rows with Paragraph objects for multi-line support
    processed_data_rows = []
    for row in data_rows:
        processed_row = [Paragraph(shape_text_for_arabic(str(cell)), data_style) for cell in row]
        processed_data_rows.append(processed_row)
    
    table_data = [processed_columns] + processed_data_rows

    # Optional footer totals row for PDF
    footer_totals_cols = header_config.get("tableFooterTotals", []) or []
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        name_to_index = {str(col): idx for idx, col in enumerate(columns)}
        totals_row = [Paragraph("", data_style)] * len(columns)
        totals_row[0] = Paragraph(shape_text_for_arabic("Total"), data_style)
        for col_name in footer_totals_cols:
            if str(col_name) in name_to_index:
                idx = name_to_index[str(col_name)]
                s = 0.0
                for r in data_rows:
                    try:
                        val = r[idx]
                        if val is None or val == "":
                            continue
                        s += float(str(val).replace(',', ''))
                    except Exception:
                        pass
                totals_row[idx] = Paragraph(shape_text_for_arabic(f"{s:,.2f}"), data_style)
        table_data.append(totals_row)
    
    # Create table with configuration-based styling and column widths
    # Calculate column widths to leave some margin
    num_cols = len(table_data[0]) if table_data else 1
    available_width = page_size[0] - (margin * 2)  # Subtract left and right margins
    col_width = available_width / num_cols if num_cols > 0 else available_width
    col_widths = [col_width] * num_cols
    
    table = Table(table_data, repeatRows=1, colWidths=col_widths)
    
    # Build table style based on configuration with Arabic font support
    table_style = [
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), header_bg_color_rl),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), ARABIC_FONT_NAME or 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        
        # Data rows styling
        ('BACKGROUND', (0, 1), (-1, -1), body_bg_color_rl),
        ('FONTNAME', (0, 1), (-1, -1), ARABIC_FONT_NAME or 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        
        # Grid lines
        ('GRID', (0, 0), (-1, -1), border_width, border_color_rl),
        
        # Alternating row colors if needed
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
    ]

    # If totals row exists, style it like a bold footer
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        last_row_index = len(table_data) - 1
        table_style += [
            ('BACKGROUND', (0, last_row_index), (-1, last_row_index), header_bg_color_rl),
            ('TEXTCOLOR', (0, last_row_index), (-1, last_row_index), colors.whitesmoke),
            ('FONTNAME', (0, last_row_index), (-1, last_row_index), 'Helvetica-Bold'),
        ]
    
    table.setStyle(TableStyle(table_style))
    story.append(table)
    
    # Add watermark if enabled
    if watermark_enabled:
        try:
            from export_utils import add_watermark_to_pdf
            add_watermark_to_pdf(story, header_config)
        except Exception as e:
            pass  # Continue without watermark if there's an error
    
    # Add bank information at bottom if configured
    if show_bank_info and bank_name and bank_info_location == "bottom":
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"🏦 {bank_name}", bank_style))
        if bank_address:
            story.append(Paragraph(bank_address, bank_style))
        if bank_phone or bank_website:
            contact_info = []
            if bank_phone:
                contact_info.append(f"Tel: {bank_phone}")
            if bank_website:
                contact_info.append(f"Web: {bank_website}")
            story.append(Paragraph(" | ".join(contact_info), bank_style))

    # Add footer elements
    if footer_show_date or footer_show_confidentiality or footer_show_page_numbers:
        story.append(Spacer(1, 20))
        
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        
        footer_text = []
        if footer_show_date:
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            footer_text.append(f"Generated on: {current_date}")
        if footer_show_confidentiality:
            footer_text.append(footer_confidentiality_text)
        if footer_show_page_numbers:
            footer_text.append("Page &P of &N")
        
        if footer_text:
            story.append(Paragraph(" | ".join(footer_text), footer_style))
    
    # Build the PDF
    doc.build(story)
    
    # Save export record to database (best-effort)
    try:
        import pyodbc
        from config import get_database_connection_string
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            cursor.execute(
                """
                IF NOT EXISTS (
                  SELECT * FROM INFORMATION_SCHEMA.TABLES 
                  WHERE TABLE_NAME = 'report_exports' AND TABLE_SCHEMA='dbo'
                )
                BEGIN
                  CREATE TABLE dbo.report_exports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    title NVARCHAR(255) NOT NULL,
                    src NVARCHAR(1024) NULL,
                    format NVARCHAR(20) NOT NULL,
                    dashboard NVARCHAR(100) NULL,
                    created_at DATETIME2 DEFAULT GETDATE()
                  )
                END
                """
            )
            conn.commit()
            export_title = header_config.get("title", "Dynamic Report") if header_config else "Dynamic Report"
            export_title = f"{export_title} - {datetime.now().strftime('%Y-%m-%d')}"
            export_dashboard = header_config.get("dashboard", "dynamic") if header_config else "dynamic"
            cursor.execute(
                """
                INSERT INTO dbo.report_exports (title, src, format, dashboard)
                VALUES (?, ?, ?, ?)
                """,
                (export_title, file_path, 'pdf', export_dashboard)
            )
            conn.commit()
        finally:
            cursor.close()
            conn.close()
    except Exception:
        pass

    with open(file_path, 'rb') as f:
        content = f.read()
    
    return content

def generate_excel_report(columns, data_rows, header_config=None):
    """Generate Excel report from dynamic data with chart support"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image
    import io
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import os
    from datetime import datetime
    
    # Get default header config if none provided
    if not header_config:
        from export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    
    # Create workbook and worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    
    # Add header information
    row = 1
    if header_config.get("includeHeader", True):
        title = header_config.get("title", "Report")
        ws[f'A{row}'] = title
        ws[f'A{row}'].font = Font(size=16, bold=True)
        row += 1
        
        if header_config.get("subtitle"):
            ws[f'A{row}'] = header_config.get("subtitle")
            ws[f'A{row}'].font = Font(size=12)
            row += 1
    
    # Add date if requested
    if header_config.get("showDate", True):
        ws[f'A{row}'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws[f'A{row}'].font = Font(size=10, italic=True)
        row += 2
    
    # Add data table on the left side
    table_start_row = row
    if data_rows and len(data_rows) > 0:
        # Add table headers
        if isinstance(columns[0], dict):
            headers = [col['label'] for col in columns]
        else:
            headers = columns
        
        for i, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=i, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
        
        row += 1
        
        # Add data rows
        for data_row in data_rows:
            for i, value in enumerate(data_row, 1):
                ws.cell(row=row, column=i, value=value)
            row += 1
    
    # Add chart on the right side if chart_data is provided
    if header_config.get("chart_data") and header_config.get("chart_type"):
        chart_data = header_config["chart_data"]
        chart_type = header_config["chart_type"]
        
        if chart_data and len(chart_data) > 0 and isinstance(chart_data, list):
            # Create chart
            plt.figure(figsize=(10, 6))
            
            if chart_type == 'pie':
                labels = [item.get('name', 'Unknown') for item in chart_data if isinstance(item, dict)]
                values = [item.get('value', 0) for item in chart_data if isinstance(item, dict)]
                if labels and values:
                    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
                    plt.title(header_config.get("title", "Chart"))
            elif chart_type == 'bar':
                labels = [item.get('name', 'Unknown') for item in chart_data if isinstance(item, dict)]
                values = [item.get('value', 0) for item in chart_data if isinstance(item, dict)]
                if labels and values:
                    plt.bar(labels, values)
                    plt.title(header_config.get("title", "Chart"))
                    plt.xticks(rotation=45)
            elif chart_type == 'line':
                labels = [item.get('name', 'Unknown') for item in chart_data if isinstance(item, dict)]
                values = [item.get('value', 0) for item in chart_data if isinstance(item, dict)]
                if labels and values:
                    plt.plot(labels, values, marker='o')
                    plt.title(header_config.get("title", "Chart"))
                    plt.xticks(rotation=45)
            
            plt.tight_layout()
            
            # Save chart to BytesIO
            chart_buffer = io.BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight')
            chart_buffer.seek(0)
            plt.close()
            
            # Add chart to Excel on the right side (column H)
            img = Image(chart_buffer)
            img.width = 600
            img.height = 400
            ws.add_image(img, f'H{table_start_row}')
            
            # Adjust row to be the maximum of table and chart
            row = max(row, table_start_row + 25)
    
    # Set column widths
    for i in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(i)].width = 20
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return output.getvalue()

@router.post("/api/reports/schedule")
async def save_report_schedule(request: Request):
    """Save scheduled report configuration"""
    try:
        body = await request.json()
        report_config = body.get('reportConfig', {})
        schedule = body.get('schedule', {})
        
        # Save to database (you can create a scheduled_reports table)
        import pyodbc
        from config import get_database_connection_string
        
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            # Create table if not exists
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='scheduled_reports' and xtype='U')
                CREATE TABLE scheduled_reports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    report_config NVARCHAR(MAX) NOT NULL,
                    schedule_config NVARCHAR(MAX) NOT NULL,
                    is_active BIT DEFAULT 1,
                    created_at DATETIME2 DEFAULT GETDATE()
                );
            """)
            conn.commit()
            
            # Insert schedule
            import json
            cursor.execute("""
                INSERT INTO scheduled_reports (report_config, schedule_config)
                VALUES (?, ?)
            """, json.dumps(report_config), json.dumps(schedule))
            conn.commit()
            
            return {"success": True, "message": "Schedule saved successfully"}
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save schedule: {str(e)}")

@router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# Bank Check Processing: upload PDF + form and return generated files
@router.post("/api/reports/bank-check")
async def create_bank_check_report(
    request: Request,
    file: UploadFile = File(...),
    bankName: str = Form(None),
    cost: str = Form(None),
    date: str = Form(None),
    daysRemaining: str = Form(None),
    format: str = Form("excel")
):
    try:
        content = await file.read()
        excel_bytes, word_bytes = bank_check_service.process(content, {
            "bankName": bankName,
            "cost": cost,
            "date": date,
            "daysRemaining": daysRemaining,
        })

        # allow query param override: ?format=excel|word|zip
        fmt = request.query_params.get('format') or (format or 'excel')
        fmt = (fmt or 'excel').lower()

        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        if fmt == 'excel':
            filename = f"bank_check_{ts}.xlsx"
            return Response(
                content=excel_bytes,
                media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                headers={'Content-Disposition': f'attachment; filename="{filename}"'}
            )
        elif fmt == 'word':
            filename = f"bank_check_{ts}.docx"
            return Response(
                content=word_bytes,
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': f'attachment; filename="{filename}"'}
            )
        else:
            # package into a simple zip in-memory
            import zipfile
            from io import BytesIO
            buffer = BytesIO()
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(f"bank_check_{ts}.xlsx", excel_bytes)
                zf.writestr(f"bank_check_{ts}.docx", word_bytes)
            buffer.seek(0)
            filename = f"bank_check_{ts}.zip"
            return Response(content=buffer.getvalue(), media_type='application/zip', headers={'Content-Disposition': f'attachment; filename="{filename}"'})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create bank check report: {str(e)}")

# Debug: return OCR text and parsed fields without generating files
@router.post("/api/reports/bank-check/debug")
async def debug_bank_check(
    file: UploadFile = File(...)
):
    try:
        content = await file.read()
        webhook_url = os.getenv('OCR_WEBHOOK_URL')
        if webhook_url:
            async with httpx.AsyncClient(timeout=60) as client:
                files = {
                    'file': (file.filename or 'upload.bin', content, file.content_type or 'application/octet-stream')
                }
                resp = await client.post(webhook_url, files=files)
                if resp.status_code == 200:
                    data = resp.json()
                    # Best-effort mapping for UI
                    mapped_fields = {
                        'bankName': data.get('bankName') or data.get('bank') or data.get('bank_name') or '',
                        'branch': data.get('branch') or '',
                        'currency': data.get('currency') or data.get('ccy') or '',
                        'amountNumeric': data.get('amountNumeric') or data.get('amount') or data.get('amount_number') or '',
                        'amountWords': data.get('amountWords') or data.get('amount_text') or '',
                        'date': data.get('date') or data.get('checkDate') or '',
                        'payee': data.get('payee') or data.get('to') or data.get('beneficiary') or '',
                    }
                    text_value = data.get('text') or data.get('textSnippet') or ''
                    return {
                        'textSnippet': text_value[:1000] if text_value else '',
                        'textLength': len(text_value or ''),
                        'fields': mapped_fields,
                        'diagnostics': {
                            'source': 'webhook',
                            'webhookUrl': webhook_url,
                            'status': resp.status_code
                        }
                    }
                else:
                    raise HTTPException(status_code=502, detail=f"OCR webhook failed with status {resp.status_code}")

        # Fallback: local OCR if webhook not configured
        full_text = bank_check_service.extract_text(content)
        fields = bank_check_service.extract_fields_from_pdf(content)
        diag = bank_check_service.diagnose_extraction(content)
        return {
            "textSnippet": (full_text[:1000] if full_text else ""),
            "textLength": len(full_text or ""),
            "fields": fields,
            "diagnostics": diag
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Debug extraction failed: {str(e)}")

# Enhanced Bank Check Processing with professional reports
@router.post("/api/reports/enhanced-bank-check")
async def create_enhanced_bank_check_report(
    request: Request
):
    """Process bank check records and generate professional Excel/Word reports"""
    try:
        # Check if it's a file upload or JSON data
        content_type = request.headers.get('content-type', '')
        
        if 'multipart/form-data' in content_type:
            # Handle file upload
            form = await request.form()
            file = form.get('file')
            if not file:
                raise HTTPException(status_code=400, detail="No file provided")
            
            content = await file.read()
            excel_bytes, word_bytes, data = await enhanced_bank_check_service.process_check(content, file.filename or 'check.pdf')
            
            # Create ZIP with both files
            import zipfile
            from io import BytesIO
            import os
            
            buffer = BytesIO()
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            date_folder = datetime.now().strftime('%Y-%m-%d')
            
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(f"bank_check_analysis_{ts}.xlsx", excel_bytes)
                zf.writestr(f"bank_check_report_{ts}.docx", word_bytes)
            buffer.seek(0)
            
            # Save to persistent storage
            os.makedirs(f"exports/{date_folder}", exist_ok=True)
            filename = f"bank_check_reports_{ts}.zip"
            file_path = f"exports/{date_folder}/{filename}"
            
            with open(file_path, 'wb') as f:
                f.write(buffer.getvalue())
            
            return Response(
                content=buffer.getvalue(), 
                media_type='application/zip', 
                headers={
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'X-Export-Src': file_path
                }
            )
        
        else:
            # Handle JSON data with dynamic records (arbitrary headers)
            body = await request.json()
            records = body.get('records', [])
            columns = body.get('columns')  # optional explicit headers
            rows = body.get('rows')        # optional explicit rows
            format_type = body.get('format', 'excel')

            if format_type == 'word':
                # If explicit columns/rows provided, build a simple Word report with a table
                if columns and rows:
                    from io import BytesIO
                    from docx import Document
                    from docx.shared import Pt
                    from docx.oxml.ns import qn
                    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    from docx.oxml import OxmlElement

                    def shade_cell(cell, fill="E6F0FF"):
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), fill)
                        tcPr.append(shd)

                    doc = Document()
                    # Base font
                    doc.styles['Normal'].font.name = 'Arial'
                    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

                    # Title
                    title = doc.add_paragraph()
                    run = title.add_run('تقرير سجلات الشيكات / Bank Check Records Report')
                    run.font.size = Pt(16)
                    run.bold = True
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Meta
                    meta = doc.add_paragraph()
                    meta.add_run(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").font.size = Pt(10)
                    meta.add_run(f"Number of Records: {len(rows)}\n").font.size = Pt(10)

                    # Table
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(columns):
                        hdr_cells[i].text = str(col)
                        shade_cell(hdr_cells[i], 'E6F0FF')
                        for p in hdr_cells[i].paragraphs:
                            p.runs and setattr(p.runs[0].font, 'bold', True)
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    for r in rows:
                        row_cells = table.add_row().cells
                        for i, val in enumerate(r):
                            row_cells[i].text = str(val)
                            for p in row_cells[i].paragraphs:
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    # Footer summary
                    doc.add_paragraph().add_run('\n')
                    summary = doc.add_paragraph()
                    summary.add_run('ملخص التقرير / Report Summary').bold = True
                    summary.add_run(f"\nإجمالي عدد السجلات: {len(rows)}")

                    out = BytesIO()
                    doc.save(out)
                    out.seek(0)
                    
                    # Save to persistent storage
                    import os
                    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                    date_folder = datetime.now().strftime('%Y-%m-%d')
                    os.makedirs(f"exports/{date_folder}", exist_ok=True)
                    filename = f"bank_check_report_{ts}.docx"
                    file_path = f"exports/{date_folder}/{filename}"
                    
                    with open(file_path, 'wb') as f:
                        f.write(out.getvalue())
                    
                    return Response(
                        content=out.getvalue(), 
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                        headers={
                            'Content-Disposition': f'attachment; filename="{filename}"',
                            'X-Export-Src': file_path
                        }
                    )

                if not records:
                    raise HTTPException(status_code=400, detail="No records provided for Word export")
                # Fallback to service template
                _, word_bytes = await enhanced_bank_check_service.process_records(records)
                
                # Save to persistent storage
                import os
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                date_folder = datetime.now().strftime('%Y-%m-%d')
                os.makedirs(f"exports/{date_folder}", exist_ok=True)
                filename = f"bank_check_report_{ts}.docx"
                file_path = f"exports/{date_folder}/{filename}"
                
                with open(file_path, 'wb') as f:
                    f.write(word_bytes)
                
                return Response(
                    content=word_bytes, 
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                    headers={
                        'Content-Disposition': f'attachment; filename="{filename}"',
                        'X-Export-Src': file_path
                    }
                )
            else:
                # Build a well-formatted Excel dynamically from provided data (robust to arbitrary labels)
                from io import BytesIO
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
                from openpyxl.utils import get_column_letter

                wb = Workbook()
                ws = wb.active
                ws.title = 'Bank Check Records'

                # Determine headers and rows
                if rows and columns:
                    headers = columns
                    data_rows = rows
                else:
                    if not records:
                        raise HTTPException(status_code=400, detail="No records provided")
                    headers = list(records[0].keys())
                    seen = set(headers)
                    for rec in records[1:]:
                        for k in rec.keys():
                            if k not in seen:
                                headers.append(k)
                                seen.add(k)
                    data_rows = [[str(rec.get(h, '')) for h in headers] for rec in records]

                # Custom header block
                title_font = Font(name='Calibri', size=14, bold=True, color='003366')
                meta_font = Font(name='Calibri', size=11)
                ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(headers)))
                ws.cell(row=1, column=1, value='Bank Check Records Report').font = title_font
                ws.cell(row=1, column=1).alignment = Alignment(horizontal='center')
                ws.cell(row=2, column=1, value=f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font = meta_font
                ws.cell(row=3, column=1, value=f"Number of Records: {len(data_rows)}").font = meta_font

                # Table header row
                header_row_idx = 5
                thin = Side(border_style='thin', color='CCCCCC')
                border = Border(top=thin, left=thin, right=thin, bottom=thin)
                fill = PatternFill('solid', fgColor='E6F0FF')

                for idx, h in enumerate(headers, start=1):
                    c = ws.cell(row=header_row_idx, column=idx, value=str(h))
                    c.font = Font(bold=True)
                    c.fill = fill
                    c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    c.border = border

                # Data rows
                for ridx, r in enumerate(data_rows, start=header_row_idx + 1):
                    for cidx, val in enumerate(r, start=1):
                        c = ws.cell(row=ridx, column=cidx, value=str(val))
                        c.alignment = Alignment(vertical='top')
                        c.border = border

                # Column widths
                for i in range(1, len(headers) + 1):
                    max_len = max(
                        [len(str(headers[i-1]))] + [len(str(row[i-1])) for row in data_rows if len(row) >= i]
                    )
                    ws.column_dimensions[get_column_letter(i)].width = min(max(12, max_len + 2), 40)

                buf = BytesIO()
                wb.save(buf)
                buf.seek(0)
                
                # Save to persistent storage
                import os
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                date_folder = datetime.now().strftime('%Y-%m-%d')
                os.makedirs(f"exports/{date_folder}", exist_ok=True)
                filename = f"bank_check_records_{ts}.xlsx"
                file_path = f"exports/{date_folder}/{filename}"
                
                with open(file_path, 'wb') as f:
                    f.write(buf.getvalue())
                
                return Response(
                    content=buf.getvalue(), 
                    media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                    headers={
                        'Content-Disposition': f'attachment; filename="{filename}"',
                        'X-Export-Src': file_path
                    }
                )
                
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Enhanced bank check processing failed: {str(e)}")

@router.post("/api/reports/enhanced-bank-check/preview")
async def preview_enhanced_bank_check(
    file: UploadFile = File(...)
):
    """Preview extracted data from bank check without generating files"""
    try:
        content = await file.read()
        data = await enhanced_bank_check_service.extract_check_data(content, file.filename or 'check.pdf')
        return {
            "success": True,
            "data": data,
            "extraction_summary": {
                "total_fields": 10,
                "extracted_fields": sum(1 for v in data.values() if v and v != ''),
                "missing_fields": sum(1 for v in data.values() if not v or v == ''),
                "success_rate": f"{(sum(1 for v in data.values() if v and v != '') / 10 * 100):.1f}%"
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview failed: {str(e)}")

@router.post("/api/reports/enhanced-bank-check/extract-headers")
async def extract_excel_headers(
    file: UploadFile = File(...)
):
    """Extract headers from Excel template file"""
    try:
        content = await file.read()
        headers = await enhanced_bank_check_service.extract_excel_headers(content, file.filename or 'template.xlsx')
        return {
            "success": True,
            "headers": headers,
            "count": len(headers)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Header extraction failed: {str(e)}")

@router.post("/api/reports/enhanced-bank-check/insert")
async def insert_check_record(
    request: Request
):
    """Insert check record into database"""
    try:
        import pyodbc
        from config import get_database_connection_string
        
        body = await request.json()
        record = body.get('record', {})
        table_name = body.get('table_name', 'bank_checks')
        
        if not record:
            return {
                "success": False,
                "error": "No record data provided"
            }
        
        # Use existing SQL Server connection
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            # Get table columns to validate the record
            columns_query = """
                SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE
                FROM INFORMATION_SCHEMA.COLUMNS 
                WHERE TABLE_NAME = ? AND COLUMN_NAME != 'id'
                ORDER BY ORDINAL_POSITION
            """
            cursor.execute(columns_query, table_name)
            columns = cursor.fetchall()
            
            if not columns:
                return {
                    "success": False,
                    "error": f"Table '{table_name}' not found"
                }
            
            # Prepare insert statement
            column_names = [col[0] for col in columns]
            placeholders = ', '.join(['?' for _ in column_names])
            column_list = ', '.join([f'[{col}]' for col in column_names])
            
            insert_query = f"""
                INSERT INTO [{table_name}] ({column_list})
                VALUES ({placeholders})
            """
            
            # Helper to coerce python value to SQL Server expected type
            from datetime import datetime, date, time
            def coerce_value(sql_type: str, val):
                if val is None:
                    return None
                if isinstance(val, str) and val.strip() == "":
                    return None
                t = (sql_type or '').lower()
                try:
                    if t in ['int', 'bigint', 'smallint', 'tinyint']:
                        # Extract digits if possible
                        if isinstance(val, (int, float)):
                            return int(val)
                        if isinstance(val, str):
                            return int(val.strip())
                    if t in ['float', 'real']:
                        return float(val)
                    if t in ['decimal', 'numeric', 'money', 'smallmoney']:
                        from decimal import Decimal
                        return Decimal(str(val))
                    if t in ['bit']:
                        if isinstance(val, bool):
                            return 1 if val else 0
                        if isinstance(val, (int, float)):
                            return 1 if int(val) != 0 else 0
                        if isinstance(val, str):
                            return 1 if val.strip().lower() in ['1','true','yes','y'] else 0
                    if t in ['date']:
                        if isinstance(val, (datetime, date)):
                            return val if isinstance(val, date) and not isinstance(val, datetime) else val.date()
                        if isinstance(val, str):
                            # isoformat yyyy-mm-dd
                            return datetime.fromisoformat(val.strip()).date()
                    if t in ['datetime', 'datetime2', 'smalldatetime']:
                        if isinstance(val, (datetime, date)):
                            return val if isinstance(val, datetime) else datetime.combine(val, time())
                        if isinstance(val, str):
                            # try parse ISO; append time if date-only
                            s = val.strip()
                            if len(s) == 10:
                                return datetime.fromisoformat(s + 'T00:00:00')
                            return datetime.fromisoformat(s)
                    if t in ['time']:
                        if isinstance(val, time):
                            return val
                        if isinstance(val, str):
                            return time.fromisoformat(val.strip())
                    # Text-likes as NVARCHAR/VARCHAR
                    return str(val)
                except Exception:
                    # If coercion fails, return None to let NULL insert where allowed
                    return None

            # Prepare values in the correct order with type coercion
            values = []
            for col_name, _, is_nullable in columns:
                raw_value = record.get(col_name, None)
                coerced = coerce_value(_, raw_value)  # _ is DATA_TYPE from SELECT
                if coerced is None and is_nullable == 'NO':
                    return {
                        "success": False,
                        "error": f"Required field '{col_name}' is missing or invalid type (expected {_.lower()})"
                    }
                values.append(coerced)
            
            # Execute insert
            cursor.execute(insert_query, values)
            conn.commit()
            
            # Get the inserted record ID
            record_id = cursor.execute("SELECT @@IDENTITY").fetchone()[0]
            
            return {
                "success": True,
                "message": f"Record inserted successfully into {table_name}",
                "record_id": record_id,
                "table_name": table_name
            }
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        print(f"[EnhancedBankCheck] Insert error: {e}")
        return {
            "success": False,
            "error": f"Database insertion failed: {str(e)}"
        }

@router.get("/api/reports/enhanced-bank-check/tables")
async def get_database_tables():
    """Get all tables from SQL Server database"""
    try:
        import pyodbc
        from config import get_database_connection_string
        
        # Use existing SQL Server connection
        connection_string = get_database_connection_string()
        print(f"[EnhancedBankCheck] Connecting to SQL Server with: {connection_string[:50]}...")
        
        # Connect to SQL Server database
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            # Get all tables from the database
            tables_query = """
                SELECT 
                    TABLE_SCHEMA,
                    TABLE_NAME,
                    TABLE_TYPE
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_TYPE = 'BASE TABLE'
                ORDER BY TABLE_SCHEMA, TABLE_NAME;
            """
            
            cursor.execute(tables_query)
            tables_result = cursor.fetchall()
            
            tables = []
            for table in tables_result:  # Get ALL tables from database
                table_name = table[1]  # TABLE_NAME
                schema_name = table[0]  # TABLE_SCHEMA
                full_table_name = f"{schema_name}.{table_name}" if schema_name != 'dbo' else table_name
                
                # Get columns for each table
                columns_query = """
                    SELECT 
                        COLUMN_NAME,
                        DATA_TYPE,
                        IS_NULLABLE,
                        COLUMN_DEFAULT
                    FROM INFORMATION_SCHEMA.COLUMNS 
                    WHERE TABLE_SCHEMA = ? AND TABLE_NAME = ?
                    ORDER BY ORDINAL_POSITION;
                """
                
                cursor.execute(columns_query, schema_name, table_name)
                columns_result = cursor.fetchall()
                column_names = [col[0] for col in columns_result]  # COLUMN_NAME
                
                tables.append({
                    "id": f"{schema_name}_{table_name}",
                    "name": table_name,
                    "schema": schema_name,
                    "full_name": full_table_name,
                    "fields": column_names,
                    "columns": column_names,  # Add columns for frontend compatibility
                    "field_count": len(column_names)
                })
            
            print(f"[EnhancedBankCheck] Found {len(tables)} tables in SQL Server database")
            
            return {
                "success": True,
                "tables": tables,
                "count": len(tables)
            }
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        print(f"[EnhancedBankCheck] SQL Server connection error: {e}")
        # Fallback to sample data if database connection fails
        sample_tables = [
            {
                "id": "dbo_bank_checks",
                "name": "bank_checks",
                "schema": "dbo",
                "full_name": "bank_checks",
                "fields": ["id", "bank_name", "date", "payee_name", "amount_value", "amount_text", "currency", "status_note", "issuer_signature", "created_at"],
                "columns": ["id", "bank_name", "date", "payee_name", "amount_value", "amount_text", "currency", "status_note", "issuer_signature", "created_at"],
                "field_count": 10
            },
            {
                "id": "dbo_customer_records", 
                "name": "customer_records",
                "schema": "dbo",
                "full_name": "customer_records",
                "fields": ["id", "customer_name", "project_name", "building_number", "apartment_number", "check_number", "due_date", "collection_date", "remaining_days", "collection_status", "total_receivables", "created_at"],
                "columns": ["id", "customer_name", "project_name", "building_number", "apartment_number", "check_number", "due_date", "collection_date", "remaining_days", "collection_status", "total_receivables", "created_at"],
                "field_count": 12
            }
        ]
        
        return {
            "success": True,
            "tables": sample_tables,
            "count": len(sample_tables),
            "note": "Using fallback data - SQL Server connection failed",
            "error": str(e)
        }

@router.get("/api/reports/enhanced-bank-check/test-db")
async def test_database_connection():
    """Test SQL Server database connection"""
    try:
        import pyodbc
        from config import get_database_connection_string
        
        # Use existing SQL Server connection
        connection_string = get_database_connection_string()
        print(f"[DB Test] Attempting SQL Server connection...")
        
        # Test connection
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        # Test query
        cursor.execute("SELECT @@VERSION")
        result = cursor.fetchone()[0]
        
        # Get table count
        cursor.execute("SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_TYPE = 'BASE TABLE'")
        table_count = cursor.fetchone()[0]
        
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": "SQL Server connection successful",
            "version": result,
            "table_count": table_count,
            "database": "NEWDCC-V4-UAT"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "database": "NEWDCC-V4-UAT"
        }

@router.post("/api/reports/enhanced-bank-check/check-table")
async def check_table_exists(request: Request):
    """Check if a table exists in the database"""
    try:
        import pyodbc
        from config import get_database_connection_string
        
        body = await request.json()
        table_name = body.get('table_name', '').strip()
        
        if not table_name:
            return {
                "success": False,
                "error": "Table name is required"
            }
        
        # Use existing SQL Server connection
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            # Check if table exists
            check_query = """
                SELECT COUNT(*) 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_NAME = ? AND TABLE_TYPE = 'BASE TABLE'
            """
            
            cursor.execute(check_query, table_name)
            exists = cursor.fetchone()[0] > 0
            
            if exists:
                # Get table info if it exists
                info_query = """
                    SELECT 
                        TABLE_SCHEMA,
                        TABLE_NAME,
                        (SELECT COUNT(*) FROM INFORMATION_SCHEMA.COLUMNS 
                         WHERE TABLE_SCHEMA = t.TABLE_SCHEMA AND TABLE_NAME = t.TABLE_NAME) as COLUMN_COUNT
                    FROM INFORMATION_SCHEMA.TABLES t
                    WHERE TABLE_NAME = ? AND TABLE_TYPE = 'BASE TABLE'
                """
                cursor.execute(info_query, table_name)
                table_info = cursor.fetchone()
                
                return {
                    "success": True,
                    "exists": True,
                    "table_name": table_name,
                    "schema": table_info[0] if table_info else 'dbo',
                    "column_count": table_info[2] if table_info else 0,
                    "message": f"Table '{table_name}' already exists"
                }
            else:
                return {
                    "success": True,
                    "exists": False,
                    "table_name": table_name,
                    "message": f"Table '{table_name}' does not exist - can be created"
                }
                
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@router.post("/api/reports/enhanced-bank-check/create-table")
async def create_table(request: Request):
    """Create a new table in the database"""
    try:
        import pyodbc
        from config import get_database_connection_string
        
        body = await request.json()
        table_name = body.get('table_name', '').strip()
        fields = body.get('fields', [])
        record_data = body.get('record_data', {})
        
        if not table_name:
            return {
                "success": False,
                "error": "Table name is required"
            }
        
        if not fields:
            return {
                "success": False,
                "error": "At least one field is required"
            }
        
        # Use existing SQL Server connection
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            # First check if table exists
            check_query = """
                SELECT COUNT(*) 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_NAME = ? AND TABLE_TYPE = 'BASE TABLE'
            """
            
            cursor.execute(check_query, table_name)
            exists = cursor.fetchone()[0] > 0
            
            if exists:
                return {
                    "success": False,
                    "error": f"Table '{table_name}' already exists",
                    "exists": True
                }
            
            # Create table with fields
            # Add ID column as primary key
            columns = ["id INT IDENTITY(1,1) PRIMARY KEY"]
            
            for field in fields:
                field_name = field.get('name', '').strip()
                field_type = field.get('type', 'NVARCHAR(255)')
                is_required = field.get('required', False)

                if field_name:
                    # Sanitize field name: replace spaces with underscore, then remove other special chars
                    safe_name = field_name.replace(' ', '_')
                    clean_name = ''.join(c for c in safe_name if c.isalnum() or c in '_')
                    if not clean_name:
                        clean_name = f"field_{len(columns)}"

                    null_constraint = "NOT NULL" if is_required else "NULL"
                    columns.append(f"[{clean_name}] {field_type} {null_constraint}")
            
            # Add created_at timestamp
            columns.append("[created_at] DATETIME2 DEFAULT GETDATE()")
            
            create_query = f"""
                CREATE TABLE [{table_name}] (
                    {', '.join(columns)}
                )
            """
            
            cursor.execute(create_query)
            conn.commit()
            
            # Insert record data if provided
            record_id = None
            if record_data:
                try:
                    # Get table columns to validate the record
                    columns_query = """
                        SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE
                        FROM INFORMATION_SCHEMA.COLUMNS 
                        WHERE TABLE_NAME = ? AND COLUMN_NAME != 'id'
                        ORDER BY ORDINAL_POSITION
                    """
                    cursor.execute(columns_query, table_name)
                    columns = cursor.fetchall()
                    
                    if columns:
                        # Prepare insert statement
                        column_names = [col[0] for col in columns]
                        placeholders = ', '.join(['?' for _ in column_names])
                        column_list = ', '.join([f'[{col}]' for col in column_names])
                        
                        insert_query = f"""
                            INSERT INTO [{table_name}] ({column_list})
                            VALUES ({placeholders})
                        """
                        
                        # Prepare values in the correct order
                        values = []
                        for col_name in column_names:
                            # Find matching field by name (case-insensitive)
                            field_name = None
                            for field in fields:
                                _raw = field.get('name', '')
                                _safe = _raw.replace(' ', '_')
                                clean_field_name = ''.join(c for c in _safe if c.isalnum() or c in '_')
                                if clean_field_name.lower() == col_name.lower():
                                    field_name = field.get('name', '')
                                    break
                            
                            if field_name and field_name in record_data:
                                value = record_data[field_name]
                            else:
                                value = None
                            
                            if value is None and any(col[0] == col_name and col[2] == 'NO' for col in columns):
                                # Required field is missing, skip insertion
                                print(f"[EnhancedBankCheck] Warning: Required field '{col_name}' is missing, skipping record insertion")
                                record_id = None
                                break
                            values.append(value)
                        
                        if values:  # Only insert if we have valid values
                            cursor.execute(insert_query, values)
                            conn.commit()
                            
                            # Get the inserted record ID
                            record_id = cursor.execute("SELECT @@IDENTITY").fetchone()[0]
                            print(f"[EnhancedBankCheck] Record inserted with ID: {record_id}")
                        
                except Exception as insert_error:
                    print(f"[EnhancedBankCheck] Error inserting record: {insert_error}")
                    # Don't fail the table creation if record insertion fails
                    pass
            
            return {
                "success": True,
                "message": f"Table '{table_name}' created successfully" + (f" and record inserted (ID: {record_id})" if record_id else ""),
                "table_name": table_name,
                "field_count": len(fields),
                "created_fields": [f.get('name', '') for f in fields],
                "record_id": record_id,
                "record_inserted": record_id is not None
            }
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }
