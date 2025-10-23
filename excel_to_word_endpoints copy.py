# Top-level imports
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import FileResponse
from typing import Optional
import os
import json
from datetime import datetime
from io import BytesIO
import logging

from openpyxl import load_workbook
from openpyxl.styles.numbers import is_date_format
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re  # Add this import at the top with other imports

router = APIRouter()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _read_excel_cell_values(excel_bytes: bytes, sheet_name: Optional[str]) -> tuple:
    """Load workbook and return (workbook, sheet)."""
    try:
        wb = load_workbook(filename=BytesIO(excel_bytes), data_only=True)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
        return wb, ws
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed reading Excel file: {exc}")


def _get_cell_value(ws, address: str) -> str:
    try:
        return str(ws[address].value if ws[address].value is not None else "")
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Invalid cell address '{address}': {exc}")


def _build_table_from_range(document: Document, ws, table_range: str):
    """Insert a table into the document from an Excel range like 'A1:D10'."""
    try:
        start, end = table_range.split(":", 1)
        start_col = ''.join([c for c in start if c.isalpha()])
        start_row = int(''.join([c for c in start if c.isdigit()]))
        end_col = ''.join([c for c in end if c.isalpha()])
        end_row = int(''.join([c for c in end if c.isdigit()]))

        # Convert column letters to indices (A=1)
        def col_to_idx(col_letters: str) -> int:
            idx = 0
            for ch in col_letters.upper():
                idx = idx * 26 + (ord(ch) - ord('A') + 1)
            return idx

        start_col_idx = col_to_idx(start_col)
        end_col_idx = col_to_idx(end_col)

        num_rows = end_row - start_row + 1
        num_cols = end_col_idx - start_col_idx + 1

        table = document.add_table(rows=num_rows, cols=num_cols)
        _apply_table_style_safely(table)
        table.autofit = True
        for r in range(num_rows):
            for c in range(num_cols):
                xl_cell = ws.cell(row=start_row + r, column=start_col_idx + c)
                docx_cell = table.cell(r, c)
                _write_cell_with_style(docx_cell, xl_cell)
        
        return table
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Invalid table_range '{table_range}': {exc}")


def _replace_text_placeholders(doc: Document, replacements: dict) -> None:
    """Replace placeholders like {{NAME}} across paragraphs and table cells."""
    if not replacements:
        return

    def replace_in_run_text(text: str) -> str:
        for key, value in replacements.items():
            # Accept both bare and braced keys: "NAME" or "{{NAME}}"
            placeholder = key if (key.startswith("{{") and key.endswith("}}")) else f"{{{{{key}}}}}"
            text = text.replace(placeholder, value)
        return text

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = replace_in_run_text(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = replace_in_run_text(run.text)


def _normalize_placeholder_text(text: str) -> str:
    """Normalize text by removing spaces and zero-width/RTL marks for robust matching."""
    if text is None:
        return ""
    zero_width = [
        "\u200e",  # LRM
        "\u200f",  # RLM
        "\u202a",  # LRE
        "\u202b",  # RLE
        "\u202c",  # PDF
        "\u202d",  # LRO
        "\u202e",  # RLO
    ]
    for z in zero_width:
        text = text.replace(z, "")
    return "".join(text.split())


def _find_paragraph_with_token_anywhere(doc: Document, token: str):
    # Robust search in paragraphs and inside table cells using normalized text
    desired = _normalize_placeholder_text(token)
    for paragraph in doc.paragraphs:
        if desired in _normalize_placeholder_text(paragraph.text):
            return paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if desired in _normalize_placeholder_text(p.text):
                        return p
    return None


def _delete_paragraph(paragraph) -> None:
    """Remove a paragraph from the document."""
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _clear_token_in_document(doc: Document, token: str) -> int:
    """Remove all occurrences of token across paragraphs and table cells. Returns count removed.
    Works even when token is split across multiple runs by clearing entire paragraphs that contain it.
    """
    count = 0
    desired = _normalize_placeholder_text(token)
    # Normal paragraphs
    for paragraph in list(doc.paragraphs):
        if desired in _normalize_placeholder_text(paragraph.text):
            for run in list(paragraph.runs):
                run.text = ""
            count += 1
    # Paragraphs in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    if desired in _normalize_placeholder_text(p.text):
                        for run in list(p.runs):
                            run.text = ""
                        count += 1
    return count


@router.post("/api/reports/excel-to-word")
async def excel_to_word(
    request: Request,
    excel: UploadFile = File(...),
    template: UploadFile = File(...),
    mappings: str = Form(...)
):
    """
    Flexible Excel → Word generator
    Supports:
    - {{TAG}} placeholders replaced with Excel values while preserving original formatting
    - Table or text insertion from Excel ranges
    - RTL (right-to-left) table alignment when requested
    """

    try:
        excel_bytes = await excel.read()
        template_bytes = await template.read()

        wb = load_workbook(filename=BytesIO(excel_bytes), data_only=True)
        doc = Document(BytesIO(template_bytes))

        # Handle missing mappings
        if not mappings:
            mapping_dict = {}
            logger.info("No mappings provided - will only process table mappings from form data")
        else:
            try:
                mapping_dict = json.loads(mappings)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Invalid JSON mappings: {e}")

        # Merge table mappings sent as repeated form fields
        form = await request.form()
        sheet_default = form.get('sheet')
        table_tags = form.getlist('table_tag') if hasattr(form, 'getlist') else []
        table_ranges = form.getlist('table_range') if hasattr(form, 'getlist') else []
        table_sheets = form.getlist('table_sheet') if hasattr(form, 'getlist') else []
    
        # Process table mappings from form data
        for i, ttag in enumerate(table_tags):
            trange = table_ranges[i] if i < len(table_ranges) else None
            tsheet = table_sheets[i] if i < len(table_sheets) else None
            if trange:
                existing = mapping_dict.get(ttag)
                base_type = "row_tables2"
                base_align = "left"
                if isinstance(existing, dict):
                    base_type = existing.get("type", base_type)
                    base_align = existing.get("align", base_align)
                payload = {
                    "sheet": tsheet or sheet_default,
                    "range": trange,
                    "type": base_type,
                    "align": base_align,
                }
                if isinstance(existing, dict):
                    existing.update({k: v for k, v in payload.items() if k not in existing})
                    mapping_dict[ttag] = existing
                else:
                    mapping_dict[ttag] = payload

        logger.info(f"Loaded mappings: {mapping_dict}")

        # Process ALL mappings
        for tag, value in mapping_dict.items():
            # -------------------------------
            # CASE 1: Simple string value (text replacement preserving original style)
            # -------------------------------
            if isinstance(value, str):
                # Get text value from Excel cell or use direct text
                ws = None
                if sheet_default and sheet_default in wb.sheetnames:
                    ws = wb[sheet_default]
                else:
                    ws = wb.active

                if any(c.isalpha() for c in value) and any(c.isdigit() for c in value):
                    try:
                        cell_val = ws[value].value
                        text_value = str(cell_val or "")
                    except Exception:
                        text_value = f"<Invalid cell {value}>"
                else:
                    text_value = value

                # Replace tag with value while preserving original formatting
                _replace_preserve_style(doc, tag, text_value)
                logger.info(f"Applied text replacement for {tag}: {text_value}")

            # -------------------------------
            # CASE 2: Object-based mapping (table, range, or styled text)
            # -------------------------------
            elif isinstance(value, dict):
                # Check if this is a simple text replacement
                if value.get("type") == "text" or ("range" not in value and "sheet" not in value):
                    # Text replacement with optional explicit styling
                    text_value = value.get("value", "")
                    explicit_style = value.get("style", {})
                    explicit_align = value.get("align")
                    
                    # Get value from Excel cell if it's a reference
                    if any(c.isalpha() for c in text_value) and any(c.isdigit() for c in text_value):
                        sheet_name = value.get("sheet")
                        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else (
                            wb[sheet_default] if sheet_default and sheet_default in wb.sheetnames else wb.active
                        )
                        try:
                            cell_val = ws[text_value].value
                            text_value = str(cell_val or "")
                        except Exception:
                            text_value = f"<Invalid cell {text_value}>"
                    
                    # Replace with optional explicit styling (overrides original style)
                    _replace_with_optional_style(doc, tag, text_value, explicit_style, explicit_align)
                    logger.info(f"Applied styled text replacement for {tag}: {text_value}")
                    continue

                # Otherwise, process as table/range (existing table logic)
                sheet_name = value.get("sheet")
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else (
                    wb[sheet_default] if sheet_default and sheet_default in wb.sheetnames else wb.active
                )
                range_str = value.get("range")
                insert_type = value.get("type", "row_tables2")
                align = value.get("align", "left")
                style_override = value.get("style", {}) or {}
                direction = value.get("direction", "auto")
                trim_empty = bool(value.get("trimEmpty", True))
                drop_empty_any = bool(value.get("dropEmptyCells", True))

                if not range_str:
                    logger.warning(f"No range specified for tag {tag}")
                    continue

                logger.info(f"Processing {tag} as {insert_type} from {range_str}")

                # Parse range
                try:
                    start_row, end_row, start_col_idx, end_col_idx = _parse_range_flexible(range_str, ws)
                except Exception as e:
                    logger.error(f"Failed to parse range {range_str} for tag {tag}: {e}")
                    doc.add_paragraph(f"⚠️ Invalid range '{range_str}' for {tag}")
                    continue

                rtl = False
                if direction == "rtl":
                    rtl = True
                elif direction == "auto":
                    rtl = _range_has_arabic(ws, start_row, end_row, start_col_idx, end_col_idx)

                paragraph = _find_paragraph_with_token(doc, tag)
                if not paragraph:
                    logger.warning(f"Tag {tag} not found in template, adding placeholder")
                    doc.add_paragraph(f"⚠️ Tag {tag} not found in template.")
                    continue

                # Clear the tag from the paragraph
                for run in paragraph.runs:
                    run.text = run.text.replace(tag, "").strip()

                if insert_type == "row_tables":
                    _insert_row_tables(
                        doc, ws,
                        start_row, end_row,
                        start_col_idx, end_col_idx,
                        paragraph, rtl, align, style_override,
                        trim_empty=trim_empty,
                        drop_empty_any=drop_empty_any
                    )
                    logger.info(f"Inserted row-by-row tables for {tag}")

                elif insert_type == "row_tables2":
                    as_table = value.get("asTable", False)
                    spacing = value.get("spacing", 5)
                    text_align = value.get("textAlign")

                    # DEBUG: Check what we're getting
                    logger.info(f"DEBUG BEFORE FORCE: as_table={as_table}, text_align={text_align}")
                    
                    # FORCE TABLE CREATION - OVERRIDE THE MAPPING
                    as_table = True  # ← THIS SHOULD FORCE IT
                    rtl = True
                    align = "right" 
                    text_align = "right"

                    logger.info(f"DEBUG AFTER FORCE: as_table={as_table}, text_align={text_align}")

                    # Use the enhanced _insert_row_data_with_borders function
                    _insert_row_data_with_borders(
                        doc, ws,
                        start_row, end_row,
                        start_col_idx, end_col_idx,
                        paragraph, rtl, align, style_override,
                        as_table=as_table,
                        spacing=spacing,
                        text_align=text_align,
                        trim_empty=trim_empty,
                        drop_empty_any=drop_empty_any
                    )
                elif insert_type == "table":
                    print('llllllllll')
                    num_rows = end_row - start_row + 1
                    num_cols = end_col_idx - start_col_idx + 1

                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    
                    # Apply comprehensive borders to the entire table first
                    _add_comprehensive_borders_to_table(table)
                    table.autofit = True

                    # Fill rows with NEW RULE: merge-if-previous-empty
                    for r in range(num_rows):
                        c = 0
                        while c < num_cols:
                            xl_cur = ws.cell(row=start_row + r, column=start_col_idx + c)
                            cur_text = str(xl_cell.value or "").strip()
                            cur_is_empty = not cur_text

                            # Check previous cell availability
                            has_prev = c > 0
                            prev_is_empty = False
                            if has_prev:
                                xl_prev = ws.cell(row=start_row + r, column=start_col_idx + (c - 1))
                                prev_text = str(xl_prev.value or "").strip()
                                prev_is_empty = not prev_text

                            # Target col positions in Word (handle RTL)
                            target_col_cur = (num_cols - 1 - c) if rtl else c
                            target_col_prev = (num_cols - 1 - (c - 1)) if rtl else (c - 1)

                            # NEW RULE: If current is empty AND previous exists AND previous is not empty, merge backward
                            if has_prev and cur_is_empty and not prev_is_empty and _should_merge_backward(prev_text):
                                # Merge previous cell with current empty cell
                                merged_cell = table.cell(r, target_col_prev).merge(table.cell(r, target_col_cur))
                                _write_cell_with_style(
                                    merged_cell,
                                    xl_prev,  # Use the previous cell's value and style
                                    {**style_override, "forceRTL": (rtl or align.lower() == "right"), "textAlign": align}
                                )
                                c += 1  # Move to next cell (we've consumed current empty cell)
                                continue

                            # Normal cell writing
                            _write_cell_with_style(
                                table.cell(r, target_col_cur),
                                xl_cur,
                                {**style_override, "forceRTL": (rtl or align.lower() == "right"), "textAlign": align}
                            )
                            c += 1

                    # Apply borders above bold text (this adds EXTRA borders above bold rows)
                    _apply_borders_above_bold_text(table)

                    # Align the table block itself
                    if align.lower() == "right" or rtl:
                        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                    elif align.lower() == "center":
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    else:
                        table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    _set_table_rtl(table, rtl)

                    paragraph._p.addnext(table._tbl)
                    logger.info(f"Inserted table for {tag} with borders")
                elif insert_type == "text":
                    lines = []
                    for r in range(start_row, end_row + 1):
                        cols = list(range(start_col_idx, end_col_idx + 1))
                        if rtl:
                            cols = cols[::-1]
                        row_vals = []
                        for c_idx in cols:
                            val = ws.cell(row=r, column=c_idx).value
                            row_vals.append(str(val or ""))
                        lines.append("    ".join(row_vals))
                    
                    # Create styled text paragraphs
                    for line in lines:
                        new_para = paragraph.insert_paragraph_before(line)
                        new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT if (rtl or align.lower() == "right") else (
                            WD_ALIGN_PARAGRAPH.CENTER if align.lower() == "center" else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        if rtl:
                            _set_paragraph_rtl(new_para, True)
                        
                        # Apply styling to all runs in the paragraph
                        for run in new_para.runs:
                            _apply_run_styling(run, style_override)

                    logger.info(f"Inserted styled text for {tag}")

                else:
                    logger.warning(f"Unknown insert type '{insert_type}' for tag {tag}")

        # Save output
        export_dir = "exports"
        os.makedirs(export_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"excel_to_word_{timestamp}.docx"
        filepath = os.path.join(export_dir, filename)
        doc.save(filepath)

        logger.info(f"Successfully generated Word document: {filename}")
        return FileResponse(filepath, filename=filename)

    except Exception as e:
        logger.error(f"Excel to Word conversion failed: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error generating Word file: {str(e)}")

def _insert_row_data_with_borders(
    doc, ws, start_row, end_row, start_col_idx, end_col_idx, 
    anchor_paragraph, rtl, align, style_override, 
    as_table=False, spacing=5, trim_empty=True, drop_empty_any=True,
    text_align=None
):
    """
    ENHANCED VERSION - Create a SINGLE table with ALL rows and apply borders
    """
    # Force RTL for Arabic content
    rtl = True
    align = "right"
    text_align = "right"
    
    current_anchor = anchor_paragraph._p

    # Collect ALL data first
    all_rows_data = []
    max_cols = 0
    
    for r in range(start_row, end_row + 1):
        # Collect row cells
        row_cells = [ws.cell(row=r, column=c) for c in range(start_col_idx, end_col_idx + 1)]
        
        # Get visible cells for this row
        visible = [cell for cell in row_cells if cell.value is not None and str(cell.value).strip() != ""]
        if not visible:
            continue
            
        all_rows_data.append(visible)
        max_cols = max(max_cols, len(visible))

    if not all_rows_data:
        return

    # Create a SINGLE table with ALL rows
    num_rows = len(all_rows_data)
    num_cols = max_cols
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Apply safe styling (NO borders initially)
    _apply_table_style_safely(table)
    
    # Set table to right alignment
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_rtl(table, True)

    # Fill ALL rows in the single table
    for row_idx, visible_cells in enumerate(all_rows_data):
        for col_idx, xl_cell in enumerate(visible_cells):
            # Reverse column order for RTL (right-to-left)
            target_col = (num_cols - 1 - col_idx) if rtl else col_idx
            
            # Use your existing function with forced RTL
            _write_cell_with_style(
                table.cell(row_idx, target_col), 
                xl_cell, 
                {**(style_override or {}), "forceRTL": True}
            )

    # APPLY BORDERS ABOVE BOLD TEXT (this is the key addition)
     
    _debug_bold_numeric_cells(table)
    
    # Apply borders ONLY to BOLD numeric cells and center align them
    _apply_borders_to_bold_numeric_cells(table)
    

    # Insert the complete table
    current_anchor.addnext(table._tbl)
    current_anchor = table._tbl
    
    logger.info(f"Created table with {num_rows} rows and {num_cols} columns with borders")
def _replace_preserve_style(doc: Document, tag: str, replacement_text: str):
    """Replace tag with text while preserving the original formatting of the tag."""
    
    def replace_in_paragraph(paragraph, tag, replacement_text):
        if tag in paragraph.text:
            # Store the original runs and their styles
            original_runs = []
            for run in paragraph.runs:
                if tag in run.text:
                    # This run contains our tag - preserve its style
                    original_runs.append({
                        'text': run.text,
                        'font': {
                            'name': run.font.name,
                            'size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic,
                            'color': run.font.color.rgb if run.font.color else None
                        }
                    })
            
            # Clear paragraph and rebuild with replacement text
            paragraph.clear()
            
            # Recreate the text with the same styling
            for original_run in original_runs:
                run = paragraph.add_run(original_run['text'].replace(tag, replacement_text))
                
                # Apply original styling
                if original_run['font']['name']:
                    run.font.name = original_run['font']['name']
                if original_run['font']['size']:
                    run.font.size = original_run['font']['size']
                if original_run['font']['bold'] is not None:
                    run.font.bold = original_run['font']['bold']
                if original_run['font']['italic'] is not None:
                    run.font.italic = original_run['font']['italic']
                if original_run['font']['color']:
                    try:
                        run.font.color.rgb = original_run['font']['color']
                    except:
                        pass

    def replace_in_table_cell(paragraph, tag, replacement_text):
        if tag in paragraph.text:
            # Store original run styles
            styled_runs = []
            for run in paragraph.runs:
                if tag in run.text:
                    styled_runs.append({
                        'run': run,
                        'font': {
                            'name': run.font.name,
                            'size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic,
                            'color': run.font.color.rgb if run.font.color else None,
                            'underline': run.font.underline
                        }
                    })
            
            # Replace text in each run that contains the tag
            for styled_run in styled_runs:
                styled_run['run'].text = styled_run['run'].text.replace(tag, replacement_text)

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, tag, replacement_text)

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_table_cell(paragraph, tag, replacement_text)


def _replace_with_optional_style(doc: Document, tag: str, replacement_text: str, explicit_style: dict = None, explicit_align: str = None):
    """Replace tag with text, using explicit style if provided, otherwise preserving original style."""
    
    def process_paragraph(paragraph, tag, replacement_text, explicit_style, explicit_align):
        if tag in paragraph.text:
            if explicit_style or explicit_align:
                # Use explicit styling
                paragraph.clear()
                run = paragraph.add_run(paragraph.text.replace(tag, replacement_text))
                
                # Apply explicit styling
                if explicit_style:
                    _apply_run_styling(run, explicit_style)
                
                # Apply explicit alignment
                if explicit_align:
                    if explicit_align.lower() == "right":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif explicit_align.lower() == "center":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Preserve original styling
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, replacement_text)

    def process_table_cell(paragraph, tag, replacement_text, explicit_style, explicit_align):
        if tag in paragraph.text:
            if explicit_style:
                # Store the paragraph alignment first
                original_alignment = paragraph.alignment
                
                # Clear and rebuild with explicit style
                paragraph.clear()
                run = paragraph.add_run(paragraph.text.replace(tag, replacement_text))
                _apply_run_styling(run, explicit_style)
                
                # Apply alignment
                if explicit_align:
                    if explicit_align.lower() == "right":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif explicit_align.lower() == "center":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = original_alignment
            else:
                # Preserve original styling
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, replacement_text)

    # Process paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, tag, replacement_text, explicit_style, explicit_align)

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_table_cell(paragraph, tag, replacement_text, explicit_style, explicit_align)

def _apply_run_styling(run, style_override: dict):
    """Apply styling to a run."""
    try:
        if style_override.get('fontName'):
            run.font.name = style_override['fontName']
        if style_override.get('fontSize'):
            try:
                run.font.size = Pt(float(style_override['fontSize']))
            except Exception:
                pass
        if style_override.get('bold') is not None:
            run.font.bold = bool(style_override['bold'])
        if style_override.get('italic') is not None:
            run.font.italic = bool(style_override['italic'])
        if style_override.get('color'):
            try:
                run.font.color.rgb = RGBColor.from_string(style_override['color'].replace('#', ''))
            except Exception:
                pass
        if style_override.get('underline') is not None:
            run.font.underline = bool(style_override['underline'])
    except Exception:
        pass

# Keep the existing _find_all_placeholders function and other helpers...
def _find_all_placeholders(doc: Document) -> list:
    """Find all {{TAG}} placeholders in the document."""
    placeholders = set()
    
    # Regex to find {{TAG}} patterns
    import re
    pattern = r'\{\{([^}]+)\}\}'
    
    # Search in paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        placeholders.update([f"{{{{{match}}}}}" for match in matches])
    
    # Search in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(pattern, para.text)
                    placeholders.update([f"{{{{{match}}}}}" for match in matches])
    
    return list(placeholders)

def _insert_row_data(
    doc, ws, start_row, end_row, start_col_idx, end_col_idx, 
    anchor_paragraph, rtl, align, style_override, 
    as_table=False, spacing=5, trim_empty=True, drop_empty_any=True,
    text_align=None
):
    """
    ENHANCED VERSION - Create a SINGLE table with ALL rows and apply borders
    """
    # Force RTL for Arabic content
    rtl = True
    align = "right"
    text_align = "right"
    
    current_anchor = anchor_paragraph._p

    # Collect ALL data first
    all_rows_data = []
    max_cols = 0
    
    for r in range(start_row, end_row + 1):
        # Collect row cells
        row_cells = [ws.cell(row=r, column=c) for c in range(start_col_idx, end_col_idx + 1)]
        
        # Get visible cells for this row
        visible = [cell for cell in row_cells if cell.value is not None and str(cell.value).strip() != ""]
        if not visible:
            continue
            
        all_rows_data.append(visible)
        max_cols = max(max_cols, len(visible))

    if not all_rows_data:
        return

    # Create a SINGLE table with ALL rows
    num_rows = len(all_rows_data)
    num_cols = max_cols
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # ADD COMPREHENSIVE BORDERS TO THE TABLE
    _add_comprehensive_borders_to_table(table)
    
    # Set table to right alignment
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_rtl(table, True)

    # Fill ALL rows in the single table
    for row_idx, visible_cells in enumerate(all_rows_data):
        for col_idx, xl_cell in enumerate(visible_cells):
            # Reverse column order for RTL (right-to-left)
            target_col = (num_cols - 1 - col_idx) if rtl else col_idx
            
            # Use your existing function with forced RTL
            _write_cell_with_style(
                table.cell(row_idx, target_col), 
                xl_cell, 
                {**(style_override or {}), "forceRTL": True}
            )

    # APPLY BORDERS ABOVE BOLD TEXT
    _apply_borders_above_bold_text(table)

    # Insert the complete table
    current_anchor.addnext(table._tbl)
    current_anchor = table._tbl
    
    logger.info(f"Created table with {num_rows} rows and {num_cols} columns with borders")

def _add_manual_borders(table):
    """Add manual borders to table"""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove any existing borders
        for elem in tblPr:
            if elem.tag.endswith('tblBorders'):
                tblPr.remove(elem)
        
        # Create new borders element
        tblBorders = OxmlElement('w:tblBorders')
        
        # Define border types
        border_types = [
            ('top', 'single'),
            ('left', 'single'), 
            ('bottom', 'single'),
            ('right', 'single'),
            ('insideH', 'single'),
            ('insideV', 'single')
        ]
        
        for border_name, border_val in border_types:
            border_elem = OxmlElement(f'w:{border_name}')
            border_elem.set(qn('w:val'), border_val)
            border_elem.set(qn('w:sz'), '4')
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), '000000')
            tblBorders.append(border_elem)
        
        # Add borders to table
        tblPr.append(tblBorders)
        
    except Exception as e:
        print(f"Manual borders failed: {e}")
# --------------------------------------------------------------------------
# Helper functions
# --------------------------------------------------------------------------

def _remove_table_borders(table):
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove any existing borders node
        for elem in list(tblPr):
            if elem.tag.endswith('tblBorders') or elem.tag == qn('w:tblBorders'):
                tblPr.remove(elem)

        # Create borders with 'nil' on all edges (no borders)
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)

        tblPr.append(tblBorders)
    except Exception as e:
        logger.warning(f"Failed to remove table borders: {e}")

def _apply_table_style_safely(table):
    """Apply a table style but ensure NO borders are added."""
    try:
        # Use a plain style that doesn't add borders
        table.style = 'Table Grid'  # This style typically has minimal formatting
    except (ValueError, KeyError):
        try:
            table.style = 'Table Normal'
        except (ValueError, KeyError):
            # If no style available, proceed without styling
            pass

    # EXPLICITLY remove all borders regardless of style defaults
    _remove_all_table_borders(table) 
     
def _add_borders_to_table(table):
    """Manually add borders to table if no style is available"""
    try:
        # Set all borders for the table
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Add table borders
        tblBorders = OxmlElement('w:tblBorders')
        
        # Add top border
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), 'auto')
        tblBorders.append(top)
        
        # Add left border
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '4')
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), 'auto')
        tblBorders.append(left)
        
        # Add bottom border
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')
        tblBorders.append(bottom)
        
        # Add right border
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), '4')
        right.set(qn('w:space'), '0')
        right.set(qn('w:color'), 'auto')
        tblBorders.append(right)
        
        # Add inside horizontal border
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'single')
        insideH.set(qn('w:sz'), '4')
        insideH.set(qn('w:space'), '0')
        insideH.set(qn('w:color'), 'auto')
        tblBorders.append(insideH)
        
        # Add inside vertical border
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'single')
        insideV.set(qn('w:sz'), '4')
        insideV.set(qn('w:space'), '0')
        insideV.set(qn('w:color'), 'auto')
        tblBorders.append(insideV)
        
        # Add borders to table properties
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblPr.append(tblBorders)
        
    except Exception as e:
        print(f"DEBUG: Failed to add manual borders: {e}")

def _build_table_from_worksheet(doc, ws, min_row, max_row, min_col, max_col):
    """Create a Word table from Excel sheet cells."""
    table = doc.add_table(rows=max_row - min_row + 1, cols=max_col - min_col + 1)
    _apply_table_style_safely(table)
    
    for r in range(min_row, max_row + 1):
        for c in range(min_col, max_col + 1):
            val = ws.cell(row=r, column=c).value
            table.cell(r - min_row, c - min_col).text = str(val or "")
    return table


def _find_paragraph_with_token(doc, token):
    """Find a paragraph that contains a given token like {{TABLE}}."""
    for para in doc.paragraphs:
        if token in para.text:
            return para
    return None


@router.post("/api/reports/excel-to-word/preview")
async def excel_to_word_preview(
    excel: UploadFile = File(...),
    sheet: Optional[str] = Form(None)
):
    try:
        excel_bytes = await excel.read()
        wb = load_workbook(filename=BytesIO(excel_bytes), data_only=True)
        sheet_names = wb.sheetnames
        ws = wb[sheet] if sheet and sheet in sheet_names else wb.active

        # Compute used size (rows/cols having data)
        max_row_used = 0
        max_col_used = 0
        for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            last_col_with_data = 0
            row_has_data = False
            for c_idx, val in enumerate(row, start=1):
                if val not in (None, ""):
                    row_has_data = True
                    last_col_with_data = c_idx
            if row_has_data:
                max_row_used = r_idx
                if last_col_with_data > max_col_used:
                    max_col_used = last_col_with_data

        return {
            "sheetNames": sheet_names,
            "sheetPreview": {
                "size": {"rows": max_row_used, "cols": max_col_used}
            }
        }
    except Exception as exc:
        logger.error(f"Excel preview analysis failed: {exc}", exc_info=True)
        raise HTTPException(status_code=400, detail="Failed to analyze Excel file")

@router.post("/api/reports/excel-to-word/debug")
async def excel_to_word_debug(
    template: UploadFile = File(...),
    table_tag: Optional[str] = Form("{{TABLE}}"),
):
    """Debug endpoint to analyze Word template and show all text content."""
    try:
        template_bytes = await template.read()
        doc = Document(BytesIO(template_bytes))
        
        debug_info = {
            "table_tag": table_tag,
            "total_paragraphs": len(doc.paragraphs),
            "paragraphs": [],
            "total_tables": len(doc.tables),
            "table_cells": [],
            "found_table_placeholder": False,
            "found_table_placeholder_normalized": False,
        }
        
        # Analyze all paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text
            normalized_text = _normalize_placeholder_text(para_text)
            
            para_info = {
                "index": i,
                "text": para_text,
                "normalized": normalized_text,
                "length": len(para_text),
                "runs": len(paragraph.runs),
                "contains_table": table_tag in para_text,
                "contains_table_normalized": table_tag in normalized_text,
            }
            
            # Show individual runs if there are multiple
            if len(paragraph.runs) > 1:
                para_info["runs_detail"] = []
                for j, run in enumerate(paragraph.runs):
                    para_info["runs_detail"].append({
                        "index": j,
                        "text": run.text,
                        "length": len(run.text)
                    })
            
            debug_info["paragraphs"].append(para_info)
            
            if table_tag in para_text:
                debug_info["found_table_placeholder"] = True
            if table_tag in normalized_text:
                debug_info["found_table_placeholder_normalized"] = True
        
        # Analyze table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        para_text = para.text
                        normalized_text = _normalize_placeholder_text(para_text)
                        
                        cell_info = {
                            "table": table_idx,
                            "row": row_idx,
                            "cell": cell_idx,
                            "paragraph": para_idx,
                            "text": para_text,
                            "normalized": normalized_text,
                            "contains_table": table_tag in para_text,
                            "contains_table_normalized": table_tag in normalized_text,
                        }
                        
                        debug_info["table_cells"].append(cell_info)
                        
                        if table_tag in para_text:
                            debug_info["found_table_placeholder"] = True
                        if table_tag in normalized_text:
                            debug_info["found_table_placeholder_normalized"] = True
        
        return debug_info
        
    except Exception as exc:
        logger.error(f"Debug analysis failed: {exc}", exc_info=True)
        error_detail = {
            "category": "debug_error",
            "message": str(exc),
            "suggestion": "Check if Word template is valid and not corrupted.",
            "timestamp": datetime.now().isoformat()
        }
        raise HTTPException(status_code=500, detail=json.dumps(error_detail))


@router.post("/api/reports/excel-to-word/simple")
async def excel_to_word_simple(
    excel: UploadFile = File(...),
    sheet: Optional[str] = Form(None),
    table_range: Optional[str] = Form(None),
    table_sheet: Optional[str] = Form(None),
):
    """
    Simple Excel to Word conversion - creates a Word document with Excel data as a table.
    No template needed, just Excel data converted to Word table.
    """
    try:
        excel_bytes = await excel.read()
        
        # Load Excel
        wb = load_workbook(filename=BytesIO(excel_bytes), data_only=True)
        ws_mappings = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
        ws_table = wb[table_sheet] if table_sheet and table_sheet in wb.sheetnames else ws_mappings
        
        # Create new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Excel Data from {ws_table.title}', 0)
        title.alignment = 1  # Center alignment
        
        # Add info paragraph
        info = doc.add_paragraph()
        info.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
        info.add_run(f'Source Sheet: {ws_table.title}\n')
        if table_range:
            info.add_run(f'Data Range: {table_range}\n')
        info.alignment = 1  # Center alignment
        
        # Add spacing
        doc.add_paragraph()
        
        if table_range:
            # Use specific range
            try:
                start, end = table_range.split(":", 1)
                start_col = ''.join([c for c in start if c.isalpha()])
                start_row = int(''.join([c for c in start if c.isdigit()]))
                end_col = ''.join([c for c in end if c.isalpha()])
                end_row = int(''.join([c for c in end if c.isdigit()]))

                def col_to_idx(col_letters: str) -> int:
                    idx = 0
                    for ch in col_letters.upper():
                        idx = idx * 26 + (ord(ch) - ord('A') + 1)
                    return idx

                start_col_idx = col_to_idx(start_col)
                end_col_idx = col_to_idx(end_col)

                num_rows = end_row - start_row + 1
                num_cols = end_col_idx - start_col_idx + 1

                # Create table
                table = doc.add_table(rows=num_rows, cols=num_cols)
                _apply_table_style_safely(table)
                
                # Fill table with data
                for r in range(num_rows):
                    for c in range(num_cols):
                        value = ws_table.cell(row=start_row + r, column=start_col_idx + c).value
                        table.cell(r, c).text = "" if value is None else str(value)
                        
                logger.info(f"Created table with {num_rows} rows and {num_cols} columns from range {table_range}")
                
            except Exception as e:
                logger.error(f"Failed to create table from range {table_range}: {e}")
                raise HTTPException(status_code=400, detail=f"Invalid table range: {e}")
        else:
            # Use entire sheet (first 50 rows x 20 columns)
            max_rows = min(ws_table.max_row or 0, 50)
            max_cols = min(ws_table.max_column or 0, 20)
            
            if max_rows > 0 and max_cols > 0:
                # Create table
                table = doc.add_table(rows=max_rows, cols=max_cols)
                _apply_table_style_safely(table)
                
                # Fill table with data
                for r in range(max_rows):
                    for c in range(max_cols):
                        value = ws_table.cell(row=r + 1, column=c + 1).value
                        table.cell(r, c).text = "" if value is None else str(value)
                        
                logger.info(f"Created table with {max_rows} rows and {max_cols} columns from entire sheet")
            else:
                raise HTTPException(status_code=400, detail="No data found in Excel sheet")

        # Save file to exports
        export_dir = "exports"
        os.makedirs(export_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"excel_to_word_simple_{timestamp}.docx"
        filepath = os.path.join(export_dir, filename)
        doc.save(filepath)

        logger.info(f"Simple Word document saved: {filepath}")

        return FileResponse(
            path=filepath,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"X-Export-Src": filepath},
        )

    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"Simple Excel to Word generation failed: {exc}", exc_info=True)
        error_detail = {
            "category": "simple_generation_error",
            "message": str(exc),
            "suggestion": "Check Excel file format and ensure it contains data.",
            "timestamp": datetime.now().isoformat()
        }
        raise HTTPException(status_code=500, detail=json.dumps(error_detail))


@router.post("/api/reports/excel-to-word/no-replace")
async def excel_to_word_no_replace(
    template: UploadFile = File(...)
):
    try:
        template_bytes = await template.read()

        export_dir = "exports"
        os.makedirs(export_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"excel_to_word_raw_{timestamp}.docx"
        filepath = os.path.join(export_dir, filename)

        with open(filepath, "wb") as f:
            f.write(template_bytes)

        return FileResponse(
            path=filepath,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"X-Export-Src": filepath},
        )
    except Exception as exc:
        logger.error(f"Pass-through Word return failed: {exc}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to return template")
        raise HTTPException(status_code=500, detail=f"Error generating Word file: {e}")

# Helper functions
def _contains_arabic(text: str) -> bool:
    if not text:
        return False
    for ch in text:
        code = ord(ch)
        if (0x0600 <= code <= 0x06FF) or (0x0750 <= code <= 0x077F) or (0x08A0 <= code <= 0x08FF):
            return True
    return False

def _set_paragraph_rtl(paragraph, rtl: bool = True):
    # Enable bidirectional rendering for Arabic
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if rtl else '0')

def _set_table_rtl(table, rtl: bool = True):
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._element.insert(0, tblPr)
        bidi = tblPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            tblPr.append(bidi)
        bidi.set(qn('w:val'), '1' if rtl else '0')
    except Exception:
        # Non-critical if not supported
        pass

# ---- NEW: Flexible range parsing helpers ----
def _col_to_idx(col_letters: str) -> int:
    """Convert Excel column letters to 1-based index (A=1)."""
    idx = 0
    for ch in (col_letters or "").upper():
        if 'A' <= ch <= 'Z':
            idx = idx * 26 + (ord(ch) - ord('A') + 1)
    return idx

def _detect_used_row_span(ws, start_col_idx: int, end_col_idx: int, start_row_min: int = 1) -> tuple[int, int]:
    """
    Scan the worksheet to find the first and last non-empty row across the given column range.
    Returns (start_row, end_row). Falls back to (start_row_min, ws.max_row or start_row_min) if no data.
    """
    max_row = ws.max_row or start_row_min
    first = None
    last = None

    for r in range(start_row_min, max_row + 1):
        row_has_data = False
        for c in range(start_col_idx, end_col_idx + 1):
            v = ws.cell(row=r, column=c).value
            if v is not None and str(v).strip() != "":
                row_has_data = True
                break
        if row_has_data:
            if first is None:
                first = r
            last = r

    if first is None or last is None:
        # No data found: return a safe span
        return start_row_min, max_row
    return first, last

def _parse_range_flexible(range_str: str, ws):
    """
    Parse ranges like:
      - 'A:G'        (columns only, detect used rows)
      - 'A5:G'       (start row fixed, detect end row)
      - 'A:G60'      (detect start row, end row fixed)
      - 'A1:G60'     (fully specified)
      - 'C7'         (single cell)
    Returns: (start_row, end_row, start_col_idx, end_col_idx)
    """
    s = (range_str or "").replace(" ", "")
    if not s:
        raise HTTPException(status_code=400, detail="Empty range string")

    def split_part(part: str) -> tuple[str, int | None]:
        letters = "".join(ch for ch in part if ch.isalpha())
        digits = "".join(ch for ch in part if ch.isdigit())
        return letters, (int(digits) if digits else None)

    if ":" in s:
        left, right = s.split(":", 1)
        left_col_letters, left_row = split_part(left)
        right_col_letters, right_row = split_part(right)

        if not left_col_letters or not right_col_letters:
            raise HTTPException(status_code=400, detail=f"Invalid range '{s}': missing column letters")

        start_col_idx = _col_to_idx(left_col_letters)
        end_col_idx = _col_to_idx(right_col_letters)
        if start_col_idx < 1 or end_col_idx < 1:
            raise HTTPException(status_code=400, detail=f"Invalid range '{s}': bad column letters")

        # Ensure start_col_idx <= end_col_idx
        if start_col_idx > end_col_idx:
            start_col_idx, end_col_idx = end_col_idx, start_col_idx

        # Determine start_row and end_row
        if left_row is not None and right_row is not None:
            start_row = max(1, left_row)
            end_row = max(start_row, right_row)
        elif left_row is not None and right_row is None:
            # 'A5:G' → detect end_row from left_row to bottom
            start_row = max(1, left_row)
            _, end_row = _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=start_row)
        elif left_row is None and right_row is not None:
            # 'A:G60' → detect start_row from top, end_row fixed
            start_row, _ = _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=1)
            end_row = max(start_row, right_row)
        else:
            # 'A:G' → detect both
            start_row, end_row = _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=1)

        return start_row, end_row, start_col_idx, end_col_idx
    else:
        # Single address like 'C7' or 'AA12'
        col_letters, row = split_part(s)
        if not col_letters or row is None:
            raise HTTPException(status_code=400, detail=f"Invalid range/address '{s}'")
        col_idx = _col_to_idx(col_letters)
        start_row = end_row = max(1, row)
        return start_row, end_row, col_idx, col_idx

def _range_has_arabic(ws, start_row, end_row, start_col_idx, end_col_idx) -> bool:
    for r in range(start_row, end_row + 1):
        for c in range(start_col_idx, end_col_idx + 1):
            v = ws.cell(row=r, column=c).value
            if isinstance(v, str) and _contains_arabic(v):
                return True
    return False

def _format_excel_value(xl_cell) -> str:
    val = xl_cell.value
    if val is None:
        return ""
    fmt = xl_cell.number_format or ""

    # Date/time formats
    try:
        if is_date_format(fmt):
            # Return like 30/Jun/2025 to match Excel-style preview
            if hasattr(val, 'strftime'):
                return val.strftime('%d/%b/%Y')
    except Exception:
        pass

    # Percent
    if isinstance(val, (int, float)) and '%' in fmt:
        try:
            # Excel typically stores 0.12 for 12%
            return f"{val:.2%}"
        except Exception:
            return str(val)

    # Thousands separators and decimals
    if isinstance(val, (int, float)):
        try:
            if '.00' in fmt or '.##' in fmt or '.0' in fmt or '.##' in fmt:
                # 2 decimal places if format implies decimals
                return format(val, ',.2f')
            if ',' in fmt:
                return format(val, ',.0f')
            return str(val)
        except Exception:
            return str(val)

    return str(val)

def _is_excel_cell_empty(xl_cell) -> bool:
    """Return True if the Excel cell is effectively empty (None or blank string)."""
    try:
        v = xl_cell.value
    except Exception:
        return True
    if v is None:
        return True
    if isinstance(v, str) and v.strip() == "":
        return True
    return False

def _should_merge_to_one_line(cur_text: str) -> bool:
    """Business rule: merge current cell with the next empty cell when current has non-empty text."""
    return bool((cur_text or "").strip())

def _apply_cell_shading(docx_cell, xl_cell, style: Optional[dict]):
    # Shading OFF by default; only apply if explicitly requested
    if not style or not style.get('applyShading'):
        return
    try:
        fill = xl_cell.fill
        fill_type = getattr(fill, 'fill_type', None) or getattr(fill, 'patternType', None)
        if str(fill_type or '').lower() != 'solid':
            return
        # Prefer fgColor when available
        color = getattr(fill, 'fgColor', None) or getattr(fill, 'start_color', None)
        rgb = getattr(color, 'rgb', None)
        if not rgb or rgb in ('00000000', 'FFFFFFFF'):
            return

        hexrgb = rgb[-6:]  # strip alpha if ARGB
        tc = docx_cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hexrgb)
        tcPr.append(shd)
    except Exception:
        # Non-critical
        pass


def _debug_bold_numeric_cells(table):
    """Debug function to identify bold numeric cells."""
    logger.info("=== DEBUG BOLD NUMERIC CELLS ===")
    
    bold_numeric_cells = []
    
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            is_numeric = _is_numeric_value(cell_text)
            is_bold = _is_cell_bold(cell)
            
            if is_numeric and is_bold:
                bold_numeric_cells.append((row_idx, cell_idx, cell_text))
                logger.info(f"BOLD NUMERIC: Row {row_idx}, Col {cell_idx}: '{cell_text}'")
            elif is_numeric and not is_bold:
                logger.info(f"Regular numeric: Row {row_idx}, Col {cell_idx}: '{cell_text}'")
            elif is_bold and not is_numeric:
                logger.info(f"Bold text: Row {row_idx}, Col {cell_idx}: '{cell_text}'")
    
    logger.info(f"Total bold numeric cells found: {len(bold_numeric_cells)}")
    logger.info("=== END DEBUG ===")
    
    return bold_numeric_cells


def _write_cell_with_style(docx_cell, xl_cell, style: Optional[dict] = None):
    text = _format_excel_value(xl_cell)
    docx_cell.text = ""
    para = docx_cell.paragraphs[0] if docx_cell.paragraphs else docx_cell.add_paragraph()
    run = para.add_run(text)

    # DEBUG: Check if Excel cell has bold formatting
    excel_bold = getattr(xl_cell.font, 'bold', None)
    if excel_bold:
        logger.info(f"Excel cell bold: {excel_bold}, value: '{text}'")
        run.font.bold = excel_bold  # Force preserve Excel bold

    # ... rest of your alignment and styling code ...

    # Force bold if specified in style
    if style and style.get('bold'):
        run.font.bold = True
        logger.info(f"Force bold from style: '{text}'")

def _insert_row_tables(doc, ws, start_row, end_row, start_col_idx, end_col_idx, anchor_paragraph, rtl, align, style_override, trim_empty=True, drop_empty_any=True):
    """
    Insert a sequence of 1-row tables, one per Excel row.
    """
    current_anchor = anchor_paragraph._p

    for r in range(start_row, end_row + 1):
        # Collect row cells
        row_cells = [ws.cell(row=r, column=c) for c in range(start_col_idx, end_col_idx + 1)]

        # Visible span
        if drop_empty_any:
            visible = [cell for cell in row_cells if cell.value is not None and str(cell.value).strip() != ""]
            if not visible:
                continue
        else:
            left = 0
            right = len(row_cells) - 1
            if trim_empty:
                while left <= right and (row_cells[left].value is None or str(row_cells[left].value).strip() == ""):
                    left += 1
                while right >= left and (row_cells[right].value is None or str(row_cells[right].value).strip() == ""):
                    right -= 1
            if right < left:
                continue
            visible = row_cells[left:right + 1]

        num_cols = len(visible)

        # Create 1-row table
        table = doc.add_table(rows=1, cols=num_cols)
        _apply_table_style_safely(table)

        # Force right alignment and RTL
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        _set_table_rtl(table, True)

        # Fill cells with forced RTL
        for i, xl_cell in enumerate(visible):
            target_col = (num_cols - 1 - i)  # Reverse for RTL
            _write_cell_with_style(
                table.cell(0, target_col), 
                xl_cell, 
                {**(style_override or {}), "forceRTL": True}  # Force RTL
            )

        # Insert table
        current_anchor.addnext(table._tbl)
        current_anchor = table._tbl



def _apply_simple_borders_above_bold_text(table):
    """DEPRECATED – kept for reference; use _apply_borders_above_bold_text instead."""
    pass
    """Simple approach: add borders above rows with bold text using paragraph borders."""
    try:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.bold:
                            # Add border to the paragraph above this cell
                            if row_idx > 0:
                                above_cell = table.rows[row_idx - 1].cells[cell._index]
                                for above_para in above_cell.paragraphs:
                                    p_pr = above_para._element.get_or_add_pPr()
                                    p_borders = OxmlElement('w:pBorders')
                                    
                                    bottom = OxmlElement('w:bottom')
                                    bottom.set(qn('w:val'), 'single')
                                    bottom.set(qn('w:sz'), '8')
                                    bottom.set(qn('w:space'), '0')
                                    bottom.set(qn('w:color'), '000000')
                                    p_borders.append(bottom)
                                    
                                    p_pr.append(p_borders)
                            break
    except Exception as e:
        logger.warning(f"Failed to apply simple borders: {e}")


def _find_paragraph_with_token(doc, token):
    """Find a paragraph that contains a given token like {{TABLE}}."""
    for para in doc.paragraphs:
        if token in para.text:
            return para
    
    # Also search in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if token in para.text:
                        return para
    return None



def _add_comprehensive_borders_to_table(table):
    """Add borders to entire table for better visibility."""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove any existing borders
        for elem in list(tblPr):
            if elem.tag.endswith('tblBorders'):
                tblPr.remove(elem)
        
        # Create comprehensive borders
        tblBorders = OxmlElement('w:tblBorders')
        
        border_types = [
            ('top', 'single'),
            ('left', 'single'), 
            ('bottom', 'single'),
            ('right', 'single'),
            ('insideH', 'single'),
            ('insideV', 'single')
        ]
        
        for border_name, border_val in border_types:
            border_elem = OxmlElement(f'w:{border_name}')
            border_elem.set(qn('w:val'), border_val)
            border_elem.set(qn('w:sz'), '8')
            border_elem.set(qn('w:space'), '0')
            border_elem.set(qn('w:color'), '000000')
            tblBorders.append(border_elem)
        
        tblPr.append(tblBorders)
        
    except Exception as e:
        logger.warning(f"Failed to add comprehensive borders: {e}")



def _debug_table_content(table):
    """Debug function to see table structure and styling"""
    logger.info("=== DEBUG TABLE CONTENT ===")
    logger.info(f"Table has {len(table.rows)} rows and {len(table.columns) if table.columns else 'unknown'} columns")
    
    for row_idx, row in enumerate(table.rows):
        row_info = []
        bold_cells = []
        numeric_cells = []
        
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            is_bold = False
            is_numeric = False
            
            # Check each run in the cell for bold formatting
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.bold:
                        is_bold = True
                        break
                if is_bold:
                    break
            
            # Check if numeric
            is_numeric = _is_numeric_value(cell_text)
            
            row_info.append(f"'{cell_text}'")
            if is_bold:
                bold_cells.append(cell_idx)
            if is_numeric:
                numeric_cells.append(cell_idx)
        
        logger.info(f"Row {row_idx}: {', '.join(row_info)}")
        if bold_cells:
            logger.info(f"  → Bold cells: {bold_cells}")
        if numeric_cells:
            logger.info(f"  → Numeric cells: {numeric_cells}")
        if bold_cells and numeric_cells:
            logger.info(f"  → BOLD NUMERIC FOUND! Row {row_idx} has both bold and numeric")
    
    logger.info("=== END DEBUG TABLE ===")



def _apply_borders_above_bold_text(table):
    """Add top borders ONLY above rows that contain BOLD cells - NO other borders."""
    try:
        # First remove ALL borders from the table
        _remove_all_table_borders(table)
        
        # Track which rows have bold content
        bold_rows = []
        
        for row_idx, row in enumerate(table.rows):
            has_bold = False
            
            # Check if any cell in this row has bold text
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.bold and run.text.strip():
                            has_bold = True
                            break
                    if has_bold:
                        break
                if has_bold:
                    break
            
            if has_bold:
                bold_rows.append(row_idx)
        
        logger.info(f"Bold rows found: {bold_rows}")
        
        # Add borders ONLY above each bold row (on the previous row)
        for bold_row_idx in bold_rows:
            if bold_row_idx > 0:  # Can't add border above first row
                target_row = table.rows[bold_row_idx - 1]
                _add_top_border_to_row(target_row)
                logger.info(f"ADDED BORDER: Above row {bold_row_idx} (row index {bold_row_idx - 1})")
                
    except Exception as e:
        logger.warning(f"Failed to apply borders above bold text: {e}")

def _add_top_border_to_row(row):
    """Add top border ONLY to all cells in a row - no other borders."""
    try:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Get or create borders element
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            # Remove ALL existing borders first
            for elem in list(tcBorders):
                tcBorders.remove(elem)
            
            # Add ONLY top border
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '8')  # Normal thickness
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), '000000')  # Black
            tcBorders.append(top)
            
            # Explicitly set all other borders to "nil" (no border)
            for side in ['left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            
    except Exception as e:
        logger.warning(f"Failed to add top border to cell: {e}")




def _apply_borders_to_numeric_cells(table):
    """Add top and bottom borders ONLY to cells with numeric values."""
    try:
        # First remove ALL borders from the table
        _remove_all_table_borders(table)
        
        numeric_cells_count = 0
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # Check if cell contains numeric value
                if _is_numeric_value(cell_text):
                    _add_top_bottom_borders_to_cell(cell)
                    _center_align_cell(cell)  # Center align numeric cells
                    numeric_cells_count += 1
                    logger.info(f"Added borders to numeric cell: '{cell_text}' at [{row_idx},{cell_idx}]")
        
        logger.info(f"Total numeric cells with borders: {numeric_cells_count}")
                
    except Exception as e:
        logger.warning(f"Failed to apply borders to numeric cells: {e}")



def _is_numeric_value(text: str) -> bool:
    """Check if text represents a numeric value (with commas, currency, etc.)."""
    if not text or not text.strip():
        return False
    
    cleaned = text.strip()
    
    # Remove common currency symbols and text
    cleaned = cleaned.replace('جنيه', '').replace('مصري', '').replace('ج', '').replace('م', '').strip()
    
    # Remove commas for numeric checking
    cleaned_no_commas = cleaned.replace(',', '')
    
    # Check for various numeric patterns
    numeric_patterns = [
        r'^\d+$',                    # 123456
        r'^\d+\.\d+$',               # 123.456
        r'^\d+\.\d+%$',              # 123.456%
        r'^[\d,]+$',                 # 123,456
        r'^[\d,]+\.\d+$',            # 123,456.78
        r'^[\d,]+\.\d+%$',           # 123,456.78%
        r'^-?\d+$',                  # -123456
        r'^-?[\d,]+$',               # -123,456
    ]
    
    for pattern in numeric_patterns:
        if re.match(pattern, cleaned_no_commas):
            return True
    
    # Also check if it's a large financial number (like your table examples)
    if (any(char.isdigit() for char in cleaned) and 
        len([c for c in cleaned if c.isdigit()]) >= 6):  # At least 6 digits
        # Make sure it's not a date or other numeric-looking text
        if not any(keyword in text for keyword in ['/', 'يونيو', 'يناير', 'فبراير', 'مارس', 'أبريل', 'مايو', 'يونيه', 'يوليو', 'أغسطس', 'سبتمبر', 'أكتوبر', 'نوفمبر', 'ديسمبر']):
            return True
        
    return False

def _should_merge_backward(prev_text: str) -> bool:
    """Business rule: merge backward when current cell is empty and previous has text."""
    return bool((prev_text or "").strip())


def _remove_all_table_borders(table):
    """Remove ALL borders from the entire table first."""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove table-level borders
        for elem in list(tblPr):
            if elem.tag.endswith('tblBorders'):
                tblPr.remove(elem)
        
        # Remove cell-level borders from ALL cells
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Remove any existing cell borders
                for elem in list(tcPr):
                    if elem.tag.endswith('tcBorders'):
                        tcPr.remove(elem)
                        
        logger.info("Removed all table borders")
    except Exception as e:
        logger.warning(f"Failed to remove table borders: {e}")

def _apply_borders_to_bold_numeric_cells(table):
    """Add top and bottom borders ONLY to cells with BOLD numeric values."""
    try:
        # First remove ALL borders from the table
        _remove_all_table_borders(table)
        
        bold_numeric_cells_count = 0
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # Check if cell contains numeric value AND has bold formatting
                if _is_numeric_value(cell_text) and _is_cell_bold(cell):
                    _add_top_bottom_borders_to_cell(cell)
                    _center_align_cell(cell)  # Center align bold numeric cells
                    bold_numeric_cells_count += 1
                    logger.info(f"Added borders to BOLD numeric cell: '{cell_text}' at [{row_idx},{cell_idx}]")
        
        logger.info(f"Total BOLD numeric cells with borders: {bold_numeric_cells_count}")
                
    except Exception as e:
        logger.warning(f"Failed to apply borders to bold numeric cells: {e}")

def _is_cell_bold(cell):
    """Check if any text in the cell has bold formatting."""
    try:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.font.bold:
                    return True
        return False
    except Exception:
        return False

def _add_top_bottom_borders_to_cell(cell):
    """Add top and bottom borders ONLY to a cell."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create borders element
        tcBorders = OxmlElement('w:tcBorders')
        
        # Add TOP border
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')  # Medium thickness
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')  # Black
        tcBorders.append(top)
        
        # Add BOTTOM border
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Medium thickness
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')  # Black
        tcBorders.append(bottom)
        
        # Set all other borders to "nil" (no border)
        for side in ['left', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
            
    except Exception as e:
        logger.warning(f"Failed to add top/bottom borders to cell: {e}")

def _center_align_cell(cell):
    """Center align all paragraphs in a cell."""
    try:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning(f"Failed to center align cell: {e}")