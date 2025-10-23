# Top-level imports
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import FileResponse
from typing import Optional
import os
import json
from datetime import datetime
from io import BytesIO
import logging

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles.numbers import is_date_format
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re  # Add this import at the top with other imports
import urllib.parse

router = APIRouter()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@router.post("/api/reports/excel-to-word/preview")
async def excel_to_word_preview(
    excel: UploadFile = File(...),
    sheet: Optional[str] = Form(None)
):
    try:
        excel_bytes = await excel.read()
        wb = load_workbook(filename=BytesIO(excel_bytes), data_only=True)
        sheet_names = wb.sheetnames
        selected = sheet if sheet and sheet in sheet_names else wb.active.title

        sheetsData = {}

        for name in sheet_names:
            ws = wb[name]
            min_row = None
            max_row = 0
            min_col = None
            max_col = 0

            # Detect bounding box of used cells for this sheet
            for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                row_has_data = False
                first_col_in_row = None
                last_col_in_row = None
                for c_idx, val in enumerate(row, start=1):
                    if val not in (None, ""):
                        row_has_data = True
                        if first_col_in_row is None:
                            first_col_in_row = c_idx
                        last_col_in_row = c_idx
                if row_has_data:
                    if min_row is None:
                        min_row = r_idx
                    max_row = r_idx
                    if first_col_in_row is not None:
                        if min_col is None or first_col_in_row < min_col:
                            min_col = first_col_in_row
                        if last_col_in_row > max_col:
                            max_col = last_col_in_row

            if min_row is None:
                # Empty sheet
                range_str = "A1:A1"
                size = {"rows": 0, "cols": 0}
                samples = {}
            else:
                start_col_letter = get_column_letter(min_col or 1)
                end_col_letter = get_column_letter(max_col or 1)
                range_str = f"{start_col_letter}{min_row}:{end_col_letter}{max_row}"
                size = {"rows": max_row, "cols": max_col}

                # Collect a few sample non-empty cells (for UI compatibility)
                samples = {}
                count = 0
                for r in range(min_row, max_row + 1):
                    for c in range(min_col or 1, (max_col or 1) + 1):
                        val = ws.cell(row=r, column=c).value
                        if val not in (None, ""):
                            addr = f"{get_column_letter(c)}{r}"
                            samples[addr] = {
                                "value": str(val),
                                "row": r,
                                "col": c,
                                "type": type(val).__name__,
                            }
                            count += 1
                            if count >= 12:
                                break
                    if count >= 12:
                        break

            sheetsData[name] = {
                "size": size,
                "cell_samples": samples,
                "range": range_str
            }

        # Backward-compatible selected sheet preview
        sel_size = sheetsData.get(selected, {"size": {"rows": 0, "cols": 0}})["size"]
        return {
            "sheetNames": sheet_names,
            "sheetsData": sheetsData,
            "selectedSheet": selected,
            "sheetPreview": {"size": sel_size}
        }
    except Exception as exc:
        logger.error(f"Excel preview analysis failed: {exc}", exc_info=True)
        raise HTTPException(status_code=400, detail="Failed to analyze Excel file")

@router.get("/api/reports/files/templates")
async def list_template_files():
    base_dir = os.path.dirname(__file__)
    templates_dir = os.path.join(base_dir, "template")
    if not os.path.isdir(templates_dir):
        raise HTTPException(status_code=404, detail="Templates directory not found")
    files = []
    for name in os.listdir(templates_dir):
        path = os.path.join(templates_dir, name)
        if os.path.isfile(path):
            ext = os.path.splitext(name)[1].lower()
            if ext in [".docx", ".doc", ".dotx"]:
                stat = os.stat(path)
                files.append({
                    "name": name,
                    "size": stat.st_size,
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })
    return {"files": files}

@router.get("/api/reports/files/disclosures")
async def list_disclosure_files():
    base_dir = os.path.dirname(__file__)
    disclosures_dir = os.path.join(base_dir, "Disclosures")
    if not os.path.isdir(disclosures_dir):
        raise HTTPException(status_code=404, detail="Disclosures directory not found")
    files = []
    for name in os.listdir(disclosures_dir):
        path = os.path.join(disclosures_dir, name)
        if os.path.isfile(path):
            ext = os.path.splitext(name)[1].lower()
            if ext in [".docx", ".doc", ".dotx"]:
                stat = os.stat(path)
                files.append({
                    "name": name,
                    "size": stat.st_size,
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })
    return {"files": files}

@router.post("/api/reports/excel-to-word/debug")
async def excel_to_word_debug(
    template: UploadFile = File(...),
    table_tag: Optional[str] = Form("{{TABLE}}"),
):
    """Debug endpoint to analyze Word template and show all text content."""
    try:
        template_bytes = await template.read()
        doc = Document(BytesIO(template_bytes))
        
        # Save uploaded template to 'template' folder (debug endpoint) once per day using original name + date
        try:
            template_dir = os.path.join(os.getcwd(), "template")
            os.makedirs(template_dir, exist_ok=True)
            date_stamp = datetime.now().strftime("%Y%m%d")
            orig_name = os.path.basename(template.filename) if getattr(template, 'filename', None) else "template.docx"
            stem, ext = os.path.splitext(orig_name)
            if not ext:
                ext = ".docx"
            template_path = os.path.join(template_dir, f"{stem}_{date_stamp}{ext}")
            if not os.path.exists(template_path):
                with open(template_path, "wb") as f:
                    f.write(template_bytes)
                logger.info(f"Saved uploaded template (debug) to {template_path}")
            else:
                logger.info(f"Template already saved today: {template_path}")
        except Exception as e:
            logger.warning(f"Failed to save uploaded template (debug): {e}")
        
        debug_info = {
            "table_tag": table_tag,
            "total_paragraphs": len(doc.paragraphs),
            "paragraphs": [],
            "total_tables": len(doc.tables),
            "table_cells": [],
            "found_table_placeholder": False,
            "found_table_placeholder_normalized": False,
        }
        
        # Analyze all paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text
            normalized_text = _normalize_placeholder_text(para_text)
            
            para_info = {
                "index": i,
                "text": para_text,
                "normalized": normalized_text,
                "length": len(para_text),
                "runs": len(paragraph.runs),
                "contains_table": table_tag in para_text,
                "contains_table_normalized": table_tag in normalized_text,
            }
            
            # Show individual runs if there are multiple
            if len(paragraph.runs) > 1:
                para_info["runs_detail"] = []
                for j, run in enumerate(paragraph.runs):
                    para_info["runs_detail"].append({
                        "index": j,
                        "text": run.text,
                        "length": len(run.text)
                    })
            
            debug_info["paragraphs"].append(para_info)
            
            if table_tag in para_text:
                debug_info["found_table_placeholder"] = True
            if table_tag in normalized_text:
                debug_info["found_table_placeholder_normalized"] = True
        
        # Analyze table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        para_text = para.text
                        normalized_text = _normalize_placeholder_text(para_text)
                        
                        cell_info = {
                            "table": table_idx,
                            "row": row_idx,
                            "cell": cell_idx,
                            "paragraph": para_idx,
                            "text": para_text,
                            "normalized": normalized_text,
                            "contains_table": table_tag in para_text,
                            "contains_table_normalized": table_tag in normalized_text,
                        }
                        
                        debug_info["table_cells"].append(cell_info)
                        
                        if table_tag in para_text:
                            debug_info["found_table_placeholder"] = True
                        if table_tag in normalized_text:
                            debug_info["found_table_placeholder_normalized"] = True
        
        return debug_info
        
    except Exception as exc:
        logger.error(f"Debug analysis failed: {exc}", exc_info=True)
        error_detail = {
            "category": "debug_error",
            "message": str(exc),
            "suggestion": "Check if Word template is valid and not corrupted.",
            "timestamp": datetime.now().isoformat()
        }
        raise HTTPException(status_code=500, detail=json.dumps(error_detail))


@router.post("/api/reports/excel-to-word")
async def excel_to_word(
    request: Request,
    excel: UploadFile = File(...),
    template: UploadFile = File(...),
    mappings: str = Form(...)
):
    """
    Flexible Excel → Word generator
    Supports:
    - {{TAG}} placeholders replaced with Excel values while preserving original formatting
    - Table or text insertion from Excel ranges
    - RTL (right-to-left) table alignment when requested
    """

    try:
        excel_bytes = await excel.read()
        template_bytes = await template.read()

        wb = load_workbook(filename=BytesIO(excel_bytes), data_only=True)
        doc = Document(BytesIO(template_bytes))
        
        # Removed template saving in main endpoint to avoid duplicates; debug endpoint still saves uploads.

        # Handle missing mappings
        if not mappings:
            mapping_dict = {}
            logger.info("No mappings provided - will only process table mappings from form data")
        else:
            try:
                mapping_dict = json.loads(mappings)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Invalid JSON mappings: {e}")

        # Merge table mappings sent as repeated form fields
        form = await request.form()
        sheet_default = form.get('sheet')
        table_tags = form.getlist('table_tag') if hasattr(form, 'getlist') else []
        table_ranges = form.getlist('table_range') if hasattr(form, 'getlist') else []
        table_sheets = form.getlist('table_sheet') if hasattr(form, 'getlist') else []
    
        # Only fill missing sheet/range; do not override type/align/direction from frontend
        for i, ttag in enumerate(table_tags):
            trange = table_ranges[i] if i < len(table_ranges) else None
            tsheet = table_sheets[i] if i < len(table_sheets) else None
            if not trange:
                continue

            existing = mapping_dict.get(ttag)

            # Respect explicit value/text mappings from frontend
            if isinstance(existing, str) or (isinstance(existing, dict) and existing.get("type") in ("text", "value")):
                logger.info(f"Skipping legacy table merge for {ttag} (explicit value/text mapping)")
                continue

            if isinstance(existing, dict):
                # Fill only missing fields; don't override type/align/direction from frontend JSON
                if "range" not in existing:
                    existing["range"] = trange
                if "sheet" not in existing:
                    existing["sheet"] = tsheet or sheet_default
                mapping_dict[ttag] = existing
            else:
                # Legacy-only tag: let downstream defaults apply type/align/direction
                mapping_dict[ttag] = {
                    "sheet": tsheet or sheet_default,
                    "range": trange
                }

        logger.info(f"Loaded mappings: {mapping_dict}")

        # Process ALL mappings
        for tag, value in mapping_dict.items():
            # -------------------------------
            # CASE 1: Simple string value (text replacement preserving original style)
            # -------------------------------
            if isinstance(value, str):
                # Get text value from Excel cell or use direct text
                ws = None
                if sheet_default and sheet_default in wb.sheetnames:
                    ws = wb[sheet_default]
                else:
                    ws = wb.active

                if any(c.isalpha() for c in value) and any(c.isdigit() for c in value):
                    try:
                        cell_val = ws[value].value
                        text_value = str(cell_val or "")
                    except Exception:
                        text_value = f"<Invalid cell {value}>"
                else:
                    text_value = value

                # Replace tag with value while preserving original formatting
                _replace_preserve_style(doc, tag, text_value)
                logger.info(f"Applied text replacement for {tag}: {text_value}")

            # -------------------------------
            # CASE 2: Object-based mapping (table, range, or styled text)
            # -------------------------------
            elif isinstance(value, dict):
                # Check if this is a simple text replacement
                if value.get("type") == "text" or ("range" not in value and "sheet" not in value):
                    # Text replacement with optional explicit styling
                    text_value = value.get("value", "")
                    explicit_style = value.get("style", {})
                    explicit_align = value.get("align")
                    
                    # Get value from Excel cell if it's a reference
                    if any(c.isalpha() for c in text_value) and any(c.isdigit() for c in text_value):
                        sheet_name = value.get("sheet")
                        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else (
                            wb[sheet_default] if sheet_default and sheet_default in wb.sheetnames else wb.active
                        )
                        try:
                            cell_val = ws[text_value].value
                            text_value = str(cell_val or "")
                        except Exception:
                            text_value = f"<Invalid cell {text_value}>"
                    
                    # Replace with optional explicit styling (overrides original style)
                    _replace_with_optional_style(doc, tag, text_value, explicit_style, explicit_align)
                    logger.info(f"Applied styled text replacement for {tag}: {text_value}")
                    continue

                # Otherwise, process as table/range (existing table logic)
                sheet_name = value.get("sheet")
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else (
                    wb[sheet_default] if sheet_default and sheet_default in wb.sheetnames else wb.active
                )
                range_str = value.get("range")
                insert_type = value.get("type", "row_tables2")
                align = value.get("align", "left")
                style_override = value.get("style", {}) or {}
                direction = value.get("direction", "auto")
                trim_empty = bool(value.get("trimEmpty", True))
                drop_empty_any = bool(value.get("dropEmptyCells", True))

                if not range_str:
                    logger.warning(f"No range specified for tag {tag}")
                    continue

                logger.info(f"Processing {tag} as {insert_type} from {range_str}")

                # Parse range
                try:
                    start_row, end_row, start_col_idx, end_col_idx = _parse_range_flexible(range_str, ws)
                except Exception as e:
                    logger.error(f"Failed to parse range {range_str} for tag {tag}: {e}")
                    doc.add_paragraph(f"⚠️ Invalid range '{range_str}' for {tag}")
                    continue

                rtl = False
                if direction == "rtl":
                    rtl = True
                elif direction == "auto":
                    rtl = _range_has_arabic(ws, start_row, end_row, start_col_idx, end_col_idx)
                    # Enhanced RTL logic with numeric protection
                    if rtl:
                        # Additional check: if range is mostly numeric, don't use RTL
                        numeric_cell_count = 0
                        total_cells = 0
                        for r in range(start_row, end_row + 1):
                            for c in range(start_col_idx, end_col_idx + 1):
                                cell_value = ws.cell(row=r, column=c).value
                                if cell_value is not None:
                                    total_cells += 1
                                    if _is_numeric_value(str(cell_value)):
                                        numeric_cell_count += 1
                        
                        # If most cells are numeric, don't use RTL to preserve number alignment
                        if total_cells > 0 and (numeric_cell_count / total_cells) > 0.6:  # 60% threshold
                            rtl = False
                            logger.info(f"Auto-disabled RTL for {tag} - {numeric_cell_count}/{total_cells} cells are numeric")

                paragraph = _find_paragraph_with_token(doc, tag)
                if not paragraph:
                    logger.warning(f"Tag {tag} not found in template, adding placeholder")
                    doc.add_paragraph(f"⚠️ Tag {tag} not found in template.")
                    continue

                # Clear the tag from the paragraph
                for run in paragraph.runs:
                    run.text = run.text.replace(tag, "").strip()

                if insert_type == "row_tables":
                    _insert_row_tables(
                        doc, ws,
                        start_row, end_row,
                        start_col_idx, end_col_idx,
                        paragraph, rtl, align, style_override,
                        trim_empty=trim_empty,
                        drop_empty_any=drop_empty_any
                    )
                    logger.info(f"Inserted row-by-row tables for {tag}")

                elif insert_type == "row_tables2":
                    table = _create_table_with_excel_styles(doc, ws, start_row, end_row, start_col_idx, end_col_idx, rtl=rtl)
                    # Apply forward merges across all rows where the next Excel cell is empty
                    _merge_forward_when_next_empty(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl=rtl)
                    paragraph._p.addnext(table._tbl)
                    logger.info(f"Inserted row_tables2 table with Excel styles for {tag}")
                    _debug_table_alignment_summary(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl, tag)
               
                elif insert_type == "table":
                    table = _insert_table_with_smart_merging(
                        doc, ws, start_row, end_row, start_col_idx, end_col_idx, 
                        paragraph, rtl, align, style_override
                    )
                    paragraph._p.addnext(table._tbl)
                    logger.info(f"Inserted table for {tag} with smart merging")
                    _debug_table_alignment_summary(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl, tag)
               
                elif insert_type == "text":
                    lines = []
                    for r in range(start_row, end_row + 1):
                        cols = list(range(start_col_idx, end_col_idx + 1))
                        if rtl:
                            cols = cols[::-1]
                        row_vals = []
                        for c_idx in cols:
                            val = ws.cell(row=r, column=c_idx).value
                            row_vals.append(str(val or ""))
                        lines.append("    ".join(row_vals))
                    
                    # Create styled text paragraphs
                    for line in lines:
                        new_para = paragraph.insert_paragraph_before(line)
                        new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT if (rtl or align.lower() == "right") else (
                            WD_ALIGN_PARAGRAPH.CENTER if align.lower() == "center" else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        if rtl:
                            _set_paragraph_rtl(new_para, True)
                        
                        # Apply styling to all runs in the paragraph
                        for run in new_para.runs:
                            _apply_run_styling(run, style_override)

                    logger.info(f"Inserted styled text for {tag}")

                else:
                    logger.warning(f"Unknown insert type '{insert_type}' for tag {tag}")

        # Save output (exports) using original name + date (YYYYMMDD)
        export_dir = "exports"
        os.makedirs(export_dir, exist_ok=True)
        date_stamp = datetime.now().strftime("%Y%m%d")
        orig_name = os.path.basename(template.filename) if getattr(template, 'filename', None) else "document.docx"
        stem, ext = os.path.splitext(orig_name)
        if not ext:
            ext = ".docx"
        filename = f"{stem}_{date_stamp}{ext}"
        filepath = os.path.join(export_dir, filename)
        doc.save(filepath)

        # Also save in 'Disclosures' with date as YYYY_MM_DD (e.g., 2025_10_23)
        try:
            disclosures_dir = "Disclosures"
            os.makedirs(disclosures_dir, exist_ok=True)
            date_stamp_disclosures = datetime.now().strftime("%Y_%m_%d")
            disclosures_filename = f"{stem}_{date_stamp_disclosures}{ext}"
            disclosures_path = os.path.join(disclosures_dir, disclosures_filename)
            doc.save(disclosures_path)
            logger.info(f"Saved disclosure copy: {disclosures_path}")
        except Exception as e:
            logger.warning(f"Failed to save to Disclosures: {e}")

        logger.info(f"Successfully generated Word document: {filename}")
        return FileResponse(filepath, filename=filename)

    except Exception as e:
        logger.error(f"Excel to Word conversion failed: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error generating Word file: {str(e)}")


def _range_has_arabic(ws, start_row, end_row, start_col_idx, end_col_idx):
    """Check if range contains Arabic text to determine RTL."""
    arabic_pattern = re.compile(r'[\u0600-\u06FF]')  # Arabic Unicode range
    
    arabic_cell_count = 0
    total_text_cells = 0
    
    for row in range(start_row, end_row + 1):
        for col in range(start_col_idx, end_col_idx + 1):
            cell_value = ws.cell(row=row, column=col).value
            if cell_value and str(cell_value).strip():
                str_value = str(cell_value)
                
                # Skip if it's purely numeric (with or without commas, decimals, spaces)
                if re.match(r'^[0-9,\.\s\%\(\)\-]+$', str_value):
                    continue
                
                total_text_cells += 1
                if arabic_pattern.search(str_value):
                    arabic_cell_count += 1
    
    # Only consider as RTL if significant portion contains Arabic text
    if total_text_cells > 0:
        arabic_ratio = arabic_cell_count / total_text_cells
        return arabic_ratio > 0.3  # At least 30% of text cells should be Arabic
    
    return False


def _normalize_placeholder_text(text: str) -> str:
    """Normalize text by removing spaces and zero-width/RTL marks for robust matching."""
    if text is None:
        return ""
    zero_width = [
        "\u200e",  # LRM
        "\u200f",  # RLM
        "\u202a",  # LRE
        "\u202b",  # RLE
        "\u202c",  # PDF
        "\u202d",  # LRO
        "\u202e",  # RLO
    ]
    for z in zero_width:
        text = text.replace(z, "")
    return "".join(text.split())


def _parse_range_flexible(range_str: str, ws):
    """
    Parse ranges like:
      - 'A:G'        (columns only, detect used rows)
      - 'A5:G'       (start row fixed, detect end row)
      - 'A:G60'      (detect start row, end row fixed)
      - 'A1:G60'     (fully specified)
      - 'C7'         (single cell)
    Returns: (start_row, end_row, start_col_idx, end_col_idx)
    """
    s = (range_str or "").replace(" ", "")
    if not s:
        raise HTTPException(status_code=400, detail="Empty range string")

    def split_part(part: str) -> tuple[str, int | None]:
        letters = "".join(ch for ch in part if ch.isalpha())
        digits = "".join(ch for ch in part if ch.isdigit())
        return letters, (int(digits) if digits else None)

    if ":" in s:
        left, right = s.split(":", 1)
        left_col_letters, left_row = split_part(left)
        right_col_letters, right_row = split_part(right)

        if not left_col_letters or not right_col_letters:
            raise HTTPException(status_code=400, detail=f"Invalid range '{s}': missing column letters")

        start_col_idx = _col_to_idx(left_col_letters)
        end_col_idx = _col_to_idx(right_col_letters)
        if start_col_idx < 1 or end_col_idx < 1:
            raise HTTPException(status_code=400, detail=f"Invalid range '{s}': bad column letters")

        # Ensure start_col_idx <= end_col_idx
        if start_col_idx > end_col_idx:
            start_col_idx, end_col_idx = end_col_idx, start_col_idx

        # Determine start_row and end_row
        if left_row is not None and right_row is not None:
            start_row = max(1, left_row)
            end_row = max(start_row, right_row)
        elif left_row is not None and right_row is None:
            # 'A5:G' → detect end_row from left_row to bottom
            start_row = max(1, left_row)
            _, end_row = _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=start_row)
        elif left_row is None and right_row is not None:
            # 'A:G60' → detect start_row from top, end_row fixed
            start_row, _ = _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=1)
            end_row = max(start_row, right_row)
        else:
            # 'A:G' → detect both
            start_row, end_row = _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=1)

        return start_row, end_row, start_col_idx, end_col_idx
    else:
        # Single address like 'C7' or 'AA12'
        col_letters, row = split_part(s)
        if not col_letters or row is None:
            raise HTTPException(status_code=400, detail=f"Invalid range/address '{s}'")
        col_idx = _col_to_idx(col_letters)
        start_row = end_row = max(1, row)
        return start_row, end_row, col_idx, col_idx


def _col_to_idx(col_letters: str) -> int:
    """Convert Excel column letters to 1-based index (A=1)."""
    idx = 0
    for ch in (col_letters or "").upper():
        if 'A' <= ch <= 'Z':
            idx = idx * 26 + (ord(ch) - ord('A') + 1)
    return idx


def _detect_used_row_span(ws, start_col_idx, end_col_idx, start_row_min=1):
    """Detect the actual used row span in the worksheet."""
    max_row = ws.max_row
    start_row_used = None
    end_row_used = start_row_min
    
    # Find first row with data
    for row in range(start_row_min, max_row + 1):
        has_data = False
        for col in range(start_col_idx, end_col_idx + 1):
            cell_value = ws.cell(row=row, column=col).value
            if cell_value not in (None, ""):
                has_data = True
                break
        if has_data:
            start_row_used = row
            break
    
    if start_row_used is None:
        return start_row_min, start_row_min
    
    # Find last row with data
    for row in range(max_row, start_row_used - 1, -1):
        has_data = False
        for col in range(start_col_idx, end_col_idx + 1):
            cell_value = ws.cell(row=row, column=col).value
            if cell_value not in (None, ""):
                has_data = True
                break
        if has_data:
            end_row_used = row
            break
    
    return start_row_used, end_row_used


def _find_paragraph_with_token(doc, token):
    """Find a paragraph that contains a given token like {{TABLE}}."""
    for para in doc.paragraphs:
        if token in para.text:
            return para
    return None


def _set_table_rtl(table, rtl: bool = True):
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._element.insert(0, tblPr)
        bidi = tblPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            tblPr.append(bidi)
        bidi.set(qn('w:val'), '1' if rtl else '0')
    except Exception:
        # Non-critical if not supported
        pass


def _set_paragraph_rtl(paragraph, rtl=True):
    """Set paragraph to RTL direction."""
    try:
        pPr = paragraph._element.get_or_add_pPr()
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
        bidi.set(qn('w:val'), '1' if rtl else '0')
    except Exception:
        pass


def _format_excel_value(xl_cell) -> str:
    val = xl_cell.value
    if val is None:
        return ""
    fmt = xl_cell.number_format or ""

    # Date/time formats
    try:
        if is_date_format(fmt):
            # Return like 30/Jun/2025 to match Excel-style preview
            if hasattr(val, 'strftime'):
                return val.strftime('%d/%b/%Y')
    except Exception:
        pass

    # Percent
    if isinstance(val, (int, float)) and '%' in fmt:
        try:
            # Excel typically stores 0.12 for 12%
            return f"{val:.2%}"
        except Exception:
            return str(val)

    # Thousands separators and decimals
    if isinstance(val, (int, float)):
        try:
            if '.00' in fmt or '.##' in fmt or '.0' in fmt or '.##' in fmt:
                # 2 decimal places if format implies decimals
                return format(val, ',.2f')
            if ',' in fmt:
                return format(val, ',.0f')
            return str(val)
        except Exception:
            return str(val)

    return str(val)


def _is_numeric_value(text):
    """Check if text represents a numeric value."""
    if text is None:
        return False
    # Remove commas and spaces for checking
    clean_text = str(text).replace(',', '').replace(' ', '').strip()
    try:
        float(clean_text)
        return True
    except ValueError:
        return False


def _replace_preserve_style(doc, tag, replacement_text):
    """Replace tag with text while preserving original formatting - FIXED VERSION."""
    found = False
    
    # Normalize the tag for matching
    normalized_tag = tag.strip()
    
    logger.info(f"🔍 Searching for tag: '{normalized_tag}' to replace with: '{replacement_text}'")
    
    # Search in paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        if normalized_tag in original_text:
            found = True
            logger.info(f"✅ Found '{normalized_tag}' in paragraph: '{original_text[:100]}...'")
            
            # SIMPLE APPROACH: Clear all runs and rebuild with replacement
            # Store the formatting of the first run for the replacement
            first_run_formatting = None
            if paragraph.runs:
                first_run = paragraph.runs[0]
                first_run_formatting = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'font_size': first_run.font.size,
                    'font_name': first_run.font.name,
                    'font_color': first_run.font.color.rgb if first_run.font.color.rgb else None
                }
            
            # Clear all content
            for run in paragraph.runs:
                run.text = ""
            
            # Add the replacement text with original formatting
            if first_run_formatting:
                new_run = paragraph.add_run(original_text.replace(normalized_tag, replacement_text))
                new_run.bold = first_run_formatting['bold']
                new_run.italic = first_run_formatting['italic']
                new_run.underline = first_run_formatting['underline']
                if first_run_formatting['font_size']:
                    new_run.font.size = first_run_formatting['font_size']
                if first_run_formatting['font_name']:
                    new_run.font.name = first_run_formatting['font_name']
                if first_run_formatting['font_color']:
                    new_run.font.color.rgb = first_run_formatting['font_color']
            else:
                # Fallback: just add the text
                paragraph.add_run(original_text.replace(normalized_tag, replacement_text))
            
            logger.info(f"✅ Replaced '{normalized_tag}' with '{replacement_text}' in paragraph")
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    if normalized_tag in original_text:
                        found = True
                        logger.info(f"✅ Found '{normalized_tag}' in table cell: '{original_text[:100]}...'")
                        
                        # Store formatting
                        first_run_formatting = None
                        if paragraph.runs:
                            first_run = paragraph.runs[0]
                            first_run_formatting = {
                                'bold': first_run.bold,
                                'italic': first_run.italic,
                                'underline': first_run.underline,
                                'font_size': first_run.font.size,
                                'font_name': first_run.font.name,
                                'font_color': first_run.font.color.rgb if first_run.font.color.rgb else None
                            }
                        
                        # Clear all content
                        for run in paragraph.runs:
                            run.text = ""
                        
                        # Add replacement with formatting
                        if first_run_formatting:
                            new_run = paragraph.add_run(original_text.replace(normalized_tag, replacement_text))
                            new_run.bold = first_run_formatting['bold']
                            new_run.italic = first_run_formatting['italic']
                            new_run.underline = first_run_formatting['underline']
                            if first_run_formatting['font_size']:
                                new_run.font.size = first_run_formatting['font_size']
                            if first_run_formatting['font_name']:
                                new_run.font.name = first_run_formatting['font_name']
                            if first_run_formatting['font_color']:
                                new_run.font.color.rgb = first_run_formatting['font_color']
                        else:
                            paragraph.add_run(original_text.replace(normalized_tag, replacement_text))
                        
                        logger.info(f"✅ Replaced '{normalized_tag}' with '{replacement_text}' in table cell")
    
    if not found:
        logger.warning(f"❌ Tag '{normalized_tag}' not found anywhere in the document")
    else:
        logger.info(f"✅ Successfully replaced all occurrences of '{normalized_tag}'")
    
    return found

def _replace_with_optional_style(doc, tag, replacement_text, explicit_style=None, explicit_align=None):
    """Replace tag with optional explicit styling."""
    explicit_style = explicit_style or {}
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            for run in paragraph.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, replacement_text)
                    # Apply explicit styling if provided
                    if explicit_style:
                        _apply_run_styling(run, explicit_style)
            
            # Apply paragraph alignment if specified
            if explicit_align:
                if explicit_align.lower() == "right":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif explicit_align.lower() == "center":
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _apply_run_styling(run, style_dict):
    """Apply styling to a run from a style dictionary."""
    if style_dict.get("bold"):
        run.font.bold = True
    if style_dict.get("italic"):
        run.font.italic = True
    if style_dict.get("underline"):
        run.font.underline = True
    if style_dict.get("fontSize"):
        try:
            run.font.size = Pt(float(style_dict["fontSize"]))
        except (ValueError, TypeError):
            pass
    if style_dict.get("fontFamily"):
        run.font.name = style_dict["fontFamily"]
    if style_dict.get("color"):
        try:
            run.font.color.rgb = RGBColor.from_string(style_dict["color"].lstrip('#'))
        except Exception:
            pass


def _write_cell_with_style(cell, xl_cell, style_override):
    """Write content to cell with proper styling."""
    try:
        cell_address = f"{_get_column_letter(xl_cell.column)}{xl_cell.row}"
        
        # Clear existing content
        for paragraph in cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Add new paragraph
        paragraph = cell.add_paragraph()
        
        # Get cell value
        cell_value = _format_excel_value(xl_cell)
        
        # Determine alignment
        force_rtl = style_override.get("forceRTL", False)
        text_align = style_override.get("textAlign", "left")

        # DEBUG: Log alignment decision
        logger.debug(f"📝 Cell {cell_address} alignment: "
                    f"forceRTL={force_rtl}, textAlign={text_align}, "
                    f"value='{cell_value}', value_type={type(xl_cell.value).__name__}")
        
        # Apply alignment
        if force_rtl or text_align.lower() == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logger.debug(f"   → Setting RIGHT alignment")
        elif text_align.lower() == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logger.debug(f"   → Setting CENTER alignment")
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logger.debug(f"   → Setting LEFT alignment")
        
        # Add text
        if cell_value:
            run = paragraph.add_run(cell_value)
            _apply_run_styling(run, style_override)
            
    except Exception as e:
        logger.warning(f"Failed to write cell with style: {e}")


def _should_merge_backward(text):
    """Determine if cell should be merged backward."""
    return text and len(text.strip()) > 0


def _add_comprehensive_borders_to_table(table):
    """Add comprehensive borders to entire table."""
    try:
        for row in table.rows:
            for cell in row.cells:
                # Add borders to each cell
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Remove existing borders
                for elem in list(tcPr):
                    if elem.tag.endswith('tcBorders'):
                        tcPr.remove(elem)
                
                # Add new borders
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_elem = OxmlElement(f'w:{border_name}')
                    border_elem.set(qn('w:val'), 'single')
                    border_elem.set(qn('w:sz'), '4')
                    border_elem.set(qn('w:color'), '000000')
                    border_elem.set(qn('w:space'), '0')
                    tcBorders.append(border_elem)
                
                tcPr.append(tcBorders)
    except Exception as e:
        logger.warning(f"Failed to add comprehensive borders: {e}")


def _apply_borders_above_bold_text(table):
    """Apply extra borders above rows with bold text."""
    try:
        for i, row in enumerate(table.rows):
            if i > 0:  # Skip first row
                has_bold = False
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.bold:
                                has_bold = True
                                break
                        if has_bold:
                            break
                    if has_bold:
                        break
                
                if has_bold:
                    # Add extra border above this row
                    prev_row = table.rows[i-1]
                    for cell in prev_row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is not None:
                            bottom_border = tcBorders.find(qn('w:bottom'))
                            if bottom_border is not None:
                                bottom_border.set(qn('w:val'), 'double')
                                bottom_border.set(qn('w:sz'), '8')
    except Exception as e:
        logger.warning(f"Failed to apply borders above bold text: {e}")


def _insert_row_tables(doc, ws, start_row, end_row, start_col_idx, end_col_idx, paragraph, rtl, align, style_override, trim_empty=True, drop_empty_any=True):
    """Insert data as row-by-row tables."""
    try:
        for r in range(start_row, end_row + 1):
            # Check if row should be skipped (if empty and trim_empty is True)
            if trim_empty:
                row_empty = True
                for c in range(start_col_idx, end_col_idx + 1):
                    cell_value = ws.cell(row=r, column=c).value
                    if cell_value not in (None, ""):
                        row_empty = False
                        break
                if row_empty:
                    continue
            
            # Create a single-row table
            num_cols = end_col_idx - start_col_idx + 1
            table = doc.add_table(rows=1, cols=num_cols)
            
            # Apply basic styling
            table.style = None
            table.autofit = False
            
            # Fill the row
            for c in range(num_cols):
                excel_col = start_col_idx + c
                xl_cell = ws.cell(row=r, column=excel_col)
                
                # Determine target column (handle RTL)
                target_col = (num_cols - 1 - c) if rtl else c
                
                cell_value = _format_excel_value(xl_cell)
                cell = table.cell(0, target_col)
                
                # Write cell content
                _write_cell_with_style(
                    cell,
                    xl_cell,
                    {**style_override, "forceRTL": (rtl or align.lower() == "right"), "textAlign": align}
                )
            
            # Set table alignment
            if align.lower() == "right" or rtl:
                table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            elif align.lower() == "center":
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            _set_table_rtl(table, rtl)
            
            # Insert after the paragraph
            paragraph._p.addnext(table._tbl)
            
    except Exception as e:
        logger.error(f"Failed to insert row tables: {e}")


def _apply_excel_styles_to_table(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl=True):
    """Apply ALL Excel styles (borders, colors, fonts, column widths, row heights) to Word table with RTL support."""
    try:
        # Apply column widths from Excel
        _apply_excel_column_widths(table, ws, start_col_idx, end_col_idx)
        
        # Apply row heights from Excel
        _apply_excel_row_heights(table, ws, start_row, end_row)
        
        # Apply cell styles (borders, colors, fonts) from Excel
        for r_idx in range(len(table.rows)):
            for c_idx in range(len(table.columns)):
                excel_row = start_row + r_idx
                excel_col = start_col_idx + c_idx
                
                if excel_row <= end_row and excel_col <= end_col_idx:
                    xl_cell = ws.cell(row=excel_row, column=excel_col)
                    
                    # For RTL, reverse the column order
                    if rtl:
                        target_col = len(table.columns) - 1 - c_idx
                    else:
                        target_col = c_idx
                    
                    docx_cell = table.cell(r_idx, target_col)
                    
                    # Apply ALL Excel styles to Word cell
                    _apply_excel_cell_styles(docx_cell, xl_cell)
        
        logger.info(f"Applied Excel styles to table: {len(table.rows)}x{len(table.columns)} (RTL: {rtl})")
        
    except Exception as e:
        logger.warning(f"Failed to apply Excel styles: {e}")


def _apply_excel_column_widths(table, ws, start_col_idx, end_col_idx):
    """Apply Excel column widths to Word table columns."""
    try:
        for col_idx in range(len(table.columns)):
            excel_col = start_col_idx + col_idx
            if excel_col <= end_col_idx:
                # Get Excel column width
                col_letter = _get_column_letter(excel_col)
                col_dim = ws.column_dimensions[col_letter]
                excel_width = col_dim.width
                
                # Use default width if not set
                if not excel_width or excel_width == 0:
                    excel_width = 8.43  # Excel default column width
                
                # Convert Excel width to Word width (in dxa - twentieths of a point)
                word_width = int(excel_width * 115)  # Adjusted conversion factor
                
                # Set column width in Word table
                for row in table.rows:
                    if col_idx < len(row.cells):
                        cell = row.cells[col_idx]
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # Remove existing width
                        for elem in list(tcPr):
                            if elem.tag.endswith('tcW'):
                                tcPr.remove(elem)
                        
                        # Add new width
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(word_width))
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.append(tcW)
        
        logger.info(f"Applied Excel column widths to {len(table.columns)} columns")
        
    except Exception as e:
        logger.warning(f"Failed to apply Excel column widths: {e}")


def _apply_excel_row_heights(table, ws, start_row, end_row):
    """Apply Excel row heights to Word table rows."""
    try:
        for row_idx in range(len(table.rows)):
            excel_row = start_row + row_idx
            if excel_row <= end_row:
                # Get Excel row height
                row_dim = ws.row_dimensions[excel_row]
                excel_height = row_dim.height
                
                if excel_height and excel_height > 0:
                    # Convert Excel height to Word height (in dxa)
                    word_height = int(excel_height * 20)  # Points to twentieths of points
                    
                    # Set row height in Word table
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        
                        # Remove existing height
                        for elem in list(trPr):
                            if elem.tag.endswith('trHeight'):
                                trPr.remove(elem)
                        
                        # Add new height
                        trHeight = OxmlElement('w:trHeight')
                        trHeight.set(qn('w:val'), str(word_height))
                        trHeight.set(qn('w:hRule'), 'atLeast')
                        trPr.append(trHeight)
        
        logger.info(f"Applied Excel row heights to {len(table.rows)} rows")
        
    except Exception as e:
        logger.warning(f"Failed to apply Excel row heights: {e}")


def _get_column_letter(col_idx):
    """Convert column index to Excel column letter (1->A, 2->B, etc.)."""
    letters = ""
    while col_idx > 0:
        col_idx, remainder = divmod(col_idx - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters


def _apply_excel_cell_styles(docx_cell, xl_cell):
    """Apply ALL Excel cell styles to Word cell."""
    try:
        # Clear existing content but preserve cell structure
        for paragraph in docx_cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Add new paragraph
        paragraph = docx_cell.add_paragraph()
        
        # Apply borders from Excel
        _apply_excel_borders(docx_cell, xl_cell)
        
        # Apply background color from Excel
        _apply_excel_fill_color(docx_cell, xl_cell)
        
        # Apply font styles from Excel
        _apply_excel_font_styles(paragraph, xl_cell)
        
        # Apply alignment from Excel
        _apply_excel_alignment(paragraph, xl_cell)
        
        # Apply cell value
        cell_value = _format_excel_value(xl_cell)
        if cell_value:
            run = paragraph.add_run(cell_value)
            # Re-apply font styles to the run
            _apply_excel_font_styles_to_run(run, xl_cell)
        
    except Exception as e:
        logger.warning(f"Failed to apply Excel cell styles: {e}")


def _apply_excel_borders(docx_cell, xl_cell):
    """Apply Excel cell borders to Word cell."""
    try:
        xl_border = xl_cell.border
        if not xl_border:
            return
            
        tc = docx_cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing borders
        for elem in list(tcPr):
            if elem.tag.endswith('tcBorders'):
                tcPr.remove(elem)
        
        tcBorders = OxmlElement('w:tcBorders')
        
        # Define border sides to check
        border_sides = [
            ('top', xl_border.top),
            ('left', xl_border.left),
            ('bottom', xl_border.bottom),
            ('right', xl_border.right),
        ]
        
        has_any_border = False
        
        for side_name, xl_side in border_sides:
            border_elem = OxmlElement(f'w:{side_name}')
            
            if xl_side and xl_side.style and xl_side.style != 'none':
                # Map Excel border style to Word border style
                border_style = _map_excel_border_style(xl_side.style)
                border_elem.set(qn('w:val'), border_style)
                
                # Set border size based on Excel style
                border_size = _map_excel_border_size(xl_side.style)
                border_elem.set(qn('w:sz'), str(border_size))
                
                # Set border color if available - FIXED RGB HANDLING
                border_color = '000000'  # Default black
                if xl_side.color:
                    try:
                        # Handle different color types
                        if hasattr(xl_side.color, 'rgb'):
                            rgb_value = xl_side.color.rgb
                            if rgb_value and isinstance(rgb_value, str) and len(rgb_value) >= 6:
                                border_color = rgb_value[-6:]  # Get last 6 chars for RGB
                        elif hasattr(xl_side.color, 'index'):
                            # Handle indexed colors
                            border_color = '000000'  # Default for indexed colors
                    except Exception:
                        border_color = '000000'
                
                border_elem.set(qn('w:color'), border_color)
                border_elem.set(qn('w:space'), '0')
                has_any_border = True
            else:
                # No border for this side
                border_elem.set(qn('w:val'), 'nil')
            
            tcBorders.append(border_elem)
        
        # Add inside borders as nil
        for inside_side in ['insideH', 'insideV']:
            border_elem = OxmlElement(f'w:{inside_side}')
            border_elem.set(qn('w:val'), 'nil')
            tcBorders.append(border_elem)
        
        if has_any_border:
            tcPr.append(tcBorders)
        
    except Exception as e:
        logger.warning(f"Failed to apply Excel borders: {e}")


def _map_excel_border_style(excel_style):
    """Map Excel border style to Word border style."""
    style_map = {
        'thin': 'single',
        'medium': 'single',
        'thick': 'single',
        'dashed': 'dashed',
        'dotted': 'dotted',
        'double': 'double',
        'hair': 'single',
        'mediumDashed': 'dashed',
        'dashDot': 'dashDot',
        'mediumDashDot': 'dashDot',
        'dashDotDot': 'dashDotDot',
        'mediumDashDotDot': 'dashDotDot',
        'slantDashDot': 'dashDot',
    }
    return style_map.get(excel_style, 'single')


def _map_excel_border_size(excel_style):
    """Map Excel border style to Word border size."""
    size_map = {
        'hair': 2,
        'thin': 4,
        'medium': 8,
        'thick': 16,
        'dashed': 4,
        'dotted': 4,
        'double': 8,
        'mediumDashed': 8,
        'dashDot': 4,
        'mediumDashDot': 8,
        'dashDotDot': 4,
        'mediumDashDotDot': 8,
        'slantDashDot': 4,
    }
    return size_map.get(excel_style, 4)


def _apply_excel_fill_color(docx_cell, xl_cell):
    """Apply Excel cell background color to Word cell."""
    try:
        xl_fill = xl_cell.fill
        if not xl_fill:
            return
            
        fill_type = getattr(xl_fill, 'fill_type', None) or getattr(xl_fill, 'patternType', None)
        if fill_type != 'solid':
            return
            
        color = getattr(xl_fill, 'fgColor', None) or getattr(xl_fill, 'start_color', None)
        if not color:
            return
            
        rgb_hex = 'FFFFFF'  # Default white
        
        # Handle different color types - FIXED RGB HANDLING
        try:
            if hasattr(color, 'rgb') and color.rgb:
                rgb_value = color.rgb
                if isinstance(rgb_value, str) and len(rgb_value) >= 6:
                    rgb_hex = rgb_value[-6:]  # Get last 6 chars for RGB
                # Skip if transparent or white
                if rgb_hex in ['00000000', 'FFFFFFFF']:
                    return
            elif hasattr(color, 'index'):
                # Handle indexed colors - use default white
                return
        except Exception:
            return  # Skip if there's any issue with color processing
        
        tc = docx_cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing shading
        for elem in list(tcPr):
            if elem.tag.endswith('shd'):
                tcPr.remove(elem)
        
        # Add new shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), rgb_hex)
        tcPr.append(shd)
        
    except Exception as e:
        logger.warning(f"Failed to apply Excel fill color: {e}")


def _apply_excel_font_styles(paragraph, xl_cell):
    """Apply Excel font styles to Word paragraph."""
    try:
        xl_font = xl_cell.font
        if not xl_font:
            return
            
        # Paragraph-level formatting
        if xl_font.name:
            paragraph.style.font.name = xl_font.name
        
        # Font size affects paragraph properties too
        if xl_font.size:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
                
    except Exception as e:
        logger.warning(f"Failed to apply Excel font styles to paragraph: {e}")


def _apply_excel_font_styles_to_run(run, xl_cell):
    """Apply Excel font styles to Word run."""
    try:
        xl_font = xl_cell.font
        if not xl_font:
            return
        
        # Font name
        if xl_font.name:
            run.font.name = xl_font.name
        
        # Font size
        if xl_font.size:
            try:
                run.font.size = Pt(float(xl_font.size))
            except (ValueError, TypeError):
                run.font.size = Pt(11)  # Default size
        
        # Bold
        if xl_font.bold:
            run.font.bold = True
        
        # Italic
        if xl_font.italic:
            run.font.italic = True
        
        # Underline
        if xl_font.underline and xl_font.underline != 'none':
            run.font.underline = True
        
        # Font color - FIXED RGB HANDLING
        if xl_font.color:
            try:
                if hasattr(xl_font.color, 'rgb') and xl_font.color.rgb:
                    rgb_value = xl_font.color.rgb
                    if isinstance(rgb_value, str) and len(rgb_value) >= 6:
                        rgb_hex = rgb_value[-6:]
                        if rgb_hex not in ['00000000', 'FFFFFFFF']:
                            run.font.color.rgb = RGBColor.from_string(rgb_hex)
                # Skip indexed colors for now
            except Exception:
                pass  # Skip color if there's an issue
                
        # Strike through
        if xl_font.strike:
            run.font.strike = True
            
    except Exception as e:
        logger.warning(f"Failed to apply Excel font styles to run: {e}")


def _apply_excel_alignment(paragraph, xl_cell):
    """Apply Excel cell alignment to Word paragraph."""
    try:
        xl_alignment = xl_cell.alignment
        cell_address = f"{_get_column_letter(xl_cell.column)}{xl_cell.row}"
        cell_value = xl_cell.value
        
        if not xl_alignment:
            # hereeeeeeeeeeeeeeeeeeeeee - Check if number and apply consistent alignment
            if cell_value is not None:
                # Check if it's a numeric value (integer, float, or numeric string)
                is_numeric = isinstance(cell_value, (int, float)) or _is_numeric_value(str(cell_value))
                
                if is_numeric:
                    # Apply right alignment for numbers (common practice for numeric data)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Debug log for numeric cells
                    logger.debug(f"🔢 Numeric cell {cell_address}: '{cell_value}' → RIGHT alignment (no Excel alignment)")
                else:
                    # For non-numeric cells without Excel alignment, keep default (left)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return
            
        # Horizontal alignment
        if xl_alignment.horizontal:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'general': WD_ALIGN_PARAGRAPH.LEFT,
                'fill': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'centerContinuous': WD_ALIGN_PARAGRAPH.CENTER,
                'distributed': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            word_align = align_map.get(xl_alignment.horizontal, WD_ALIGN_PARAGRAPH.LEFT)
            
            # Special handling for 'general' alignment with numeric values
            if xl_alignment.horizontal == 'general' and cell_value is not None:
                is_numeric = isinstance(cell_value, (int, float)) or _is_numeric_value(str(cell_value))
                if is_numeric:
                    # For 'general' alignment with numbers, use right alignment
                    word_align = WD_ALIGN_PARAGRAPH.RIGHT
                    logger.debug(f"🔢 Numeric cell {cell_address} with 'general' alignment: '{cell_value}' → RIGHT alignment")
            
            paragraph.alignment = word_align
            paragraph.paragraph_format.alignment = word_align

        else:
            # hereeeeeeeeeeeeeeeeeeeeee - Check if number and apply consistent alignment
            if cell_value is not None:
                # Check if it's a numeric value
                is_numeric = isinstance(cell_value, (int, float)) or _is_numeric_value(str(cell_value))
                
                if is_numeric:
                    # Apply right alignment for numbers
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Debug log for numeric cells
                    logger.debug(f"🔢 Numeric cell {cell_address}: '{cell_value}' → RIGHT alignment (no horizontal alignment)")
                else:
                    # For non-numeric cells, keep default (left)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Text rotation (vertical text)
        if xl_alignment.textRotation and xl_alignment.textRotation != 0:
            # This would require more complex XML manipulation for vertical text
            pass
            
        # Wrap text
        if xl_alignment.wrap_text:
            paragraph.paragraph_format.widow_control = True
                
    except Exception as e:
        logger.warning(f"Failed to apply Excel alignment: {e}")


def _create_table_with_excel_styles(doc, ws, start_row, end_row, start_col_idx, end_col_idx, rtl=True):
    """Create a Word table with full Excel style preservation and RTL support."""
    num_rows = end_row - start_row + 1
    num_cols = end_col_idx - start_col_idx + 1

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Remove any default styling
    table.style = None
    
    # Apply RTL settings to the entire table
    if rtl:
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        _set_table_rtl(table, True)
    
    # Apply ALL Excel styles
    _apply_excel_styles_to_table(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl)
    
    # Merge header row if first Excel row has exactly one non-empty cell
    _merge_header_row_if_single_cell(table, ws, start_row, start_col_idx, end_col_idx, rtl)
    
    # NEW: Merge forward across all rows when the next Excel cell is empty
    _merge_forward_when_next_empty(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl=rtl)
    
    return table


def _merge_header_row_if_single_cell(table, ws, start_row, start_col_idx, end_col_idx, rtl=True):
    """If the first Excel row within the range has exactly one non-empty cell, merge the entire first table row
    into a single cell and apply that Excel cell's styles and content."""
    try:
        # Determine non-empty cells in first Excel row
        non_empty_cols = []
        for c in range(start_col_idx, end_col_idx + 1):
            val = ws.cell(row=start_row, column=c).value
            if val is not None and str(val).strip():
                non_empty_cols.append(c)

        # Only proceed if exactly one cell has content
        if len(non_empty_cols) != 1:
            return

        # Use that Excel cell's content/style for the merged header cell
        xl_col = non_empty_cols[0]
        xl_cell = ws.cell(row=start_row, column=xl_col)

        # Merge across the entire first row in Word
        try:
            merged_cell = table.cell(0, 0).merge(table.cell(0, len(table.columns) - 1))
        except Exception:
            # Fallback: incremental merges if direct full merge fails
            merged_cell = table.cell(0, 0)
            for col in range(1, len(table.columns)):
                try:
                    merged_cell = merged_cell.merge(table.cell(0, col))
                except Exception:
                    pass

        # Apply Excel styles and content to the merged header cell
        _apply_excel_cell_styles(merged_cell, xl_cell)

        # Center-align the merged header content for better presentation
        try:
            for para in merged_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        logger.info("Merged entire header row into a single cell based on single non-empty Excel header cell")

    except Exception as e:
        logger.warning(f"Header merge failed: {e}")


def _enhanced_format_excel_value(xl_cell):
    """Enhanced Excel value formatting with number format preservation."""
    val = xl_cell.value
    if val is None:
        return ""
    
    fmt = xl_cell.number_format or "General"
    
    # Handle dates
    try:
        if is_date_format(fmt) and hasattr(val, 'strftime'):
            # Use Excel's date format or default to sensible format
            if 'yy' in fmt or 'yyyy' in fmt:
                return val.strftime('%d/%m/%Y')
            else:
                return val.strftime('%d/%b/%Y')
    except Exception:
        pass
    
    # Handle percentages
    if isinstance(val, (int, float)) and '%' in fmt:
        try:
            return f"{val:.1%}".replace('.0%', '%')
        except Exception:
            return str(val)
    
    # Handle currency and numbers with formatting
    if isinstance(val, (int, float)):
        try:
            # Check for currency symbols
            if '[$' in fmt or '£' in fmt or '$' in fmt:
                # Format as currency
                return f"{val:,.2f}"
            
            # Check for decimal places in format
            if '.00' in fmt or '0.0' in fmt:
                return f"{val:,.2f}"
            elif '.0' in fmt:
                return f"{val:,.1f}"
            elif '.#' in fmt or '0' in fmt:
                return f"{val:,.0f}"
            
            # Default number formatting
            return f"{val:,.2f}"
            
        except Exception:
            return str(val)
    
    # Handle booleans
    if isinstance(val, bool):
        return "نعم" if val else "لا"  # Arabic for Yes/No
    
    return str(val)


def _debug_excel_alignment_enhanced(xl_cell, paragraph, rtl_context=False, align_setting="left", tag_name=""):
    """Enhanced debug function to log Excel alignment details with comprehensive analysis."""
    try:
        xl_alignment = xl_cell.alignment
        cell_value = xl_cell.value
        cell_address = f"{_get_column_letter(xl_cell.column)}{xl_cell.row}"
        
        # Get alignment details
        excel_horizontal = getattr(xl_alignment, 'horizontal', None) if xl_alignment else None
        excel_vertical = getattr(xl_alignment, 'vertical', None) if xl_alignment else None
        excel_wrap_text = getattr(xl_alignment, 'wrap_text', None) if xl_alignment else None
        excel_text_rotation = getattr(xl_alignment, 'textRotation', None) if xl_alignment else None
        
        # Determine content type
        is_numeric = isinstance(cell_value, (int, float))
        is_numeric_text = _is_numeric_value(str(cell_value or ""))
        is_date = isinstance(cell_value, datetime) or (hasattr(cell_value, 'strftime'))
        is_arabic = bool(re.search(r'[\u0600-\u06FF]', str(cell_value or "")))
        
        # Get final alignment
        final_alignment = paragraph.alignment
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER", 
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY"
        }
        final_align_name = alignment_map.get(final_alignment, "UNKNOWN")
        
        # Get all runs in paragraph
        run_details = []
        for i, run in enumerate(paragraph.runs):
            run_details.append({
                'index': i,
                'text': run.text,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'font_size': run.font.size,
                'font_name': run.font.name
            })
        
        # Log all details
        logger.info(f"🔍 ALIGNMENT DEBUG - Cell {cell_address} (Tag: {tag_name})")
        logger.info(f"   📊 Value: '{cell_value}' (type: {type(cell_value).__name__})")
        logger.info(f"   🔢 Type Analysis: Numeric={is_numeric}, NumericText={is_numeric_text}, Date={is_date}, Arabic={is_arabic}")
        logger.info(f"   📐 Excel Alignment: horizontal='{excel_horizontal}', vertical='{excel_vertical}'")
        logger.info(f"   📋 Excel Formatting: wrap_text={excel_wrap_text}, text_rotation={excel_text_rotation}")
        logger.info(f"   🌍 Context: RTL={rtl_context}, AlignSetting='{align_setting}'")
        logger.info(f"   ✅ Final Alignment: {final_align_name}")
        logger.info(f"   📝 Number Format: '{xl_cell.number_format}'")
        logger.info(f"   🏷️  Paragraph Runs: {len(paragraph.runs)}")
        
        # Show run details if multiple runs
        if len(paragraph.runs) > 1:
            for run_detail in run_details:
                logger.info(f"        Run {run_detail['index']}: '{run_detail['text']}' (bold={run_detail['bold']})")
        
        # Decision process analysis
        logger.info(f"   🤔 DECISION ANALYSIS:")
        
        if rtl_context:
            if is_numeric or is_numeric_text:
                logger.info(f"        RTL + Numeric → Should be RIGHT alignment")
            elif is_date:
                logger.info(f"        RTL + Date → Should be RIGHT alignment") 
            elif is_arabic:
                logger.info(f"        RTL + Arabic → Should be RIGHT alignment")
            else:
                logger.info(f"        RTL + Non-Arabic Text → Should be RIGHT alignment")
        else:
            if excel_horizontal:
                logger.info(f"        LTR + Excel horizontal='{excel_horizontal}' → Should follow Excel")
            else:
                logger.info(f"        LTR + No Excel alignment → Should use '{align_setting}'")
        
        # Check for potential issues
        issues = []
        if rtl_context and final_align_name == "LEFT" and (is_arabic or not (is_numeric or is_numeric_text)):
            issues.append("❌ RTL context but got LEFT alignment for non-numeric content")
        if not rtl_context and final_align_name == "RIGHT" and not is_arabic and not excel_horizontal == 'right':
            issues.append("❌ LTR context but got RIGHT alignment without explicit setting")
        if excel_horizontal and excel_horizontal != 'general' and final_align_name != alignment_map.get(WD_ALIGN_PARAGRAPH.RIGHT if excel_horizontal == 'right' else WD_ALIGN_PARAGRAPH.LEFT, "UNKNOWN"):
            issues.append(f"❌ Excel alignment '{excel_horizontal}' not respected")
        
        if issues:
            logger.info(f"   ⚠️  POTENTIAL ISSUES:")
            for issue in issues:
                logger.info(f"        {issue}")
        else:
            logger.info(f"   ✅ No alignment issues detected")
                
        logger.info("   " + "="*60)
        
        return {
            'cell_address': cell_address,
            'value': cell_value,
            'value_type': type(cell_value).__name__,
            'excel_horizontal': excel_horizontal,
            'final_alignment': final_align_name,
            'rtl_context': rtl_context,
            'is_numeric': is_numeric,
            'is_arabic': is_arabic,
            'issues': issues
        }
        
    except Exception as e:
        logger.warning(f"Alignment debug failed: {e}")
        return None


def _debug_table_alignment_summary(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl, tag_name):
    """Debug alignment for entire table and provide summary."""
    logger.info(f"📊 TABLE ALIGNMENT SUMMARY - {tag_name}")
    logger.info(f"   Range: {start_row}-{end_row}, cols {start_col_idx}-{end_col_idx}")
    logger.info(f"   RTL Setting: {rtl}")
    logger.info(f"   Table Size: {len(table.rows)} rows x {len(table.columns)} columns")
    
    alignment_stats = {
        'LEFT': 0,
        'RIGHT': 0, 
        'CENTER': 0,
        'JUSTIFY': 0,
        'UNKNOWN': 0
    }
    
    cell_issues = []
    
    for r_idx in range(len(table.rows)):
        for c_idx in range(len(table.columns)):
            excel_row = start_row + r_idx
            excel_col = start_col_idx + c_idx
            
            if excel_row <= end_row and excel_col <= end_col_idx:
                xl_cell = ws.cell(row=excel_row, column=excel_col)
                docx_cell = table.cell(r_idx, c_idx)
                
                if docx_cell.paragraphs:
                    paragraph = docx_cell.paragraphs[0]
                    alignment_map = {
                        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
                        WD_ALIGN_PARAGRAPH.CENTER: "CENTER", 
                        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
                        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY"
                    }
                    align_name = alignment_map.get(paragraph.alignment, "UNKNOWN")
                    alignment_stats[align_name] += 1
                    
                    # Debug specific cells that might have issues
                    cell_value = xl_cell.value
                    if cell_value and str(cell_value).strip():
                        is_numeric = _is_numeric_value(str(cell_value))
                        is_arabic = bool(re.search(r'[\u0600-\u06FF]', str(cell_value)))
                        
                        # Check for potential alignment issues
                        if rtl and align_name == "LEFT" and not is_numeric and is_arabic:
                            cell_issues.append(f"Cell {_get_column_letter(excel_col)}{excel_row}: '{cell_value}' - RTL context but LEFT aligned Arabic text")
                        elif not rtl and align_name == "RIGHT" and not is_arabic and not is_numeric:
                            cell_issues.append(f"Cell {_get_column_letter(excel_col)}{excel_row}: '{cell_value}' - LTR context but RIGHT aligned non-Arabic text")
    
    logger.info(f"   📈 Alignment Distribution:")
    for align, count in alignment_stats.items():
        if count > 0:
            percentage = (count / (len(table.rows) * len(table.columns))) * 100
            logger.info(f"        {align}: {count} cells ({percentage:.1f}%)")
    
    if cell_issues:
        logger.info(f"   ⚠️  Alignment Issues Found:")
        for issue in cell_issues[:10]:  # Show first 10 issues
            logger.info(f"        {issue}")
        if len(cell_issues) > 10:
            logger.info(f"        ... and {len(cell_issues) - 10} more issues")
    else:
        logger.info(f"   ✅ No alignment issues detected in table")
    
    logger.info("   " + "="*60)
    
    return {
        'alignment_stats': alignment_stats,
        'cell_issues': cell_issues,
        'total_cells': len(table.rows) * len(table.columns)
    }


def _should_merge_with_next_cell(ws, current_row, current_col, max_cols):
    """
    Check if current cell should merge with next cell:
    NEW RULE: Merge whenever the next Excel cell is empty and current has content.
    """
    try:
        # No next cell at the end of the range
        if current_col >= max_cols:
            return False

        current_cell = ws.cell(row=current_row, column=current_col)
        next_cell = ws.cell(row=current_row, column=current_col + 1)

        current_value = str(current_cell.value or "").strip()
        next_value = str(next_cell.value or "").strip()

        # Unconditional forward merge when next is empty (and current has content)
        next_is_empty = not next_value
        should_merge = next_is_empty and bool(current_value)

        if should_merge:
            logger.debug(
                f"🔄 Simple forward merge: Cell {_get_column_letter(current_col)}{current_row} "
                f"→ {_get_column_letter(current_col + 1)}{current_row} | Reason: empty-next"
            )

        return should_merge

    except Exception as e:
        logger.warning(f"Failed to check simple forward merge: {e}")
        return False


def _merge_forward_when_next_empty(table, ws, start_row, end_row, start_col_idx, end_col_idx, rtl=True, align=None):
    """
    Merge forward across all rows:
    If the next Excel cell(s) are empty (and current has content), merge current Word cell with the chain.
    """
    try:
        num_rows = end_row - start_row + 1
        num_cols = end_col_idx - start_col_idx + 1

        for r_idx in range(num_rows):
            c_idx = 0
            while c_idx < num_cols:
                excel_row = start_row + r_idx
                excel_col_cur = start_col_idx + c_idx
                xl_cur = ws.cell(row=excel_row, column=excel_col_cur)

                cur_text = str(xl_cur.value or "").strip()
                if not cur_text:
                    c_idx += 1
                    continue

                # Find a chain of empty adjacent cells to the right
                chain_end = c_idx + 1
                while chain_end < num_cols:
                    excel_col_next = start_col_idx + chain_end
                    next_text = str(ws.cell(row=excel_row, column=excel_col_next).value or "").strip()
                    if next_text:
                        break
                    chain_end += 1

                if chain_end > c_idx + 1:
                    # Compute Word target columns respecting RTL
                    target_col_start = (num_cols - 1 - c_idx) if rtl else c_idx
                    target_col_last = (num_cols - 1 - (chain_end - 1)) if rtl else (chain_end - 1)

                    # Merge across the entire empty chain
                    try:
                        merged_cell = table.cell(r_idx, target_col_start).merge(table.cell(r_idx, target_col_last))
                    except Exception:
                        # Fallback: stepwise merge if direct span fails
                        merged_cell = table.cell(r_idx, target_col_start)
                        for mc in range(c_idx + 1, chain_end):
                            tc_next = (num_cols - 1 - mc) if rtl else mc
                            try:
                                merged_cell = merged_cell.merge(table.cell(r_idx, tc_next))
                            except Exception:
                                pass

                    # Apply explicit alignment setting if provided
                    if align:
                        try:
                            for para in merged_cell.paragraphs:
                                if align.lower() == "right":
                                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                elif align.lower() == "center":
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        except Exception:
                            pass

                    # Default to RIGHT alignment when no explicit align is provided
                    if not align:
                        try:
                            for para in merged_cell.paragraphs:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        except Exception:
                            pass

                    # Skip past merged chain
                    c_idx = chain_end
                else:
                    c_idx += 1

        logger.info("Applied forward chain merges across rows where next cells are empty")

    except Exception as e:
        logger.warning(f"Failed forward merging pass: {e}")


def _insert_table_with_smart_merging(doc, ws, start_row, end_row, start_col_idx, end_col_idx, paragraph, rtl, align, style_override):
    """Insert table with smart merging for multi-line and long text."""
    num_rows = end_row - start_row + 1
    num_cols = end_col_idx - start_col_idx + 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Apply comprehensive borders
    _add_comprehensive_borders_to_table(table)
    table.autofit = True

    # Fill rows with smart merging
    for r in range(num_rows):
        c = 0
        while c < num_cols:
            excel_row = start_row + r
            excel_col = start_col_idx + c
            
            xl_cur = ws.cell(row=excel_row, column=excel_col)
            cur_text = str(xl_cur.value or "").strip()
            cur_is_empty = not cur_text

            # Check previous cell for backward merging
            has_prev = c > 0
            prev_is_empty = False
            if has_prev:
                xl_prev = ws.cell(row=excel_row, column=start_col_idx + (c - 1))
                prev_text = str(xl_prev.value or "").strip()
                prev_is_empty = not prev_text

            # Target col positions in Word (handle RTL)
            target_col_cur = (num_cols - 1 - c) if rtl else c
            target_col_prev = (num_cols - 1 - (c - 1)) if rtl else (c - 1)

            # RULE 1: Merge backward if current is empty AND previous exists AND previous is not empty
            if has_prev and cur_is_empty and not prev_is_empty and _should_merge_backward(prev_text):
                merged_cell = table.cell(r, target_col_prev).merge(table.cell(r, target_col_cur))
                _write_cell_with_style(
                    merged_cell,
                    xl_prev,
                    {**style_override, "forceRTL": (rtl or align.lower() == "right"), "textAlign": align}
                )
                c += 1
                continue

            # RULE 2: Merge forward if current cell has content AND next cell is empty
            should_merge_forward = _should_merge_with_next_cell(ws, excel_row, excel_col, end_col_idx)
            
            if should_merge_forward and c < num_cols - 1:
                # Find a chain of empty cells to the right
                chain_end = c + 1
                while chain_end < num_cols:
                    excel_col_next = start_col_idx + chain_end
                    next_text = str(ws.cell(row=excel_row, column=excel_col_next).value or "").strip()
                    if next_text:
                        break
                    chain_end += 1

                # Merge from current to the last empty in the chain
                target_col_last = (num_cols - 1 - (chain_end - 1)) if rtl else (chain_end - 1)
                merged_cell = table.cell(r, target_col_cur).merge(table.cell(r, target_col_last))
                
                # Enforce RIGHT alignment for forward-merged content
                _write_cell_with_style(
                    merged_cell,
                    xl_cur,
                    {**style_override, "forceRTL": (rtl or align.lower() == "right"), "textAlign": "left"}
                )

                # Advance past the merged chain
                c = chain_end
                continue

            # Normal cell writing
            _write_cell_with_style(
                table.cell(r, target_col_cur),
                xl_cur,
                {**style_override, "forceRTL": (rtl or align.lower() == "right"), "textAlign": align}
            )
            c += 1

    # Apply additional formatting
    _apply_borders_above_bold_text(table)

    # Align the table block itself
    if align.lower() == "right" or rtl:
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    elif align.lower() == "center":
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    else:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_rtl(table, rtl)

    # Merge header row if first Excel row has exactly one non-empty cell
    _merge_header_row_if_single_cell(table, ws, start_row, start_col_idx, end_col_idx, rtl)

    return table


@router.get("/api/reports/files/templates/download/{filename}")
async def download_template_file(filename: str):
    """Download a template file."""
    try:
        # URL decode the filename to handle spaces and special characters
        decoded_filename = urllib.parse.unquote(filename)
        base_dir = os.path.dirname(__file__)
        templates_dir = os.path.join(base_dir, "template")
        file_path = os.path.join(templates_dir, decoded_filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        return FileResponse(file_path, filename=filename)
    except Exception as e:
        logger.error(f"Error downloading template file {filename}: {e}")
        raise HTTPException(status_code=500, detail="Failed to download file")

@router.get("/api/reports/files/disclosures/download/{filename}")
async def download_disclosure_file(filename: str):
    """Download a disclosure file."""
    try:
        # URL decode the filename to handle spaces and special characters
        decoded_filename = urllib.parse.unquote(filename)
        base_dir = os.path.dirname(__file__)
        disclosures_dir = os.path.join(base_dir, "Disclosures")
        file_path = os.path.join(disclosures_dir, decoded_filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        return FileResponse(file_path, filename=filename)
    except Exception as e:
        logger.error(f"Error downloading disclosure file {filename}: {e}")
        raise HTTPException(status_code=500, detail="Failed to download file")

@router.delete("/api/reports/files/templates/delete/{filename}")
async def delete_template_file(filename: str):
    """Delete a template file."""
    try:
        # URL decode the filename to handle spaces and special characters
        decoded_filename = urllib.parse.unquote(filename)
        base_dir = os.path.dirname(__file__)
        templates_dir = os.path.join(base_dir, "template")
        file_path = os.path.join(templates_dir, decoded_filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        os.remove(file_path)
        logger.info(f"Deleted template file: {filename}")
        return {"message": "File deleted successfully"}
    except HTTPException:
        # Re-raise HTTP exceptions (like 404) as-is
        raise
    except Exception as e:
        logger.error(f"Error deleting template file {filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")

@router.delete("/api/reports/files/disclosures/delete/{filename}")
async def delete_disclosure_file(filename: str):
    """Delete a disclosure file."""
    try:
        # URL decode the filename to handle spaces and special characters
        decoded_filename = urllib.parse.unquote(filename)
        base_dir = os.path.dirname(__file__)
        disclosures_dir = os.path.join(base_dir, "Disclosures")
        file_path = os.path.join(disclosures_dir, decoded_filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        os.remove(file_path)
        logger.info(f"Deleted disclosure file: {filename}")
        return {"message": "File deleted successfully"}
    except HTTPException:
        # Re-raise HTTP exceptions (like 404) as-is
        raise
    except Exception as e:
        logger.error(f"Error deleting disclosure file {filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")