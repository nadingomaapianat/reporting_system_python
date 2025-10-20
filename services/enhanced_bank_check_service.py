"""
Enhanced Bank Check Service for Arabic/English check processing
Extracts data using webhook and generates formatted Excel/Word reports
"""
import os
import json
import httpx
from datetime import datetime
from typing import Dict, Any, Tuple, List
from pathlib import Path
from io import BytesIO

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class EnhancedBankCheckService:
    def __init__(self):
        self.webhook_url = "https://n8nio.pianat.ai/webhook/ocr-check"
        self.exports_dir = Path("exports")
        self.exports_dir.mkdir(exist_ok=True)

    async def extract_check_data(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract data from check image using webhook"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                files = {
                    'file': (filename, file_bytes, 'application/octet-stream')
                }
                response = await client.post(self.webhook_url, files=files)
                
                if response.status_code == 200:
                    data = response.json()
                    print(f"[EnhancedBankCheck] Webhook response: {data}")
                    return self._normalize_webhook_data(data)
                else:
                    raise Exception(f"Webhook failed with status {response.status_code}")
        except Exception as e:
            print(f"[EnhancedBankCheck] Webhook error: {e}")
            return self._get_empty_data()

    def _normalize_webhook_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Normalize webhook response to our standard format"""
        # Handle the specific webhook response format
        if isinstance(data, list) and len(data) > 0:
            output = data[0].get('output', '{}')
            if isinstance(output, str):
                try:
                    parsed_output = json.loads(output)
                    data = parsed_output
                except:
                    pass
        
        # Extract amount information
        amount_info = data.get('amount', {})
        if isinstance(amount_info, dict):
            amount_value = amount_info.get('value', '')
            amount_text = amount_info.get('text', '')
            currency = amount_info.get('currency', 'EGP')
        else:
            amount_value = data.get('amount_value', '')
            amount_text = data.get('amount_text', '')
            currency = data.get('currency', 'EGP')
        
        # Extract issuer information
        issuer_info = data.get('issuer', {})
        if isinstance(issuer_info, dict):
            issuer_name = issuer_info.get('name', '')
            issuer_signature = issuer_info.get('signature_present', False)
        else:
            issuer_name = data.get('issuer_name', '')
            issuer_signature = data.get('issuer_signature', False)
        
        return {
            "bank_name": data.get('bank_name', ''),
            "date": data.get('date', ''),
            "payee_name": data.get('payee_name', ''),
            "amount_value": amount_value,
            "amount_text": amount_text,
            "currency": currency,
            "status_note": data.get('status_note', ''),
            "issuer_signature": issuer_signature,
            "issuer_name": issuer_name,
            "serial_number": data.get('serial_number', ''),
            "raw_text": data.get('raw_text', '')
        }

    def _get_empty_data(self) -> Dict[str, Any]:
        """Return empty data structure"""
        return {
            "bank_name": "",
            "date": "",
            "payee_name": "",
            "amount_value": "",
            "amount_text": "",
            "currency": "EGP",
            "status_note": "",
            "issuer_signature": False,
            "issuer_name": "",
            "serial_number": "",
            "raw_text": ""
        }

    def generate_excel_report(self, records: List[Dict[str, Any]]) -> bytes:
        """Generate Excel report matching the professional template structure"""
        wb = Workbook()
        ws = wb.active
        ws.title = "تقرير الشيكات المصرفية"

        # Define styles
        title_font = Font(name='Arial', size=16, bold=True, color='000000')
        header_font = Font(name='Arial', size=12, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        data_font = Font(name='Arial', size=10)
        total_font = Font(name='Arial', size=11, bold=True)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Main title
        ws.merge_cells('A1:L1')
        title_cell = ws.cell(row=1, column=1, value="تقرير الشيكات المصرفية - Bank Check Report")
        title_cell.font = title_font
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Date and time
        ws.cell(row=2, column=1, value=f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        ws.cell(row=2, column=1).font = Font(name='Arial', size=10)
        
        # Section 1: Main Receivables Table
        ws.cell(row=4, column=1, value="اجمالي المستحقات - Total Receivables").font = Font(name='Arial', size=14, bold=True)
        
        # Headers matching the Excel template exactly
        headers = [
            "اسم المشروع", "رقم العمارة", "رقم الشقة", "اسم العميل", 
            "رقم الشيك", "البنك المسحوب على مبلغ الشيك", "تاريخ استحقاق الشيك",
            "تاريخ اليوم المحدث", "تاريخ التحصيل", "عدد الايام المتبقية",
            "حالة التحصيل", "اجمالي المستحقات"
        ]
        
        # Write headers (starting from row 6)
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=6, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border

        # Write data rows
        for row_idx, record in enumerate(records, 7):
            row_data = [
                record.get('project_name', ''),
                record.get('building_no', ''),
                record.get('apartment_no', ''),
                record.get('customer_name', ''),
                record.get('check_no', ''),
                record.get('bank_name', ''),
                record.get('check_due_date', ''),
                record.get('updated_date', ''),
                record.get('collection_date', ''),
                record.get('remaining_days', 0),
                record.get('collection_status', ''),
                record.get('total_receivables', 0)
            ]
            
            for col, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.font = data_font
                cell.border = border
                if col == 12:  # Amount column
                    cell.number_format = '#,##0'

        # Add total row
        total_row = len(records) + 7
        total_amount = sum(record.get('total_receivables', 0) for record in records)
        
        ws.cell(row=total_row, column=11, value="الاجمالي").font = total_font
        ws.cell(row=total_row, column=12, value=total_amount).font = total_font
        ws.cell(row=total_row, column=12).number_format = '#,##0'
        
        # Style total row
        for col in range(1, 13):
            cell = ws.cell(row=total_row, column=col)
            cell.fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
            cell.border = border

        # Section 2: Summary Statistics
        summary_start_row = total_row + 3
        ws.cell(row=summary_start_row, column=1, value="ملخص الإحصائيات - Summary Statistics").font = Font(name='Arial', size=14, bold=True)
        
        # Summary data
        summary_data = [
            ("إجمالي عدد الشيكات", "Total Number of Checks", len(records)),
            ("إجمالي المبلغ", "Total Amount", f"{total_amount:,} EGP"),
            ("متوسط المبلغ", "Average Amount", f"{total_amount // len(records) if records else 0:,} EGP"),
            ("تاريخ التقرير", "Report Date", datetime.now().strftime('%Y-%m-%d')),
            ("وقت التقرير", "Report Time", datetime.now().strftime('%H:%M:%S'))
        ]
        
        for i, (arabic, english, value) in enumerate(summary_data):
            row = summary_start_row + 2 + i
            ws.cell(row=row, column=1, value=arabic).font = data_font
            ws.cell(row=row, column=2, value=english).font = data_font
            ws.cell(row=row, column=3, value=value).font = Font(name='Arial', size=10, bold=True)

        # Section 3: Collection Status Breakdown
        status_start_row = summary_start_row + 8
        ws.cell(row=status_start_row, column=1, value="توزيع حالة التحصيل - Collection Status Breakdown").font = Font(name='Arial', size=14, bold=True)
        
        # Count by status
        status_counts = {}
        for record in records:
            status = record.get('collection_status', 'غير محدد')
            status_counts[status] = status_counts.get(status, 0) + 1
        
        status_row = status_start_row + 2
        for status, count in status_counts.items():
            ws.cell(row=status_row, column=1, value=status).font = data_font
            ws.cell(row=status_row, column=2, value=count).font = data_font
            status_row += 1

        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 25)
            ws.column_dimensions[column_letter].width = adjusted_width

        # Save to bytes
        buffer = BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    def generate_word_report(self, records: List[Dict[str, Any]]) -> bytes:
        """Generate Word report with check records table"""
        doc = Document()

        # Title
        title = doc.add_heading('تقرير الشيكات المصرفية', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_heading('Bank Check Records Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report info
        info_para = doc.add_paragraph()
        info_para.add_run(f"تاريخ التقرير / Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"عدد السجلات / Number of Records: {len(records)}\n")
        info_para.add_run(f"إجمالي المبلغ / Total Amount: {sum(record.get('total_receivables', 0) for record in records):,} EGP")

        # Main data table
        doc.add_heading('سجل الشيكات / Check Records', level=3)
        
        table = doc.add_table(rows=1, cols=12)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Table headers
        hdr_cells = table.rows[0].cells
        headers = [
            "اسم المشروع", "رقم العمارة", "رقم الشقة", "اسم العميل", 
            "رقم الشيك", "البنك", "تاريخ الاستحقاق", "تاريخ التحديث",
            "تاريخ التحصيل", "الأيام المتبقية", "الحالة", "المبلغ"
        ]
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Table data
        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.get('project_name', '')
            row_cells[1].text = record.get('building_no', '')
            row_cells[2].text = record.get('apartment_no', '')
            row_cells[3].text = record.get('customer_name', '')
            row_cells[4].text = record.get('check_no', '')
            row_cells[5].text = record.get('bank_name', '')
            row_cells[6].text = record.get('check_due_date', '')
            row_cells[7].text = record.get('updated_date', '')
            row_cells[8].text = record.get('collection_date', '')
            row_cells[9].text = str(record.get('remaining_days', 0))
            row_cells[10].text = record.get('collection_status', '')
            row_cells[11].text = f"{record.get('total_receivables', 0):,}"

        # Summary section
        doc.add_heading('ملخص التقرير / Report Summary', level=3)
        summary_para = doc.add_paragraph()
        total_amount = sum(record.get('total_receivables', 0) for record in records)
        summary_para.add_run(f"إجمالي عدد السجلات: {len(records)}\n")
        summary_para.add_run(f"إجمالي المبلغ: {total_amount:,} EGP\n")
        summary_para.add_run(f"متوسط المبلغ لكل شيك: {total_amount / len(records) if records else 0:,.2f} EGP")

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run("تم إنشاء هذا التقرير تلقائياً بواسطة نظام تحليل الشيكات المصرفية")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    async def process_check(self, file_bytes: bytes, filename: str) -> Tuple[bytes, bytes, Dict[str, Any]]:
        """Complete check processing workflow"""
        # Extract data
        data = await self.extract_check_data(file_bytes, filename)
        
        # Create a single record from the extracted data
        record = {
            'project_name': 'تاج سيتي',
            'building_no': 'A-101',
            'apartment_no': '101',
            'customer_name': data.get('payee_name', ''),
            'check_no': data.get('serial_number', ''),
            'bank_name': data.get('bank_name', ''),
            'check_due_date': data.get('date', ''),
            'updated_date': datetime.now().strftime('%Y-%m-%d'),
            'collection_date': '',
            'remaining_days': 0,
            'collection_status': 'لم يحن موعد السداد',
            'total_receivables': data.get('amount_value', 0)
        }
        
        # Generate reports
        excel_bytes = self.generate_excel_report([record])
        word_bytes = self.generate_word_report([record])
        
        return excel_bytes, word_bytes, data

    async def process_records(self, records: List[Dict[str, Any]]) -> Tuple[bytes, bytes]:
        """Process multiple records and generate reports"""
        # Generate reports
        excel_bytes = self.generate_excel_report(records)
        word_bytes = self.generate_word_report(records)
        
        return excel_bytes, word_bytes

    async def extract_excel_headers(self, file_bytes: bytes, filename: str) -> List[str]:
        """Extract headers from Excel template file"""
        try:
            from openpyxl import load_workbook
            from io import BytesIO
            
            # Load the Excel file
            wb = load_workbook(BytesIO(file_bytes))
            ws = wb.active
            
            # Extract headers from the first row
            headers = []
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                if cell_value:
                    headers.append(str(cell_value).strip())
            
            print(f"[EnhancedBankCheck] Extracted {len(headers)} headers from {filename}")
            return headers
            
        except Exception as e:
            print(f"[EnhancedBankCheck] Error extracting headers: {e}")
            # Return default headers if extraction fails
            return [
                "اسم المشروع", "رقم العمارة", "رقم الشقة", "اسم العميل", 
                "رقم الشيك", "البنك المسحوب على مبلغ الشيك", "تاريخ استحقاق الشيك",
                "تاريخ اليوم المحدث", "تاريخ التحصيل", "عدد الايام المتبقية",
                "حالة التحصيل", "اجمالي المستحقات"
            ]