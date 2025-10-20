from io import BytesIO
from datetime import datetime
from typing import Dict, Any, Tuple
import os
import re
import glob
import unicodedata
import httpx
import httpx

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
try:
    import cv2  # optional
    import numpy as np  # optional
    OPENCV_AVAILABLE = True
except Exception:
    OPENCV_AVAILABLE = False


class BankCheckService:
    def __init__(self) -> None:
        pass

    def _guess_poppler_path(self) -> str | None:
        # 1) Env
        env_path = os.getenv('POPPLER_PATH')
        if env_path and os.path.isfile(os.path.join(env_path, 'pdfinfo.exe')):
            return env_path
        # 2) Common windows installs within project
        cwd = os.getcwd()
        candidates = []
        candidates += glob.glob(os.path.join(cwd, 'poppler-*', 'Library', 'bin'))
        candidates += glob.glob(os.path.join(cwd, 'poppler-*', 'library', 'bin'))
        for cand in candidates:
            if os.path.isfile(os.path.join(cand, 'pdfinfo.exe')):
                return cand
        return None

    def _guess_tesseract_cmd(self) -> str | None:
        env_cmd = os.getenv('TESSERACT_CMD')
        if env_cmd and os.path.isfile(env_cmd):
            return env_cmd
        defaults = [
            r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
            r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
        ]
        for p in defaults:
            if os.path.isfile(p):
                return p
        return None

    def extract_text(self, pdf_bytes: bytes) -> str:
        # Prefer external OCR webhook if configured
        webhook_url = os.getenv('OCR_WEBHOOK_URL')
        if webhook_url:
            try:
                with httpx.Client(timeout=60) as client:
                    files = {'file': ('upload.pdf', pdf_bytes, 'application/pdf')}
                    resp = client.post(webhook_url, files=files)
                    if resp.status_code == 200:
                        data = resp.json()
                        text = data.get('text') or data.get('textSnippet') or ''
                        if text:
                            print(f"[BankCheck][Webhook] Used external OCR, text length: {len(text)}")
                            return text
            except Exception as e:
                print(f"[BankCheck][Webhook] OCR webhook failed: {e}")
        reader = PdfReader(BytesIO(pdf_bytes))
        text_chunks = []
        for page in reader.pages:
            try:
                text_chunks.append(page.extract_text() or "")
            except Exception:
                continue
        full_text = "\n".join(text_chunks)
        # If no text extracted (likely scanned PDF), try OCR
        if len(full_text.strip()) == 0:
            try:
                # Allow configuring paths via environment variables for Windows installs
                tess_cmd = self._guess_tesseract_cmd()
                if tess_cmd:
                    pytesseract.pytesseract.tesseract_cmd = tess_cmd
                poppler_path = self._guess_poppler_path()

                images = convert_from_bytes(pdf_bytes, dpi=400, poppler_path=poppler_path if poppler_path else None)
                ocr_texts = []
                for img in images:
                    # optional OpenCV preprocessing; fallback to PIL if not available
                    if OPENCV_AVAILABLE:
                        np_img = np.array(img)
                        if len(np_img.shape) == 3:
                            gray = cv2.cvtColor(np_img, cv2.COLOR_BGR2GRAY)
                        else:
                            gray = np_img
                        gray = cv2.bilateralFilter(gray, 9, 75, 75)
                        bw = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 15)
                        kernel = np.ones((2,2), np.uint8)
                        bw = cv2.morphologyEx(bw, cv2.MORPH_CLOSE, kernel, iterations=1)
                        text_img = Image.fromarray(bw)
                    else:
                        try:
                            gray = img.convert('L')
                            bw = gray.point(lambda x: 0 if x < 180 else 255, '1')
                            text_img = bw
                        except Exception:
                            text_img = img
                    lang = os.getenv('TESS_LANG', 'ara+eng')
                    # try two PSMs and pick the longer result
                    t1 = pytesseract.image_to_string(text_img, lang=lang, config='--oem 1 --psm 6')
                    t2 = pytesseract.image_to_string(text_img, lang=lang, config='--oem 1 --psm 4')
                    ocr_texts.append(t1 if len(t1) >= len(t2) else t2)
                full_text = "\n".join(ocr_texts)
                print(f"[BankCheck][OCR] Used OCR, text length: {len(full_text)}")
            except Exception as ocr_err:
                print(f"[BankCheck][OCR] Failed OCR: {ocr_err}")
        return full_text

    def diagnose_extraction(self, pdf_bytes: bytes) -> Dict[str, Any]:
        diagnostics: Dict[str, Any] = {
            "popplerPath": self._guess_poppler_path() or '',
            "tesseractCmd": self._guess_tesseract_cmd() or '',
            "ocrUsed": False,
            "ocrError": '',
            "textLength": 0,
            "textSnippet": '',
        }
        # Try normal text
        reader = PdfReader(BytesIO(pdf_bytes))
        text_chunks = []
        for page in reader.pages:
            try:
                text_chunks.append(page.extract_text() or "")
            except Exception:
                continue
        full_text = "\n".join(text_chunks)
        if len(full_text.strip()) == 0:
            try:
                tess_cmd = self._guess_tesseract_cmd()
                if tess_cmd:
                    pytesseract.pytesseract.tesseract_cmd = tess_cmd
                poppler_path = self._guess_poppler_path()
                images = convert_from_bytes(pdf_bytes, dpi=400, poppler_path=poppler_path if poppler_path else None)
                ocr_texts = []
                for img in images:
                    if OPENCV_AVAILABLE:
                        np_img = np.array(img)
                        if len(np_img.shape) == 3:
                            gray = cv2.cvtColor(np_img, cv2.COLOR_BGR2GRAY)
                        else:
                            gray = np_img
                        gray = cv2.bilateralFilter(gray, 9, 75, 75)
                        bw = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 15)
                        kernel = np.ones((2,2), np.uint8)
                        bw = cv2.morphologyEx(bw, cv2.MORPH_CLOSE, kernel, iterations=1)
                        text_img = Image.fromarray(bw)
                    else:
                        try:
                            gray = img.convert('L')
                            bw = gray.point(lambda x: 0 if x < 180 else 255, '1')
                            text_img = bw
                        except Exception:
                            text_img = img
                    lang = os.getenv('TESS_LANG', 'ara+eng')
                    t1 = pytesseract.image_to_string(text_img, lang=lang, config='--oem 1 --psm 6')
                    t2 = pytesseract.image_to_string(text_img, lang=lang, config='--oem 1 --psm 4')
                    ocr_texts.append(t1 if len(t1) >= len(t2) else t2)
                full_text = "\n".join(ocr_texts)
                diagnostics["ocrUsed"] = True
            except Exception as e:
                diagnostics["ocrError"] = str(e)
        diagnostics["textLength"] = len(full_text or '')
        diagnostics["textSnippet"] = (full_text[:1000] if full_text else '')
        return diagnostics

    def extract_fields_from_pdf(self, pdf_bytes: bytes) -> Dict[str, Any]:
        # If webhook is configured and returns fields, use it
        webhook_url = os.getenv('OCR_WEBHOOK_URL')
        if webhook_url:
            try:
                with httpx.Client(timeout=60) as client:
                    files = {'file': ('upload.pdf', pdf_bytes, 'application/pdf')}
                    resp = client.post(webhook_url, files=files)
                    if resp.status_code == 200:
                        data = resp.json()
                        # Try to normalize keys to our schema
                        mapped = {
                            'bankName': data.get('bankName') or data.get('bank') or data.get('bank_name') or '',
                            'branch': data.get('branch') or '',
                            'currency': data.get('currency') or data.get('ccy') or '',
                            'amountNumeric': data.get('amountNumeric') or data.get('amount') or data.get('amount_number') or '',
                            'amountWords': data.get('amountWords') or data.get('amount_text') or '',
                            'date': data.get('date') or data.get('checkDate') or '',
                            'payee': data.get('payee') or data.get('to') or data.get('beneficiary') or '',
                            'rawText': data.get('text') or data.get('textSnippet') or ''
                        }
                        print(f"[BankCheck][Webhook] Fields from webhook: {mapped}")
                        return mapped
            except Exception as e:
                print(f"[BankCheck][Webhook] Field extraction via webhook failed: {e}")

        full_text = self.extract_text(pdf_bytes)
        # Debug: show length and a snippet of text extracted
        try:
            print(f"[BankCheck] Extracted text length: {len(full_text)}")
            snippet = full_text[:500].replace("\n", " ")
            print(f"[BankCheck] Text snippet: {snippet}")
        except Exception:
            pass

        # Normalize Arabic text for robust matching
        def normalize_text(s: str) -> str:
            if not s:
                return ""
            # NFKC normalize
            s = unicodedata.normalize('NFKC', s)
            # unify Arabic forms (yeh/aleph/teh marbuta)
            arabic_map = {
                '\u0649': '\u064A',  # alef maksura -> yeh
                '\u0629': '\u0647',  # teh marbuta -> heh (approx)
                '\u0640': '',        # tatweel
            }
            for k, v in arabic_map.items():
                s = s.replace(k, v)
            # convert Arabic-Indic digits to Latin
            digits_ar = '٠١٢٣٤٥٦٧٨٩'
            for i, d in enumerate(digits_ar):
                s = s.replace(d, str(i))
            # collapse spaces
            s = re.sub(r"[\u0610-\u061A\u064B-\u065F\u0670]", "", s)  # remove harakat/diacritics
            s = re.sub(r"\s+", " ", s).strip()
            return s

        ntext = normalize_text(full_text)

        # naive parsing; improved with Arabic-aware rules
        def find_between(label: str) -> str:
            idx = ntext.lower().find(label.lower())
            if idx == -1:
                return ""
            segment = ntext[idx: idx + 200]
            parts = segment.split(":", 1)
            if len(parts) == 2:
                return parts[1].split("\n", 1)[0].strip()
            return ""

        # Regex helpers for typical check patterns
        date_match = (
            re.search(r"\b(\d{4})\s*/\s*(\d{1,2})\s*/\s*(\d{1,2})\b", ntext)
            or re.search(r"\b(\d{2}[\-/]\d{2}[\-/]\d{4})\b", ntext)
        )
        ccy_match = re.search(r"\b(EGP|USD|EUR|SAR|AED|GBP)\b", ntext, flags=re.I)
        # Numeric amount commonly shown with thousands commas and optional decimals
        amount_num_match = None
        if ccy_match:
            # try capture number near currency
            currency = ccy_match.group(1)
            near = ntext[ccy_match.start():ccy_match.start() + 80]
            amount_num_match = re.search(r"([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)", near)
        if not amount_num_match:
            amount_num_match = re.search(r"#\s*([0-9,\.]+)\s*#", ntext) or re.search(r"([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)", ntext)

        # Amount in words (Arabic): "فقط ... لا غير"
        amount_words_match_ar = re.search(r"فقط\s+(.+?)\s+لا\s+غير", ntext)
        # English fallback: '... Only'
        amount_words_match_en = re.search(r"([A-Z][A-Za-z\s-]+ Only)\**", ntext)

        # Payee detection: after 'PAY TO THE ORDER OF' or between ** **
        payee_match = re.search(r"PAY\s+TO\s+THE\s+ORDER\s+OF\s+\**([^\n*]+)\**", ntext, flags=re.I)
        if not payee_match:
            # Arabic payee often near the long line area; heuristic: pick a long Arabic word group before amount words
            payee_match = re.search(r"([\u0621-\u064A][\u0621-\u064A\s'.-]{2,})", ntext)

        # Bank & branch
        bank_name = None
        if re.search(r"Commercial\s+International\s+Bank", ntext, flags=re.I):
            bank_name = "Commercial International Bank"
        elif re.search(r"\bCIB\b", ntext):
            bank_name = "CIB"
        elif re.search(r"بنك\s*مصر|BANQUE\s*MISR", ntext, flags=re.I):
            bank_name = "Banque Misr"
        branch_match = re.search(r"([A-Z][A-Z\s]+)\s+BRANCH", ntext)

        extracted: Dict[str, Any] = {
            "bankName": bank_name or find_between("Bank Name"),
            "branch": branch_match.group(1).title().strip() if branch_match else "",
            "currency": (ccy_match.group(1).upper() if ccy_match else ""),
            "amountNumeric": amount_num_match.group(1) if amount_num_match else "",
            "amountWords": (amount_words_match_ar.group(1).strip() if amount_words_match_ar else (amount_words_match_en.group(1).strip() if amount_words_match_en else "")),
            "date": ("/".join(date_match.groups()) if (date_match and len(date_match.groups())==3) else (date_match.group(1) if date_match else find_between("Date"))),
            "payee": (payee_match.group(1).strip() if payee_match else ""),
            "rawText": ntext[:1000]
        }

        try:
            print(f"[BankCheck] Extracted fields: {extracted}")
        except Exception:
            pass
        return extracted

    def generate_excel(self, extracted: Dict[str, Any], form: Dict[str, Any], full_text: str | None = None) -> bytes:
        data = [{
            "Bank Name": form.get("bankName") or extracted.get("bankName"),
            "Branch": extracted.get("branch") or "",
            "Currency": extracted.get("currency") or "",
            "Amount (Numeric)": form.get("cost") or extracted.get("amountNumeric"),
            "Amount (Words)": extracted.get("amountWords") or "",
            "Date": form.get("date") or extracted.get("date"),
            "Payee": extracted.get("payee") or "",
            "Days Remaining": form.get("daysRemaining"),
            "Raw Text (first 500)": (full_text[:500].replace("\n", " ") if full_text else ""),
        }]
        df = pd.DataFrame(data)
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Bank Check")
        return buffer.getvalue()

    def fill_word_template(self, extracted: Dict[str, Any], form: Dict[str, Any]) -> bytes:
        # minimal template synthesis without external .docx template
        doc = Document()
        doc.add_heading("Bank Check Report", level=1)
        doc.add_paragraph(f"Bank Name: {form.get('bankName') or extracted.get('bankName') or ''}")
        if extracted.get('branch'):
            doc.add_paragraph(f"Branch: {extracted.get('branch')}")
        if extracted.get('currency'):
            doc.add_paragraph(f"Currency: {extracted.get('currency')}")
        # Prefer explicit cost, else numeric amount from PDF
        doc.add_paragraph(f"Amount: {form.get('cost') or extracted.get('amountNumeric') or ''}")
        if extracted.get('amountWords'):
            doc.add_paragraph(f"Amount (Words): {extracted.get('amountWords')}")
        doc.add_paragraph(f"Date: {form.get('date') or extracted.get('date') or ''}")
        if extracted.get('payee'):
            doc.add_paragraph(f"Payee: {extracted.get('payee')}")
        doc.add_paragraph(f"Days Remaining: {form.get('daysRemaining') or ''}")
        doc.add_paragraph(f"Generated At: {datetime.now().isoformat(timespec='seconds')}")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def process(self, pdf_bytes: bytes, form: Dict[str, Any]) -> Tuple[bytes, bytes]:
        try:
            print(f"[BankCheck] Incoming PDF bytes: {len(pdf_bytes)}")
            print(f"[BankCheck] Incoming form: {form}")
        except Exception:
            pass
        extracted = self.extract_fields_from_pdf(pdf_bytes)
        # Recompute full_text snippet for Excel debug column
        try:
            # simple reuse: extract text again for snippet without OCR to avoid heavy cost
            reader = PdfReader(BytesIO(pdf_bytes))
            snippet_text = "\n".join([(p.extract_text() or "") for p in reader.pages])
        except Exception:
            snippet_text = ""
        excel_bytes = self.generate_excel(extracted, form, snippet_text if snippet_text else None)
        word_bytes = self.fill_word_template(extracted, form)
        try:
            print(f"[BankCheck] Excel size: {len(excel_bytes)} bytes, Word size: {len(word_bytes)} bytes")
        except Exception:
            pass
        return excel_bytes, word_bytes


