from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import FileResponse
import pyodbc
from config import get_database_connection_string
from export_utils import get_default_header_config

# Create router for dynamic report endpoints
router = APIRouter()

@router.post("/api/reports/dynamic")
async def generate_dynamic_report(request: Request):
    """Generate dynamic report based on table selection, joins, columns, and conditions"""
    try:
        body = await request.json()
        print(f"[DynamicReport] Request body: {body}")
        
        tables = body.get('tables', [])
        joins = body.get('joins', [])
        columns = body.get('columns', [])
        where_conditions = body.get('whereConditions', [])
        time_filter = body.get('timeFilter')
        format_type = body.get('format', 'excel')
        header_config = body.get('headerConfig', {})
        
        print(f"[DynamicReport] Tables: {tables}, Columns: {columns}, Format: {format_type}")
        print(f"[DynamicReport] Header Config: {header_config}")
        print(f"[DynamicReport] Header Config Type: {type(header_config)}")
        print(f"[DynamicReport] Header Config Keys: {header_config.keys() if isinstance(header_config, dict) else 'Not a dict'}")
        print(f"[DynamicReport] Header Config includeHeader: {header_config.get('includeHeader') if isinstance(header_config, dict) else 'N/A'}")
        print(f"[DynamicReport] Header Config title: {header_config.get('title') if isinstance(header_config, dict) else 'N/A'}")
        print(f"[DynamicReport] Header Config subtitle: {header_config.get('subtitle') if isinstance(header_config, dict) else 'N/A'}")
        print(f"[DynamicReport] Header Config fontColor: {header_config.get('fontColor') if isinstance(header_config, dict) else 'N/A'}")
        
        if not tables or not columns:
            raise HTTPException(status_code=400, detail="Tables and columns are required")
        
        # Build SQL query
        sql_query = build_dynamic_sql_query(tables, joins, columns, where_conditions, time_filter)
        print(f"[DynamicReport] SQL Query: {sql_query}")
        print(f"[DynamicReport] Columns: {columns}")
        print(f"[DynamicReport] Tables: {tables}")
        
        # Execute query and get data
        connection_string = get_database_connection_string()
        print(f"[DynamicReport] Connecting to database...")
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            print(f"[DynamicReport] Executing query...")
            cursor.execute(sql_query)
            rows = cursor.fetchall()
            print(f"[DynamicReport] Found {len(rows)} rows")
            
            # Convert to list of dictionaries
            data_rows = []
            for row in rows:
                data_rows.append([str(cell) if cell is not None else '' for cell in row])
            
            print(f"[DynamicReport] Generating {format_type} report...")
            # Get default header config and merge with user config
            default_config = get_default_header_config("dynamic")
            print(f"[DynamicReport] Default config: {default_config}")
            print(f"[DynamicReport] User config: {header_config}")
            merged_config = {**default_config, **header_config}
            print(f"[DynamicReport] Merged config: {merged_config}")
            
            # Generate report based on format
            if format_type == 'excel':
                return generate_excel_report(columns, data_rows, merged_config)
            elif format_type == 'word':
                return generate_word_report(columns, data_rows, merged_config)
            elif format_type == 'pdf':
                return generate_pdf_report(columns, data_rows, merged_config)
            else:
                raise HTTPException(status_code=400, detail="Unsupported format")
                
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        print(f"[DynamicReport] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate dynamic report: {str(e)}")

def build_dynamic_sql_query(tables, joins, columns, where_conditions, time_filter):
    """Build SQL query based on dynamic report configuration"""
    # Start with SELECT clause
    select_columns = []
    for col in columns:
        if '.' in col:
            # Split table.column and quote each part separately
            table_part, column_part = col.split('.', 1)
            select_columns.append(f"[{table_part}].[{column_part}]")
        else:
            # Add to first table
            select_columns.append(f"[{tables[0]}].[{col}]")
    
    query = f"SELECT {', '.join(select_columns)}"
    
    # Add FROM clause
    query += f" FROM [{tables[0]}]"
    
    # Add JOINs
    for join in joins:
        if join.get('leftTable') and join.get('rightTable') and join.get('leftColumn') and join.get('rightColumn'):
            join_type = join.get('type', 'INNER')
            query += f" {join_type} JOIN [{join['rightTable']}] ON [{join['leftTable']}].[{join['leftColumn']}] = [{join['rightTable']}].[{join['rightColumn']}]"
    
    # Add WHERE conditions
    where_clauses = []
    for condition in where_conditions:
        if condition.get('column') and condition.get('operator') and condition.get('value'):
            col_name = condition['column']
            if '.' in col_name:
                # Split table.column and quote each part separately
                table_part, column_part = col_name.split('.', 1)
                col_name = f"[{table_part}].[{column_part}]"
            else:
                col_name = f"[{col_name}]"
            where_clauses.append(f"{col_name} {condition['operator']} '{condition['value']}'")
    
    # Add time filter
    if time_filter and time_filter.get('column') and time_filter.get('startDate') and time_filter.get('endDate'):
        col_name = time_filter['column']
        if '.' in col_name:
            # Split table.column and quote each part separately
            table_part, column_part = col_name.split('.', 1)
            col_name = f"[{table_part}].[{column_part}]"
        else:
            col_name = f"[{col_name}]"
        where_clauses.append(f"{col_name} BETWEEN '{time_filter['startDate']}' AND '{time_filter['endDate']}'")
    
    if where_clauses:
        query += " WHERE " + " AND ".join(where_clauses)
    
    return query

default_header_config = {
    "includeHeader": True,
    "title": "ADIB GRC - Risk & Control Report",
    "subtitle": "Confidential Internal Report",
    "fontColor": "#1F4E79",
    "showDate": True,
    "showBankInfo": True,
    "bankName": "Abu Dhabi Islamic Bank Egypt",
    "bankAddress": "Head Office: Corniche El Nil, Cairo, Egypt",
    "bankPhone": "+20 2 1234 5678",
    "bankWebsite": "https://www.adib.eg",
    "tableHeaderBgColor": "#1F4E79",
    "tableBodyBgColor": "#FFFFFF",
    "excelZebraStripes": True,
    "excelAutoFitColumns": True,
    "excelFreezeTopRow": True,
    "footerShowDate": True,
    "footerShowConfidentiality": True,
    "footerConfidentialityText": "Confidential Report - For Internal Use Only"
}

def generate_excel_report(columns, data, header_config=None):
    wb = Workbook()
    ws = wb.active
    ws.title = "Dynamic Report 777"

    # --- Header config ---
    title = header_config.get("title", "Dynamic Report8888") if header_config else "Dynamic Report55"
    subtitle = header_config.get("subtitle", "") if header_config else "rr"
    show_date = header_config.get("showDate", True) if header_config else True
    bank_name = header_config.get("bankName", "") if header_config else ""
    bank_address = header_config.get("bankAddress", "") if header_config else ""

    # --- Title ---
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(columns))
    ws["A1"] = title
    ws["A1"].font = Font(size=16, bold=True, color="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")

    # --- Subtitle / Bank Info ---
    if subtitle:
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(columns))
        ws["A2"] = subtitle
        ws["A2"].font = Font(size=12, italic=True, color="1F4E79")
        ws["A2"].alignment = Alignment(horizontal="center")

    info_row = 3
    if bank_name:
        ws.merge_cells(start_row=info_row, start_column=1, end_row=info_row, end_column=len(columns))
        ws[f"A{info_row}"] = bank_name
        ws[f"A{info_row}"].alignment = Alignment(horizontal="center")
        info_row += 1
    if bank_address:
        ws.merge_cells(start_row=info_row, start_column=1, end_row=info_row, end_column=len(columns))
        ws[f"A{info_row}"] = bank_address
        ws[f"A{info_row}"].alignment = Alignment(horizontal="center")
        info_row += 1
    if show_date:
        ws.merge_cells(start_row=info_row, start_column=1, end_row=info_row, end_column=len(columns))
        ws[f"A{info_row}"] = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws[f"A{info_row}"].alignment = Alignment(horizontal="center")
        info_row += 1

    start_data_row = info_row + 1

    # --- Table Header ---
    for col_index, col_name in enumerate(columns, start=1):
        cell = ws.cell(row=start_data_row, column=col_index, value=col_name)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E79", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    # --- Table Data ---
    for row_idx, row_data in enumerate(data, start=start_data_row + 1):
        for col_idx, col_name in enumerate(columns, start=1):
            ws.cell(row=row_idx, column=col_idx, value=row_data.get(col_name, ""))

    # Auto column width
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 3

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def generate_word_report(columns, data_rows, header_config=None):
    """Generate Word report with custom headers and data"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        import os
        from datetime import datetime
        
        doc = Document()
        
        # Create exports directory if it doesn't exist
        export_dir = "exports"
        if not os.path.exists(export_dir):
            os.makedirs(export_dir)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"dynamic_report_{timestamp}.docx"
        filepath = os.path.join(export_dir, filename)
        
        # Add custom header if configured
        if header_config and header_config.get('includeHeader', True):
            # Add title
            title_text = header_config.get('title', 'Dynamic Repor999yyyt')
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle
            subtitle_text = header_config.get('subtitle', 'Generated Report')
            subtitle = doc.add_paragraph(subtitle_text)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(12)
            
            # Add report info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            info_para.add_run(f"Total Records: {len(data_rows)}\n")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Add basic title
            title = doc.add_heading('Dynamic Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add report info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            info_para.add_run(f"Total Records: {len(data_rows)}\n")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add table
        table = doc.add_table(rows=1, cols=len(columns))
        
        # Apply table style safely
        try:
            table.style = 'Table Grid'
        except (ValueError, KeyError):
            try:
                table.style = 'Light Grid'
            except (ValueError, KeyError):
                try:
                    table.style = 'Table Normal'
                except (ValueError, KeyError):
                    pass  # Use default style
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, column in enumerate(columns):
            header_cells[i].text = column
            # Style header cells
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Add data rows
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                row_cells[i].text = str(cell_value)
        
        # Save file
        doc.save(filepath)
        
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"X-Export-Src": filepath}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Word report: {str(e)}")

def generate_pdf_report(columns, data_rows, header_config=None):
    """Generate PDF report with custom headers and data"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        import os
        from datetime import datetime
        
        # Create exports directory if it doesn't exist
        export_dir = "exports"
        if not os.path.exists(export_dir):
            os.makedirs(export_dir)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"dynamic_report_{timestamp}.pdf"
        filepath = os.path.join(export_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        # Add custom header if configured
        if header_config and header_config.get('includeHeader', True):
            # Add title
            title_text = header_config.get('title', 'Dynamic Report')
            title = Paragraph(title_text, title_style)
            story.append(title)
            
            # Add subtitle
            subtitle_text = header_config.get('subtitle', 'Generated Report')
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=20,
                alignment=1  # Center alignment
            )
            subtitle = Paragraph(subtitle_text, subtitle_style)
            story.append(subtitle)
            
            # Add report info
            info_text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>Total Records: {len(data_rows)}"
            info = Paragraph(info_text, styles['Normal'])
            story.append(info)
            story.append(Spacer(1, 20))
        else:
            # Add basic title
            title = Paragraph("Dynamic Report", title_style)
            story.append(title)
            
            # Add report info
            info_text = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>Total Records: {len(data_rows)}"
            info = Paragraph(info_text, styles['Normal'])
            story.append(info)
            story.append(Spacer(1, 20))
        
        # Prepare table data
        table_data = [columns] + data_rows
        
        # Create table
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        
        # Build PDF
        doc.build(story)
        
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type="application/pdf",
            headers={"X-Export-Src": filepath}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF report: {str(e)}")

@router.post("/api/reports/schedule")
async def save_report_schedule(request: Request):
    """Save scheduled report configuration"""
    try:
        body = await request.json()
        report_config = body.get('reportConfig', {})
        schedule = body.get('schedule', {})
        
        # Save to database (you can create a scheduled_reports table)
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        
        try:
            # Create table if not exists
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='scheduled_reports' and xtype='U')
                CREATE TABLE scheduled_reports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    report_config NVARCHAR(MAX) NOT NULL,
                    schedule_config NVARCHAR(MAX) NOT NULL,
                    is_active BIT DEFAULT 1,
                    created_at DATETIME2 DEFAULT GETDATE()
                );
            """)
            conn.commit()
            
            # Insert schedule
            import json
            cursor.execute("""
                INSERT INTO scheduled_reports (report_config, schedule_config)
                VALUES (?, ?)
            """, json.dumps(report_config), json.dumps(schedule))
            conn.commit()
            
            return {"success": True, "message": "Schedule saved successfully"}
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save schedule: {str(e)}")
