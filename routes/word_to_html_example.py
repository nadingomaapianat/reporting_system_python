"""
مثال توضيحي: كيف يتم تحويل Word إلى HTML
Example: How to convert Word document to HTML
"""

from docx import Document
from docx.shared import RGBColor, Pt
from io import BytesIO
import html

def paragraph_to_html_example(paragraph):
	"""
	مثال: تحويل فقرة من Word إلى HTML
	Example: Convert Word paragraph to HTML
	"""
	html_parts = []
	
	# 1. معالجة كل run في الفقرة
	# Process each run in the paragraph
	for run in paragraph.runs:
		text = run.text if hasattr(run, 'text') else ''
		
		if not text:
			continue
		
		# 2. Escape HTML entities (مثل &, <, >)
		# Escape HTML entities (like &, <, >)
		text = html.escape(text)
		
		# 3. بناء style attribute مع جميع التنسيقات
		# Build style attribute with all formatting
		styles = []
		
		# أ) اسم الخط / Font name
		if hasattr(run, 'font') and run.font:
			if hasattr(run.font, 'name') and run.font.name:
				styles.append(f"font-family: '{run.font.name}'")
			
			# ب) حجم الخط / Font size
			if hasattr(run.font, 'size') and run.font.size:
				size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 12700
				styles.append(f"font-size: {size_pt:.1f}pt")
			
			# ج) لون الخط / Font color
			if hasattr(run.font, 'color') and run.font.color:
				if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
					rgb = run.font.color.rgb
					if isinstance(rgb, RGBColor):
						# RGBColor stores as hex string like "FF0000"
						# Extract R, G, B values
						rgb_str = str(rgb)
						# Convert hex string to integer
						rgb_int = int(rgb_str, 16) if len(rgb_str) == 6 else int(rgb_str, 16) if rgb_str.startswith('0x') else int(rgb_str)
						r = (rgb_int >> 16) & 0xFF
						g = (rgb_int >> 8) & 0xFF
						b = rgb_int & 0xFF
						styles.append(f"color: rgb({r}, {g}, {b})")
		
		# 4. بناء HTML tags للتنسيقات
		# Build HTML tags for formatting
		formatting_tags = []
		if hasattr(run, 'bold') and run.bold:
			formatting_tags.append('strong')
		if hasattr(run, 'italic') and run.italic:
			formatting_tags.append('em')
		if hasattr(run, 'underline') and run.underline:
			formatting_tags.append('u')
		
		# 5. بناء HTML النهائي
		# Build final HTML
		if styles:
			style_attr = '; '.join(styles)
			text = f'<span style="{style_attr}">{text}</span>'
		
		# Wrap with formatting tags
		for tag in reversed(formatting_tags):
			text = f'<{tag}>{text}</{tag}>'
		
		html_parts.append(text)
	
	result = ''.join(html_parts)
	
	# 6. إضافة المحاذاة / Add alignment
	if paragraph.paragraph_format and hasattr(paragraph.paragraph_format, 'alignment'):
		alignment = paragraph.paragraph_format.alignment
		if alignment and 'RIGHT' in str(alignment):
			result = f'<div style="text-align: right;">{result}</div>'
	
	return result


# مثال عملي / Practical example
if __name__ == "__main__":
	# إنشاء مثال Word
	# Create example Word document
	doc = Document()
	
	# فقرة 1: نص عادي
	# Paragraph 1: Normal text
	para1 = doc.add_paragraph()
	run1 = para1.add_run("هذا نص عادي")
	run1.font.size = Pt(12)
	
	# فقرة 2: نص عريض وأزرق
	# Paragraph 2: Bold and blue text
	para2 = doc.add_paragraph()
	run2 = para2.add_run("هذا نص عريض وأزرق")
	run2.bold = True
	run2.font.color.rgb = RGBColor(0, 0, 255)
	run2.font.size = Pt(14)
	run2.font.name = 'Arial'
	
	# فقرة 3: نص مائل وأحمر
	# Paragraph 3: Italic and red text
	para3 = doc.add_paragraph()
	run3 = para3.add_run("هذا نص مائل وأحمر")
	run3.italic = True
	run3.font.color.rgb = RGBColor(255, 0, 0)
	
	# التحويل إلى HTML
	# Convert to HTML
	print("=" * 60)
	print("مثال: تحويل Word إلى HTML")
	print("Example: Word to HTML conversion")
	print("=" * 60)
	
	print("\n1. Paragraph 1 (Normal):")
	html1 = paragraph_to_html_example(para1)
	print(f"   Word: {para1.text}")
	print(f"   HTML: {html1}")
	
	print("\n2. Paragraph 2 (Bold, Blue, Arial, 14pt):")
	html2 = paragraph_to_html_example(para2)
	print(f"   Word: {para2.text}")
	print(f"   HTML: {html2}")
	
	print("\n3. Paragraph 3 (Italic, Red):")
	html3 = paragraph_to_html_example(para3)
	print(f"   Word: {para3.text}")
	print(f"   HTML: {html3}")
	
	print("\n" + "=" * 60)
	print("النتيجة النهائية في Rich Text Editor:")
	print("Final result in Rich Text Editor:")
	print("=" * 60)
	print(f"<p>{html1}</p>")
	print(f"<p>{html2}</p>")
	print(f"<p>{html3}</p>")

