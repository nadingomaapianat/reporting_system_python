"""
API routes for the reporting system
"""
import asyncio
import json
from datetime import datetime
import os
import httpx
from fastapi import APIRouter, Query, HTTPException, Request, UploadFile, File, Form
from fastapi.responses import Response, FileResponse, StreamingResponse
from typing import Optional

from services import APIService, PDFService, ExcelService, ControlService, IncidentService
from services.bank_check_service import BankCheckService
from services.enhanced_bank_check_service import EnhancedBankCheckService
# Dashboard activity service moved to Node.js backend (NestJS)
DashboardActivityService = None  # type: ignore
from utils.export_utils import get_default_header_config
from models import ExportRequest, ExportResponse
from routes.route_utils import write_debug, parse_header_config, merge_header_config, convert_to_boolean

# Initialize services
api_service = APIService()
pdf_service = PDFService()
excel_service = ExcelService()
control_service = ControlService()
incident_service = IncidentService() if IncidentService else None
dashboard_activity_service = DashboardActivityService() if DashboardActivityService else None
bank_check_service = BankCheckService()
enhanced_bank_check_service = EnhancedBankCheckService()

# db_service points to control_service for backward compatibility
db_service = control_service

# Create router
router = APIRouter()






# Bank Check Processing: upload PDF + form and return generated files
@router.post("/api/reports/bank-check")
async def create_bank_check_report(
    request: Request,
    file: UploadFile = File(...),
    bankName: str = Form(None),
    cost: str = Form(None),
    date: str = Form(None),
    daysRemaining: str = Form(None),
    format: str = Form("excel")
):
    try:
        content = await file.read()
        excel_bytes, word_bytes = bank_check_service.process(content, {
            "bankName": bankName,
            "cost": cost,
            "date": date,
            "daysRemaining": daysRemaining,
        })

        # allow query param override: ?format=excel|word|zip
        fmt = request.query_params.get('format') or (format or 'excel')
        fmt = (fmt or 'excel').lower()

        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        if fmt == 'excel':
            filename = f"bank_check_{ts}.xlsx"
            return Response(
                content=excel_bytes,
                media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                headers={'Content-Disposition': f'attachment; filename="{filename}"'}
            )
        elif fmt == 'word':
            filename = f"bank_check_{ts}.docx"
            return Response(
                content=word_bytes,
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': f'attachment; filename="{filename}"'}
            )
        else:
            # package into a simple zip in-memory
            import zipfile
            from io import BytesIO
            buffer = BytesIO()
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(f"bank_check_{ts}.xlsx", excel_bytes)
                zf.writestr(f"bank_check_{ts}.docx", word_bytes)
            buffer.seek(0)
            filename = f"bank_check_{ts}.zip"
            return Response(content=buffer.getvalue(), media_type='application/zip', headers={'Content-Disposition': f'attachment; filename="{filename}"'})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create bank check report: {str(e)}")

# Debug: return OCR text and parsed fields without generating files
@router.post("/api/reports/bank-check/debug")
async def debug_bank_check(
    file: UploadFile = File(...)
):
    try:
        content = await file.read()
        webhook_url = os.getenv('OCR_WEBHOOK_URL')
        if webhook_url:
            async with httpx.AsyncClient(timeout=60) as client:
                files = {
                    'file': (file.filename or 'upload.bin', content, file.content_type or 'application/octet-stream')
                }
                resp = await client.post(webhook_url, files=files)
                if resp.status_code == 200:
                    data = resp.json()
                    # Best-effort mapping for UI
                    mapped_fields = {
                        'bankName': data.get('bankName') or data.get('bank') or data.get('bank_name') or '',
                        'branch': data.get('branch') or '',
                        'currency': data.get('currency') or data.get('ccy') or '',
                        'amountNumeric': data.get('amountNumeric') or data.get('amount') or data.get('amount_number') or '',
                        'amountWords': data.get('amountWords') or data.get('amount_text') or '',
                        'date': data.get('date') or data.get('checkDate') or '',
                        'payee': data.get('payee') or data.get('to') or data.get('beneficiary') or '',
                    }
                    text_value = data.get('text') or data.get('textSnippet') or ''
                    return {
                        'textSnippet': text_value[:1000] if text_value else '',
                        'textLength': len(text_value or ''),
                        'fields': mapped_fields,
                        'diagnostics': {
                            'source': 'webhook',
                            'webhookUrl': webhook_url,
                            'status': resp.status_code
                        }
                    }
                else:
                    raise HTTPException(status_code=502, detail=f"OCR webhook failed with status {resp.status_code}")

        # Fallback: local OCR if webhook not configured
        full_text = bank_check_service.extract_text(content)
        fields = bank_check_service.extract_fields_from_pdf(content)
        diag = bank_check_service.diagnose_extraction(content)
        return {
            "textSnippet": (full_text[:1000] if full_text else ""),
            "textLength": len(full_text or ""),
            "fields": fields,
            "diagnostics": diag
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Debug extraction failed: {str(e)}")

# Enhanced Bank Check Processing with professional reports
@router.post("/api/reports/enhanced-bank-check")
async def create_enhanced_bank_check_report(
    request: Request
):
    """Process bank check records and generate professional Excel/Word reports"""
    try:
        # Check if it's a file upload or JSON data
        content_type = request.headers.get('content-type', '')
        
        if 'multipart/form-data' in content_type:
            # Handle file upload
            form = await request.form()
            file = form.get('file')
            if not file:
                raise HTTPException(status_code=400, detail="No file provided")
            
            content = await file.read()
            excel_bytes, word_bytes, data = await enhanced_bank_check_service.process_check(content, file.filename or 'check.pdf')
            
            # Create ZIP with both files
            import zipfile
            from io import BytesIO
            import os
            
            buffer = BytesIO()
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            date_folder = datetime.now().strftime('%Y-%m-%d')
            
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(f"bank_check_analysis_{ts}.xlsx", excel_bytes)
                zf.writestr(f"bank_check_report_{ts}.docx", word_bytes)
            buffer.seek(0)
            
            # Save to persistent storage
            os.makedirs(f"exports/{date_folder}", exist_ok=True)
            filename = f"bank_check_reports_{ts}.zip"
            file_path = f"exports/{date_folder}/{filename}"
            
            with open(file_path, 'wb') as f:
                f.write(buffer.getvalue())
            
            return Response(
                content=buffer.getvalue(), 
                media_type='application/zip', 
                headers={
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'X-Export-Src': file_path
                }
            )
        
        else:
            # Handle JSON data with dynamic records (arbitrary headers)
            body = await request.json()
            records = body.get('records', [])
            columns = body.get('columns')  # optional explicit headers
            rows = body.get('rows')        # optional explicit rows
            format_type = body.get('format', 'excel')

            if format_type == 'word':
                # If explicit columns/rows provided, build a simple Word report with a table
                if columns and rows:
                    from io import BytesIO
                    from docx import Document
                    from docx.shared import Pt
                    from docx.oxml.ns import qn
                    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    from docx.oxml import OxmlElement

                    def shade_cell(cell, fill="E6F0FF"):
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), fill)
                        tcPr.append(shd)

                    doc = Document()
                    # Base font
                    doc.styles['Normal'].font.name = 'Arial'
                    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

                    # Title
                    title = doc.add_paragraph()
                    run = title.add_run('تقرير سجلات الشيكات / Bank Check Records Report')
                    run.font.size = Pt(16)
                    run.bold = True
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Meta
                    meta = doc.add_paragraph()
                    meta.add_run(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").font.size = Pt(10)
                    meta.add_run(f"Number of Records: {len(rows)}\n").font.size = Pt(10)

                    # Table
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(columns):
                        hdr_cells[i].text = str(col)
                        shade_cell(hdr_cells[i], 'E6F0FF')
                        for p in hdr_cells[i].paragraphs:
                            p.runs and setattr(p.runs[0].font, 'bold', True)
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    for r in rows:
                        row_cells = table.add_row().cells
                        for i, val in enumerate(r):
                            row_cells[i].text = str(val)
                            for p in row_cells[i].paragraphs:
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    # Footer summary
                    doc.add_paragraph().add_run('\n')
                    summary = doc.add_paragraph()
                    summary.add_run('ملخص التقرير / Report Summary').bold = True
                    summary.add_run(f"\nإجمالي عدد السجلات: {len(rows)}")

                    out = BytesIO()
                    doc.save(out)
                    out.seek(0)
                    
                    # Save to persistent storage
                    import os
                    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                    date_folder = datetime.now().strftime('%Y-%m-%d')
                    os.makedirs(f"exports/{date_folder}", exist_ok=True)
                    filename = f"bank_check_report_{ts}.docx"
                    file_path = f"exports/{date_folder}/{filename}"
                    
                    with open(file_path, 'wb') as f:
                        f.write(out.getvalue())
                    
                    return Response(
                        content=out.getvalue(), 
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                        headers={
                            'Content-Disposition': f'attachment; filename="{filename}"',
                            'X-Export-Src': file_path
                        }
                    )

                if not records:
                    raise HTTPException(status_code=400, detail="No records provided for Word export")
                # Fallback to service template
                _, word_bytes = await enhanced_bank_check_service.process_records(records)
                
                # Save to persistent storage
                import os
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                date_folder = datetime.now().strftime('%Y-%m-%d')
                os.makedirs(f"exports/{date_folder}", exist_ok=True)
                filename = f"bank_check_report_{ts}.docx"
                file_path = f"exports/{date_folder}/{filename}"
                
                with open(file_path, 'wb') as f:
                    f.write(word_bytes)
                
                return Response(
                    content=word_bytes, 
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                    headers={
                        'Content-Disposition': f'attachment; filename="{filename}"',
                        'X-Export-Src': file_path
                    }
                )
            else:
                # Build a well-formatted Excel dynamically from provided data (robust to arbitrary labels)
                from io import BytesIO
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
                from openpyxl.utils import get_column_letter

                wb = Workbook()
                ws = wb.active
                ws.title = 'Bank Check Records'

                # Determine headers and rows
                if rows and columns:
                    headers = columns
                    data_rows = rows
                else:
                    if not records:
                        raise HTTPException(status_code=400, detail="No records provided")
                    headers = list(records[0].keys())
                    seen = set(headers)
                    for rec in records[1:]:
                        for k in rec.keys():
                            if k not in seen:
                                headers.append(k)
                                seen.add(k)
                    data_rows = [[str(rec.get(h, '')) for h in headers] for rec in records]

                # Custom header block
                title_font = Font(name='Calibri', size=14, bold=True, color='003366')
                meta_font = Font(name='Calibri', size=11)
                ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(headers)))
                ws.cell(row=1, column=1, value='Bank Check Records Report').font = title_font
                ws.cell(row=1, column=1).alignment = Alignment(horizontal='center')
                ws.cell(row=2, column=1, value=f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font = meta_font
                ws.cell(row=3, column=1, value=f"Number of Records: {len(data_rows)}").font = meta_font

                # Table header row
                header_row_idx = 5
                thin = Side(border_style='thin', color='CCCCCC')
                border = Border(top=thin, left=thin, right=thin, bottom=thin)
                fill = PatternFill('solid', fgColor='E6F0FF')

                for idx, h in enumerate(headers, start=1):
                    c = ws.cell(row=header_row_idx, column=idx, value=str(h))
                    c.font = Font(bold=True)
                    c.fill = fill
                    c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    c.border = border

                # Data rows
                for ridx, r in enumerate(data_rows, start=header_row_idx + 1):
                    for cidx, val in enumerate(r, start=1):
                        c = ws.cell(row=ridx, column=cidx, value=str(val))
                        c.alignment = Alignment(vertical='top')
                        c.border = border

                # Column widths
                for i in range(1, len(headers) + 1):
                    max_len = max(
                        [len(str(headers[i-1]))] + [len(str(row[i-1])) for row in data_rows if len(row) >= i]
                    )
                    ws.column_dimensions[get_column_letter(i)].width = min(max(12, max_len + 2), 40)

                buf = BytesIO()
                wb.save(buf)
                buf.seek(0)
                
                # Save to persistent storage
                import os
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                date_folder = datetime.now().strftime('%Y-%m-%d')
                os.makedirs(f"exports/{date_folder}", exist_ok=True)
                filename = f"bank_check_records_{ts}.xlsx"
                file_path = f"exports/{date_folder}/{filename}"
                
                with open(file_path, 'wb') as f:
                    f.write(buf.getvalue())
                
                return Response(
                    content=buf.getvalue(), 
                    media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                    headers={
                        'Content-Disposition': f'attachment; filename="{filename}"',
                        'X-Export-Src': file_path
                    }
                )
                
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Enhanced bank check processing failed: {str(e)}")

@router.post("/api/reports/enhanced-bank-check/preview")
async def preview_enhanced_bank_check(
    file: UploadFile = File(...)
):
    """Preview extracted data from bank check without generating files"""
    try:
        content = await file.read()
        data = await enhanced_bank_check_service.extract_check_data(content, file.filename or 'check.pdf')
        return {
            "success": True,
            "data": data,
            "extraction_summary": {
                "total_fields": 10,
                "extracted_fields": sum(1 for v in data.values() if v and v != ''),
                "missing_fields": sum(1 for v in data.values() if not v or v == ''),
                "success_rate": f"{(sum(1 for v in data.values() if v and v != '') / 10 * 100):.1f}%"
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview failed: {str(e)}")

@router.post("/api/reports/enhanced-bank-check/extract-headers")
async def extract_excel_headers(
    file: UploadFile = File(...)
):
    """Extract headers from Excel template file"""
    try:
        content = await file.read()
        headers = await enhanced_bank_check_service.extract_excel_headers(content, file.filename or 'template.xlsx')
        return {
            "success": True,
            "headers": headers,
            "count": len(headers)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Header extraction failed: {str(e)}")

@router.post("/api/reports/enhanced-bank-check/insert")
async def insert_check_record(
    request: Request
):
    """Insert check record into database"""
    try:
        from config import get_db_connection
        
        body = await request.json()
        record = body.get('record', {})
        table_name = body.get('table_name', 'bank_checks')
        
        if not record:
            return {
                "success": False,
                "error": "No record data provided"
            }
        
        # Use existing SQL Server connection
        conn = get_db_connection()
        cursor = conn.cursor()
        
        try:
            # Get table columns to validate the record
            columns_query = """
                SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE
                FROM INFORMATION_SCHEMA.COLUMNS 
                WHERE TABLE_NAME = %s AND COLUMN_NAME != 'id'
                ORDER BY ORDINAL_POSITION
            """
            cursor.execute(columns_query, (table_name,))
            columns = cursor.fetchall()
            
            if not columns:
                return {
                    "success": False,
                    "error": f"Table '{table_name}' not found"
                }
            
            # Prepare insert statement
            column_names = [col[0] for col in columns]
            placeholders = ', '.join(['?' for _ in column_names])
            column_list = ', '.join([f'[{col}]' for col in column_names])
            
            insert_query = f"""
                INSERT INTO [{table_name}] ({column_list})
                VALUES ({placeholders})
            """
            
            # Helper to coerce python value to SQL Server expected type
            from datetime import datetime, date, time
            def coerce_value(sql_type: str, val):
                if val is None:
                    return None
                if isinstance(val, str) and val.strip() == "":
                    return None
                t = (sql_type or '').lower()
                try:
                    if t in ['int', 'bigint', 'smallint', 'tinyint']:
                        # Extract digits if possible
                        if isinstance(val, (int, float)):
                            return int(val)
                        if isinstance(val, str):
                            return int(val.strip())
                    if t in ['float', 'real']:
                        return float(val)
                    if t in ['decimal', 'numeric', 'money', 'smallmoney']:
                        from decimal import Decimal
                        return Decimal(str(val))
                    if t in ['bit']:
                        if isinstance(val, bool):
                            return 1 if val else 0
                        if isinstance(val, (int, float)):
                            return 1 if int(val) != 0 else 0
                        if isinstance(val, str):
                            return 1 if val.strip().lower() in ['1','true','yes','y'] else 0
                    if t in ['date']:
                        if isinstance(val, (datetime, date)):
                            return val if isinstance(val, date) and not isinstance(val, datetime) else val.date()
                        if isinstance(val, str):
                            # isoformat yyyy-mm-dd
                            return datetime.fromisoformat(val.strip()).date()
                    if t in ['datetime', 'datetime2', 'smalldatetime']:
                        if isinstance(val, (datetime, date)):
                            return val if isinstance(val, datetime) else datetime.combine(val, time())
                        if isinstance(val, str):
                            # try parse ISO; append time if date-only
                            s = val.strip()
                            if len(s) == 10:
                                return datetime.fromisoformat(s + 'T00:00:00')
                            return datetime.fromisoformat(s)
                    if t in ['time']:
                        if isinstance(val, time):
                            return val
                        if isinstance(val, str):
                            return time.fromisoformat(val.strip())
                    # Text-likes as NVARCHAR/VARCHAR
                    return str(val)
                except Exception:
                    # If coercion fails, return None to let NULL insert where allowed
                    return None

            # Prepare values in the correct order with type coercion
            values = []
            for col_name, _, is_nullable in columns:
                raw_value = record.get(col_name, None)
                coerced = coerce_value(_, raw_value)  # _ is DATA_TYPE from SELECT
                if coerced is None and is_nullable == 'NO':
                    return {
                        "success": False,
                        "error": f"Required field '{col_name}' is missing or invalid type (expected {_.lower()})"
                    }
                values.append(coerced)
            
            # Execute insert
            cursor.execute(insert_query, values)
            conn.commit()
            
            # Get the inserted record ID
            record_id = cursor.execute("SELECT @@IDENTITY").fetchone()[0]
            
            return {
                "success": True,
                "message": f"Record inserted successfully into {table_name}",
                "record_id": record_id,
                "table_name": table_name
            }
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        print(f"[EnhancedBankCheck] Insert error: {e}")
        return {
            "success": False,
            "error": f"Database insertion failed: {str(e)}"
        }

@router.get("/api/reports/enhanced-bank-check/tables")
async def get_database_tables():
    """Get all tables from SQL Server database"""
    try:
        from config import get_db_connection
        
        print(f"[EnhancedBankCheck] Connecting to SQL Server...")
        conn = get_db_connection()
        cursor = conn.cursor()
        
        try:
            # Get all tables from the database
            tables_query = """
                SELECT 
                    TABLE_SCHEMA,
                    TABLE_NAME,
                    TABLE_TYPE
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_TYPE = 'BASE TABLE'
                ORDER BY TABLE_SCHEMA, TABLE_NAME;
            """
            
            cursor.execute(tables_query)
            tables_result = cursor.fetchall()
            
            tables = []
            for table in tables_result:  # Get ALL tables from database
                table_name = table[1]  # TABLE_NAME
                schema_name = table[0]  # TABLE_SCHEMA
                full_table_name = f"{schema_name}.{table_name}" if schema_name != 'dbo' else table_name
                
                # Get columns for each table
                columns_query = """
                    SELECT 
                        COLUMN_NAME,
                        DATA_TYPE,
                        IS_NULLABLE,
                        COLUMN_DEFAULT
                    FROM INFORMATION_SCHEMA.COLUMNS 
                    WHERE TABLE_SCHEMA = %s AND TABLE_NAME = %s
                    ORDER BY ORDINAL_POSITION;
                """
                
                cursor.execute(columns_query, (schema_name, table_name))
                columns_result = cursor.fetchall()
                column_names = [col[0] for col in columns_result]  # COLUMN_NAME
                
                tables.append({
                    "id": f"{schema_name}_{table_name}",
                    "name": table_name,
                    "schema": schema_name,
                    "full_name": full_table_name,
                    "fields": column_names,
                    "columns": column_names,  # Add columns for frontend compatibility
                    "field_count": len(column_names)
                })
            
            print(f"[EnhancedBankCheck] Found {len(tables)} tables in SQL Server database")
            
            return {
                "success": True,
                "tables": tables,
                "count": len(tables)
            }
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        print(f"[EnhancedBankCheck] SQL Server connection error: {e}")
        # Fallback to sample data if database connection fails
        sample_tables = [
            {
                "id": "dbo_bank_checks",
                "name": "bank_checks",
                "schema": "dbo",
                "full_name": "bank_checks",
                "fields": ["id", "bank_name", "date", "payee_name", "amount_value", "amount_text", "currency", "status_note", "issuer_signature", "created_at"],
                "columns": ["id", "bank_name", "date", "payee_name", "amount_value", "amount_text", "currency", "status_note", "issuer_signature", "created_at"],
                "field_count": 10
            },
            {
                "id": "dbo_customer_records", 
                "name": "customer_records",
                "schema": "dbo",
                "full_name": "customer_records",
                "fields": ["id", "customer_name", "project_name", "building_number", "apartment_number", "check_number", "due_date", "collection_date", "remaining_days", "collection_status", "total_receivables", "created_at"],
                "columns": ["id", "customer_name", "project_name", "building_number", "apartment_number", "check_number", "due_date", "collection_date", "remaining_days", "collection_status", "total_receivables", "created_at"],
                "field_count": 12
            }
        ]
        
        return {
            "success": True,
            "tables": sample_tables,
            "count": len(sample_tables),
            "note": "Using fallback data - SQL Server connection failed",
            "error": str(e)
        }

@router.get("/api/reports/enhanced-bank-check/test-db")
async def test_database_connection():
    """Test SQL Server database connection"""
    try:
        from config import get_db_connection
        
        print(f"[DB Test] Attempting SQL Server connection...")
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Test query
        cursor.execute("SELECT @@VERSION")
        result = cursor.fetchone()[0]
        
        # Get table count
        cursor.execute("SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_TYPE = 'BASE TABLE'")
        table_count = cursor.fetchone()[0]
        
        cursor.close()
        conn.close()
        
        return {
            "success": True,
            "message": "SQL Server connection successful",
            "version": result,
            "table_count": table_count,
            "database": "NEWDCC-V4-UAT"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "database": "NEWDCC-V4-UAT"
        }

@router.post("/api/reports/enhanced-bank-check/check-table")
async def check_table_exists(request: Request):
    """Check if a table exists in the database"""
    try:
        from config import get_db_connection
        
        body = await request.json()
        table_name = body.get('table_name', '').strip()
        
        if not table_name:
            return {
                "success": False,
                "error": "Table name is required"
            }
        
        # Use existing SQL Server connection
        conn = get_db_connection()
        cursor = conn.cursor()
        
        try:
            # Check if table exists
            check_query = """
                SELECT COUNT(*) 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_NAME = %s AND TABLE_TYPE = 'BASE TABLE'
            """
            
            cursor.execute(check_query, (table_name,))
            exists = cursor.fetchone()[0] > 0
            
            if exists:
                # Get table info if it exists
                info_query = """
                    SELECT 
                        TABLE_SCHEMA,
                        TABLE_NAME,
                        (SELECT COUNT(*) FROM INFORMATION_SCHEMA.COLUMNS 
                         WHERE TABLE_SCHEMA = t.TABLE_SCHEMA AND TABLE_NAME = t.TABLE_NAME) as COLUMN_COUNT
                    FROM INFORMATION_SCHEMA.TABLES t
                    WHERE TABLE_NAME = %s AND TABLE_TYPE = 'BASE TABLE'
                """
                cursor.execute(info_query, (table_name,))
                table_info = cursor.fetchone()
                
                return {
                    "success": True,
                    "exists": True,
                    "table_name": table_name,
                    "schema": table_info[0] if table_info else 'dbo',
                    "column_count": table_info[2] if table_info else 0,
                    "message": f"Table '{table_name}' already exists"
                }
            else:
                return {
                    "success": True,
                    "exists": False,
                    "table_name": table_name,
                    "message": f"Table '{table_name}' does not exist - can be created"
                }
                
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@router.post("/api/reports/enhanced-bank-check/create-table")
async def create_table(request: Request):
    """Create a new table in the database"""
    try:
        from config import get_db_connection
        
        body = await request.json()
        table_name = body.get('table_name', '').strip()
        fields = body.get('fields', [])
        record_data = body.get('record_data', {})
        
        if not table_name:
            return {
                "success": False,
                "error": "Table name is required"
            }
        
        if not fields:
            return {
                "success": False,
                "error": "At least one field is required"
            }
        
        # Use existing SQL Server connection
        conn = get_db_connection()
        cursor = conn.cursor()
        
        try:
            # First check if table exists
            check_query = """
                SELECT COUNT(*) 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_NAME = %s AND TABLE_TYPE = 'BASE TABLE'
            """
            
            cursor.execute(check_query, (table_name,))
            exists = cursor.fetchone()[0] > 0
            
            if exists:
                return {
                    "success": False,
                    "error": f"Table '{table_name}' already exists",
                    "exists": True
                }
            
            # Create table with fields
            # Add ID column as primary key
            columns = ["id INT IDENTITY(1,1) PRIMARY KEY"]
            
            for field in fields:
                field_name = field.get('name', '').strip()
                field_type = field.get('type', 'NVARCHAR(255)')
                is_required = field.get('required', False)

                if field_name:
                    # Sanitize field name: replace spaces with underscore, then remove other special chars
                    safe_name = field_name.replace(' ', '_')
                    clean_name = ''.join(c for c in safe_name if c.isalnum() or c in '_')
                    if not clean_name:
                        clean_name = f"field_{len(columns)}"

                    null_constraint = "NOT NULL" if is_required else "NULL"
                    columns.append(f"[{clean_name}] {field_type} {null_constraint}")
            
            # Add created_at timestamp
            columns.append("[created_at] DATETIME2 DEFAULT GETDATE()")
            
            create_query = f"""
                CREATE TABLE [{table_name}] (
                    {', '.join(columns)}
                )
            """
            
            cursor.execute(create_query)
            conn.commit()
            
            # Insert record data if provided
            record_id = None
            if record_data:
                try:
                    # Get table columns to validate the record
                    columns_query = """
                        SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE
                        FROM INFORMATION_SCHEMA.COLUMNS 
                        WHERE TABLE_NAME = %s AND COLUMN_NAME != 'id'
                        ORDER BY ORDINAL_POSITION
                    """
                    cursor.execute(columns_query, (table_name,))
                    columns = cursor.fetchall()
                    
                    if columns:
                        # Prepare insert statement (use %s for pymssql)
                        column_names = [col[0] for col in columns]
                        placeholders = ', '.join(['%s' for _ in column_names])
                        column_list = ', '.join([f'[{col}]' for col in column_names])
                        
                        insert_query = f"""
                            INSERT INTO [{table_name}] ({column_list})
                            VALUES ({placeholders})
                        """
                        
                        # Prepare values in the correct order
                        values = []
                        for col_name in column_names:
                            # Find matching field by name (case-insensitive)
                            field_name = None
                            for field in fields:
                                _raw = field.get('name', '')
                                _safe = _raw.replace(' ', '_')
                                clean_field_name = ''.join(c for c in _safe if c.isalnum() or c in '_')
                                if clean_field_name.lower() == col_name.lower():
                                    field_name = field.get('name', '')
                                    break
                            
                            if field_name and field_name in record_data:
                                value = record_data[field_name]
                            else:
                                value = None
                            
                            if value is None and any(col[0] == col_name and col[2] == 'NO' for col in columns):
                                # Required field is missing, skip insertion
                                print(f"[EnhancedBankCheck] Warning: Required field '{col_name}' is missing, skipping record insertion")
                                record_id = None
                                break
                            values.append(value)
                        
                        if values:  # Only insert if we have valid values
                            cursor.execute(insert_query, values)
                            conn.commit()
                            
                            # Get the inserted record ID
                            record_id = cursor.execute("SELECT @@IDENTITY").fetchone()[0]
                            print(f"[EnhancedBankCheck] Record inserted with ID: {record_id}")
                        
                except Exception as insert_error:
                    print(f"[EnhancedBankCheck] Error inserting record: {insert_error}")
                    # Don't fail the table creation if record insertion fails
                    pass
            
            return {
                "success": True,
                "message": f"Table '{table_name}' created successfully" + (f" and record inserted (ID: {record_id})" if record_id else ""),
                "table_name": table_name,
                "field_count": len(fields),
                "created_fields": [f.get('name', '') for f in fields],
                "record_id": record_id,
                "record_inserted": record_id is not None
            }
            
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

