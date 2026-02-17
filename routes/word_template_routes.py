from fastapi import APIRouter, UploadFile, File, Form, Body, HTTPException
from fastapi.responses import Response, JSONResponse
from io import BytesIO
import logging
import traceback
from docx import Document
from docx.document import Document as DocType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
import json
import base64
import openpyxl
import pandas as pd
from lxml import etree
import pyodbc
from config import get_db_connection

router = APIRouter(prefix="/word-template", tags=["word-template"])
logger = logging.getLogger(__name__)


def extract_sections_from_docx(doc: DocType):
	"""Extract sections (headings, paragraphs, tables) from Word document with formatting"""
	sections = []
	section_id = 0
	
	current_section = None
	
	# Store original document reference for later use
	template_doc = doc
	
	for element in doc.element.body:
		if isinstance(element, CT_P):
			paragraph = Paragraph(element, doc)
			text = paragraph.text.strip()
			
			if not text:
				continue
			
			# Check if it's a heading (based on style)
			is_heading = False
			heading_level = 1
			
			try:
				if paragraph.style and hasattr(paragraph.style, 'name') and paragraph.style.name:
					style_name = paragraph.style.name.lower()
					if 'heading' in style_name:
						is_heading = True
						try:
							# Extract heading level (e.g., "Heading 1" -> 1)
							level_str = style_name.replace('heading', '').strip()
							if level_str:
								heading_level = int(level_str)
						except (ValueError, AttributeError):
							heading_level = 1
			except (AttributeError, Exception):
				pass
			
			# Check if paragraph has bold formatting (likely heading)
			if not is_heading:
				try:
					for run in paragraph.runs:
						if hasattr(run, 'bold') and run.bold and len(text) < 100:  # Short bold text is likely a heading
							is_heading = True
							break
				except (AttributeError, Exception):
					pass
			
			if is_heading:
				# Save previous section if exists
				if current_section:
					sections.append(current_section)
				
				# Start new section
				section_id += 1
				# Store formatting info
				style_name = paragraph.style.name if paragraph.style and hasattr(paragraph.style, 'name') else None
				current_section = {
					'id': f'section-{section_id}',
					'title': text,
					'content': '',
					'type': 'heading',
					'order': section_id,
					'level': heading_level,
					'style': style_name,
					'formatting': _extract_paragraph_formatting(paragraph)
				}
			else:
				# Add to current section or create new one
				if not current_section:
					section_id += 1
					style_name = paragraph.style.name if paragraph.style and hasattr(paragraph.style, 'name') else None
					current_section = {
						'id': f'section-{section_id}',
						'title': f'Section {section_id}',
						'content': '',
						'type': 'paragraph',
						'order': section_id,
						'style': style_name,
						'formatting': _extract_paragraph_formatting(paragraph)
					}
				
				# Convert paragraph to HTML
				para_html = paragraph_to_html(paragraph)
				if current_section['content']:
					current_section['content'] += '<p>' + para_html + '</p>'
				else:
					current_section['content'] = '<p>' + para_html + '</p>'
		
		elif isinstance(element, CT_Tbl):
			table = Table(element, doc)
			
			# Save previous section if exists
			if current_section:
				sections.append(current_section)
				current_section = None
			
			# Create table section
			section_id += 1
			table_html = table_to_html(table)
			sections.append({
				'id': f'section-{section_id}',
				'title': f'Table {section_id}',
				'content': table_html,
				'type': 'table',
				'order': section_id
			})
	
	# Add last section
	if current_section:
		sections.append(current_section)
	
	return sections


def _extract_paragraph_formatting(paragraph: Paragraph) -> dict:
	"""Extract formatting information from paragraph"""
	formatting = {}
	try:
		if paragraph.style and hasattr(paragraph.style, 'name'):
			formatting['style'] = paragraph.style.name
		
		if paragraph.paragraph_format:
			if hasattr(paragraph.paragraph_format, 'alignment') and paragraph.paragraph_format.alignment:
				formatting['alignment'] = str(paragraph.paragraph_format.alignment)
		
		# Extract run formatting
		formatting['runs'] = []
		for run in paragraph.runs:
			run_format = {}
			if hasattr(run, 'font'):
				if hasattr(run.font, 'name') and run.font.name:
					run_format['font_name'] = run.font.name
				if hasattr(run.font, 'size') and run.font.size:
					run_format['font_size'] = str(run.font.size)
				if hasattr(run.font, 'color') and run.font.color and hasattr(run.font.color, 'rgb'):
					run_format['font_color'] = str(run.font.color.rgb)
				if hasattr(run.font, 'bold'):
					run_format['bold'] = run.font.bold
				if hasattr(run.font, 'italic'):
					run_format['italic'] = run.font.italic
				if hasattr(run.font, 'underline'):
					run_format['underline'] = run.font.underline
			run_format['text'] = run.text if hasattr(run, 'text') else ''
			formatting['runs'].append(run_format)
	except Exception as e:
		logger.warning(f"Error extracting formatting: {e}")
	
	return formatting


def paragraph_to_html(paragraph: Paragraph) -> str:
	"""Convert paragraph to HTML preserving ALL formatting (colors, fonts, sizes, etc.)"""
	from docx.shared import RGBColor
	import html
	
	html_parts = []
	
	try:
		for run in paragraph.runs:
			text = run.text if hasattr(run, 'text') else ''
			
			if not text:
				continue
			
			# Escape HTML entities
			text = html.escape(text)
			
			# Build style attribute with all formatting
			styles = []
			
			# Font properties
			if hasattr(run, 'font') and run.font:
				# Font name
				if hasattr(run.font, 'name') and run.font.name:
					styles.append(f"font-family: '{run.font.name}'")
				
				# Font size
				if hasattr(run.font, 'size') and run.font.size:
					# Convert from EMU to points (1 point = 12700 EMU)
					size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 12700
					styles.append(f"font-size: {size_pt:.1f}pt")
				
				# Font color
				if hasattr(run.font, 'color') and run.font.color:
					if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
						rgb = run.font.color.rgb
						if isinstance(rgb, RGBColor):
							# RGBColor stores as hex string like "FF0000"
							# Extract R, G, B values
							rgb_str = str(rgb)
							try:
								# Convert hex string to integer
								rgb_int = int(rgb_str, 16) if len(rgb_str) == 6 else int(rgb_str, 16) if rgb_str.startswith('0x') else int(rgb_str)
								r = (rgb_int >> 16) & 0xFF
								g = (rgb_int >> 8) & 0xFF
								b = rgb_int & 0xFF
								styles.append(f"color: rgb({r}, {g}, {b})")
							except (ValueError, AttributeError):
								# Fallback: try to parse as tuple if available
								pass
				
				# Background/highlight color
				if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
					# python-docx has limited highlight support, but we can try
					pass
			
			# Text formatting
			formatting_tags = []
			if hasattr(run, 'bold') and run.bold:
				formatting_tags.append('strong')
			if hasattr(run, 'italic') and run.italic:
				formatting_tags.append('em')
			if hasattr(run, 'underline') and run.underline:
				formatting_tags.append('u')
			if hasattr(run, 'font') and run.font and hasattr(run.font, 'strike') and run.font.strike:
				formatting_tags.append('s')
			if hasattr(run, 'font') and run.font and hasattr(run.font, 'subscript') and run.font.subscript:
				formatting_tags.append('sub')
			if hasattr(run, 'font') and run.font and hasattr(run.font, 'superscript') and run.font.superscript:
				formatting_tags.append('sup')
			
			# Build HTML with style and formatting tags
			if styles:
				style_attr = '; '.join(styles)
				text = f'<span style="{style_attr}">{text}</span>'
			
			# Apply formatting tags (wrap from outside to inside)
			for tag in reversed(formatting_tags):
				text = f'<{tag}>{text}</{tag}>'
			
			html_parts.append(text)
			
	except (AttributeError, Exception) as e:
		logger.warning(f"Error processing paragraph runs: {e}")
		# Fallback to plain text with HTML escape
		if hasattr(paragraph, 'text'):
			return html.escape(paragraph.text)
		return ''
	
	result = ''.join(html_parts) if html_parts else ''
	
	# If paragraph has alignment, wrap in div with alignment
	if paragraph.paragraph_format and hasattr(paragraph.paragraph_format, 'alignment'):
		alignment = paragraph.paragraph_format.alignment
		if alignment:
			if 'RIGHT' in str(alignment):
				result = f'<div style="text-align: right;">{result}</div>'
			elif 'CENTER' in str(alignment):
				result = f'<div style="text-align: center;">{result}</div>'
			elif 'LEFT' in str(alignment):
				result = f'<div style="text-align: left;">{result}</div>'
	
	return result if result else (html.escape(paragraph.text) if hasattr(paragraph, 'text') else '')


def table_to_html(table: Table) -> str:
	"""Convert table to HTML preserving formatting"""
	from docx.shared import RGBColor
	import html
	
	html_parts = ['<table border="1" cellspacing="0" cellpadding="8" style="border-collapse: collapse; width: 100%; direction: rtl; text-align: right; margin: 10px 0;">']
	
	try:
		for row_idx, row in enumerate(table.rows):
			html_parts.append('<tr>')
			for cell_idx, cell in enumerate(row.cells):
				# Extract cell content with formatting
				cell_content = []
				for para in cell.paragraphs:
					para_html = paragraph_to_html(para)
					if para_html:
						# Remove outer div if paragraph_to_html wrapped it
						if para_html.startswith('<div'):
							# Extract inner content
							import re
							match = re.match(r'<div[^>]*>(.*?)</div>', para_html, re.DOTALL)
							if match:
								para_html = match.group(1)
						cell_content.append(para_html)
				
				cell_text = ' '.join(cell_content) if cell_content else html.escape(cell.text.strip() if cell.text else '')
				if not cell_text.strip():
					cell_text = '&nbsp;'
				
				# Get cell formatting
				cell_styles = ['direction: rtl;', 'text-align: right;', 'vertical-align: top;']
				if cell.paragraphs:
					first_para = cell.paragraphs[0]
					if first_para.paragraph_format and hasattr(first_para.paragraph_format, 'alignment'):
						alignment = first_para.paragraph_format.alignment
						if 'RIGHT' in str(alignment):
							cell_styles.append('text-align: right;')
						elif 'CENTER' in str(alignment):
							cell_styles.append('text-align: center;')
						elif 'LEFT' in str(alignment):
							cell_styles.append('text-align: left;')
				
				# Add border styles
				cell_styles.append('border: 1px solid #000;')
				cell_styles.append('padding: 8px;')
				
				style_attr = ' '.join(cell_styles)
				html_parts.append(f'<td style="{style_attr}">{cell_text}</td>')
			html_parts.append('</tr>')
	except Exception as e:
		logger.warning(f"Error processing table: {e}")
		# Fallback to simple table
		html_parts = ['<table border="1" cellspacing="0" cellpadding="8" style="border-collapse: collapse; width: 100%; direction: rtl; text-align: right; margin: 10px 0;">']
		for row in table.rows:
			html_parts.append('<tr>')
			for cell in row.cells:
				cell_text = html.escape(cell.text.strip() if cell.text else '')
				if not cell_text:
					cell_text = '&nbsp;'
				html_parts.append(f'<td style="direction: rtl; text-align: right; border: 1px solid #000; padding: 8px;">{cell_text}</td>')
			html_parts.append('</tr>')
	
	html_parts.append('</table>')
	return ''.join(html_parts)


@router.post("/analyze")
async def analyze_template(template: UploadFile = File(...)):
	"""Analyze Word template and extract sections"""
	try:
		# Validate file
		if not template:
			return JSONResponse(
				status_code=400,
				content={
					'success': False,
					'error': 'No file provided'
				}
			)
		
		# Check file extension
		file_name = template.filename or ''
		if not file_name.lower().endswith(('.docx', '.doc')):
			return JSONResponse(
				status_code=400,
				content={
					'success': False,
					'error': f'Invalid file type. Expected .docx or .doc, got: {file_name}'
				}
			)
		
		# Read uploaded file
		file_content = await template.read()
		
		if len(file_content) == 0:
			return JSONResponse(
				status_code=400,
				content={
					'success': False,
					'error': 'File is empty'
				}
			)
		
		# Try to open document
		try:
			doc = Document(BytesIO(file_content))
		except Exception as doc_error:
			logger.error(f"Error opening document: {doc_error}")
			return JSONResponse(
				status_code=400,
				content={
					'success': False,
					'error': f'Failed to open Word document: {str(doc_error)}. Please ensure the file is a valid .docx file.'
				}
			)
		
		# Extract sections
		try:
			sections = extract_sections_from_docx(doc)
		except Exception as extract_error:
			logger.error(f"Error extracting sections: {extract_error}")
			traceback_str = traceback.format_exc()
			return JSONResponse(
				status_code=500,
				content={
					'success': False,
					'error': f'Failed to extract sections: {str(extract_error)}',
					'traceback': traceback_str
				}
			)
		
		logger.info(f"Extracted {len(sections)} sections from template")
		
		# Ensure at least one section exists
		if not sections:
			sections = [{
				'id': 'section-1',
				'title': 'Document Content',
				'content': '<p>No structured sections found. Document content will be available for editing.</p>',
				'type': 'paragraph',
				'order': 1
			}]
		
		return JSONResponse(
			status_code=200,
			content={
				'success': True,
				'sections': sections,
				'total_sections': len(sections)
			}
		)
		
	except Exception as e:
		error_msg = str(e)
		traceback_str = traceback.format_exc()
		logger.error(f"ERROR analyzing template: {error_msg}")
		logger.error(traceback_str)
		return JSONResponse(
			status_code=500,
			content={
				'success': False,
				'error': f'Unexpected error: {error_msg}',
				'traceback': traceback_str
			}
		)


def _add_table_from_html(doc, table_html):
	"""Extract and add a table from HTML to Word document"""
	import re
	from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
	
	rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
	if not rows:
		return
	
	first_row = rows[0]
	cols = len(re.findall(r'<t[dh][^>]*>', first_row, re.IGNORECASE))
	if cols == 0:
		return
	
	table = doc.add_table(rows=len(rows), cols=cols)
	# Apply table style safely with fallbacks
	try:
		table.style = 'Table Grid'
	except:
		try:
			table.style = 'Light Grid'
		except:
			try:
				table.style = 'Light Grid Accent 1'
			except:
				try:
					table.style = 'Table Normal'
				except:
					pass
	
	for row_idx, row_html in enumerate(rows):
		cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row_html, re.DOTALL | re.IGNORECASE)
		for col_idx, cell_html in enumerate(cells):
			if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
				cell = table.rows[row_idx].cells[col_idx]
				# Remove HTML tags and get text, handle HTML entities
				cell_text = re.sub(r'<[^>]+>', '', cell_html).strip()
				cell_text = cell_text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"')
				cell.text = cell_text
				for para in cell.paragraphs:
					para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def html_to_word_paragraph(doc, html_content):
	"""Convert HTML content to Word paragraph with full formatting support"""
	import re
	from docx.shared import Pt, RGBColor
	from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
	
	def hex_to_rgb(hex_color):
		"""Convert hex color to RGB tuple"""
		hex_color = hex_color.lstrip('#')
		if len(hex_color) == 6:
			return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
		return (0, 0, 0)
	
	def parse_style(style_attr):
		"""Parse style attribute to extract color, background-color, etc."""
		styles = {}
		if not style_attr:
			return styles
		
		for prop in style_attr.split(';'):
			prop = prop.strip()
			if ':' in prop:
				key, value = prop.split(':', 1)
				key = key.strip()
				value = value.strip()
				styles[key] = value
		
		return styles
	
	def apply_formatting(run, tag_name, tag_attrs, stack):
		"""Apply formatting to a run based on tag and attributes"""
		# Parse style attribute FIRST (for span tags and other elements with style)
		style_attr = tag_attrs.get('style', '')
		if style_attr:
			styles = parse_style(style_attr)
			
			# Apply style properties directly (these take precedence)
			if 'color' in styles:
				color = styles['color']
				if color.startswith('rgb'):
					rgb_match = re.search(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color)
					if rgb_match:
						r, g, b = map(int, rgb_match.groups())
						run.font.color.rgb = RGBColor(r, g, b)
				elif color.startswith('#'):
					r, g, b = hex_to_rgb(color)
					run.font.color.rgb = RGBColor(r, g, b)
			
			if 'font-size' in styles:
				size_str = styles['font-size']
				size_match = re.search(r'(\d+\.?\d*)', size_str)
				if size_match:
					size = float(size_match.group(1))
					if 'px' in size_str:
						size = size * 0.75
					run.font.size = Pt(size)
			
			if 'font-family' in styles:
				font_family = styles['font-family']
				font_family = font_family.strip("'\"")
				if ',' in font_family:
					font_family = font_family.split(',')[0].strip()
				try:
					run.font.name = font_family
				except:
					pass
		
		# Basic formatting tags (only if not already applied via style)
		if tag_name in ['b', 'strong']:
			run.bold = True
		if tag_name in ['i', 'em']:
			run.italic = True
		if tag_name == 'u':
			run.underline = True
		if tag_name in ['s', 'strike', 'del']:
			run.font.strike = True
		if tag_name == 'sub':
			run.font.subscript = True
		if tag_name == 'sup':
			run.font.superscript = True
		
		# Background color (if available)
		if style_attr:
			styles = parse_style(style_attr)
			if 'background-color' in styles:
				bg_color = styles['background-color']
				if bg_color.startswith('rgb'):
					rgb_match = re.search(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', bg_color)
					if rgb_match:
						r, g, b = map(int, rgb_match.groups())
						run.font.highlight_color = 1  # Yellow highlight as closest option
						# Note: python-docx has limited highlight color support
	
	def extract_tag_attrs(tag):
		"""Extract attributes from HTML tag"""
		attrs = {}
		# Match style="..." or other attributes
		style_match = re.search(r'style=["\']([^"\']+)["\']', tag)
		if style_match:
			attrs['style'] = style_match.group(1)
		
		class_match = re.search(r'class=["\']([^"\']+)["\']', tag)
		if class_match:
			attrs['class'] = class_match.group(1)
		
		return attrs
	
	def add_formatted_text(para, html_text):
		"""Add formatted text to paragraph with full HTML support"""
		if not html_text or not html_text.strip():
			return
		
		stack = []  # Stack of (tag_name, attributes)
		pattern = r'(<[^>]+>)|([^<]+)'
		matches = list(re.finditer(pattern, html_text))
		
		if not matches:
			# Plain text
			run = para.add_run(html_text)
			run.font.size = Pt(11)
			return
		
		for match in matches:
			tag = match.group(1)
			text = match.group(2)
			
			if tag:
				tag_match = re.match(r'</?(\w+)', tag)
				tag_name = tag_match.group(1).lower() if tag_match else None
				
				if tag.startswith('</'):
					# Closing tag
					if stack and stack[-1][0] == tag_name:
						stack.pop()
				elif tag.startswith('<') and not tag.endswith('/>'):
					# Opening tag
					if tag_name:
						attrs = extract_tag_attrs(tag)
						stack.append((tag_name, attrs))
			
			if text and text.strip():
				# Add text run
				run = para.add_run(text)
				run.font.size = Pt(11)
				
				# Apply all formatting from stack (process span tags first for style attributes)
				# Sort stack: process span tags with style first, then other formatting
				sorted_stack = sorted(stack, key=lambda x: (x[0] != 'span', x[0]))
				
				for tag_name, attrs in sorted_stack:
					apply_formatting(run, tag_name, attrs, stack)
	
	# Handle lists first
	if '<ul' in html_content or '<ol' in html_content:
		list_items = re.findall(r'<li[^>]*>(.*?)</li>', html_content, re.DOTALL)
		for item in list_items:
			para = doc.add_paragraph()
			para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
			para.style = 'List Bullet' if '<ul' in html_content else 'List Number'
			add_formatted_text(para, item.strip())
		return doc
	
	# Handle headings
	heading_match = re.search(r'<h([1-6])[^>]*>(.*?)</h[1-6]>', html_content, re.DOTALL)
	if heading_match:
		level = int(heading_match.group(1))
		heading_text = re.sub(r'<[^>]+>', '', heading_match.group(2))
		heading = doc.add_heading(heading_text, level=min(level, 9))
		heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
		return doc
	
	# Handle paragraphs and line breaks
	if '<p' in html_content or '<br' in html_content:
		# Split by <p> tags or <br> tags
		parts = re.split(r'<p[^>]*>|</p>|<br\s*/?>', html_content)
		for part in parts:
			if part.strip():
				para = doc.add_paragraph()
				para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
				add_formatted_text(para, part.strip())
	else:
		# Single paragraph
		para = doc.add_paragraph()
		para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
		add_formatted_text(para, html_content)
	
	return doc


def _add_bookmark(paragraph, bookmark_name):
	"""Add a bookmark to a paragraph"""
	from docx.oxml import OxmlElement
	from docx.oxml.ns import qn
	
	# Create bookmark start
	bookmark_start = OxmlElement('w:bookmarkStart')
	bookmark_start.set(qn('w:id'), '0')
	bookmark_start.set(qn('w:name'), bookmark_name)
	
	# Create bookmark end
	bookmark_end = OxmlElement('w:bookmarkEnd')
	bookmark_end.set(qn('w:id'), '0')
	
	# Insert bookmark start at the beginning of paragraph
	paragraph._element.insert(0, bookmark_start)
	# Insert bookmark end at the end of paragraph
	paragraph._element.append(bookmark_end)


def _add_pageref_field(paragraph, bookmark_name):
	"""Add a PAGEREF field to get page number of a bookmark"""
	from docx.oxml import OxmlElement
	from docx.oxml.ns import qn
	
	# Clear existing runs
	paragraph.clear()
	
	# Create field code run
	field_code_run = paragraph.add_run()
	field_code_run._element.append(OxmlElement('w:rPr'))  # Run properties
	
	# Create field code
	field_code = OxmlElement('w:instrText')
	field_code.set(qn('w:xml:space'), 'preserve')
	field_code.text = f' PAGEREF {bookmark_name} \\h '
	
	# Create field begin
	field_begin = OxmlElement('w:fldChar')
	field_begin.set(qn('w:fldCharType'), 'begin')
	
	# Create field separator
	field_separator = OxmlElement('w:fldChar')
	field_separator.set(qn('w:fldCharType'), 'separate')
	
	# Create field end
	field_end = OxmlElement('w:fldChar')
	field_end.set(qn('w:fldCharType'), 'end')
	
	# Add field elements to run
	field_run = paragraph._element.get_or_add_r()
	field_run.append(field_begin)
	field_run.append(field_code_run._element)
	field_code_run._element.append(field_code)
	field_run.append(field_separator)
	
	# Add result run (will be updated by Word)
	result_run = paragraph.add_run('0')
	result_run._element.append(field_end)


def _add_schedule_table(doc, sections):
	"""Add a Schedule table (Table of Contents) at the beginning of the document"""
	from docx.shared import Pt, RGBColor
	from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
	from docx.oxml import OxmlElement
	from docx.oxml.ns import qn
	
	# Add title for Schedule at the beginning
	title_para = doc.add_paragraph()
	title_run = title_para.add_run('جدول المحتويات (Schedule)')
	title_run.font.size = Pt(16)
	title_run.font.bold = True
	title_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
	title_para.paragraph_format.space_after = Pt(12)
	
	# Add table with sections (only Section and Page columns)
	table = doc.add_table(rows=1, cols=2)
	# Apply table style safely with fallbacks
	try:
		table.style = 'Table Grid'
	except:
		try:
			table.style = 'Light Grid'
		except:
			try:
				table.style = 'Light Grid Accent 1'
			except:
				try:
					table.style = 'Table Normal'
				except:
					# Use default table style (no style)
					pass
	
	# Header row
	header_cells = table.rows[0].cells
	header_cells[0].text = 'القسم'
	header_cells[1].text = 'الصفحة'
	
	# Style header
	for cell in header_cells:
		for paragraph in cell.paragraphs:
			paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
			for run in paragraph.runs:
				run.font.bold = True
				run.font.size = Pt(12)
		# Set cell background color
		try:
			shading_elm = OxmlElement('w:shd')
			shading_elm.set(qn('w:fill'), 'D3D3D3')
			cell._element.get_or_add_tcPr().append(shading_elm)
		except Exception as e:
			logger.warning(f"Could not set cell shading: {e}")
	
	# Add sections to table (only Section and Page)
	sorted_sections = sorted(sections, key=lambda x: x.get('order', 0))
	for idx, section in enumerate(sorted_sections, start=1):
		row = table.add_row()
		# Section name
		row.cells[0].text = section.get('title', f'Section {idx}')
		
		# Add PAGEREF field for page number
		section_id = section.get('id', f'section_{idx}')
		bookmark_name = f'Sec_{section_id.replace("-", "_")}'
		page_cell = row.cells[1]  # Page column
		page_para = page_cell.paragraphs[0]
		page_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
		
		# Add PAGEREF field
		try:
			_add_pageref_field(page_para, bookmark_name)
		except Exception as e:
			logger.warning(f"Could not add PAGEREF field: {e}")
			# Fallback: empty cell
			page_para.clear()
		
		# Style cells
		for cell in row.cells:
			for paragraph in cell.paragraphs:
				paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
				for run in paragraph.runs:
					run.font.size = Pt(11)
	
	# Add spacing after table
	spacing_para = doc.add_paragraph()
	spacing_para.paragraph_format.space_after = Pt(12)
	
	logger.info(f"Added Schedule table with {len(sorted_sections)} sections")


@router.get("/get-users")
async def get_users():
	"""Get all active users from the database"""
	try:
		conn = get_db_connection()
		cursor = conn.cursor()
		
		try:
			# Query Users table - get active users with their names
			# Using COALESCE to prefer 'name' field, fallback to 'username' if name is NULL
			# Note: In SQL Server, BIT fields need special handling
			query = """
				SELECT DISTINCT 
					COALESCE(name, username) as display_name
				FROM Users
				WHERE active = 1 
					AND (deletedAt IS NULL OR deletedAt = '')
					AND COALESCE(name, username) IS NOT NULL
					AND COALESCE(name, username) != ''
				ORDER BY display_name
			"""
			
			logger.info("Executing query to fetch users")
			cursor.execute(query)
			rows = cursor.fetchall()
			
			# Extract user names
			users = [str(row[0]).strip() for row in rows if row[0] and str(row[0]).strip()]
			
			logger.info(f"Found {len(users)} users: {users}")
			
			return JSONResponse({
				"success": True,
				"users": users
			})
			
		except Exception as db_error:
			logger.error(f"Database error fetching users: {db_error}", exc_info=True)
			# Try a simpler query without filters
			try:
				simple_query = """
					SELECT DISTINCT 
						COALESCE(name, username) as display_name
					FROM Users
					WHERE COALESCE(name, username) IS NOT NULL
						AND COALESCE(name, username) != ''
					ORDER BY display_name
				"""
				cursor.execute(simple_query)
				rows = cursor.fetchall()
				users = [str(row[0]).strip() for row in rows if row[0] and str(row[0]).strip()]
				logger.info(f"Fallback query found {len(users)} users")
				return JSONResponse({
					"success": True,
					"users": users
				})
			except Exception as fallback_error:
				logger.error(f"Fallback query also failed: {fallback_error}", exc_info=True)
				return JSONResponse({
					"success": False,
					"users": [],
					"error": str(fallback_error)
				})
		finally:
			conn.close()
			
	except Exception as e:
		logger.error(f"Error fetching users: {e}", exc_info=True)
		# Return empty list on error
		return JSONResponse({
			"success": False,
			"users": [],
			"error": str(e)
		})


def _apply_html_to_paragraph(paragraph, html_content, original_formatting=None):
	"""Apply HTML content to existing paragraph while preserving original formatting"""
	import re
	from docx.shared import Pt, RGBColor
	
	# Clear existing runs but keep style
	paragraph.clear()
	
	# Restore original style if available
	if original_formatting and original_formatting.get('style'):
		try:
			paragraph.style = original_formatting.get('style')
		except:
			pass
	
	# Apply HTML content
	if not html_content or not html_content.strip():
		return
	
	# Parse HTML and add formatted runs
	stack = []
	pattern = r'(<[^>]+>)|([^<]+)'
	matches = re.finditer(pattern, html_content)
	
	for match in matches:
		tag = match.group(1)
		text = match.group(2)
		
		if tag:
			tag_match = re.match(r'</?(\w+)', tag)
			tag_name = tag_match.group(1).lower() if tag_match else None
			
			if tag.startswith('</'):
				if stack and stack[-1][0] == tag_name:
					stack.pop()
			elif tag.startswith('<') and not tag.endswith('/>'):
				if tag_name:
					attrs = {}
					style_match = re.search(r'style=["\']([^"\']+)["\']', tag)
					if style_match:
						attrs['style'] = style_match.group(1)
					stack.append((tag_name, attrs))
		
		if text and text.strip():
			run = paragraph.add_run(text)
			
			# Apply formatting from stack
			for tag_name, attrs in stack:
				if tag_name in ['b', 'strong']:
					run.bold = True
				if tag_name in ['i', 'em']:
					run.italic = True
				if tag_name == 'u':
					run.underline = True
				if tag_name in ['s', 'strike', 'del']:
					run.font.strike = True
				
				# Apply colors and sizes from style
				if attrs.get('style'):
					styles = {}
					for prop in attrs['style'].split(';'):
						if ':' in prop:
							key, value = prop.split(':', 1)
							styles[key.strip()] = value.strip()
					
					if 'color' in styles:
						color = styles['color']
						if color.startswith('rgb'):
							rgb_match = re.search(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color)
							if rgb_match:
								r, g, b = map(int, rgb_match.groups())
								run.font.color.rgb = RGBColor(r, g, b)
						elif color.startswith('#'):
							hex_color = color.lstrip('#')
							if len(hex_color) == 6:
								r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
								run.font.color.rgb = RGBColor(r, g, b)
					
					if 'font-size' in styles:
						size_str = styles['font-size']
						size_match = re.search(r'(\d+)', size_str)
						if size_match:
							size = int(size_match.group(1))
							if 'px' in size_str:
								size = int(size * 0.75)
							run.font.size = Pt(size)
			
			# Apply original formatting if available
			if original_formatting and original_formatting.get('runs'):
				# Use first run's formatting as default
				first_run = original_formatting['runs'][0] if original_formatting['runs'] else {}
				if first_run.get('font_name'):
					try:
						run.font.name = first_run['font_name']
					except:
						pass
				if first_run.get('font_size'):
					try:
						run.font.size = Pt(int(first_run['font_size'].replace('pt', '')))
					except:
						pass


@router.post("/generate")
async def generate_report(payload: dict = Body(...)):
	"""Generate Word report from template and data"""
	try:
		from docx import Document
		from docx.shared import Pt, RGBColor, Inches
		from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
		import re
		
		logger.info("Starting report generation")
		
		# Get template if provided
		template_base64 = payload.get('template')
		sections = payload.get('sections', [])
		data_sources = payload.get('dataSources', [])
		data_mappings = payload.get('dataMappings', [])
		include_schedule = payload.get('includeSchedule', True)  # Default to True
		
		logger.info(f"Report generation params: template={'present' if template_base64 else 'missing'}, sections={len(sections)}, include_schedule={include_schedule}")
		
		if not sections:
			logger.warning("No sections provided, creating empty document")
		
		# Fetch data from sources
		data_cache = {}
		for mapping in data_mappings:
			source_id = mapping.get('dataSource', {}).get('name')
			if source_id and source_id not in data_cache:
				# Fetch data
				source = next((s for s in data_sources if s.get('name') == source_id), None)
				if source:
					source_type = source.get('type')
					source_config = source.get('config', {})
					
					if source_type == 'database':
						# Fetch from database
						import pyodbc
						from config import get_db_connection
						
						try:
							table_name = source_config.get('table', '')
							columns = source_config.get('columns', [])
							where_clause = source_config.get('where', '')
							
							if table_name:
								conn = get_db_connection()
								cursor = conn.cursor()
								
								query = f"SELECT {', '.join(columns) if columns else '*'} FROM {table_name}"
								if where_clause:
									query += f" WHERE {where_clause}"
								
								cursor.execute(query)
								rows = cursor.fetchall()
								col_names = [desc[0] for desc in cursor.description] if cursor.description else []
								
								conn.close()
								
								# Convert to dict format for easy access
								if rows and col_names:
									# Take first row as variables
									first_row = rows[0]
									data_cache[source_id] = {col_names[i]: str(first_row[i]) for i in range(len(col_names))}
								else:
									data_cache[source_id] = {}
						except Exception as e:
							logger.error(f"Error fetching database data: {e}")
							data_cache[source_id] = {}
					
					elif source_type == 'xbrl':
						# TODO: Parse XBRL file
						data_cache[source_id] = {}
					
					elif source_type == 'excel':
						# TODO: Parse Excel file
						data_cache[source_id] = {}
					
					else:
						data_cache[source_id] = {}
		
		# Create new document
		if template_base64:
			try:
				# Use template as base - PRESERVE ALL DESIGN
				logger.info("Loading template from base64")
				template_bytes = base64.b64decode(template_base64.split(',')[1] if ',' in template_base64 else template_base64)
				doc = Document(BytesIO(template_bytes))
				logger.info(f"Template loaded successfully, {len(doc.paragraphs)} paragraphs found")
			except Exception as template_error:
				logger.error(f"Error loading template: {template_error}", exc_info=True)
				raise Exception(f"Failed to load template: {str(template_error)}")
			
			# Schedule table removed - no longer adding it to the document
			# Also remove any existing Schedule table from template
			# Remove paragraphs and tables that contain "Schedule" or "جدول المحتويات"
			paragraphs_to_remove = []
			for para in doc.paragraphs:
				para_text = para.text.strip()
				if 'Schedule' in para_text or 'جدول المحتويات' in para_text or 'جدول' in para_text and 'المحتويات' in para_text:
					paragraphs_to_remove.append(para)
			
			# Remove identified paragraphs
			for para in paragraphs_to_remove:
				try:
					para._element.getparent().remove(para._element)
				except:
					pass
			
			# Remove tables that might be Schedule tables (tables with "القسم" and "الصفحة" headers)
			# Also check for "فهرس" or "رقم الصفحة" tables
			tables_to_remove = []
			for table in doc.tables:
				if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
					header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
					# Check for various Schedule/Table of Contents indicators
					schedule_keywords = [
						('القسم', 'الصفحة'),
						('رقم الصفحة',),
						('فهرس',),
						('القوائم المالية', 'رقم الصفحة'),
						('Table of Contents',),
						('Contents',)
					]
					
					# Check if header matches any schedule pattern
					is_schedule = False
					for keywords in schedule_keywords:
						if all(keyword in header_text for keyword in keywords):
							is_schedule = True
							break
					
					# Also check if it's a 2-column table with page numbers in second column
					if not is_schedule and len(table.rows[0].cells) == 2:
						# Check first few data rows for page number pattern
						has_page_numbers = False
						for row in table.rows[1:min(4, len(table.rows))]:
							if len(row.cells) >= 2:
								second_cell = row.cells[1].text.strip()
								# If second cell contains only numbers or numbers with dash (e.g., "1 - 2")
								if second_cell and (second_cell.isdigit() or 
									(all(c.isdigit() or c in ['-', ' ', '–', '—'] for c in second_cell) and 
									 any(c.isdigit() for c in second_cell))):
									has_page_numbers = True
									break
						
						if has_page_numbers:
							is_schedule = True
					
					if is_schedule:
						tables_to_remove.append(table)
			
			# Remove identified tables
			for table in tables_to_remove:
				try:
					parent = table._element.getparent()
					if parent is not None:
						parent.remove(table._element)
						logger.info(f"Removed Schedule table from template: {header_text[:50] if 'header_text' in locals() else 'unknown'}")
				except Exception as e:
					logger.warning(f"Could not remove table from template: {e}")
			
			# Also check last few tables in template (they might be Schedule tables at the end)
			if len(doc.tables) > 0:
				all_tables = list(doc.tables)
				last_tables = all_tables[-2:] if len(all_tables) >= 2 else all_tables
				
				for table in last_tables:
					if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
						header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
						
						# Check if it's a Schedule table
						if ('رقم الصفحة' in header_text or 'Page' in header_text or 
							'فهرس' in header_text or 
							('القسم' in header_text and 'الصفحة' in header_text)) and len(table.rows[0].cells) == 2:
							try:
								parent = table._element.getparent()
								if parent is not None:
									parent.remove(table._element)
									logger.info(f"Removed Schedule table from end of template: {header_text[:50]}")
							except Exception as e:
								logger.warning(f"Could not remove end table from template: {e}")
			
			# Map sections to update content while preserving formatting
			section_map = {s.get('id'): s for s in sections}
			
			# Process sections and replace content in template while preserving formatting
			for paragraph in doc.paragraphs:
				para_text = paragraph.text.strip()
				if not para_text:
					continue
				
				# Try to match with sections by title
				for section in sections:
					section_title = section.get('title', '')
					if section_title and section_title in para_text:
						from docx.shared import Pt
						from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
						
						section_id = section.get('id')
						section_content = section.get('content', '')
						
						# Replace variables in content
						mapping = next((m for m in data_mappings if m.get('sectionId') == section_id), None)
						if mapping:
							source_data = data_cache.get(mapping.get('dataSource', {}).get('name'), {})
							for var_name, var_value in source_data.items():
								section_content = section_content.replace(f'{{{var_name}}}', str(var_value))
						
						# Replace variables
						var_pattern = r'\{([^}]+)\}'
						remaining_vars = re.findall(var_pattern, section_content)
						for var_name in remaining_vars:
							for source_id, source_data in data_cache.items():
								if var_name in source_data:
									section_content = section_content.replace(f'{{{var_name}}}', str(source_data[var_name]))
									break
						
						# Check if content contains a table
						if '<table' in section_content.lower():
							# Extract all tables from content
							table_matches = list(re.finditer(r'<table[^>]*>(.*?)</table>', section_content, re.DOTALL | re.IGNORECASE))
							if table_matches:
								# Clear paragraph first but keep the title if it exists
								original_style = paragraph.style
								para_text_before_clear = paragraph.text.strip()
								
								# Check if paragraph contains section title
								has_title = section_title and section_title in para_text_before_clear
								
								paragraph.clear()
								paragraph.style = original_style
								
								# Always add section title first (if it exists) before any content or tables
								if section_title:
									if not has_title:
										# Title wasn't in the paragraph, add it now
										title_run = paragraph.add_run(section_title)
										title_run.bold = True
										title_run.font.size = Pt(14)
										# Add bookmark for page reference
										bookmark_name = f'Sec_{section_id.replace("-", "_")}'
										_add_bookmark(paragraph, bookmark_name)
										paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
										paragraph = doc.add_paragraph()
										paragraph.style = original_style
									else:
										# Title was already in paragraph, restore it
										title_run = paragraph.add_run(section_title)
										title_run.bold = True
										title_run.font.size = Pt(14)
										# Add bookmark for page reference
										bookmark_name = f'Sec_{section_id.replace("-", "_")}'
										_add_bookmark(paragraph, bookmark_name)
										paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
										paragraph = doc.add_paragraph()
										paragraph.style = original_style
								
								# Process content, extracting tables
								last_pos = 0
								for table_match in table_matches:
									# Add content before table
									before_table = section_content[last_pos:table_match.start()].strip()
									if before_table:
										_apply_html_to_paragraph(paragraph, before_table, section.get('formatting', {}))
										paragraph = doc.add_paragraph()
										paragraph.style = original_style
									
									# Extract and add table
									table_html = table_match.group(0)
									_add_table_from_html(doc, table_html)
									
									last_pos = table_match.end()
								
								# Add remaining content after last table
								after_table = section_content[last_pos:].strip()
								if after_table:
									para = doc.add_paragraph()
									para.style = original_style
									_apply_html_to_paragraph(para, after_table, section.get('formatting', {}))
							else:
								# No table found, proceed normally
								original_style = paragraph.style
								paragraph.clear()
								paragraph.style = original_style
								_apply_html_to_paragraph(paragraph, section_content, section.get('formatting', {}))
						else:
							# No table, proceed normally
							original_style = paragraph.style
							paragraph.clear()
							paragraph.style = original_style
							
							# Add new content with HTML formatting
							_apply_html_to_paragraph(paragraph, section_content, section.get('formatting', {}))
						break
			
			# Also add any new sections that don't exist in template
			template_sections = set()
			for para in doc.paragraphs:
				if para.text.strip():
					template_sections.add(para.text.strip())
			
			for section in sorted(sections, key=lambda x: x.get('order', 0)):
				section_title = section.get('title', '')
				if section_title not in template_sections:
					# Add new section
					section_content = section.get('content', '')
					section_type = section.get('type', 'paragraph')
					
					# Replace variables
					mapping = next((m for m in data_mappings if m.get('sectionId') == section.get('id')), None)
					if mapping:
						source_data = data_cache.get(mapping.get('dataSource', {}).get('name'), {})
						for var_name, var_value in source_data.items():
							section_content = section_content.replace(f'{{{var_name}}}', str(var_value))
					
					# Add new section with preserved formatting or default
					section_id = section.get('id', f'section_{section.get("order", 0)}')
					bookmark_name = f'Sec_{section_id.replace("-", "_")}'
					
					if section_type == 'heading':
						heading = doc.add_heading(section_title, level=min(section.get('level', 1), 9))
						if section.get('formatting', {}).get('alignment'):
							heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
						# Add bookmark for page reference
						_add_bookmark(heading, bookmark_name)
						if section_content:
							html_to_word_paragraph(doc, section_content)
					else:
						para = doc.add_paragraph()
						if section.get('style'):
							try:
								para.style = section.get('style')
							except:
								pass
						if section_title:
							para.add_run(section_title).bold = True
							# Add bookmark for page reference
							_add_bookmark(para, bookmark_name)
						if section_content:
							_apply_html_to_paragraph(para, section_content, section.get('formatting', {}))
		else:
			# Create new document from sections (no template)
			doc = Document()
			
			# Schedule table removed - no longer adding it to the document
			# if include_schedule and sections:
			# 	_add_schedule_table(doc, sections)
			
			# Process sections and add content
			for section in sorted(sections, key=lambda x: x.get('order', 0)):
				section_id = section.get('id')
				section_title = section.get('title', '')
				section_content = section.get('content', '')
				section_type = section.get('type', 'paragraph')
				
				# Replace variables in content
				mapping = next((m for m in data_mappings if m.get('sectionId') == section_id), None)
				if mapping:
					source_data = data_cache.get(mapping.get('dataSource', {}).get('name'), {})
					for var_name, var_value in source_data.items():
						section_content = section_content.replace(f'{{{var_name}}}', str(var_value))
				
				# Replace variables
				var_pattern = r'\{([^}]+)\}'
				remaining_vars = re.findall(var_pattern, section_content)
				for var_name in remaining_vars:
					for source_id, source_data in data_cache.items():
						if var_name in source_data:
							section_content = section_content.replace(f'{{{var_name}}}', str(source_data[var_name]))
							break
				
				# Add section to document
				if section_type == 'heading':
					heading = doc.add_heading(section_title, level=min(section.get('level', 1), 9))
					heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
					# Add bookmark for page reference
					bookmark_name = f'Sec_{section_id.replace("-", "_")}'
					_add_bookmark(heading, bookmark_name)
					if section_content:
						html_to_word_paragraph(doc, section_content)
				elif section_type == 'table':
					if '<table' in section_content:
						rows = re.findall(r'<tr[^>]*>(.*?)</tr>', section_content, re.DOTALL)
						if rows:
							first_row = rows[0]
							cols = len(re.findall(r'<td[^>]*>', first_row))
							if cols > 0:
								table = doc.add_table(rows=len(rows), cols=cols)
								# Apply table style safely with fallbacks
								try:
									table.style = 'Table Grid'
								except:
									try:
										table.style = 'Light Grid'
									except:
										try:
											table.style = 'Light Grid Accent 1'
										except:
											try:
												table.style = 'Table Normal'
											except:
												# Use default table style (no style)
												pass
								
								for row_idx, row_html in enumerate(rows):
									cells = re.findall(r'<td[^>]*>(.*?)</td>', row_html, re.DOTALL)
									for col_idx, cell_html in enumerate(cells):
										if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
											cell = table.rows[row_idx].cells[col_idx]
											cell.text = re.sub(r'<[^>]+>', '', cell_html).strip()
											for paragraph in cell.paragraphs:
												paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
					else:
						para = doc.add_paragraph(section_title)
						para.style = 'Heading 2'
						para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
						# Add bookmark for page reference
						bookmark_name = f'Sec_{section_id.replace("-", "_")}'
						_add_bookmark(para, bookmark_name)
						html_to_word_paragraph(doc, section_content)
				else:
					# Regular paragraph section
					if section_title:
						para = doc.add_paragraph(section_title)
						para.style = 'Heading 2'
						para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
						# Add bookmark for page reference
						bookmark_name = f'Sec_{section_id.replace("-", "_")}'
						_add_bookmark(para, bookmark_name)
					
					if section_content:
						# Check if content contains a table
						if '<table' in section_content.lower():
							# Extract all tables from content
							table_matches = list(re.finditer(r'<table[^>]*>(.*?)</table>', section_content, re.DOTALL | re.IGNORECASE))
							if table_matches:
								last_pos = 0
								for table_match in table_matches:
									# Add content before table
									before_table = section_content[last_pos:table_match.start()].strip()
									if before_table:
										html_to_word_paragraph(doc, before_table)
									
									# Extract and add table
									table_html = table_match.group(0)
									_add_table_from_html(doc, table_html)
									
									last_pos = table_match.end()
								
								# Add remaining content after last table
								after_table = section_content[last_pos:].strip()
								if after_table:
									html_to_word_paragraph(doc, after_table)
							else:
								html_to_word_paragraph(doc, section_content)
						else:
							html_to_word_paragraph(doc, section_content)
		
		# Final cleanup: Remove any Schedule tables or paragraphs that might have been added
		# Check all paragraphs and tables from the end backwards
		logger.info("Performing final cleanup: removing any Schedule-related content")
		
		# More aggressive cleanup: Remove any content that looks like a table of contents
		# First, collect all elements (paragraphs and tables) with their positions
		from docx.oxml import parse_xml
		from docx.oxml.ns import qn
		
		# Remove paragraphs containing Schedule/Table of Contents keywords
		paragraphs_to_remove = []
		for para in doc.paragraphs:
			para_text = para.text.strip()
			# Check for various forms of table of contents
			schedule_keywords = [
				'Schedule', 'جدول المحتويات', 'جدول', 'المحتويات',
				'فهرس', 'القوائم المالية', 'رقم الصفحة',
				'Table of Contents', 'Contents'
			]
			if any(keyword in para_text for keyword in schedule_keywords):
				# Additional check: if paragraph is short and contains numbers, likely TOC
				if len(para_text) < 200 or any(char.isdigit() for char in para_text):
					paragraphs_to_remove.append(para)
		
		for para in paragraphs_to_remove:
			try:
				para._element.getparent().remove(para._element)
				logger.info(f"Removed Schedule paragraph: {para.text[:50]}")
			except Exception as e:
				logger.warning(f"Could not remove paragraph: {e}")
		
		# Remove tables that might be Schedule/Table of Contents tables
		# Check ALL tables, not just from the end
		tables_to_remove = []
		for table in doc.tables:
			if len(table.rows) > 0:
				# Get header row text
				header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
				
				# Check for Schedule/Table of Contents indicators
				schedule_indicators = [
					'القسم', 'الصفحة', 'Section', 'Page', 'رقم الصفحة',
					'القوائم المالية', 'تقرير', 'قائمة', 'إيضاح'
				]
				
				# If header contains schedule indicators
				if any(indicator in header_text for indicator in schedule_indicators):
					# Additional checks:
					# 1. Table has 2 columns (typical for TOC)
					# 2. OR table has many rows with numbers (page numbers)
					# 3. OR table is at the end of document
					if len(table.rows[0].cells) == 2:
						# Check if it looks like a TOC (has page numbers or section names)
						has_page_numbers = False
						has_section_names = False
						for row in table.rows[1:min(6, len(table.rows))]:  # Check first 5 data rows
							if len(row.cells) >= 2:
								# Check if second column has numbers (page numbers)
								if any(char.isdigit() for char in row.cells[1].text):
									has_page_numbers = True
								# Check if first column has section-like text
								if len(row.cells[0].text.strip()) > 5:
									has_section_names = True
						
						if has_page_numbers or (has_section_names and len(table.rows) > 3):
							tables_to_remove.append(table)
							logger.info(f"Identified Schedule table: {header_text[:50]}")
		
		# Remove identified tables
		for table in tables_to_remove:
			try:
				table._element.getparent().remove(table._element)
				logger.info("Removed Schedule table from document")
			except Exception as e:
				logger.warning(f"Could not remove table: {e}")
		
		# Also check for tables at the very end of document (last 3 tables)
		# These are more likely to be Schedule tables
		# Use a more aggressive approach: iterate through document body elements
		try:
			from docx.oxml import OxmlElement
			from docx.oxml.ns import qn
			
			# Get document body
			body = doc._body._body
			
			# Iterate through all elements in reverse order (from end to start)
			# This ensures we catch tables at the end
			elements_to_remove = []
			
			# First, collect all table elements
			for element in body:
				if element.tag.endswith('tbl'):  # Table element
					# Try to find the table object
					for table in doc.tables:
						if table._element == element:
							# Check if this is a Schedule table
							if len(table.rows) > 0:
								header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
								
								# Check for Schedule indicators
								schedule_checks = [
									'رقم الصفحة' in header_text,
									'Page' in header_text,
									'فهرس' in header_text,
									'القوائم المالية' in header_text and 'رقم الصفحة' in header_text,
									('القسم' in header_text and 'الصفحة' in header_text),
									len(table.rows[0].cells) == 2 and any(
										any(char.isdigit() for char in row.cells[1].text) 
										for row in table.rows[1:min(4, len(table.rows))] 
										if len(row.cells) >= 2
									)
								]
								
								if any(schedule_checks):
									elements_to_remove.append(element)
									logger.info(f"Marked Schedule table for removal: {header_text[:50]}")
									break
			
			# Remove identified elements
			for element in elements_to_remove:
				try:
					parent = element.getparent()
					if parent is not None:
						parent.remove(element)
						logger.info("Removed Schedule table element from document body")
				except Exception as e:
					logger.warning(f"Could not remove table element: {e}")
			
			# Also check last few tables directly (backup method)
			if len(doc.tables) > 0:
				# Get last few tables
				all_tables = list(doc.tables)
				last_tables = all_tables[-3:] if len(all_tables) >= 3 else all_tables
				
				for table in last_tables:
					if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
						header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
						
						# More aggressive check: if it has "رقم الصفحة" or looks like TOC, remove it
						if ('رقم الصفحة' in header_text or 'Page' in header_text or 'فهرس' in header_text) and len(table.rows[0].cells) == 2:
							try:
								table._element.getparent().remove(table._element)
								logger.info("Removed Schedule table from end (backup method)")
							except Exception as e:
								logger.warning(f"Could not remove end table (backup): {e}")
		
		except Exception as e:
			logger.warning(f"Error in aggressive table removal: {e}")
		
		# Save to BytesIO
		try:
			logger.info("Saving document to BytesIO")
			output = BytesIO()
			doc.save(output)
			output.seek(0)
			doc_bytes = output.getvalue()
			logger.info(f"Document saved successfully, size: {len(doc_bytes)} bytes")
			
			return Response(
				content=doc_bytes,
				media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
				headers={
					'Content-Disposition': 'attachment; filename="generated_report.docx"',
					'Content-Length': str(len(doc_bytes)),
					'Cache-Control': 'no-cache, no-store, must-revalidate',
					'X-Content-Type-Options': 'nosniff'
				},
				status_code=200
			)
		except Exception as save_error:
			logger.error(f"Error saving document: {save_error}", exc_info=True)
			raise Exception(f"Failed to save document: {str(save_error)}")
		
	except Exception as e:
		error_msg = str(e)
		traceback_str = traceback.format_exc()
		logger.error(f"ERROR generating report: {error_msg}")
		logger.error(traceback_str)
		# Return JSON error response for better frontend handling
		return JSONResponse(
			content={
				"success": False,
				"error": error_msg,
				"traceback": traceback_str.split('\n')[:10]  # First 10 lines of traceback
			},
			status_code=500
		)


@router.post("/fetch-data")
async def fetch_data_from_source(payload: dict = Body(...)):
	"""Fetch data from database/XBRL/Excel source"""
	try:
		source_type = payload.get('type')  # 'database', 'xbrl', 'excel'
		source_config = payload.get('config', {})
		
		if source_type == 'database':
			# Fetch from database table
			import pyodbc
			from config import get_db_connection
			
			table_name = source_config.get('table')
			columns = source_config.get('columns', [])
			where_clause = source_config.get('where', '')
			
			conn = get_db_connection()
			cursor = conn.cursor()
			
			query = f"SELECT {', '.join(columns) if columns else '*'} FROM {table_name}"
			if where_clause:
				query += f" WHERE {where_clause}"
			
			cursor.execute(query)
			rows = cursor.fetchall()
			columns = [desc[0] for desc in cursor.description]
			
			conn.close()
			
			return {
				'success': True,
				'data': {
					'columns': columns,
					'rows': [list(row) for row in rows]
				}
			}
		
		elif source_type == 'xbrl':
			# Parse XBRL file
			# TODO: Implement XBRL parsing
			return {
				'success': True,
				'data': {'columns': [], 'rows': []}
			}
		
		elif source_type == 'excel':
			# Parse Excel file
			# TODO: Implement Excel parsing
			return {
				'success': True,
				'data': {'columns': [], 'rows': []}
			}
		
		else:
			return {
				'success': False,
				'error': f'Unknown source type: {source_type}'
			}
			
	except Exception as e:
		error_msg = str(e)
		traceback_str = traceback.format_exc()
		logger.error(f"ERROR fetching data: {error_msg}")
		logger.error(traceback_str)
		return {
			'success': False,
			'error': error_msg,
			'traceback': traceback_str
		}


@router.post("/get-excel-columns")
async def get_excel_columns(file: UploadFile = File(...), sheet: str = Form(None)):
	"""Get column names from Excel file"""
	try:
		# Read Excel file
		file_content = await file.read()
		file_stream = BytesIO(file_content)
		
		# Try to read with openpyxl first (for .xlsx)
		try:
			wb = openpyxl.load_workbook(file_stream, read_only=True)
			if sheet:
				ws = wb[sheet]
			else:
				ws = wb.active
			
			# Get first row as headers
			headers = []
			for cell in ws[1]:
				if cell.value:
					headers.append(str(cell.value))
			
			wb.close()
			
			return JSONResponse({
				'success': True,
				'columns': headers
			})
		except Exception as e1:
			# Fallback to pandas
			try:
				file_stream.seek(0)
				if sheet:
					df = pd.read_excel(file_stream, sheet_name=sheet, nrows=0)
				else:
					df = pd.read_excel(file_stream, nrows=0)
				
				columns = df.columns.tolist()
				return JSONResponse({
					'success': True,
					'columns': [str(col) for col in columns]
				})
			except Exception as e2:
				logger.error(f"Error reading Excel: {e1}, {e2}")
				raise HTTPException(status_code=400, detail=f"Failed to read Excel file: {str(e2)}")
		
	except Exception as e:
		error_msg = str(e)
		traceback_str = traceback.format_exc()
		logger.error(f"ERROR getting Excel columns: {error_msg}")
		logger.error(traceback_str)
		raise HTTPException(status_code=500, detail=f"Error reading Excel columns: {error_msg}")


@router.post("/get-excel-data")
async def get_excel_data(file: UploadFile = File(...), sheet: str = Form(None), range: str = Form(None), columns: str = Form(...)):
	"""Get data from Excel file for selected columns"""
	try:
		import json
		selected_columns = json.loads(columns) if columns else []
		
		# Read Excel file
		file_content = await file.read()
		file_stream = BytesIO(file_content)
		
		# Read Excel with pandas
		try:
			if sheet:
				df = pd.read_excel(file_stream, sheet_name=sheet)
			else:
				df = pd.read_excel(file_stream)
			
			# Apply range if specified
			if range:
				# Parse range like "A1:D10" - simple parsing
				try:
					# Basic range parsing (e.g., "A1:D10")
					import re
					match = re.match(r'([A-Z]+)(\d+):([A-Z]+)(\d+)', range.upper())
					if match:
						start_col_letter = match.group(1)
						start_row = int(match.group(2))
						end_col_letter = match.group(3)
						end_row = int(match.group(4))
						
						# Convert column letters to numbers
						def col_letter_to_num(letters):
							result = 0
							for char in letters:
								result = result * 26 + (ord(char) - ord('A') + 1)
							return result
						
						start_col = col_letter_to_num(start_col_letter) - 1
						end_col = col_letter_to_num(end_col_letter)
						
						# Slice dataframe (pandas uses 0-based indexing)
						df = df.iloc[start_row-1:end_row, start_col:end_col]
				except Exception as range_error:
					logger.warning(f"Could not parse range {range}: {range_error}")
					pass
			
			# Filter to selected columns only
			if selected_columns:
				# Find matching columns (case-insensitive)
				available_columns = df.columns.tolist()
				matching_columns = []
				for col in selected_columns:
					# Try exact match first
					if col in available_columns:
						matching_columns.append(col)
					else:
						# Try case-insensitive match
						found = False
						for avail_col in available_columns:
							if str(avail_col).lower() == str(col).lower():
								matching_columns.append(avail_col)
								found = True
								break
						if not found:
							matching_columns.append(col)  # Keep original if not found
				
				# Filter dataframe to selected columns
				existing_cols = [c for c in matching_columns if c in df.columns]
				if existing_cols:
					df = df[existing_cols]
			
			# Convert to list format
			columns_list = df.columns.tolist()
			rows_list = df.values.tolist()
			
			return JSONResponse({
				'success': True,
				'data': {
					'columns': [str(col) for col in columns_list],
					'rows': [[str(val) if val is not None else '' for val in row] for row in rows_list]
				}
			})
			
		except Exception as e1:
			logger.error(f"Error reading Excel data: {e1}")
			raise HTTPException(status_code=400, detail=f"Failed to read Excel data: {str(e1)}")
		
	except Exception as e:
		error_msg = str(e)
		traceback_str = traceback.format_exc()
		logger.error(f"ERROR getting Excel data: {error_msg}")
		logger.error(traceback_str)
		raise HTTPException(status_code=500, detail=f"Error reading Excel data: {error_msg}")


@router.post("/get-xbrl-concepts")
async def get_xbrl_concepts(file: UploadFile = File(...)):
	"""Get XBRL taxonomy concepts from XBRL file"""
	try:
		file_content = await file.read()
		
		# Parse XBRL XML
		try:
			root = etree.fromstring(file_content)
			
			# Get namespace map
			nsmap = root.nsmap
			
			# Extract concepts from facts
			concepts = set()
			excluded_tags = {'context', 'unit', 'schemaRef', 'linkbaseRef', 'xbrl', 'xbrli'}
			
			# Iterate through all elements
			for elem in root.iter():
				if elem.tag and '}' in elem.tag:
					# Extract namespace URI and local name
					ns_uri, local_name = elem.tag[1:].split('}', 1)
					
					# Skip excluded tags
					if local_name in excluded_tags:
						continue
					
					# Find namespace prefix from nsmap
					ns_prefix = None
					for prefix, uri in nsmap.items():
						if uri == ns_uri:
							ns_prefix = prefix
							break
					
					# If no prefix found, try to extract from common namespaces
					if not ns_prefix:
						if 'ifrs-full' in ns_uri or 'ifrs' in ns_uri:
							ns_prefix = 'ifrs-full'
						elif 'egx-eas' in ns_uri or 'egx' in ns_uri:
							ns_prefix = 'egx-eas'
						elif 'us-gaap' in ns_uri:
							ns_prefix = 'us-gaap'
						else:
							# Use last part of URI as prefix
							ns_prefix = ns_uri.split('/')[-1] if ns_uri else 'unknown'
					
					full_concept = f"{ns_prefix}:{local_name}"
					concepts.add(full_concept)
			
			# Convert to sorted list
			concepts_list = sorted(list(concepts))
			
			return JSONResponse({
				'success': True,
				'concepts': concepts_list
			})
			
		except etree.XMLSyntaxError as e:
			logger.error(f"XML parsing error: {e}")
			raise HTTPException(status_code=400, detail=f"Invalid XBRL/XML file: {str(e)}")
		
	except Exception as e:
		error_msg = str(e)
		traceback_str = traceback.format_exc()
		logger.error(f"ERROR getting XBRL concepts: {error_msg}")
		logger.error(traceback_str)
		raise HTTPException(status_code=500, detail=f"Error reading XBRL concepts: {error_msg}")


@router.post("/get-xbrl-data")
async def get_xbrl_data(
	file: UploadFile = File(...),
	concepts: str = Form("")
):
	"""
	Extract values for specified XBRL concepts from an XBRL/XML file.
	Returns concept-value pairs.
	"""
	import tempfile
	import os
	
	try:
		# Save uploaded file temporarily
		with tempfile.NamedTemporaryFile(delete=False, suffix='.xml') as tmp_file:
			content = await file.read()
			tmp_file.write(content)
			tmp_path = tmp_file.name
		
		try:
			# Parse XML/XBRL file
			tree = etree.parse(tmp_path)
			root = tree.getroot()
			
			# Parse concepts list
			concept_list = [c.strip() for c in concepts.split(',') if c.strip()]
			if not concept_list:
				return JSONResponse(content={"success": False, "error": "No concepts specified"}, status_code=400)
			
			# Extract namespace map
			nsmap = root.nsmap.copy()
			# Common XBRL namespaces
			if None in nsmap:
				nsmap['xbrli'] = nsmap[None]  # Default namespace often is xbrli
			if 'xbrli' not in nsmap:
				nsmap['xbrli'] = 'http://www.xbrl.org/2003/instance'
			if 'xbrl' not in nsmap:
				nsmap['xbrl'] = 'http://www.xbrl.org/2003/instance'
			
			# Find all facts (elements with contextRef)
			concept_values = {}
			
			for concept_qname in concept_list:
				# Try to find the concept in the document
				# Concept might be prefixed (e.g., "ifrs-full:Revenue") or full QName
				concept_found = False
				
				# Try different namespace prefixes
				for prefix, uri in nsmap.items():
					if prefix:
						try:
							# Try with prefix
							if ':' in concept_qname:
								parts = concept_qname.split(':', 1)
								if parts[0] == prefix:
									tag = f"{{{uri}}}{parts[1]}"
								else:
									tag = f"{{{uri}}}{concept_qname.split(':')[-1]}"
							else:
								tag = f"{{{uri}}}{concept_qname}"
							
							elements = root.findall(f".//{tag}", namespaces=nsmap)
							if elements:
								# Get the first fact's value
								for elem in elements:
									if elem.text:
										concept_values[concept_qname] = elem.text.strip()
										concept_found = True
										break
							if concept_found:
								break
						except Exception:
							continue
				
				# Fallback: search by local name only
				if not concept_found:
					local_name = concept_qname.split(':')[-1]
					for elem in root.iter():
						if elem.tag.endswith(f"}}{local_name}") or elem.tag == local_name:
							if elem.text:
								concept_values[concept_qname] = elem.text.strip()
								break
			
			return JSONResponse(content={
				"success": True,
				"data": {
					"concepts": list(concept_values.keys()),
					"values": concept_values
				}
			})
			
		finally:
			# Clean up temp file
			if os.path.exists(tmp_path):
				os.unlink(tmp_path)
				
	except Exception as e:
		logger.error(f"Error reading XBRL file: {e}", exc_info=True)
		return JSONResponse(content={"success": False, "error": str(e)}, status_code=500)

