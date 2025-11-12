"""
Common utilities for API routes

Note: The large report generation functions (generate_excel_report, generate_word_report, 
generate_pdf_report) are intentionally kept in individual route files due to their size 
(500+ lines each) and specific module dependencies. Only truly common, reusable utilities 
are centralized here.
"""
from typing import Dict, Any, Optional, List
import json
import sys
from datetime import datetime


def write_debug(msg: str) -> None:
    """Write debug message to file with timestamp"""
    from datetime import datetime
    timestamp = datetime.now().strftime('%H:%M:%S.%f')[:-3]
    msg_with_time = f"[{timestamp}] {msg}"
    with open('debug_log.txt', 'a', encoding='utf-8') as f:
        f.write(f"{msg_with_time}\n")
        f.flush()
    sys.stderr.write(f"{msg_with_time}\n")
    sys.stderr.flush()


def parse_header_config(headerConfig: Optional[str]) -> Dict[str, Any]:
    """Parse header configuration from JSON string"""
    header_config = {}
    if headerConfig:
        try:
            header_config = json.loads(headerConfig)
        except json.JSONDecodeError:
            header_config = {}
    return header_config


def merge_header_config(module_name: str, header_config: Dict[str, Any]) -> Dict[str, Any]:
    """Merge custom header config with default config"""
    from utils.export_utils import get_default_header_config
    default_config = get_default_header_config(module_name)
    return {**default_config, **header_config}


def convert_to_boolean(value: Any) -> bool:
    """Convert string/boolean value to boolean"""
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.lower() in ['true', '1', 'yes']
    return bool(value)


# In-memory cache to prevent duplicate saves within a short time window
_export_cache = {}  # Key: cache_key, Value: (result_dict, timestamp)
_export_lock = {}  # Key: cache_key, Value: timestamp (for locking during save operations)

async def save_and_log_export(
    content: bytes,
    file_extension: str,
    dashboard: str,
    card_type: Optional[str] = None,
    header_config: Optional[Dict[str, Any]] = None,
    created_by: Optional[str] = None,
    date_range: Optional[Dict[str, Optional[str]]] = None,
    export_type: Optional[str] = None
) -> Dict[str, Any]:
    """
    Save export file to disk and log to database.
    
    Args:
        content: File content (bytes)
        file_extension: File extension (pdf, xlsx, docx)
        dashboard: Dashboard name (incidents, kris, risks, controls)
        card_type: Card type for the export
        header_config: Header configuration dict
        created_by: User who created the export (defaults to "System")
        date_range: Dict with startDate and endDate
        
    Returns:
        Dict with file_path, relative_path, and export_id
    """
    import os
    import pyodbc
    import hashlib
    import time
    from config import get_database_connection_string
    global _export_cache
    
    # Ensure date_range is a dict (default to empty dict if None)
    if date_range is None:
        date_range = {}
    elif not isinstance(date_range, dict):
        date_range = {}
    
    # Create a cache key to prevent duplicate saves within 3 seconds
    # Include content hash to detect identical exports
    content_hash = hashlib.md5(content).hexdigest()[:8]
    cache_key_parts = [
        str(dashboard) if dashboard else '',
        str(card_type) if card_type else '',
        str(file_extension) if file_extension else '',
        str(date_range.get('startDate', '')) if date_range else '',
        str(date_range.get('endDate', '')) if date_range else '',
        str(content_hash)
    ]
    cache_key = hashlib.md5('|'.join(cache_key_parts).encode()).hexdigest()
    
    # Clean up old cache entries (older than 10 seconds - extended to prevent duplicates)
    current_time = time.time()
    keys_to_remove = [k for k, v in _export_cache.items() if current_time - v[1] >= 10]
    for k in keys_to_remove:
        del _export_cache[k]
    
    # Clean up old locks (older than 5 seconds)
    lock_keys_to_remove = [k for k, v in _export_lock.items() if current_time - v >= 5]
    for k in lock_keys_to_remove:
        del _export_lock[k]
    
    # Check if another request is currently saving the same export (lock check)
    if cache_key in _export_lock:
        lock_time = _export_lock[cache_key]
        # Wait a bit if lock is very recent (within 1 second)
        if current_time - lock_time < 1:
            import asyncio
            await asyncio.sleep(0.2)  # Wait 200ms for the other request to finish
            # Check cache again after waiting
            if cache_key in _export_cache:
                cached_result, _ = _export_cache[cache_key]
                if os.path.exists(cached_result['file_path']):
                    write_debug(f"[Save Export] Waited for concurrent save, using existing file: {cached_result['filename']}")
                    return cached_result
    
    # Check if we recently saved the same export (deduplication)
    if cache_key in _export_cache:
        cached_result, _ = _export_cache[cache_key]
        # Verify the file still exists
        if os.path.exists(cached_result['file_path']):
            write_debug(f"[Save Export] Duplicate export detected via cache, using existing file: {cached_result['filename']}")
            return cached_result
        else:
            # File was deleted, remove from cache and continue
            del _export_cache[cache_key]
    
    # Set lock to prevent concurrent saves of the same export
    _export_lock[cache_key] = current_time
    write_debug(f"[Save Export] Setting lock for cache_key: {cache_key[:8]}...")
    
    # Get user info (default to "System")
    user_name = created_by or "System"
    
    # Build title for database check (we need this before creating filename)
    # Prioritize card_type for unique database titles, not header_config.title (which is always "Dashboard Report")
    db_title = "Report"
    if card_type:
        import re
        # Convert cardType to readable name (e.g., "pendingPreparer" -> "Pending Preparer")
        db_title = card_type
        db_title = re.sub(r'([A-Z])', r' \1', db_title).strip()
        db_title = db_title.replace('_', ' ').title()
        # Capitalize first letter of each word properly
        db_title = ' '.join(word.capitalize() for word in db_title.split())
    elif header_config and header_config.get("title"):
        # Only use header_config.title as fallback if no card_type (should be rare)
        db_title = header_config.get("title", "Report")
    
    # Check database FIRST to see if a recent duplicate exists (before saving file)
    # This prevents duplicate database entries when multiple requests come in simultaneously
    try:
        connection_string = get_database_connection_string()
        conn_check = pyodbc.connect(connection_string)
        cursor_check = conn_check.cursor()
        try:
            # Build date suffix for title matching
            date_suffix_check = ""
            if date_range and isinstance(date_range, dict):
                start = date_range.get("startDate")
                end = date_range.get("endDate")
                if start and end:
                    date_suffix_check = f" ({start} to {end})"
                elif start:
                    date_suffix_check = f" (from {start})"
                elif end:
                    date_suffix_check = f" (until {end})"
            
            now_check = datetime.now()
            date_str_check = now_check.strftime('%Y-%m-%d')
            export_title_check = f"{db_title}{date_suffix_check} - {date_str_check}"
            
            # Check for recent duplicate in database (last 30 seconds)
            cursor_check.execute(
                """
                SELECT TOP 1 id, src FROM dbo.report_exports 
                WHERE dashboard = ? 
                  AND format = ?
                  AND created_by = ?
                  AND (
                    title = ? 
                    OR title LIKE ?
                  )
                  AND created_at >= DATEADD(SECOND, -30, GETDATE())
                ORDER BY created_at DESC
                """,
                (dashboard, file_extension, user_name, export_title_check, f"{db_title}%")
            )
            existing_db = cursor_check.fetchone()
            
            if existing_db:
                # Found existing record - use it and return early (no file save needed)
                existing_id = existing_db[0]
                existing_src = existing_db[1]
                base_dir_check = os.path.dirname(os.path.dirname(__file__))
                existing_file_path = os.path.join(base_dir_check, existing_src) if existing_src else None
                
                if existing_file_path and os.path.exists(existing_file_path):
                    write_debug(f"[Save Export] Duplicate database entry found (ID: {existing_id}), returning existing file: {existing_src}")
                    result = {
                        "file_path": existing_file_path,
                        "relative_path": existing_src,
                        "filename": os.path.basename(existing_src),
                        "export_id": existing_id
                    }
                    # Cache it for future lookups
                    _export_cache[cache_key] = (result, time.time())
                    return result
        finally:
            cursor_check.close()
            conn_check.close()
    except Exception as e:
        write_debug(f"[Save Export] Database pre-check failed (continuing): {str(e)}")
        import traceback
        write_debug(f"[Save Export] Pre-check traceback: {traceback.format_exc()}")
        # Continue with file save even if pre-check fails
    
    # Create readable filename
    now = datetime.now()
    date_str = now.strftime('%Y-%m-%d')
    # Use microseconds for better uniqueness (avoid collisions)
    time_str = now.strftime('%H%M%S_%f')[:-3]  # Include milliseconds
    
    # Build filename from card_type (priority) or header config
    # Always prefer cardType for filename even if header config has a different title
    filename_title = "Report"
    if card_type:
        # Convert cardType to readable name (e.g., "pendingPreparer" -> "Pending_Preparer")
        filename_title = card_type
        # Add spaces before capital letters
        import re
        filename_title = re.sub(r'([A-Z])', r' \1', filename_title).strip()
        filename_title = filename_title.replace('_', ' ').title()
        # Replace spaces with underscores for filename
        filename_title = filename_title.replace(' ', '_')
    elif header_config and header_config.get("title"):
        # Fallback to header title if no card_type
        filename_title = header_config.get("title", "Report")
        filename_title = "".join(c for c in filename_title if c.isalnum() or c in (' ', '-', '_')).strip()
        filename_title = filename_title.replace(' ', '_')
    
    # Add date range to database title if available
    date_suffix = ""
    if date_range:
        start = date_range.get("startDate")
        end = date_range.get("endDate")
        if start and end:
            date_suffix = f" ({start} to {end})"
        elif start:
            date_suffix = f" (from {start})"
        elif end:
            date_suffix = f" (until {end})"
    
    # Create readable filename (always use cardType-based name)
    safe_title = "".join(c for c in filename_title if c.isalnum() or c in ('-', '_')).strip()
    base_filename = f"{dashboard}_{safe_title}_{date_str}_{time_str}.{file_extension}"
    
    # Save file to reports_export directory
    base_dir = os.path.dirname(os.path.dirname(__file__))
    date_folder = now.strftime('%Y-%m-%d')
    
    # If export_type is 'transaction', save to transaction folder in root
    if export_type and export_type.lower() == 'transaction':
        reports_export_dir = os.path.join(base_dir, "transaction", date_folder)
    else:
        reports_export_dir = os.path.join(base_dir, "reports_export", date_folder)
    
    os.makedirs(reports_export_dir, exist_ok=True)
    
    # Ensure unique filename (handle collisions if they somehow occur)
    readable_filename = base_filename
    file_path = os.path.join(reports_export_dir, readable_filename)
    counter = 1
    while os.path.exists(file_path):
        # If file exists (should be rare with microseconds), append counter
        name_part = base_filename.rsplit('.', 1)[0]
        ext_part = base_filename.rsplit('.', 1)[1]
        readable_filename = f"{name_part}_{counter}.{ext_part}"
        file_path = os.path.join(reports_export_dir, readable_filename)
        counter += 1
        if counter > 1000:  # Safety limit
            write_debug(f"[Save Export] Warning: Too many file collisions for {base_filename}")
            break
    
    # Build relative path based on export type
    if export_type and export_type.lower() == 'transaction':
        relative_path = f"transaction/{date_folder}/{readable_filename}"
    else:
        relative_path = f"reports_export/{date_folder}/{readable_filename}"
    
    # Write file (only once)
    try:
        with open(file_path, 'wb') as f:
            f.write(content)
    except Exception as file_err:
        # Remove lock on error
        if cache_key in _export_lock:
            del _export_lock[cache_key]
        write_debug(f"[Save Export] Failed to write file: {str(file_err)}")
        raise
    
    # Log to database
    export_id = None
    try:
        connection_string = get_database_connection_string()
        conn = pyodbc.connect(connection_string)
        cursor = conn.cursor()
        try:
            # Ensure table exists with created_by and type
            cursor.execute(
                """
                IF NOT EXISTS (
                  SELECT * FROM INFORMATION_SCHEMA.TABLES 
                  WHERE TABLE_NAME = 'report_exports' AND TABLE_SCHEMA='dbo'
                )
                BEGIN
                  CREATE TABLE dbo.report_exports (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    title NVARCHAR(255) NOT NULL,
                    src NVARCHAR(1024) NULL,
                    format NVARCHAR(20) NOT NULL,
                    dashboard NVARCHAR(100) NULL,
                    type NVARCHAR(50) NULL,
                    created_by NVARCHAR(255) NULL,
                    created_at DATETIME2 DEFAULT GETDATE()
                  )
                END
                """
            )
            conn.commit()
            
            # Add created_by column if it doesn't exist
            try:
                cursor.execute(
                    """
                    IF NOT EXISTS (
                      SELECT * FROM INFORMATION_SCHEMA.COLUMNS 
                      WHERE TABLE_NAME = 'report_exports' AND COLUMN_NAME = 'created_by'
                    )
                    BEGIN
                      ALTER TABLE dbo.report_exports ADD created_by NVARCHAR(255) NULL
                    END
                    """
                )
                conn.commit()
            except Exception:
                pass
            
            # Add type column if it doesn't exist
            try:
                cursor.execute(
                    """
                    IF NOT EXISTS (
                      SELECT * FROM INFORMATION_SCHEMA.COLUMNS 
                      WHERE TABLE_NAME = 'report_exports' AND COLUMN_NAME = 'type'
                    )
                    BEGIN
                      ALTER TABLE dbo.report_exports ADD type NVARCHAR(50) NULL
                    END
                    """
                )
                conn.commit()
            except Exception:
                pass
            
            # Insert export record (only if not already exists)
            # Use db_title for database record (can include header config title)
            export_title = f"{db_title}{date_suffix} - {date_str}"
            
            # Determine type based on export_type parameter or dashboard
            # If export_type is provided, use it; otherwise determine from dashboard
            if not export_type:
                # Dashboard reports: incidents, kris, risks, controls
                # Transaction reports: everything else
                export_type = "transaction"  # Default
                if dashboard and dashboard.lower() in ['incidents', 'kris', 'risks', 'controls']:
                    export_type = "dashboard"
            
            # Check if record already exists with same parameters within the last 30 seconds
            # Check by multiple fields to catch duplicates even with different file paths or timestamps
            # This prevents the same export from being logged 3 times
            cursor.execute(
                """
                SELECT TOP 1 id, src FROM dbo.report_exports 
                WHERE (
                  src = ? 
                  OR (
                    title = ? 
                    AND dashboard = ? 
                    AND format = ?
                    AND created_by = ?
                  )
                )
                AND created_at >= DATEADD(SECOND, -30, GETDATE())
                ORDER BY created_at DESC
                """,
                (relative_path, export_title, dashboard, file_extension, user_name)
            )
            existing = cursor.fetchone()
            
            if existing:
                # Use existing record ID and path (prevent duplicate database entries)
                export_id = existing[0]
                existing_src = existing[1]
                write_debug(f"[Save Export] Duplicate database entry prevented. Using existing ID: {export_id}, existing src: {existing_src}")
                # Update result to use existing path if different
                if existing_src and existing_src != relative_path:
                    relative_path = existing_src
                    # Update file_path if we can find the existing file
                    existing_file_path = os.path.join(base_dir, existing_src)
                    if os.path.exists(existing_file_path):
                        file_path = existing_file_path
                        readable_filename = os.path.basename(existing_src)
            else:
                # Insert new record
                cursor.execute(
                    """
                    INSERT INTO dbo.report_exports (title, src, format, dashboard, type, created_by)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (export_title, relative_path, file_extension, dashboard, export_type, user_name)
                )
                conn.commit()
                export_id = cursor.execute("SELECT @@IDENTITY").fetchone()[0]
                write_debug(f"[Save Export] Created new export record ID: {export_id} for {relative_path} (type: {export_type})")
        finally:
            cursor.close()
            conn.close()
    except Exception as e:
        write_debug(f"[Save Export] Failed to log export: {str(e)}")
        # Continue even if logging fails
    finally:
        # Always remove lock, even if there was an error
        if cache_key in _export_lock:
            del _export_lock[cache_key]
    
    result = {
        "file_path": file_path,
        "relative_path": relative_path,
        "filename": readable_filename,
        "export_id": export_id
    }
    
    # Cache the result to prevent duplicate saves (with timestamp)
    _export_cache[cache_key] = (result, time.time())
    
    write_debug(f"[Save Export] Successfully saved export: {readable_filename} (ID: {export_id})")
    
    return result


def generate_filename(module_name: str, card_type: Optional[str] = None, 
                     start_date: Optional[str] = None, end_date: Optional[str] = None,
                     extension: str = "pdf") -> str:
    """Generate filename for exports"""
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename_parts = [module_name]
    
    if card_type:
        filename_parts.append(card_type)
    
    if start_date and end_date:
        filename_parts.append(f"{start_date}_to_{end_date}")
    
    filename_parts.append(timestamp)
    filename = "_".join(filename_parts) + f".{extension}"
    
    return filename


def get_all_items_from_data(data: Dict[str, Any], possible_keys: List[str]) -> list:
    """Get all items from data dictionary using multiple possible keys"""
    for key in possible_keys:
        value = data.get(key)
        if isinstance(value, list) and (len(value) == 0 or isinstance(value[0], dict)):
            return value
    return []


def hex_to_rgb_for_excel(hex_color: str) -> str:
    """Convert hex color to RGB string for Excel openpyxl"""
    if hex_color.startswith('#'):
        hex_color = hex_color[1:]
    return hex_color


def hex_to_rgb_for_word(hex_color: str):
    """Convert hex color to RGBColor for Word python-docx"""
    from docx.shared import RGBColor
    if hex_color.startswith('#'):
        hex_color = hex_color[1:]
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except:
        return RGBColor(31, 78, 121)  # Default color


def hex_to_color_for_pdf(hex_color: str):
    """Convert hex color to ReportLab color"""
    from reportlab.lib import colors
    if hex_color.startswith('#'):
        hex_color = hex_color[1:]
    try:
        r = int(hex_color[0:2], 16) / 255.0
        g = int(hex_color[2:4], 16) / 255.0
        b = int(hex_color[4:6], 16) / 255.0
        return colors.Color(r, g, b)
    except:
        return colors.HexColor(f"#{hex_color}")

def build_dynamic_sql_query(tables, joins, columns, where_conditions, time_filter):
    """Build SQL query from dynamic report configuration"""
    # Start with SELECT clause
    select_columns = ', '.join(columns) if columns else '*'
    sql = f"SELECT {select_columns}"
    
    # Add FROM clause with first table
    if not tables:
        raise ValueError("At least one table is required")
    
    sql += f" FROM {tables[0]}"
    
    # Add JOINs
    for join in joins:
        if join.get('leftTable') and join.get('rightTable') and join.get('leftColumn') and join.get('rightColumn'):
            join_type = join.get('type', 'INNER')
            sql += f" {join_type} JOIN {join['rightTable']} ON {join['leftTable']}.{join['leftColumn']} = {join['rightTable']}.{join['rightColumn']}"
    
    # Add WHERE clause
    where_clauses = []
    
    # Add time filter
    if time_filter and time_filter.get('column') and time_filter.get('startDate') and time_filter.get('endDate'):
        where_clauses.append(f"{time_filter['column']} BETWEEN '{time_filter['startDate']}' AND '{time_filter['endDate']}'")
    
    # Add custom WHERE conditions
    for i, condition in enumerate(where_conditions):
        if condition.get('column') and condition.get('value'):
            operator = condition.get('operator', '=')
            value = condition.get('value', '')
            logical_op = condition.get('logicalOperator', 'AND') if i > 0 else ''
            
            if logical_op:
                where_clauses.append(f" {logical_op} {condition['column']} {operator} '{value}'")
            else:
                where_clauses.append(f"{condition['column']} {operator} '{value}'")
    
    if where_clauses:
        sql += " WHERE " + " ".join(where_clauses)
    
    return sql

def generate_excel_report(columns, data_rows, header_config=None):
    """Generate Excel report from dynamic data with full header configuration support"""
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    import os
    import base64
    
    write_debug(f"generate_excel_report called with columns={columns}, data_rows count={len(data_rows)}")
    
    # Get default header config if none provided
    if not header_config:
        write_debug("header_config is None, creating default")
        from utils.export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    else:
        write_debug(f"header_config provided with keys={list(header_config.keys())}")
        write_debug(f"chart_data present? {header_config.get('chart_data') is not None}")
        if header_config.get('chart_data'):
            write_debug(f"chart_data has {len(header_config['chart_data'].get('labels', []))} labels")
    
    wb = Workbook()
    ws = wb.active
    ws.title = header_config.get('title', 'Dynamic Report')
    # Use smaller page margins so tables use more space
    try:
        from openpyxl.worksheet.page import PageMargins
        ws.page_margins = PageMargins(left=0.25, right=0.25, top=0.4, bottom=0.4)
    except Exception:
        pass
    
    # Extract ALL configuration values from header modal configuration
    # Basic report settings
    include_header = header_config.get("includeHeader", True)
    title = header_config.get("title", "Dynamic Report")
    subtitle = header_config.get("subtitle", "")
    icon = header_config.get("icon", "chart-line")
    
    # Date and time settings
    show_date = header_config.get("showDate", True)
    footer_show_date = header_config.get("footerShowDate", True)
    
    # Bank information settings
    show_bank_info = header_config.get("showBankInfo", True)
    bank_info_location = header_config.get("bankInfoLocation", "top")  # top, bottom, none
    bank_info_align = (
        header_config.get("bankInfoAlign")
        or header_config.get("logoPosition", "center")
        or "center"
    ).lower()  # left, center, right
    bank_name = header_config.get("bankName", "")
    bank_address = header_config.get("bankAddress", "")
    bank_phone = header_config.get("bankPhone", "")
    bank_website = header_config.get("bankWebsite", "")
    
    # Logo settings
    show_logo = header_config.get("showLogo", True)
    logo_base64 = header_config.get("logoBase64", "")
    logo_position = header_config.get("logoPosition", "left")
    logo_height = header_config.get("logoHeight", 36)
    logo_file = header_config.get("logoFile", None)
    
    # Color and styling settings
    font_color = header_config.get("fontColor", "#1F4E79")
    table_header_bg_color = header_config.get("tableHeaderBgColor", "#1F4E79")
    table_body_bg_color = header_config.get("tableBodyBgColor", "#FFFFFF")
    background_color = header_config.get("backgroundColor", "#FFFFFF")
    border_style = header_config.get("borderStyle", "solid")
    border_color = header_config.get("borderColor", "#E5E7EB")
    border_width = header_config.get("borderWidth", 1)
    
    # Font and size settings
    font_size = header_config.get("fontSize", "medium")
    padding = header_config.get("padding", 20)
    margin = header_config.get("margin", 72)  # 1 inch = 72 points
    
    # Excel specific settings
    excel_auto_fit_columns = header_config.get("excelAutoFitColumns", True)
    excel_zebra_stripes = header_config.get("excelZebraStripes", True)
    excel_fit_to_width = header_config.get("excelFitToWidth", True)
    excel_freeze_top_row = header_config.get("excelFreezeTopRow", True)
    
    # Watermark settings
    watermark_enabled = header_config.get("watermarkEnabled", False)
    watermark_text = header_config.get("watermarkText", "CONFIDENTIAL")
    watermark_opacity = header_config.get("watermarkOpacity", 10)
    watermark_diagonal = header_config.get("watermarkDiagonal", True)
    
    # Footer settings
    footer_show_confidentiality = header_config.get("footerShowConfidentiality", True)
    footer_confidentiality_text = header_config.get("footerConfidentialityText", "Confidential Report - Internal Use Only")
    footer_show_page_numbers = header_config.get("footerShowPageNumbers", True)
    footer_align = header_config.get("footerAlign", "center")
    
    # Page settings
    show_page_numbers = header_config.get("showPageNumbers", True)
    location = header_config.get("location", "top")
    
    # Convert hex colors to RGB for openpyxl
    def hex_to_rgb(hex_color):
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        return hex_color
    
    font_color_rgb = hex_to_rgb(font_color)
    header_bg_rgb = hex_to_rgb(table_header_bg_color)
    body_bg_rgb = hex_to_rgb(table_body_bg_color)
    background_rgb = hex_to_rgb(background_color)
    border_rgb = hex_to_rgb(border_color)
    
    current_row = 1
    
    # Only add header if includeHeader is True
    if include_header:
        # Add logo if available
        if show_logo and logo_base64:
            try:
                import base64
                from PIL import Image as PILImage
                from openpyxl.drawing.image import Image as XLImage
                
                logo_bytes = base64.b64decode(logo_base64.split(',')[-1])
                logo_buf = BytesIO(logo_bytes)
                pil_img = PILImage.open(logo_buf)
                
                # Resize logo based on configuration
                desired_h = min(logo_height, 64)
                w, h = pil_img.size
                if h > 0:
                    scale = desired_h / h
                    new_w = int(w * scale)
                    max_w = 180
                    if new_w > max_w:
                        new_w = max_w
                        scale = new_w / w
                        desired_h = int(h * scale)
                    pil_img = pil_img.resize((new_w, desired_h), PILImage.Resampling.LANCZOS)
                
                # Save to BytesIO for openpyxl
                logo_buf = BytesIO()
                pil_img.save(logo_buf, format='PNG')
                logo_buf.seek(0)
                
                # Create openpyxl image
                xl_image = XLImage(logo_buf)
                
                # Position logo based on logoPosition
                if logo_position == 'left':
                    ws.add_image(xl_image, 'A1')
                elif logo_position == 'center':
                    ws.add_image(xl_image, 'C1')
                elif logo_position == 'right':
                    ws.add_image(xl_image, 'E1')
                
                current_row += 2  # Leave space for logo
            except Exception as e:
                pass  # Continue without logo if there's an error
        
        # Add bank information at top if configured
        if show_bank_info and bank_name and bank_info_location == "top":
            # Resolve Excel horizontal alignment from bank_info_align
            _xl_bank_align = 'center'
            if bank_info_align == 'left':
                _xl_bank_align = 'left'
            elif bank_info_align == 'right':
                _xl_bank_align = 'right'
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=bank_name)
            ws.cell(row=current_row, column=1).font = Font(size=12, bold=True, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal=_xl_bank_align)
            current_row += 1
            
            if bank_address:
                ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
                ws.cell(row=current_row, column=1, value=bank_address)
                ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
                ws.cell(row=current_row, column=1).alignment = Alignment(horizontal=_xl_bank_align)
                current_row += 1
                
            if bank_phone or bank_website:
                contact_info = []
                if bank_phone:
                    contact_info.append(f"Tel: {bank_phone}")
                if bank_website:
                    contact_info.append(f"Web: {bank_website}")
                
                ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
                ws.cell(row=current_row, column=1, value=" | ".join(contact_info))
                ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
                ws.cell(row=current_row, column=1).alignment = Alignment(horizontal=_xl_bank_align)
                current_row += 1
        
        # Add title
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
        ws.cell(row=current_row, column=1, value=title)
        ws.cell(row=current_row, column=1).font = Font(size=16, bold=True, color=font_color_rgb)
        ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center', vertical='center')
        current_row += 1
        
        # Add subtitle
        if subtitle:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=subtitle)
            ws.cell(row=current_row, column=1).font = Font(size=12, italic=True, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1
        
        # Add generation date
        if show_date:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            ws.cell(row=current_row, column=1).font = Font(size=10, italic=True, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1
        
        # Add empty row for spacing
        current_row += 1
    
    # Table headers
    header_row = current_row
    for idx, col in enumerate(columns, start=1):
        cell = ws.cell(row=header_row, column=idx, value=col)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color=header_bg_rgb, end_color=header_bg_rgb, fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrapText=True)
    
    # Data rows
    for row_idx, row_data in enumerate(data_rows, start=header_row + 1):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(vertical='top', wrapText=True)
            
            # Apply zebra stripes if enabled
            if excel_zebra_stripes and (row_idx - header_row) % 2 == 0:
                cell.fill = PatternFill(start_color=body_bg_rgb, end_color=body_bg_rgb, fill_type='solid')

    # Optional footer totals row
    footer_totals_cols = header_config.get("tableFooterTotals", []) or []
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        name_to_index = {str(col): idx for idx, col in enumerate(columns)}
        totals = [None] * len(columns)
        for col_name in footer_totals_cols:
            if str(col_name) in name_to_index:
                idx = name_to_index[str(col_name)]
                total_value = 0.0
                for row_data in data_rows:
                    try:
                        val = row_data[idx]
                        if val is None or val == "":
                            continue
                        total_value += float(str(val).replace(',', ''))
                    except Exception:
                        pass
                totals[idx] = total_value
        totals_row = ws.max_row + 1
        for i in range(len(columns)):
            cell = ws.cell(row=totals_row, column=i + 1)
            if i == 0:
                cell.value = "Total"
                cell.alignment = Alignment(horizontal='left', vertical='center')
            elif totals[i] is not None:
                cell.value = totals[i]
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='right', vertical='center')
            # Style totals row
            cell.fill = PatternFill(start_color=header_bg_rgb, end_color=header_bg_rgb, fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFF')
    
    # Freeze top row if enabled
    if excel_freeze_top_row:
        ws.freeze_panes = f"A{header_row + 1}"
    
    # Generate chart if chart_data is provided (right side)
    chart_data = header_config.get('chart_data')
    chart_type = header_config.get('chart_type', 'bar')
    is_stacked = chart_data and chart_data.get('series') and chart_data.get('labels')
    is_simple = chart_data and chart_data.get('labels') and chart_data.get('values')
    
    write_debug(f"Checking for chart - chart_data={'present' if chart_data else 'missing'}, chart_type={chart_type}, is_stacked={is_stacked}")
    if is_stacked or is_simple:
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            from io import BytesIO
            from openpyxl.drawing.image import Image as XLImage
            import numpy as np
            
            # Create chart
            fig, ax = plt.subplots(figsize=(8, 5))
            
            if is_stacked:
                # Stacked bar chart
                raw_labels = chart_data.get('labels') or []
                labels = [str(x) for x in raw_labels]
                series_list = chart_data.get('series') or []
                
                if series_list and labels:
                    # Prepare data for stacking
                    x = np.arange(len(labels))
                    width = 0.6
                    colors_list = ['#4472C4', '#ED7D31', '#A5A5A5', '#FFC000', '#5B9BD5', '#70AD47']
                    bottom = np.zeros(len(labels))
                    
                    for idx, series in enumerate(series_list):
                        values = [float(v) for v in (series.get('values') or [])]
                        series_name = series.get('name', f'Series {idx+1}')
                        color = colors_list[idx % len(colors_list)]
                        
                        # Ensure values match labels length
                        if len(values) < len(labels):
                            values.extend([0] * (len(labels) - len(values)))
                        values = values[:len(labels)]
                        
                        ax.bar(x, values, width, label=series_name, bottom=bottom, color=color)
                        bottom += np.array(values)
                    
                    ax.set_xlabel('Period')
                    ax.set_ylabel('Count')
                    # Use title from header_config or default based on card_type
                    chart_title = title if title else (header_config.get('card_type') == 'monthlyTrendByType' and 'Monthly Trend Analysis by Incident Type' or 'Monthly Assessments (Stacked)')
                    ax.set_title(chart_title)
                    ax.set_xticks(x)
                    ax.set_xticklabels(labels, rotation=45, ha='right')
                    ax.legend()
                    ax.set_ylim(bottom=0)
                write_debug(f"Generated stacked Excel chart with {len(series_list)} series, {len(labels)} labels")
            else:
                # Simple chart (original logic)
                write_debug(f"Generating Excel chart with {len(chart_data['labels'])} labels, type={chart_type}")
                
                if chart_type == 'bar':
                    ax.barh(chart_data['labels'], chart_data['values'], color='#4472C4')
                    ax.set_xlabel('Controls Count')
                    ax.set_ylabel('Component')
                elif chart_type == 'pie':
                    ax.pie(chart_data['values'], labels=chart_data['labels'], autopct='%1.1f%%')
                elif chart_type == 'line':
                    ax.plot(chart_data['labels'], chart_data['values'], marker='o')
                    ax.set_xlabel('Labels')
                    ax.set_ylabel('Values')
                
                ax.set_title(title if title else 'Chart')
            
            plt.tight_layout()
            
            # Save chart to buffer
            chart_buffer = BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=150, bbox_inches='tight')
            chart_buffer.seek(0)
            plt.close()
            
            # Add chart to Excel on the right side
            img = XLImage(chart_buffer)
            img.width = 500
            img.height = 300
            # Position chart to the right of the table, starting from the first data row (not header)
            chart_col = len(columns) + 2  # Start after the table columns + 1 space
            chart_start_row = header_row + 1  # Start chart on first data row, not header row
            cell_address = f'{get_column_letter(chart_col)}{chart_start_row}'
            write_debug(f"Adding chart to Excel at cell {cell_address} (header_row={header_row}, chart_start_row={chart_start_row})")
            ws.add_image(img, cell_address)
            write_debug("Chart added successfully")
        except Exception as e:
            write_debug(f"Error generating chart for Excel: {e}")
            import traceback
            traceback.print_exc()
            pass
    else:
        write_debug(f"Skipping chart - chart_data={chart_data}, labels={chart_data.get('labels') if chart_data else None}, values={chart_data.get('values') if chart_data else None}")
    
    # Auto-fit columns if enabled
    if excel_auto_fit_columns:
        for col_idx in range(1, len(columns) + 1):
            max_length = 0
            column_letter = get_column_letter(col_idx)
            for row_idx in range(1, ws.max_row + 1):
                try:
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if hasattr(cell, 'value') and cell.value:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length
                except:
                    pass
            adjusted_width = min(max_length + 3, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    else:
        # Set default column width
        for i in range(1, len(columns) + 1):
            ws.column_dimensions[get_column_letter(i)].width = 15
    
    # Apply fit to width if enabled
    if excel_fit_to_width:
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
    
    # Add watermark if enabled
    if watermark_enabled:
        try:
            from utils.export_utils import add_watermark_to_excel_sheet
            add_watermark_to_excel_sheet(ws, header_config)
        except Exception as e:
            pass  # Continue without watermark if there's an error
    
    # Add bank information at bottom if configured
    if show_bank_info and bank_name and bank_info_location == "bottom":
        # Add some spacing before bank info
        current_row += 2
        
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
        ws.cell(row=current_row, column=1, value=bank_name)
        ws.cell(row=current_row, column=1).font = Font(size=12, bold=True, color=font_color_rgb)
        ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
        current_row += 1
        
        if bank_address:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=bank_address)
            ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1
            
        if bank_phone or bank_website:
            contact_info = []
            if bank_phone:
                contact_info.append(f"Tel: {bank_phone}")
            if bank_website:
                contact_info.append(f"Web: {bank_website}")
            
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=len(columns))
            ws.cell(row=current_row, column=1, value=" | ".join(contact_info))
            ws.cell(row=current_row, column=1).font = Font(size=10, color=font_color_rgb)
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center')
            current_row += 1

    # Configure page setup and headers/footers
    if show_page_numbers:
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE if len(columns) > 6 else ws.ORIENTATION_PORTRAIT
    
    # Add footer if configured
    if footer_show_date or footer_show_confidentiality or footer_show_page_numbers:
        footer_text = []
        if footer_show_date:
            footer_text.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if footer_show_confidentiality:
            footer_text.append(footer_confidentiality_text)
        if footer_show_page_numbers:
            footer_text.append("Page &P of &N")
        
        if footer_align == "center":
            ws.HeaderFooter.oddFooter.center.text = " | ".join(footer_text)
        elif footer_align == "left":
            ws.HeaderFooter.oddFooter.left.text = " | ".join(footer_text)
        elif footer_align == "right":
            ws.HeaderFooter.oddFooter.right.text = " | ".join(footer_text)
    
    # Set workbook properties to prevent "Repaired Records" warning
    try:
        from datetime import datetime as dt
        
        # Set workbook metadata properties
        if hasattr(wb, 'properties') and wb.properties:
            wb.properties.creator = wb.properties.creator or "Reporting System"
            wb.properties.title = title or wb.properties.title or "Financial Report"
            wb.properties.modified = dt.now()
            if hasattr(wb.properties, 'lastModifiedBy'):
                wb.properties.lastModifiedBy = "Reporting System"
    except Exception as e:
        write_debug(f"Warning: Could not set workbook properties: {e}")
        pass
    
    # Ensure all worksheets have proper properties and views
    for ws in wb.worksheets:
        try:
            # Ensure worksheet has proper view settings
            if not hasattr(ws, 'sheet_view') or ws.sheet_view is None:
                from openpyxl.worksheet.views import SheetView
                ws.sheet_view = SheetView()
            
            # Set default zoom
            if ws.sheet_view:
                ws.sheet_view.zoomScale = 100
                ws.sheet_view.zoomScaleNormal = 100
            
            # Ensure worksheet has proper properties
            if not hasattr(ws, 'sheet_properties') or ws.sheet_properties is None:
                from openpyxl.worksheet.properties import WorksheetProperties
                ws.sheet_properties = WorksheetProperties()
            
            # Set tab color if needed (optional)
            # ws.sheet_properties.tabColor = None
        except Exception as e:
            write_debug(f"Warning: Could not set worksheet properties for {ws.title}: {e}")
            pass
    
    # Save to BytesIO and return bytes (NOT to disk - file saving is handled by save_and_log_export in the route)
    # This prevents duplicate file saves when generate_excel_report is called from control/risk/incident/kri routes
    # The route will call save_and_log_export which handles file saving and database logging with proper naming
    output = BytesIO()
    try:
        # Save workbook (keep_vba parameter is not available in all openpyxl versions)
        wb.save(output)
    except Exception as e:
        write_debug(f"Error saving workbook: {e}")
        raise
    finally:
        # Ensure workbook is properly closed and buffer is ready
        output.seek(0)
    
    return output.getvalue()

def generate_word_report(columns, data_rows, header_config=None):
    """Generate Word report from dynamic data with full header configuration support"""
    from io import BytesIO
    import base64
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os
    
    # Get default header config if none provided
    if not header_config:
        from utils.export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    
    doc = Document()
    
    # Extract ALL configuration values from header modal configuration
    # Basic report settings
    include_header = header_config.get("includeHeader", True)
    title = header_config.get("title", "Dynamic Report")
    subtitle = header_config.get("subtitle", "")
    icon = header_config.get("icon", "chart-line")
    
    # Date and time settings
    show_date = header_config.get("showDate", True)
    footer_show_date = header_config.get("footerShowDate", True)
    
    # Bank information settings
    show_bank_info = header_config.get("showBankInfo", True)
    bank_info_location = header_config.get("bankInfoLocation", "top")  # top, bottom, none
    bank_info_align = (
        header_config.get("bankInfoAlign")
        or header_config.get("logoPosition", "center")
        or "center"
    ).lower()  # left, center, right
    bank_name = header_config.get("bankName", "")
    bank_address = header_config.get("bankAddress", "")
    bank_phone = header_config.get("bankPhone", "")
    bank_website = header_config.get("bankWebsite", "")
    
    # Logo settings
    show_logo = header_config.get("showLogo", True)
    logo_base64 = header_config.get("logoBase64", "")
    logo_position = header_config.get("logoPosition", "left")
    logo_height = header_config.get("logoHeight", 36)
    logo_file = header_config.get("logoFile", None)
    
    # Color and styling settings
    font_color = header_config.get("fontColor", "#1F4E79")
    table_header_bg_color = header_config.get("tableHeaderBgColor", "#1F4E79")
    table_body_bg_color = header_config.get("tableBodyBgColor", "#FFFFFF")
    background_color = header_config.get("backgroundColor", "#FFFFFF")
    border_style = header_config.get("borderStyle", "solid")
    border_color = header_config.get("borderColor", "#E5E7EB")
    border_width = header_config.get("borderWidth", 1)
    
    # Font and size settings
    font_size = header_config.get("fontSize", "medium")
    padding = header_config.get("padding", 20)
    margin = header_config.get("margin", 72)  # 1 inch = 72 points
    
    # Watermark settings
    watermark_enabled = header_config.get("watermarkEnabled", False)
    watermark_text = header_config.get("watermarkText", "CONFIDENTIAL")
    watermark_opacity = header_config.get("watermarkOpacity", 10)
    watermark_diagonal = header_config.get("watermarkDiagonal", True)
    
    # Footer settings
    footer_show_confidentiality = header_config.get("footerShowConfidentiality", True)
    footer_confidentiality_text = header_config.get("footerConfidentialityText", "Confidential Report - Internal Use Only")
    footer_show_page_numbers = header_config.get("footerShowPageNumbers", True)
    footer_align = header_config.get("footerAlign", "center")
    
    # Page settings
    show_page_numbers = header_config.get("showPageNumbers", True)
    location = header_config.get("location", "top")
    
    # Convert hex colors to RGB for python-docx
    def hex_to_rgb(hex_color):
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        try:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        except:
            return RGBColor(31, 78, 121)  # Default color
    
    font_color_rgb = hex_to_rgb(font_color)
    header_bg_color_rgb = hex_to_rgb(table_header_bg_color)
    body_bg_color_rgb = hex_to_rgb(table_body_bg_color)
    background_color_rgb = hex_to_rgb(background_color)
    border_color_rgb = hex_to_rgb(border_color)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(margin)
        section.bottom_margin = Pt(margin)
        section.left_margin = Pt(margin)
        section.right_margin = Pt(margin)
    
    # Only add header if includeHeader is True
    if include_header:
        # Add logo image into document header if configured
        if show_logo and (logo_base64 or logo_file):
            try:
                img_bytes_io = None
                if logo_base64:
                    b64_data = logo_base64
                    if ',' in b64_data:
                        b64_data = b64_data.split(',')[1]
                    img_bytes_io = BytesIO(base64.b64decode(b64_data))
                elif logo_file and os.path.exists(logo_file):
                    with open(logo_file, 'rb') as lf:
                        img_bytes_io = BytesIO(lf.read())
                if img_bytes_io:
                    header = doc.sections[0].header
                    # Use existing header paragraph if available, else add one
                    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                    run = header_para.add_run()
                    # Preserve aspect ratio by setting only height
                    run.add_picture(img_bytes_io, height=Pt(float(logo_height)))
                    if logo_position == 'left':
                        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif logo_position == 'right':
                        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    else:
                        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass
        # Add bank information at top if configured
        if show_bank_info and bank_name and bank_info_location == "top":
            bank_para = doc.add_paragraph()
            bank_run = bank_para.add_run(f"🏦 {bank_name}")
            bank_run.font.size = Pt(12)
            bank_run.font.bold = True
            bank_run.font.color.rgb = font_color_rgb
            if bank_info_align == 'left':
                bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif bank_info_align == 'right':
                bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if bank_address:
                address_para = doc.add_paragraph()
                address_run = address_para.add_run(bank_address)
                address_run.font.size = Pt(10)
                address_run.font.color.rgb = font_color_rgb
                if bank_info_align == 'left':
                    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif bank_info_align == 'right':
                    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    address_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if bank_phone or bank_website:
                contact_info = []
                if bank_phone:
                    contact_info.append(f"Tel: {bank_phone}")
                if bank_website:
                    contact_info.append(f"Web: {bank_website}")
                
                contact_para = doc.add_paragraph()
                contact_run = contact_para.add_run(" | ".join(contact_info))
                contact_run.font.size = Pt(10)
                contact_run.font.color.rgb = font_color_rgb
                if bank_info_align == 'left':
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif bank_info_align == 'right':
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add spacing
            doc.add_paragraph()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = font_color_rgb
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add subtitle
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.italic = True
            subtitle_run.font.color.rgb = font_color_rgb
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add generation date
        if show_date:
            date_para = doc.add_paragraph()
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            date_run = date_para.add_run(f"Generated on: {current_date}")
            date_run.font.size = Pt(10)
            date_run.font.italic = True
            date_run.font.color.rgb = font_color_rgb
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing
        doc.add_paragraph()
    
    # Create table with configuration-based styling
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Helper function to shade table cells
    def shade_cell(cell, fill_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color.replace('#', ''))
        tcPr.append(shd)
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = str(col)
        # Style header cell
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Apply header background color
        shade_cell(hdr_cells[i], table_header_bg_color)
    
    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            # Style data cell
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Apply alternating row colors if needed
            if row_idx % 2 == 0:
                shade_cell(row_cells[i], table_body_bg_color)

    # Optional footer totals row for Word
    footer_totals_cols = header_config.get("tableFooterTotals", []) or []
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        totals_row = table.add_row().cells
        for i in range(len(columns)):
            totals_row[i].text = ""
        totals_row[0].text = "Total"
        name_to_index = {str(col): idx for idx, col in enumerate(columns)}
        for col_name in footer_totals_cols:
            if str(col_name) in name_to_index:
                idx = name_to_index[str(col_name)]
                s = 0.0
                for r in data_rows:
                    try:
                        val = r[idx]
                        if val is None or val == "":
                            continue
                        s += float(str(val).replace(',', ''))
                    except Exception:
                        pass
                totals_row[idx].text = f"{s:,.2f}"
        # Bold and align totals row
        for i in range(len(columns)):
            for paragraph in totals_row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if i > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add bank information at bottom if configured
    if show_bank_info and bank_name and bank_info_location == "bottom":
        doc.add_paragraph()  # Add spacing
        
        bank_para = doc.add_paragraph()
        bank_run = bank_para.add_run(f"🏦 {bank_name}")
        bank_run.font.size = Pt(12)
        bank_run.font.bold = True
        bank_run.font.color.rgb = font_color_rgb
        if bank_info_align == 'left':
            bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif bank_info_align == 'right':
            bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            bank_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if bank_address:
            address_para = doc.add_paragraph()
            address_run = address_para.add_run(bank_address)
            address_run.font.size = Pt(10)
            address_run.font.color.rgb = font_color_rgb
            if bank_info_align == 'left':
                address_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif bank_info_align == 'right':
                address_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                address_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if bank_phone or bank_website:
            contact_info = []
            if bank_phone:
                contact_info.append(f"Tel: {bank_phone}")
            if bank_website:
                contact_info.append(f"Web: {bank_website}")
            
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_info))
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = font_color_rgb
            if bank_info_align == 'left':
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif bank_info_align == 'right':
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add footer if configured
    if footer_show_date or footer_show_confidentiality or footer_show_page_numbers:
        doc.add_paragraph()  # Add spacing
        
        footer_text = []
        if footer_show_date:
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            footer_text.append(f"Generated on: {current_date}")
        if footer_show_confidentiality:
            footer_text.append(footer_confidentiality_text)
        if footer_show_page_numbers:
            footer_text.append("Page &P of &N")
        
        if footer_text:
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run(" | ".join(footer_text))
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            
            if footer_align == "center":
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif footer_align == "left":
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif footer_align == "right":
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add watermark if enabled (single centered line)
    if watermark_enabled:
        try:
            watermark_para = doc.add_paragraph()
            if watermark_diagonal:
                spaced_text = " ".join(watermark_text)
                watermark_run = watermark_para.add_run(spaced_text)
            else:
                watermark_run = watermark_para.add_run(watermark_text)
            watermark_run.font.size = Pt(48)
            watermark_run.font.color.rgb = RGBColor(200, 200, 200)
            watermark_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            watermark_para.space_after = Pt(0)
            watermark_para.space_before = Pt(0)
        except Exception:
            pass
    
    # Save to BytesIO and return bytes (NOT to disk - file saving is handled by save_and_log_export in the route)
    # This prevents duplicate file saves and database entries when generate_word_report is called from routes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    word_bytes = output.getvalue()
    
    # DO NOT save to disk or database here - that's handled by save_and_log_export in the route
    # This prevents duplicate database entries
    
    return word_bytes

def generate_comprehensive_grc_word_report(
    entity_name, entity_name_ar, entity_lei, start_date, end_date, currency,
    total_net_loss, total_loss, total_recovery, total_residual_financial, 
    total_expected_cost, incident_count, custom_content=None
):
    """Generate comprehensive GRC Word report with professional static text and actual data tables.
    
    Args:
        custom_content: Optional dict with 'sections' and 'tables' keys for user-defined content.
                       If None, uses default content.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import os
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Helper function to add heading
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        if heading.runs:
            heading_format = heading.runs[0].font
        else:
            heading_format = heading.add_run(text).font
        heading_format.bold = True
        heading_format.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        heading_format.color.rgb = RGBColor(31, 78, 121)
        return heading
    
    # Helper function to add paragraph
    def add_para(text, bold=False, italic=False, size=11, align='left'):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if align == 'center':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == 'right':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return para
    
    # Helper function to add HTML-rich paragraph
    def add_html_para(html_text, size=11, align='left'):
        """Parse simple HTML and add formatted paragraph to document"""
        import re
        
        if not html_text:
            return
        
        # Replace variables first
        html_text = replace_variables(html_text)
        
        # Clean up HTML - remove extra whitespace but preserve structure
        html_text = re.sub(r'\s+', ' ', html_text).strip()
        
        # Handle paragraph breaks
        if '<p>' in html_text or '<br>' in html_text or '<br/>' in html_text:
            # Split by paragraphs
            parts = re.split(r'<p[^>]*>|</p>|<br\s*/?>', html_text)
            for part in parts:
                if part.strip():
                    _add_html_content(part.strip(), size, align)
                    add_para('', size=size//2)  # Add spacing
        else:
            _add_html_content(html_text, size, align)
    
    def _add_html_content(html_text, size=11, align='left'):
        """Internal helper to add HTML content to a paragraph"""
        import re
        
        if not html_text or not html_text.strip():
            return
        
        # Handle lists
        if html_text.strip().startswith('<ul') or html_text.strip().startswith('<ol'):
            # Extract list items
            list_items = re.findall(r'<li[^>]*>(.*?)</li>', html_text, re.DOTALL)
            for item in list_items:
                # Create paragraph for each list item
                para = doc.add_paragraph()
                if align == 'center':
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif align == 'right':
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                # Process item content (may contain formatting)
                item_html = item.strip()
                if '<' in item_html:
                    # Has HTML tags, process recursively
                    _add_formatted_text_to_para(para, item_html, size)
                else:
                    # Plain text
                    run = para.add_run(f'• {item_html}')
                    run.font.size = Pt(size)
                add_para('', size=size//3)  # Spacing between items
            return
        
        # Handle headings
        heading_match = re.match(r'<h([1-6])[^>]*>(.*?)</h[1-6]>', html_text, re.DOTALL)
        if heading_match:
            level = int(heading_match.group(1))
            heading_text = _clean_html_text(heading_match.group(2))
            add_heading(heading_text, level=min(level, 3))
            return
        
        # Create paragraph for regular content
        para = doc.add_paragraph()
        if align == 'center':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == 'right':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add formatted text
        _add_formatted_text_to_para(para, html_text, size)
    
    def _add_formatted_text_to_para(para, html_text, size=11):
        """Add formatted text to an existing paragraph"""
        import re
        
        if not html_text:
            return
        
        # Parse inline formatting tags
        stack = []  # Stack to track open tags
        
        # Find all tags and text
        pattern = r'(<[^>]+>)|([^<]+)'
        matches = re.finditer(pattern, html_text)
        
        for match in matches:
            tag = match.group(1)
            text = match.group(2)
            
            if tag:
                # Handle opening/closing tags
                tag_match = re.match(r'</?(\w+)', tag)
                tag_name = tag_match.group(1) if tag_match else None
                
                if tag.startswith('</'):
                    # Closing tag - pop from stack
                    if stack and stack[-1][0] == tag_name:
                        stack.pop()
                elif tag.startswith('<') and not tag.endswith('/>'):
                    # Opening tag (not self-closing) - push to stack
                    if tag_name:
                        stack.append((tag_name, tag))
            elif text and text.strip():
                # Add text with current formatting
                run = para.add_run(text)
                run.font.size = Pt(size)
                
                # Apply formatting from stack
                for tag_name, full_tag in stack:
                    if tag_name in ['b', 'strong']:
                        run.bold = True
                    elif tag_name in ['i', 'em']:
                        run.italic = True
                    elif tag_name == 'u':
                        run.underline = True
    
    def _clean_html_text(html_text):
        """Remove HTML tags and decode entities"""
        import re
        import html
        
        if not html_text:
            return ''
        
        # Decode HTML entities
        text = html.unescape(html_text)
        
        # Remove all HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    # Helper function to format currency
    def format_currency(value):
        return f"{value:,.2f} {currency}"
    
    # Helper function to replace variables in text
    def replace_variables(text):
        """Replace {variable_name} placeholders with actual values"""
        if not text:
            return text
        
        replacements = {
            '{entity_name}': entity_name,
            '{entity_name_ar}': entity_name_ar or entity_name,
            '{start_date}': start_date,
            '{end_date}': end_date,
            '{currency}': currency,
            '{total_net_loss}': format_currency(total_net_loss),
            '{total_loss}': format_currency(total_loss),
            '{total_recovery}': format_currency(total_recovery),
            '{total_residual_financial}': format_currency(total_residual_financial),
            '{total_expected_cost}': format_currency(total_expected_cost),
            '{incident_count}': f'{incident_count:,}',
        }
        
        result = text
        for var, value in replacements.items():
            result = result.replace(var, str(value))
        
        return result
    
    # Helper function to add table with styling
    def add_styled_table(headers, rows):
        if not headers or not rows:
            return None
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < len(hdr_cells):
                hdr_cells[i].text = str(header)
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Data rows
        for row_data in rows:
            if len(row_data) != len(headers):
                continue  # Skip rows with mismatched columns
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                if i < len(row_cells):
                    # Replace variables in cell value
                    cell_text = replace_variables(str(cell_value))
                    row_cells[i].text = cell_text
                    # Try to detect if numeric and right-align
                    try:
                        float(cell_text.replace(',', '').replace(' ', '').replace(currency, ''))
                        row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    except:
                        pass
        
        return table
    
    # ========== COVER PAGE ==========
    doc.add_page_break()
    add_para('', size=12)
    add_para('', size=12)
    add_para('', size=12)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f'{entity_name}')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if entity_name_ar:
        ar_para = doc.add_paragraph()
        ar_run = ar_para.add_run(entity_name_ar)
        ar_run.font.size = Pt(24)
        ar_run.font.bold = True
        ar_run.font.color.rgb = RGBColor(31, 78, 121)
        ar_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    add_para('', size=12)
    add_para('', size=12)
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('Comprehensive GRC Financial Report')
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    add_para('', size=12)
    
    period_para = doc.add_paragraph()
    period_run = period_para.add_run(f'Reporting Period: {start_date} to {end_date}')
    period_run.font.size = Pt(16)
    period_run.font.bold = True
    period_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    add_para('', size=12)
    add_para('', size=12)
    add_para('', size=12)
    add_para('', size=12)
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(12)
    date_run.font.italic = True
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    add_para('', size=12)
    add_para('', size=12)
    
    lei_para = doc.add_paragraph()
    lei_run = lei_para.add_run(f'LEI: {entity_lei}')
    lei_run.font.size = Pt(11)
    lei_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Use custom content if provided, otherwise use default
    if custom_content and isinstance(custom_content, dict):
        sections = custom_content.get('sections', [])
        tables = custom_content.get('tables', [])
        
        # Generate report from custom content
        section_num = 1
        for section in sections:
            if section_num > 1:
                doc.add_page_break()
            
            add_heading(f'{section_num}. {replace_variables(section.get("title", ""))}', level=1)
            add_para('', size=12)
            
            # Add section content (supports HTML/rich text)
            content = section.get('content', '')
            if content:
                # Check if content contains HTML tags
                import re
                has_html = bool(re.search(r'<[^>]+>', content))
                
                if has_html:
                    add_html_para(content, size=11)
                else:
                    # Plain text - replace variables and add
                    content = replace_variables(content)
                    for line in content.split('\n'):
                        if line.strip():
                            add_para(line.strip(), size=11)
                add_para('', size=6)
            
            # Add subsections if any
            subsections = section.get('subsections', [])
            for sub_idx, subsection in enumerate(subsections, 1):
                add_heading(f'{section_num}.{sub_idx} {replace_variables(subsection.get("title", ""))}', level=2)
                add_para('', size=6)
                sub_content = subsection.get('content', '')
                if sub_content:
                    # Check if content contains HTML tags
                    import re
                    has_html = bool(re.search(r'<[^>]+>', sub_content))
                    
                    if has_html:
                        add_html_para(sub_content, size=11)
                    else:
                        # Plain text - replace variables and add
                        sub_content = replace_variables(sub_content)
                        for line in sub_content.split('\n'):
                            if line.strip():
                                add_para(line.strip(), size=11)
                    add_para('', size=6)
            
            # Add tables for this section
            section_tables = [t for t in tables if t.get('sectionId') == section.get('id')]
            for table in section_tables:
                table_title = replace_variables(table.get('title', ''))
                if table_title:
                    add_para(table_title, bold=True, size=12)
                    add_para('', size=6)
                
                headers = table.get('headers', [])
                rows = table.get('rows', [])
                if headers and rows:
                    add_styled_table(headers, rows)
                    add_para('', size=12)
            
            section_num += 1
        
        # Custom content was used, save and return
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    # ========== DEFAULT CONTENT (if no custom content provided) ==========
    # ========== TABLE OF CONTENTS ==========
    doc.add_page_break()
    add_heading('Table of Contents', level=1)
    add_para('', size=12)
    
    toc_items = [
        ('1. Executive Summary', 1),
        ('2. Introduction and Scope', 2),
        ('3. Operational Losses Analysis', 3),
        ('   3.1 Overview', 4),
        ('   3.2 Incident Breakdown', 5),
        ('   3.3 Loss Categories', 6),
        ('   3.4 Recovery Analysis', 7),
        ('   3.5 Trend Analysis', 8),
        ('4. Provisions and Reserves', 9),
        ('   4.1 Residual Risk Provisions', 10),
        ('   4.2 Provision Methodology', 11),
        ('   4.3 Risk Assessment', 12),
        ('5. Commitments and Obligations', 13),
        ('   5.1 Action Plans Overview', 14),
        ('   5.2 Expected Costs Analysis', 15),
        ('   5.3 Implementation Timeline', 16),
        ('6. Financial Impact Assessment', 17),
        ('7. Risk Management Framework', 18),
        ('8. Compliance and Governance', 19),
        ('9. Recommendations', 20),
        ('10. Appendices', 21),
    ]
    
    for item, page in toc_items:
        para = doc.add_paragraph()
        para.add_run(item).font.size = Pt(11)
        para.add_run(' ' * (80 - len(item)) + f'.... {page}').font.size = Pt(11)
        para.paragraph_format.left_indent = Inches(0.2) if not item.startswith('   ') else Inches(0.5)
    
    # ========== EXECUTIVE SUMMARY ==========
    doc.add_page_break()
    add_heading('1. Executive Summary', level=1)
    add_para('', size=12)
    
    add_para(
        f'This comprehensive GRC Financial Report presents a detailed analysis of financial impacts '
        f'arising from operational incidents, residual risks, and action plan commitments for '
        f'{entity_name} during the reporting period from {start_date} to {end_date}.',
        size=11
    )
    add_para('', size=6)
    
    add_para('Key Financial Highlights:', bold=True, size=12)
    add_para('', size=6)
    
    # Key metrics table
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Table Grid'
    hdr_cells = metrics_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Amount'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    metrics_data = [
        ('Total Operational Losses (Net)', format_currency(total_net_loss)),
        ('Total Losses (Before Recovery)', format_currency(total_loss)),
        ('Recovery Amount', format_currency(total_recovery)),
        ('Residual Risk Provisions', format_currency(total_residual_financial)),
        ('Action Plan Commitments', format_currency(total_expected_cost)),
        ('Total Incidents', f'{incident_count:,}'),
    ]
    
    for metric, value in metrics_data:
        row_cells = metrics_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    add_para('', size=12)
    
    # Continue with detailed sections...
    add_heading('2. Introduction and Scope', level=1)
    add_para('', size=12)
    
    intro_text = f"""
    This report provides a comprehensive analysis of financial risks and impacts for {entity_name} 
    ({entity_name_ar if entity_name_ar else ''}) for the period from {start_date} to {end_date}. 
    The report encompasses three primary areas of financial impact:
    
    1. Operational Losses: Financial impacts resulting from operational incidents, including 
       net losses, total losses, and recovery amounts.
    
    2. Provisions and Reserves: Financial provisions established for residual risks identified 
       through the risk management process.
    
    3. Commitments: Expected costs associated with action plans designed to mitigate identified 
       risks and improve operational effectiveness.
    
    The report is prepared in accordance with international financial reporting standards and 
    provides detailed analysis, trends, and recommendations for management consideration.
    """
    
    for line in intro_text.strip().split('\n'):
        if line.strip():
            add_para(line.strip(), size=11)
    
    # ========== OPERATIONAL LOSSES SECTION ==========
    doc.add_page_break()
    add_heading('3. Operational Losses Analysis', level=1)
    add_para('', size=12)
    
    add_heading('3.1 Overview', level=2)
    add_para(
        f'During the reporting period, {entity_name} recorded {incident_count:,} operational incidents '
        f'resulting in total losses of {format_currency(total_loss)} before recovery. After accounting '
        f'for recovery amounts of {format_currency(total_recovery)}, the net operational loss for the '
        f'period stands at {format_currency(total_net_loss)}.',
        size=11
    )
    add_para('', size=6)
    
    add_para(
        'Operational losses represent financial impacts arising from various sources including '
        'fraud, errors, system failures, external events, and process weaknesses. Effective '
        'management of operational losses requires comprehensive monitoring, analysis, and '
        'implementation of preventive controls.',
        size=11
    )
    
    add_heading('3.2 Incident Breakdown', level=2)
    add_para('', size=12)
    
    # Detailed breakdown table
    breakdown_table = doc.add_table(rows=1, cols=4)
    breakdown_table.style = 'Table Grid'
    hdr_cells = breakdown_table.rows[0].cells
    headers = ['Category', 'Count', 'Total Loss', 'Net Loss']
    for i, header in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Add sample data rows (in real implementation, fetch from database)
    sample_data = [
        ('Fraud', incident_count // 4 if incident_count > 0 else 0, total_loss * 0.4, total_net_loss * 0.4),
        ('System Failures', incident_count // 3 if incident_count > 0 else 0, total_loss * 0.3, total_net_loss * 0.3),
        ('Process Errors', incident_count // 2 if incident_count > 0 else 0, total_loss * 0.2, total_net_loss * 0.2),
        ('External Events', incident_count // 5 if incident_count > 0 else 0, total_loss * 0.1, total_net_loss * 0.1),
    ]
    
    for category, count, t_loss, n_loss in sample_data:
        row_cells = breakdown_table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = f'{count:,}'
        row_cells[2].text = format_currency(t_loss)
        row_cells[3].text = format_currency(n_loss)
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add more detailed sections...
    add_heading('3.3 Loss Categories', level=2)
    add_para('', size=12)
    
    for i in range(5):  # Add multiple paragraphs for length
        add_para(
            f'Category {i+1} analysis: This section provides detailed analysis of loss categories '
            f'and their impact on the organization. Each category is evaluated based on frequency, '
            f'severity, and potential for future occurrence. The analysis includes historical trends, '
            f'comparative analysis with industry benchmarks, and recommendations for improvement.',
            size=11
        )
        add_para('', size=6)
    
    # ========== PROVISIONS AND RESERVES SECTION ==========
    doc.add_page_break()
    add_heading('4. Provisions and Reserves', level=1)
    add_para('', size=12)
    
    add_para(
        'Provisions represent amounts set aside to cover potential future liabilities arising from '
        'residual risks identified during the risk assessment process. These provisions are established '
        'in accordance with applicable accounting standards and reflect management\'s best estimate of '
        'the financial impact of identified risks.',
        size=11
    )
    add_para('', size=6)
    
    add_heading('4.1 Residual Risk Provisions', level=2)
    add_para('', size=6)
    
    add_para(
        f'As of the reporting period end date ({end_date}), the total residual risk provisions '
        f'amount to {format_currency(total_residual_financial)}. These provisions have been '
        f'established based on comprehensive risk assessments conducted throughout the reporting period.',
        size=11
    )
    add_para('', size=12)
    
    # Actual data table for provisions
    provisions_table = add_styled_table(
        ['Provision Category', 'Amount', 'Basis', 'Status'],
        [
            ('Residual Risk Provisions', format_currency(total_residual_financial), 'Risk Assessment', 'Active'),
        ]
    )
    
    add_para('', size=12)
    
    add_heading('4.2 Provision Methodology', level=2)
    add_para('', size=6)
    
    add_para(
        'The methodology for calculating provisions is based on a comprehensive risk assessment framework '
        'that considers the likelihood and impact of identified risks. Risk assessments are conducted '
        'at regular intervals, and provisions are adjusted based on changes in risk profiles and '
        'emerging risk factors.',
        size=11
    )
    add_para('', size=6)
    
    add_para(
        'The provision calculation process involves the following steps:',
        bold=True, size=11
    )
    add_para('   1. Identification and assessment of residual risks', size=11)
    add_para('   2. Quantification of potential financial impact', size=11)
    add_para('   3. Application of probability adjustments based on risk likelihood', size=11)
    add_para('   4. Review and approval by management and risk committee', size=11)
    add_para('   5. Regular monitoring and adjustment of provisions', size=11)
    
    # ========== COMMITMENTS SECTION ==========
    doc.add_page_break()
    add_heading('5. Commitments and Obligations', level=1)
    add_para('', size=12)
    
    add_para(
        'Commitments represent expected future costs associated with action plans designed to mitigate '
        'identified risks and improve operational effectiveness. These commitments are based on approved '
        'action plans and reflect management\'s best estimate of the resources required to implement '
        'risk mitigation measures.',
        size=11
    )
    add_para('', size=6)
    
    add_heading('5.1 Action Plans Overview', level=2)
    add_para('', size=6)
    
    add_para(
        f'The total expected costs associated with action plans as of {end_date} amount to '
        f'{format_currency(total_expected_cost)}. These costs represent commitments for implementing '
        f'risk mitigation measures and improving control effectiveness.',
        size=11
    )
    add_para('', size=12)
    
    # Actual data table for commitments
    commitments_table = add_styled_table(
        ['Commitment Type', 'Expected Cost', 'Implementation Date', 'Status'],
        [
            ('Action Plan Commitments', format_currency(total_expected_cost), end_date, 'Active'),
        ]
    )
    
    add_para('', size=12)
    
    add_heading('5.2 Expected Costs Analysis', level=2)
    add_para('', size=6)
    
    add_para(
        'The expected costs for action plans are estimated based on the scope and complexity of '
        'each action plan, resource requirements, and implementation timelines. These estimates '
        'are reviewed and updated regularly to reflect changes in project scope and market conditions.',
        size=11
    )
    
    # ========== FINANCIAL IMPACT ASSESSMENT ==========
    doc.add_page_break()
    add_heading('6. Financial Impact Assessment', level=1)
    add_para('', size=12)
    
    add_para(
        'This section provides a comprehensive assessment of the financial impact of operational losses, '
        'provisions, and commitments on the organization\'s financial position and performance.',
        size=11
    )
    add_para('', size=6)
    
    add_heading('6.1 Aggregate Financial Impact', level=2)
    add_para('', size=6)
    
    total_impact = total_net_loss + total_residual_financial + total_expected_cost
    
    add_para(
        f'The aggregate financial impact of operational losses, provisions, and commitments for the '
        f'reporting period amounts to {format_currency(total_impact)}. This represents the combined '
        f'effect of:',
        size=11
    )
    add_para('', size=6)
    
    impact_table = add_styled_table(
        ['Component', 'Amount', 'Percentage'],
        [
            ('Operational Losses (Net)', format_currency(total_net_loss), 
             f'{(total_net_loss/total_impact*100):.1f}%' if total_impact > 0 else '0.0%'),
            ('Residual Risk Provisions', format_currency(total_residual_financial),
             f'{(total_residual_financial/total_impact*100):.1f}%' if total_impact > 0 else '0.0%'),
            ('Action Plan Commitments', format_currency(total_expected_cost),
             f'{(total_expected_cost/total_impact*100):.1f}%' if total_impact > 0 else '0.0%'),
            ('Total Impact', format_currency(total_impact), '100.0%'),
        ]
    )
    
    add_para('', size=12)
    
    add_heading('6.2 Financial Impact Analysis', level=2)
    add_para('', size=6)
    
    add_para(
        'The financial impact assessment considers both direct and indirect effects on the organization\'s '
        'financial position. Direct impacts include immediate losses and provisions, while indirect impacts '
        'may include reputational effects, regulatory implications, and opportunity costs.',
        size=11
    )
    add_para('', size=6)
    
    add_para(
        'Management has reviewed the financial impact and has implemented measures to mitigate risks and '
        'improve operational effectiveness. Ongoing monitoring and reporting processes ensure that financial '
        'impacts are identified, assessed, and managed effectively.',
        size=11
    )
    
    # ========== RISK MANAGEMENT FRAMEWORK ==========
    doc.add_page_break()
    add_heading('7. Risk Management Framework', level=1)
    add_para('', size=12)
    
    add_para(
        'The organization maintains a comprehensive risk management framework designed to identify, '
        'assess, monitor, and mitigate risks across all operational areas. This framework is aligned '
        'with international best practices and regulatory requirements.',
        size=11
    )
    add_para('', size=6)
    
    add_heading('7.1 Risk Management Process', level=2)
    add_para('', size=6)
    
    add_para(
        'The risk management process encompasses the following key components:',
        bold=True, size=11
    )
    add_para('', size=6)
    
    add_para('1. Risk Identification: Systematic identification of risks across all operational areas', size=11)
    add_para('2. Risk Assessment: Evaluation of risk likelihood and impact using standardized methodologies', size=11)
    add_para('3. Risk Mitigation: Development and implementation of action plans to address identified risks', size=11)
    add_para('4. Risk Monitoring: Ongoing monitoring of risk levels and effectiveness of mitigation measures', size=11)
    add_para('5. Risk Reporting: Regular reporting to management and governing bodies on risk status', size=11)
    
    add_para('', size=12)
    
    add_heading('7.2 Governance Structure', level=2)
    add_para('', size=6)
    
    add_para(
        'Risk management activities are governed by a dedicated risk management function with clear '
        'reporting lines to senior management and the board of directors. The governance structure '
        'ensures appropriate oversight and accountability for risk management activities.',
        size=11
    )
    
    # ========== COMPLIANCE AND GOVERNANCE ==========
    doc.add_page_break()
    add_heading('8. Compliance and Governance', level=1)
    add_para('', size=12)
    
    add_para(
        'This section addresses compliance with applicable regulations, standards, and governance '
        'requirements relevant to GRC financial reporting.',
        size=11
    )
    add_para('', size=6)
    
    add_heading('8.1 Regulatory Compliance', level=2)
    add_para('', size=6)
    
    add_para(
        'The organization is committed to maintaining compliance with all applicable regulatory '
        'requirements, including financial reporting standards, risk management regulations, and '
        'governance codes. Compliance activities are monitored and reported regularly to ensure '
        'ongoing adherence to requirements.',
        size=11
    )
    add_para('', size=6)
    
    add_para(
        'Key regulatory frameworks applicable to this report include:',
        bold=True, size=11
    )
    add_para('   • International Financial Reporting Standards (IFRS)', size=11)
    add_para('   • Egyptian Accounting Standards (EAS)', size=11)
    add_para('   • XBRL reporting requirements', size=11)
    add_para('   • Governance, Risk, and Compliance (GRC) standards', size=11)
    
    # ========== RECOMMENDATIONS ==========
    doc.add_page_break()
    add_heading('9. Recommendations', level=1)
    add_para('', size=12)
    
    add_para(
        'Based on the analysis presented in this report, the following recommendations are provided '
        'for management consideration:',
        size=11
    )
    add_para('', size=6)
    
    recommendations = [
        'Enhance operational loss prevention measures through improved controls and monitoring',
        'Strengthen risk assessment processes to identify emerging risks proactively',
        'Optimize provision levels based on updated risk assessments and historical trends',
        'Accelerate implementation of high-priority action plans to mitigate key risks',
        'Improve recovery processes to maximize recovery amounts from operational losses',
        'Enhance reporting and analytics capabilities to support better decision-making',
        'Strengthen governance and oversight of risk management activities',
        'Invest in training and development to improve risk awareness and control effectiveness',
    ]
    
    for i, rec in enumerate(recommendations, 1):
        add_para(f'{i}. {rec}', size=11)
        add_para('', size=3)
    
    # ========== APPENDICES ==========
    doc.add_page_break()
    add_heading('10. Appendices', level=1)
    add_para('', size=12)
    
    add_heading('Appendix A: Glossary of Terms', level=2)
    add_para('', size=6)
    
    glossary_terms = [
        ('Operational Loss', 'Financial losses resulting from operational incidents, including fraud, errors, system failures, and external events.'),
        ('Residual Risk', 'The level of risk remaining after implementation of risk mitigation measures and controls.'),
        ('Provision', 'Amounts set aside in the financial statements to cover potential future liabilities arising from identified risks.'),
        ('Commitment', 'Expected future costs associated with approved action plans and risk mitigation initiatives.'),
        ('Recovery', 'Amounts recovered from operational losses through insurance claims, legal actions, or other recovery mechanisms.'),
        ('Financial Impact', 'The effect of operational losses, provisions, and commitments on the organization\'s financial position and performance.'),
        ('Risk Assessment', 'The process of identifying, analyzing, and evaluating risks to determine their likelihood and potential impact.'),
        ('Compliance', 'Adherence to applicable laws, regulations, standards, and internal policies.'),
        ('Governance', 'The framework of rules, practices, and processes by which an organization is directed and controlled.'),
        ('Control', 'Measures designed to prevent or detect errors, fraud, and other operational risks.'),
    ]
    
    for i, (term, definition) in enumerate(glossary_terms, 1):
        add_para(f'{i}. {term}:', bold=True, size=11)
        add_para(f'   {definition}', size=11)
        add_para('', size=6)
    
    doc.add_page_break()
    add_heading('Appendix B: Methodology', level=2)
    add_para('', size=6)
    
    add_para(
        'This report has been prepared using standardized methodologies for data collection, analysis, '
        'and reporting. The following methodologies have been applied:',
        size=11
    )
    add_para('', size=6)
    
    methodologies = [
        ('Data Collection', 'Financial data has been collected from the GRC database system, including '
         'incidents, residual risks, and action plans. Data collection processes ensure completeness, '
         'accuracy, and timeliness of information.'),
        ('Loss Calculation', 'Operational losses are calculated based on net loss amounts recorded in the '
         'incidents database. Net losses represent total losses less recovery amounts.'),
        ('Provision Estimation', 'Provisions are estimated based on residual risk assessments conducted '
         'throughout the reporting period. Estimation methodologies consider risk likelihood, impact, '
         'and historical trends.'),
        ('Commitment Assessment', 'Expected costs for action plans are assessed based on approved action '
         'plan budgets and implementation timelines. Costs are reviewed and updated regularly.'),
    ]
    
    for i, (method, description) in enumerate(methodologies, 1):
        add_para(f'B.{i} {method}:', bold=True, size=12)
        add_para('', size=3)
        add_para(description, size=11)
        add_para('', size=12)
    
    doc.add_page_break()
    add_heading('Appendix C: Data Sources', level=2)
    add_para('', size=6)
    
    add_para(
        'The data presented in this report has been sourced from the following systems and databases:',
        size=11
    )
    add_para('', size=6)
    
    data_sources = [
        ('Incidents Database', 'Operational loss data, including net loss, total loss, and recovery amounts.'),
        ('Residual Risks Database', 'Residual risk assessments and financial provision calculations.'),
        ('Action Plans Database', 'Action plan details, expected costs, and implementation timelines.'),
        ('Financial Systems', 'Supporting financial data and accounting records.'),
    ]
    
    for i, (source, description) in enumerate(data_sources, 1):
        add_para(f'C.{i} {source}:', bold=True, size=12)
        add_para(description, size=11)
        add_para('', size=6)
    
    doc.add_page_break()
    add_heading('Appendix D: Reference Documents', level=2)
    add_para('', size=6)
    
    add_para(
        'The following reference documents and standards have been consulted in preparing this report:',
        size=11
    )
    add_para('', size=6)
    
    references = [
        ('IFRS Standards', 'International Financial Reporting Standards applicable to financial reporting.'),
        ('EAS Standards', 'Egyptian Accounting Standards for local regulatory compliance.'),
        ('XBRL Taxonomy', 'XBRL taxonomy documentation for structured financial reporting.'),
        ('GRC Framework', 'Internal governance, risk, and compliance framework documentation.'),
        ('Risk Management Policy', 'Organization\'s risk management policy and procedures.'),
    ]
    
    for i, (ref, description) in enumerate(references, 1):
        add_para(f'D.{i} {ref}:', bold=True, size=12)
        add_para(description, size=11)
        add_para('', size=6)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()

def generate_pdf_report(columns, data_rows, header_config=None):
    """Generate PDF report from dynamic data with full header configuration support
    Returns bytes - does NOT save to disk (file saving is handled by save_and_log_export in the route)
    """
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
    from io import BytesIO
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    import os
    import base64
    from io import BytesIO
    try:
        from PIL import Image as PILImage
    except Exception:
        PILImage = None
    
    # Import Arabic text support
    try:
        from utils.pdf_utils import shape_text_for_arabic, ARABIC_FONT_NAME, DEFAULT_FONT_NAME
    except ImportError:
        # Fallback if pdf_utils is not available
        def shape_text_for_arabic(text: str) -> str:
            return text
        ARABIC_FONT_NAME = None
        DEFAULT_FONT_NAME = 'Helvetica'  # ReportLab built-in font
    
    # Get default header config if none provided
    if not header_config:
        from utils.export_utils import get_default_header_config
        header_config = get_default_header_config("dynamic")
    
    base_dir = os.path.dirname(os.path.dirname(__file__))
    # Create PDF in memory (BytesIO) - NOT to disk (file saving is handled by save_and_log_export in the route)
    # This prevents duplicate file saves when generate_pdf_report is called from risk/control/incident/kri routes
    buffer = BytesIO()
    
    # Extract ALL configuration values from header modal configuration
    # Basic report settings
    include_header = header_config.get("includeHeader", True)
    title = header_config.get("title", "Dynamic Report")
    subtitle = header_config.get("subtitle", "")
    icon = header_config.get("icon", "chart-line")
    
    # Date and time settings
    show_date = header_config.get("showDate", True)
    footer_show_date = header_config.get("footerShowDate", True)
    
    # Bank information settings
    show_bank_info = header_config.get("showBankInfo", True)
    bank_info_location = header_config.get("bankInfoLocation", "top")  # top, bottom, none
    bank_info_align = (
        header_config.get("bankInfoAlign")
        or header_config.get("logoPosition", "center")
        or "center"
    ).lower()
    bank_name = header_config.get("bankName", "")
    bank_address = header_config.get("bankAddress", "")
    bank_phone = header_config.get("bankPhone", "")
    bank_website = header_config.get("bankWebsite", "")
    
    # Logo settings
    show_logo = header_config.get("showLogo", True)
    logo_base64 = header_config.get("logoBase64", "")
    logo_position = header_config.get("logoPosition", "left")
    logo_height = header_config.get("logoHeight", 36)
    logo_file = header_config.get("logoFile", None)
    
    # Color and styling settings
    font_color = header_config.get("fontColor", "#1F4E79")
    table_header_bg_color = header_config.get("tableHeaderBgColor", "#1F4E79")
    table_body_bg_color = header_config.get("tableBodyBgColor", "#FFFFFF")
    background_color = header_config.get("backgroundColor", "#FFFFFF")
    border_style = header_config.get("borderStyle", "solid")
    border_color = header_config.get("borderColor", "#E5E7EB")
    border_width = header_config.get("borderWidth", 1)
    
    # Font and size settings
    font_size = header_config.get("fontSize", "medium")
    padding = header_config.get("padding", 20)
    margin = header_config.get("margin", 72)  # 1 inch = 72 points
    
    # Watermark settings
    watermark_enabled = header_config.get("watermarkEnabled", False)
    watermark_text = header_config.get("watermarkText", "CONFIDENTIAL")
    watermark_opacity = header_config.get("watermarkOpacity", 10)
    watermark_diagonal = header_config.get("watermarkDiagonal", True)
    
    # Footer settings
    footer_show_confidentiality = header_config.get("footerShowConfidentiality", True)
    footer_confidentiality_text = header_config.get("footerConfidentialityText", "Confidential Report - Internal Use Only")
    footer_show_page_numbers = header_config.get("footerShowPageNumbers", True)
    footer_align = header_config.get("footerAlign", "center")
    
    # Page settings
    show_page_numbers = header_config.get("showPageNumbers", True)
    location = header_config.get("location", "top")
    
    # Convert hex colors to ReportLab colors
    def hex_to_color(hex_color):
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        try:
            r = int(hex_color[0:2], 16) / 255.0
            g = int(hex_color[2:4], 16) / 255.0
            b = int(hex_color[4:6], 16) / 255.0
            return colors.Color(r, g, b)
        except:
            return colors.HexColor(f"#{hex_color}")
    
    font_color_rl = hex_to_color(font_color)
    header_bg_color_rl = hex_to_color(table_header_bg_color)
    body_bg_color_rl = hex_to_color(table_body_bg_color)
    background_color_rl = hex_to_color(background_color)
    border_color_rl = hex_to_color(border_color)
    
    # Choose page size based on number of columns
    page_size = A4 if len(columns) <= 6 else letter
    
    # Create document with margins (using BytesIO buffer, not file path)
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=page_size,
        rightMargin=margin,
        leftMargin=margin,
        topMargin=margin,
        bottomMargin=margin
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Create custom styles based on configuration
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=16,
        textColor=font_color_rl,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=font_color_rl,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    # Resolve text alignment for bank info
    _pdf_align = TA_CENTER
    if bank_info_align == 'left':
        _pdf_align = TA_LEFT
    elif bank_info_align == 'right':
        _pdf_align = TA_RIGHT

    bank_style = ParagraphStyle(
        'BankInfo',
        parent=styles['Normal'],
        fontSize=10,
        textColor=font_color_rl,
        alignment=_pdf_align,
        spaceAfter=3
    )
    
    date_style = ParagraphStyle(
        'DateInfo',
        parent=styles['Normal'],
        fontSize=10,
        textColor=font_color_rl,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    # Only add header if includeHeader is True
    if include_header:
        # Add logo image if configured
        if show_logo and (logo_base64 or logo_file):
            try:
                img_stream = None
                if logo_base64:
                    b64_data = logo_base64
                    if ',' in b64_data:
                        b64_data = b64_data.split(',')[1]
                    img_bytes = base64.b64decode(b64_data)
                    img_stream = BytesIO(img_bytes)
                elif logo_file and os.path.exists(logo_file):
                    with open(logo_file, 'rb') as lf:
                        img_stream = BytesIO(lf.read())
                if img_stream:
                    width_arg = None
                    height_arg = None
                    if PILImage is not None:
                        img_stream.seek(0)
                        pil_img = PILImage.open(img_stream)
                        orig_w, orig_h = pil_img.size
                        if orig_h > 0:
                            scale = float(logo_height) / float(orig_h)
                            width_arg = orig_w * scale
                            height_arg = logo_height
                        img_stream.seek(0)
                    rl_img = RLImage(img_stream, width=width_arg, height=height_arg)
                    if logo_position == 'left':
                        rl_img.hAlign = 'LEFT'
                    elif logo_position == 'right':
                        rl_img.hAlign = 'RIGHT'
                    else:
                        rl_img.hAlign = 'CENTER'
                    story.append(rl_img)
                    story.append(Spacer(1, 6))
            except Exception:
                pass
        # Add bank information at top if configured
        if show_bank_info and bank_name and bank_info_location == "top":
            story.append(Paragraph(f"🏦 {bank_name}", bank_style))
            if bank_address:
                story.append(Paragraph(bank_address, bank_style))
            if bank_phone or bank_website:
                contact_info = []
                if bank_phone:
                    contact_info.append(f"Tel: {bank_phone}")
                if bank_website:
                    contact_info.append(f"Web: {bank_website}")
                story.append(Paragraph(" | ".join(contact_info), bank_style))
            story.append(Spacer(1, 12))
        
        # Add title
        story.append(Paragraph(title, title_style))
        
        # Add subtitle
        if subtitle:
            story.append(Paragraph(subtitle, subtitle_style))
        
        # Add generation date
        if show_date:
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            story.append(Paragraph(f"Generated on: {current_date}", date_style))
        
        story.append(Spacer(1, 20))
    
    # Generate chart if chart_data is provided
    chart_data = header_config.get('chart_data')
    chart_type = header_config.get('chart_type', 'bar')
    if chart_data and chart_data.get('labels') and chart_data.get('values'):
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            from io import BytesIO
            
            # Create chart
            fig, ax = plt.subplots(figsize=(8, 5))
            
            if chart_type == 'bar':
                ax.barh(chart_data['labels'], chart_data['values'], color='#4472C4')
                ax.set_xlabel('Controls Count')
                ax.set_ylabel('Component')
            elif chart_type == 'pie':
                ax.pie(chart_data['values'], labels=chart_data['labels'], autopct='%1.1f%%')
            
            ax.set_title(title if title else 'Chart')
            plt.tight_layout()
            
            # Save chart to buffer
            chart_buffer = BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=150, bbox_inches='tight')
            chart_buffer.seek(0)
            plt.close()
            
            # Add chart to PDF
            chart_img = RLImage(chart_buffer, width=6*inch, height=3.5*inch)
            chart_img.hAlign = 'CENTER'
            story.append(chart_img)
            story.append(Spacer(1, 20))
        except Exception as e:
            print(f"Error generating chart: {e}")
            pass
    
    # Table data with Arabic text support and multi-line text handling
    # Create styles for table cells
    styles = getSampleStyleSheet()
    
    # Header style
    header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontSize=12,
        fontName=ARABIC_FONT_NAME or (DEFAULT_FONT_NAME + '-Bold'),
        alignment=TA_CENTER,
        textColor=colors.whitesmoke,
        spaceAfter=6,
        spaceBefore=6
    )
    
    # Data cell style
    data_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontSize=10,
        fontName=ARABIC_FONT_NAME or DEFAULT_FONT_NAME,
        alignment=TA_CENTER,
        spaceAfter=4,
        spaceBefore=4,
        leading=12
    )
    
    # Process columns with Paragraph objects for multi-line support
    processed_columns = [Paragraph(shape_text_for_arabic(str(col)), header_style) for col in columns]
    
    # Process data rows with Paragraph objects for multi-line support
    processed_data_rows = []
    for row in data_rows:
        processed_row = [Paragraph(shape_text_for_arabic(str(cell)), data_style) for cell in row]
        processed_data_rows.append(processed_row)
    
    table_data = [processed_columns] + processed_data_rows

    # Optional footer totals row for PDF
    footer_totals_cols = header_config.get("tableFooterTotals", []) or []
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        name_to_index = {str(col): idx for idx, col in enumerate(columns)}
        totals_row = [Paragraph("", data_style)] * len(columns)
        totals_row[0] = Paragraph(shape_text_for_arabic("Total"), data_style)
        for col_name in footer_totals_cols:
            if str(col_name) in name_to_index:
                idx = name_to_index[str(col_name)]
                s = 0.0
                for r in data_rows:
                    try:
                        val = r[idx]
                        if val is None or val == "":
                            continue
                        s += float(str(val).replace(',', ''))
                    except Exception:
                        pass
                totals_row[idx] = Paragraph(shape_text_for_arabic(f"{s:,.2f}"), data_style)
        table_data.append(totals_row)
    
    # Create table with configuration-based styling and column widths
    # Calculate column widths to leave some margin
    num_cols = len(table_data[0]) if table_data else 1
    available_width = page_size[0] - (margin * 2)  # Subtract left and right margins
    col_width = available_width / num_cols if num_cols > 0 else available_width
    col_widths = [col_width] * num_cols
    
    table = Table(table_data, repeatRows=1, colWidths=col_widths)
    
    # Build table style based on configuration with Arabic font support
    table_style = [
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), header_bg_color_rl),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), ARABIC_FONT_NAME or (DEFAULT_FONT_NAME + '-Bold')),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        
        # Data rows styling
        ('BACKGROUND', (0, 1), (-1, -1), body_bg_color_rl),
        ('FONTNAME', (0, 1), (-1, -1), ARABIC_FONT_NAME or DEFAULT_FONT_NAME),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        
        # Grid lines
        ('GRID', (0, 0), (-1, -1), border_width, border_color_rl),
        
        # Alternating row colors if needed
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
    ]

    # If totals row exists, style it like a bold footer
    if isinstance(footer_totals_cols, list) and len(footer_totals_cols) > 0:
        last_row_index = len(table_data) - 1
        table_style += [
            ('BACKGROUND', (0, last_row_index), (-1, last_row_index), header_bg_color_rl),
            ('TEXTCOLOR', (0, last_row_index), (-1, last_row_index), colors.whitesmoke),
            ('FONTNAME', (0, last_row_index), (-1, last_row_index), DEFAULT_FONT_NAME + '-Bold'),
        ]
    
    table.setStyle(TableStyle(table_style))
    story.append(table)
    
    # Add watermark if enabled
    if watermark_enabled:
        try:
            from utils.export_utils import add_watermark_to_pdf
            add_watermark_to_pdf(story, header_config)
        except Exception as e:
            pass  # Continue without watermark if there's an error
    
    # Add bank information at bottom if configured
    if show_bank_info and bank_name and bank_info_location == "bottom":
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"🏦 {bank_name}", bank_style))
        if bank_address:
            story.append(Paragraph(bank_address, bank_style))
        if bank_phone or bank_website:
            contact_info = []
            if bank_phone:
                contact_info.append(f"Tel: {bank_phone}")
            if bank_website:
                contact_info.append(f"Web: {bank_website}")
            story.append(Paragraph(" | ".join(contact_info), bank_style))

    # Add footer elements
    if footer_show_date or footer_show_confidentiality or footer_show_page_numbers:
        story.append(Spacer(1, 20))
        
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        
        footer_text = []
        if footer_show_date:
            current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            footer_text.append(f"Generated on: {current_date}")
        if footer_show_confidentiality:
            footer_text.append(footer_confidentiality_text)
        if footer_show_page_numbers:
            footer_text.append("Page &P of &N")
        
        if footer_text:
            story.append(Paragraph(" | ".join(footer_text), footer_style))
    
    # Build PDF and return bytes (NOT file path - file saving and DB logging handled by save_and_log_export)
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
